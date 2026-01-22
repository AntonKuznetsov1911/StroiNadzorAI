"""
Telegram –±–æ—Ç –°—Ç—Ä–æ–π–ù–∞–¥–∑–æ—ÄAI - AI –∫–æ–Ω—Å—É–ª—å—Ç–∞–Ω—Ç –ø–æ —Å—Ç—Ä–æ–∏—Ç–µ–ª—å–Ω—ã–º –Ω–æ—Ä–º–∞—Ç–∏–≤–∞–º
–° –ø–æ–¥–¥–µ—Ä–∂–∫–æ–π —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–æ–≥–æ –∞–Ω–∞–ª–∏–∑–∞ —Ñ–æ—Ç–æ–≥—Ä–∞—Ñ–∏–π
"""

import os
import logging
import base64
import json
import re
from io import BytesIO
from datetime import datetime, timedelta
from collections import defaultdict, Counter
from pathlib import Path
from dotenv import load_dotenv

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, WebAppInfo, BotCommand, ReplyKeyboardRemove, ReplyKeyboardMarkup, KeyboardButton
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    CallbackQueryHandler,
    ConversationHandler,
    ContextTypes,
    filters
)
import asyncio

# –ó–∞–≥—Ä—É–∑–∫–∞ –ø–µ—Ä–µ–º–µ–Ω–Ω—ã—Ö –æ–∫—Ä—É–∂–µ–Ω–∏—è
load_dotenv()

# –ù–∞—Å—Ç—Ä–æ–π–∫–∞ –ª–æ–≥–∏—Ä–æ–≤–∞–Ω–∏—è
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO,
    handlers=[
        logging.FileHandler('bot.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# –ò–º–ø–æ—Ä—Ç –±–∞–∑—ã –∞–∫—Ç—É–∞–ª—å–Ω—ã—Ö –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤ 2025
try:
    from regulations_2025 import (
        FEDERAL_LAWS,
        MANDATORY_PROCEDURES_2025,
        SRO_REQUIREMENTS,
        TIM_BIM_REQUIREMENTS,
        PRICING_2025,
        INDUSTRIAL_CONSTRUCTION,
        CIVIL_CONSTRUCTION,
        COMMERCIAL_CONSTRUCTION,
        KEY_REGULATIONS_2025,
        DAILY_CHECKLIST,
        TRENDS_2025_2027,
        get_all_regulations,
        search_regulation
    )
    REGULATIONS_2025_AVAILABLE = True
    logger.info("‚úÖ –ë–∞–∑–∞ –∞–∫—Ç—É–∞–ª—å–Ω—ã—Ö –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤ 2025 –∑–∞–≥—Ä—É–∂–µ–Ω–∞")
except ImportError:
    REGULATIONS_2025_AVAILABLE = False
    logger.warning("‚ö†Ô∏è –§–∞–π–ª regulations_2025.py –Ω–µ –Ω–∞–π–¥–µ–Ω")

# –ò–º–ø–æ—Ä—Ç –ø—Ä–∞–∫—Ç–∏—á–µ—Å–∫–∏—Ö –∑–Ω–∞–Ω–∏–π 2025
try:
    from practical_knowledge_2025 import (
        HSE_REQUIREMENTS,
        CONSTRUCTION_TECHNOLOGY,
        ESTIMATING_FINANCE,
        LEGAL_ISSUES,
        PROJECT_MANAGEMENT,
        get_all_practical_knowledge,
        search_practical
    )
    PRACTICAL_KNOWLEDGE_AVAILABLE = True
    logger.info("‚úÖ –ë–∞–∑–∞ –ø—Ä–∞–∫—Ç–∏—á–µ—Å–∫–∏—Ö –∑–Ω–∞–Ω–∏–π 2025 –∑–∞–≥—Ä—É–∂–µ–Ω–∞")
except ImportError:
    PRACTICAL_KNOWLEDGE_AVAILABLE = False
    logger.warning("‚ö†Ô∏è –§–∞–π–ª practical_knowledge_2025.py –Ω–µ –Ω–∞–π–¥–µ–Ω")

# –ò–º–ø–æ—Ä—Ç —Ä–∞—Å—à–∏—Ä–µ–Ω–Ω—ã—Ö –ø—Ä–∞–∫—Ç–∏—á–µ—Å–∫–∏—Ö –∑–Ω–∞–Ω–∏–π 2025
try:
    from practical_knowledge_advanced_2025 import (
        MIGRATION_LAW,
        GEODESY,
        LOGISTICS,
        ECOLOGY,
        SPECIAL_CONDITIONS,
        ENGINEERING_NETWORKS,
        get_all_advanced_knowledge,
        search_advanced
    )
    ADVANCED_KNOWLEDGE_AVAILABLE = True
    logger.info("‚úÖ –†–∞—Å—à–∏—Ä–µ–Ω–Ω–∞—è –±–∞–∑–∞ –∑–Ω–∞–Ω–∏–π 2025 –∑–∞–≥—Ä—É–∂–µ–Ω–∞ (–∫–∞–¥—Ä—ã, –≥–µ–æ–¥–µ–∑–∏—è, –ª–æ–≥–∏—Å—Ç–∏–∫–∞, —ç–∫–æ–ª–æ–≥–∏—è, —Å–µ—Ç–∏)")
except ImportError:
    ADVANCED_KNOWLEDGE_AVAILABLE = False
    logger.warning("‚ö†Ô∏è –§–∞–π–ª practical_knowledge_advanced_2025.py –Ω–µ –Ω–∞–π–¥–µ–Ω")

# ============================================================================
# –û–ü–¢–ò–ú–ò–ó–ê–¶–ò–Ø: –£–º–Ω—ã–π –≤—ã–±–æ—Ä –º–æ–¥–µ–ª–µ–π AI
# ============================================================================
try:
    from model_selector import ModelSelector, should_use_web_search, extract_regulation_codes
    MODEL_SELECTOR_AVAILABLE = True
    logger.info("‚úÖ ModelSelector –∑–∞–≥—Ä—É–∂–µ–Ω - —É–º–Ω—ã–π –≤—ã–±–æ—Ä AI –º–æ–¥–µ–ª–µ–π –∞–∫—Ç–∏–≤–µ–Ω")
except ImportError:
    MODEL_SELECTOR_AVAILABLE = False
    logger.warning("‚ö†Ô∏è model_selector.py –Ω–µ –Ω–∞–π–¥–µ–Ω - –∏—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è —Ç–æ–ª—å–∫–æ Grok")

try:
    from optimized_prompts import (
        CLAUDE_SYSTEM_PROMPT_TECHNICAL,
        CLAUDE_DALLE_PROMPT_CREATOR,
        GROK_SYSTEM_PROMPT_GENERAL,
        GEMINI_VISION_PROMPT_DEFECTS
    )
    OPTIMIZED_PROMPTS_AVAILABLE = True
    logger.info("‚úÖ –û–ø—Ç–∏–º–∏–∑–∏—Ä–æ–≤–∞–Ω–Ω—ã–µ –ø—Ä–æ–º–ø—Ç—ã –∑–∞–≥—Ä—É–∂–µ–Ω—ã")
except ImportError:
    OPTIMIZED_PROMPTS_AVAILABLE = False
    logger.warning("‚ö†Ô∏è optimized_prompts.py –Ω–µ –Ω–∞–π–¥–µ–Ω - –∏—Å–ø–æ–ª—å–∑—É—é—Ç—Å—è —Å—Ç–∞–Ω–¥–∞—Ä—Ç–Ω—ã–µ –ø—Ä–æ–º–ø—Ç—ã")

try:
    from optimized_handlers import (
        handle_with_claude_technical,
        handle_with_gemini_vision,
        handle_with_gemini_image,
        handle_with_grok
    )
    OPTIMIZED_HANDLERS_AVAILABLE = True
    logger.info("‚úÖ –û–ø—Ç–∏–º–∏–∑–∏—Ä–æ–≤–∞–Ω–Ω—ã–µ –æ–±—Ä–∞–±–æ—Ç—á–∏–∫–∏ –∑–∞–≥—Ä—É–∂–µ–Ω—ã (Claude, Gemini Vision, Gemini Image, Grok)")
except ImportError:
    OPTIMIZED_HANDLERS_AVAILABLE = False
    logger.warning("‚ö†Ô∏è optimized_handlers.py –Ω–µ –Ω–∞–π–¥–µ–Ω - –∏—Å–ø–æ–ª—å–∑—É—é—Ç—Å—è —Å—Ç–∞–Ω–¥–∞—Ä—Ç–Ω—ã–µ –æ–±—Ä–∞–±–æ—Ç—á–∏–∫–∏")

try:
    from smart_model_wrapper import smart_model_selection_text, smart_model_selection_photo
    SMART_WRAPPER_AVAILABLE = True
    logger.info("‚úÖ Smart wrapper –∑–∞–≥—Ä—É–∂–µ–Ω - —É–º–Ω—ã–π –≤—ã–±–æ—Ä –º–æ–¥–µ–ª–µ–π –∞–∫—Ç–∏–≤–µ–Ω!")
except ImportError:
    SMART_WRAPPER_AVAILABLE = False
    logger.warning("‚ö†Ô∏è smart_model_wrapper.py –Ω–µ –Ω–∞–π–¥–µ–Ω")

# –ò–º–ø–æ—Ä—Ç —Å–ø—Ä–∞–≤–æ—á–Ω–∏–∫–∞ —Å—Ç—Ä–æ–∏—Ç–µ–ª—è
try:
    from builder_reference import (
        get_builder_skills_context,
        search_builder_reference,
        is_builder_question,
        BUILDER_TOPICS
    )
    BUILDER_REFERENCE_AVAILABLE = True
    logger.info("‚úÖ –°–ø—Ä–∞–≤–æ—á–Ω–∏–∫ —Å—Ç—Ä–æ–∏—Ç–µ–ª—è –∑–∞–≥—Ä—É–∂–µ–Ω")
except ImportError:
    BUILDER_REFERENCE_AVAILABLE = False
    logger.warning("‚ö†Ô∏è –§–∞–π–ª builder_reference.py –Ω–µ –Ω–∞–π–¥–µ–Ω")

# PDF/Word —ç–∫—Å–ø–æ—Ä—Ç
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("ReportLab not available - PDF export disabled")

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available - Word export disabled")

# PostgreSQL Database
try:
    from database import (
        init_db,
        save_message,
        get_user_messages,
        search_messages as db_search_messages,
        get_user_tags,
        clear_user_history as db_clear_history,
        get_total_messages,
        close_db
    )
    DATABASE_AVAILABLE = True
    logger.info("‚úÖ PostgreSQL –º–æ–¥—É–ª—å –∑–∞–≥—Ä—É–∂–µ–Ω")
except ImportError:
    DATABASE_AVAILABLE = False
    logger.warning("‚ö†Ô∏è PostgreSQL –Ω–µ –¥–æ—Å—Ç—É–ø–µ–Ω, –∏—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è JSON")

# –£–ª—É—á—à–µ–Ω–∏—è v3.0
try:
    from improvements_v3 import (
        create_answer_buttons,
        create_quick_actions_menu,
        create_reply_suggestions_keyboard,
        create_calculators_menu,
        create_regulations_category_menu,
        create_region_selection_menu,
        create_related_questions_buttons,
        generate_smart_related_questions_prompt,
        parse_generated_questions,
        get_improved_help_text,
        REGULATIONS_CATEGORIES,
        CALCULATORS,
        REGIONS
    )
    IMPROVEMENTS_V3_AVAILABLE = True
    logger.info("‚úÖ –ú–æ–¥—É–ª—å —É–ª—É—á—à–µ–Ω–∏–π v3.0 –∑–∞–≥—Ä—É–∂–µ–Ω")
except ImportError:
    IMPROVEMENTS_V3_AVAILABLE = False
    logger.warning("‚ö†Ô∏è –ú–æ–¥—É–ª—å improvements_v3.py –Ω–µ –Ω–∞–π–¥–µ–Ω")

# –ö–∞–ª—å–∫—É–ª—è—Ç–æ—Ä—ã
try:
    from calculators import (
        calculate_concrete,
        calculate_reinforcement,
        calculate_formwork,
        calculate_electrical,
        calculate_water,
        calculate_winter_heating,
        calculate_brick,
        calculate_tile,
        calculate_paint,
        calculate_wall_area,
        calculate_roof,
        calculate_plaster,
        calculate_wallpaper,
        calculate_laminate,
        calculate_insulation,
        calculate_foundation,
        calculate_stairs,
        calculate_drywall,
        calculate_earthwork,
        calculate_labor,
        format_calculator_result
    )
    CALCULATORS_AVAILABLE = True
    logger.info("‚úÖ –ú–æ–¥—É–ª—å –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä–æ–≤ –∑–∞–≥—Ä—É–∂–µ–Ω (21 –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä)")
except ImportError:
    CALCULATORS_AVAILABLE = False
    logger.warning("‚ö†Ô∏è –ú–æ–¥—É–ª—å calculators.py –Ω–µ –Ω–∞–π–¥–µ–Ω")

# –ò–Ω—Ç–µ—Ä–∞–∫—Ç–∏–≤–Ω—ã–µ –æ–±—Ä–∞–±–æ—Ç—á–∏–∫–∏ –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä–æ–≤ v4.0
try:
    from calculator_handlers import (
        create_concrete_calculator_handler,
        create_rebar_calculator_handler,
        create_formwork_calculator_handler,
        create_electrical_calculator_handler,
        create_water_calculator_handler,
        create_winter_calculator_handler,
        create_math_calculator_handler,
        create_brick_calculator_handler,
        create_tile_calculator_handler,
        create_paint_calculator_handler,
        create_wall_area_calculator_handler,
        create_roof_calculator_handler,
        create_plaster_calculator_handler,
        create_wallpaper_calculator_handler,
        create_laminate_calculator_handler,
        create_insulation_calculator_handler,
        create_foundation_calculator_handler,
        create_stairs_calculator_handler,
        create_drywall_calculator_handler,
        create_earthwork_calculator_handler,
        create_labor_calculator_handler,
        quick_concrete,
        quick_math
    )
    CALCULATOR_HANDLERS_AVAILABLE = True
    logger.info("‚úÖ –ò–Ω—Ç–µ—Ä–∞–∫—Ç–∏–≤–Ω—ã–µ –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä—ã v4.0 (–≤—Å–µ 21) –∑–∞–≥—Ä—É–∂–µ–Ω—ã")
except ImportError:
    CALCULATOR_HANDLERS_AVAILABLE = False
    logger.warning("‚ö†Ô∏è –ú–æ–¥—É–ª—å calculator_handlers.py –Ω–µ –Ω–∞–π–¥–µ–Ω")

# –ò–Ω—Ç–µ—Ä–∞–∫—Ç–∏–≤–Ω—ã–µ –æ–±—Ä–∞–±–æ—Ç—á–∏–∫–∏ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ v1.0
try:
    from document_handlers import (
        create_acceptance_foundation_handler,
        create_complaint_contractor_handler,
        create_safety_plan_handler,
        create_hidden_works_act_handler,
    )
    DOCUMENT_HANDLERS_AVAILABLE = True
    logger.info("‚úÖ –ò–Ω—Ç–µ—Ä–∞–∫—Ç–∏–≤–Ω—ã–µ –æ–±—Ä–∞–±–æ—Ç—á–∏–∫–∏ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ v1.0 (–≤—Å–µ 4) –∑–∞–≥—Ä—É–∂–µ–Ω—ã")
except ImportError as e:
    DOCUMENT_HANDLERS_AVAILABLE = False
    logger.warning(f"‚ö†Ô∏è –ú–æ–¥—É–ª—å document_handlers.py –Ω–µ –Ω–∞–π–¥–µ–Ω: {e}")

# –ú–æ–¥—É–ª—å –≤–µ–±-–ø–æ–∏—Å–∫–∞ –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤
try:
    from web_search import perform_web_search, should_perform_web_search
    WEB_SEARCH_AVAILABLE = True
    logger.info("‚úÖ –ú–æ–¥—É–ª—å –≤–µ–±-–ø–æ–∏—Å–∫–∞ –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤ –∑–∞–≥—Ä—É–∂–µ–Ω (docs.cntd.ru, minstroyrf.gov.ru)")
except ImportError as e:
    WEB_SEARCH_AVAILABLE = False
    logger.warning(f"‚ö†Ô∏è –ú–æ–¥—É–ª—å web_search.py –Ω–µ –Ω–∞–π–¥–µ–Ω: {e}")

# –ú–æ–¥—É–ª—å –ø–æ–≥–æ–¥—ã (–Ø–Ω–¥–µ–∫—Å –ü–æ–≥–æ–¥–∞ API)
try:
    from weather import get_weather, is_weather_query
    WEATHER_AVAILABLE = True
    logger.info("‚úÖ –ú–æ–¥—É–ª—å –ø–æ–≥–æ–¥—ã –∑–∞–≥—Ä—É–∂–µ–Ω (–Ø–Ω–¥–µ–∫—Å –ü–æ–≥–æ–¥–∞ API)")
except ImportError as e:
    WEATHER_AVAILABLE = False
    logger.warning(f"‚ö†Ô∏è –ú–æ–¥—É–ª—å weather.py –Ω–µ –Ω–∞–π–¥–µ–Ω: {e}")

# –ú–æ–¥—É–ª—å –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π
try:
    from image_generator import (
        generate_construction_image,
        should_generate_image,
        format_generation_result
    )
    IMAGE_GENERATION_AVAILABLE = True
    logger.info("‚úÖ –ú–æ–¥—É–ª—å –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π –∑–∞–≥—Ä—É–∂–µ–Ω (Gemini AI)")
except ImportError as e:
    IMAGE_GENERATION_AVAILABLE = False
    logger.warning(f"‚ö†Ô∏è –ú–æ–¥—É–ª—å image_generator.py –Ω–µ –Ω–∞–π–¥–µ–Ω: {e}")

# Gemini Vision –¥–ª—è –∞–Ω–∞–ª–∏–∑–∞ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π v4.0
try:
    from gemini_vision import (
        initialize_gemini_vision,
        is_gemini_available,
        analyze_construction_image
    )
    gemini_vision_analyzer = initialize_gemini_vision()
    GEMINI_VISION_AVAILABLE = is_gemini_available()
    if GEMINI_VISION_AVAILABLE:
        logger.info("‚úÖ Gemini 2.5 Flash Vision –∑–∞–≥—Ä—É–∂–µ–Ω (–∞–Ω–∞–ª–∏–∑ —Ñ–æ—Ç–æ)")
    else:
        logger.warning("‚ö†Ô∏è Gemini Vision –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω (–Ω—É–∂–µ–Ω GEMINI_API_KEY)")
except ImportError as e:
    GEMINI_VISION_AVAILABLE = False
    gemini_vision_analyzer = None
    logger.warning(f"‚ö†Ô∏è –ú–æ–¥—É–ª—å gemini_vision.py –Ω–µ –Ω–∞–π–¥–µ–Ω: {e}")

# –†–µ–∂–∏–º —Ä–∞–∑—Ä–∞–±–æ—Ç—á–∏–∫–∞ v3.0 - –∞–≤—Ç–æ–≤—ã–±–æ—Ä –ª–æ–∫–∞–ª—å–Ω–æ–π/–æ–±–ª–∞—á–Ω–æ–π –≤–µ—Ä—Å–∏–∏
is_developer = None
try:
    # –°–Ω–∞—á–∞–ª–∞ –ø—Ä–æ–±—É–µ–º –∑–∞–≥—Ä—É–∑–∏—Ç—å –õ–û–ö–ê–õ–¨–ù–£–Æ –≤–µ—Ä—Å–∏—é (—Å git –∞–≤—Ç–æ–ø—É—à–µ–º)
    from dev_mode_local import create_dev_mode_handler, is_developer
    DEV_MODE_AVAILABLE = True
    logger.info("‚úÖ –†–µ–∂–∏–º —Ä–∞–∑—Ä–∞–±–æ—Ç—á–∏–∫–∞ v3.0 –õ–û–ö–ê–õ–¨–ù–´–ô –∑–∞–≥—Ä—É–∂–µ–Ω (—Å git –∞–≤—Ç–æ–ø—É—à–µ–º)")
except ImportError:
    try:
        # –ï—Å–ª–∏ –Ω–µ –ø–æ–ª—É—á–∏–ª–æ—Å—å (–Ω–µ—Ç git) - –∑–∞–≥—Ä—É–∂–∞–µ–º –û–ë–õ–ê–ß–ù–£–Æ –≤–µ—Ä—Å–∏—é (—Ç–æ–ª—å–∫–æ –∞–Ω–∞–ª–∏–∑)
        from dev_mode import create_dev_mode_handler, is_developer
        DEV_MODE_AVAILABLE = True
        logger.info("‚úÖ –†–µ–∂–∏–º —Ä–∞–∑—Ä–∞–±–æ—Ç—á–∏–∫–∞ v2.0 –û–ë–õ–ê–ß–ù–´–ô –∑–∞–≥—Ä—É–∂–µ–Ω (–±–µ–∑ git)")
    except ImportError as e:
        DEV_MODE_AVAILABLE = False
        is_developer = lambda user_id: False  # –ó–∞–≥–ª—É—à–∫–∞, –µ—Å–ª–∏ –º–æ–¥—É–ª–∏ –Ω–µ –∑–∞–≥—Ä—É–∂–µ–Ω—ã
        logger.warning(f"‚ö†Ô∏è –ú–æ–¥—É–ª–∏ dev_mode –Ω–µ –Ω–∞–π–¥–µ–Ω—ã: {e}")

# –ê–≤—Ç–æ–ø—Ä–∏–º–µ–Ω–µ–Ω–∏–µ –∏–∑–º–µ–Ω–µ–Ω–∏–π v1.0 - –∫–Ω–æ–ø–∫–∞ "–ü—Ä–∏–º–µ–Ω–∏—Ç—å –∏–∑–º–µ–Ω–µ–Ω–∏—è" –ø–æ—Å–ª–µ –æ—Ç–≤–µ—Ç–æ–≤
try:
    from auto_apply import add_apply_button, should_show_apply_button, handle_apply_changes
    AUTO_APPLY_AVAILABLE = True
    logger.info("‚úÖ –ê–≤—Ç–æ–ø—Ä–∏–º–µ–Ω–µ–Ω–∏–µ –∏–∑–º–µ–Ω–µ–Ω–∏–π v1.0 –∑–∞–≥—Ä—É–∂–µ–Ω–æ")
except ImportError as e:
    AUTO_APPLY_AVAILABLE = False
    logger.warning(f"‚ö†Ô∏è –ú–æ–¥—É–ª—å auto_apply.py –Ω–µ –Ω–∞–π–¥–µ–Ω: {e}")

# –ü—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è –ø–æ —É–ª—É—á—à–µ–Ω–∏—é v1.0
try:
    from suggestions import suggestions_menu, create_suggestions_handler
    SUGGESTIONS_AVAILABLE = True
    logger.info("‚úÖ –ú–æ–¥—É–ª—å –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π v1.0 –∑–∞–≥—Ä—É–∂–µ–Ω")
except ImportError as e:
    SUGGESTIONS_AVAILABLE = False
    logger.warning(f"‚ö†Ô∏è –ú–æ–¥—É–ª—å suggestions.py –Ω–µ –Ω–∞–π–¥–µ–Ω: {e}")

# –û–±—Ä–∞–±–æ—Ç—á–∏–∫ –≥–æ–ª–æ—Å–æ–≤—ã—Ö —Å–æ–æ–±—â–µ–Ω–∏–π v3.9
try:
    from voice_handler import process_voice_message
    VOICE_HANDLER_AVAILABLE = True
    logger.info("‚úÖ –û–±—Ä–∞–±–æ—Ç—á–∏–∫ –≥–æ–ª–æ—Å–æ–≤—ã—Ö —Å–æ–æ–±—â–µ–Ω–∏–π v3.9 –∑–∞–≥—Ä—É–∂–µ–Ω")
except ImportError:
    VOICE_HANDLER_AVAILABLE = False
    logger.warning("‚ö†Ô∏è –ú–æ–¥—É–ª—å voice_handler.py –Ω–µ –Ω–∞–π–¥–µ–Ω")

# –®–∞–±–ª–æ–Ω—ã –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ v3.9
try:
    from document_templates import DOCUMENT_TEMPLATES, generate_document
    TEMPLATES_AVAILABLE = True
    logger.info(f"‚úÖ –®–∞–±–ª–æ–Ω—ã –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ v3.9 –∑–∞–≥—Ä—É–∂–µ–Ω—ã ({len(DOCUMENT_TEMPLATES)} —à–∞–±–ª–æ–Ω–æ–≤)")
except ImportError:
    TEMPLATES_AVAILABLE = False
    logger.warning("‚ö†Ô∏è –ú–æ–¥—É–ª—å document_templates.py –Ω–µ –Ω–∞–π–¥–µ–Ω")

# –£–ø—Ä–∞–≤–ª–µ–Ω–∏–µ –ø—Ä–æ–µ–∫—Ç–∞–º–∏ v3.9
try:
    from project_manager import get_user_projects, create_project, load_project, Project
    PROJECTS_AVAILABLE = True
    logger.info("‚úÖ –£–ø—Ä–∞–≤–ª–µ–Ω–∏–µ –ø—Ä–æ–µ–∫—Ç–∞–º–∏ v3.9 –∑–∞–≥—Ä—É–∂–µ–Ω–æ")
except ImportError:
    PROJECTS_AVAILABLE = False
    logger.warning("‚ö†Ô∏è –ú–æ–¥—É–ª—å project_manager.py –Ω–µ –Ω–∞–π–¥–µ–Ω")

# –†–µ–∂–∏–º—ã —Ä–∞–±–æ—Ç—ã –ø–æ —Ä–æ–ª—è–º v3.2
try:
    from role_modes import (
        role_command,
        handle_role_selection,
        get_user_role,
        get_role_system_prompt
    )
    ROLES_AVAILABLE = True
    logger.info("‚úÖ –†–µ–∂–∏–º—ã —Ä–∞–±–æ—Ç—ã –ø–æ —Ä–æ–ª—è–º v3.2 –∑–∞–≥—Ä—É–∂–µ–Ω—ã")
except ImportError:
    ROLES_AVAILABLE = False
    logger.warning("‚ö†Ô∏è –ú–æ–¥—É–ª—å role_modes.py –Ω–µ –Ω–∞–π–¥–µ–Ω")

# –£–ø—Ä–∞–≤–ª–µ–Ω–∏–µ –∏—Å—Ç–æ—Ä–∏–µ–π v3.5
try:
    from history_manager import (
        history_command,
        stats_command,
        search_command,
        export_command,
        clear_history_command,
        handle_history_callback
    )
    HISTORY_MANAGER_AVAILABLE = True
    logger.info("‚úÖ –£–ø—Ä–∞–≤–ª–µ–Ω–∏–µ –∏—Å—Ç–æ—Ä–∏–µ–π v3.5 –∑–∞–≥—Ä—É–∂–µ–Ω–æ (–ø–æ–∏—Å–∫ + —ç–∫—Å–ø–æ—Ä—Ç)")
except ImportError:
    HISTORY_MANAGER_AVAILABLE = False
    logger.warning("‚ö†Ô∏è –ú–æ–¥—É–ª—å history_manager.py –Ω–µ –Ω–∞–π–¥–µ–Ω")

# –°–æ—Ö—Ä–∞–Ω—ë–Ω–Ω—ã–µ —Ä–∞—Å—á—ë—Ç—ã v3.6
try:
    from saved_calculations import (
        saved_command,
        handle_saved_callback
    )
    SAVED_CALCS_AVAILABLE = True
    logger.info("‚úÖ –°–æ—Ö—Ä–∞–Ω—ë–Ω–Ω—ã–µ —Ä–∞—Å—á—ë—Ç—ã v3.6 –∑–∞–≥—Ä—É–∂–µ–Ω—ã")
except ImportError:
    SAVED_CALCS_AVAILABLE = False
    logger.warning("‚ö†Ô∏è –ú–æ–¥—É–ª—å saved_calculations.py –Ω–µ –Ω–∞–π–¥–µ–Ω")

# –ë–∞–∑–∞ —á–∞—Å—Ç–æ –∑–∞–¥–∞–≤–∞–µ–º—ã—Ö –≤–æ–ø—Ä–æ—Å–æ–≤ (FAQ) v3.7
try:
    from faq import (
        faq_command,
        faq_search_command,
        handle_faq_callback,
        get_total_faq_count
    )
    FAQ_AVAILABLE = True
    logger.info("‚úÖ FAQ v3.7 –∑–∞–≥—Ä—É–∂–µ–Ω")
except ImportError:
    FAQ_AVAILABLE = False
    logger.warning("‚ö†Ô∏è –ú–æ–¥—É–ª—å faq.py –Ω–µ –Ω–∞–π–¥–µ–Ω")

# –ö—ç—à–∏—Ä–æ–≤–∞–Ω–∏–µ –æ—Ç–≤–µ—Ç–æ–≤ v3.8
try:
    from cache_manager import (
        init_cache,
        close_cache,
        get_cached_answer,
        set_cached_answer,
        find_similar_cached_question,
        get_cache_stats
    )
    CACHE_AVAILABLE = True
    logger.info("‚úÖ Cache manager v3.8 –∑–∞–≥—Ä—É–∂–µ–Ω")
except ImportError:
    CACHE_AVAILABLE = False
    logger.warning("‚ö†Ô∏è –ú–æ–¥—É–ª—å cache_manager.py –Ω–µ –Ω–∞–π–¥–µ–Ω")

# –ü–ª–∞–Ω–∏—Ä–æ–≤—â–∏–∫ —Ä–∞–±–æ—Ç v3.8
try:
    from work_planner import (
        planner_command,
        plan_calc_command,
        handle_planner_callback
    )
    PLANNER_AVAILABLE = True
    logger.info("‚úÖ Work planner v3.8 –∑–∞–≥—Ä—É–∂–µ–Ω")
except ImportError:
    PLANNER_AVAILABLE = False
    logger.warning("‚ö†Ô∏è –ú–æ–¥—É–ª—å work_planner.py –Ω–µ –Ω–∞–π–¥–µ–Ω")

# Gemini Image Generation (–Ω–æ–≤—ã–π API —Å —Ä–µ–∞–ª—å–Ω–æ–π –≥–µ–Ω–µ—Ä–∞—Ü–∏–µ–π –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π)
try:
    from gemini_image_gen import (
        initialize_gemini_generator,
        GeminiImageGenerator,
        generate_construction_image_gemini,
        is_image_generation_available
    )
    GEMINI_AVAILABLE = True
    logger.info("‚úÖ Gemini Image Generator –º–æ–¥—É–ª—å –∑–∞–≥—Ä—É–∂–µ–Ω (—Å –ø–æ–¥–¥–µ—Ä–∂–∫–æ–π –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏)")
except ImportError:
    GEMINI_AVAILABLE = False
    logger.warning("‚ö†Ô∏è –ú–æ–¥—É–ª—å gemini_image_gen.py –Ω–µ –Ω–∞–π–¥–µ–Ω")

# –û–ø—Ç–∏–º–∏–∑–∏—Ä–æ–≤–∞–Ω–Ω—ã–µ –ø—Ä–æ–º–ø—Ç—ã –∏ —Å–µ–ª–µ–∫—Ç–æ—Ä –º–æ–¥–µ–ª–µ–π v5.0
try:
    from optimized_prompts import (
        CLAUDE_SYSTEM_PROMPT_TECHNICAL,
        GROK_SYSTEM_PROMPT_GENERAL,
        GEMINI_IMAGE_PROMPT_SYSTEM,
        GEMINI_VISION_PROMPT_DEFECTS,
        WEB_SEARCH_DECISION_PROMPT
    )
    from model_selector import ModelSelector
    OPTIMIZED_PROMPTS_AVAILABLE = True
    model_selector = ModelSelector()
    logger.info("‚úÖ –û–ø—Ç–∏–º–∏–∑–∏—Ä–æ–≤–∞–Ω–Ω—ã–µ –ø—Ä–æ–º–ø—Ç—ã –∏ —Å–µ–ª–µ–∫—Ç–æ—Ä –º–æ–¥–µ–ª–µ–π v5.0 –∑–∞–≥—Ä—É–∂–µ–Ω—ã")
except ImportError as e:
    OPTIMIZED_PROMPTS_AVAILABLE = False
    model_selector = None
    logger.warning(f"‚ö†Ô∏è –ú–æ–¥—É–ª–∏ optimized_prompts/model_selector –Ω–µ –Ω–∞–π–¥–µ–Ω—ã: {e}")

# –ò–Ω—Ç–µ—Ä–∞–∫—Ç–∏–≤–Ω—ã–µ –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä—ã v4.0 - –£–î–ê–õ–Å–ù –¥—É–±–ª–∏—Ä—É—é—â–∏–π –∏–º–ø–æ—Ä—Ç
# –í—Å–µ –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä—ã –∏–º–ø–æ—Ä—Ç–∏—Ä—É—é—Ç—Å—è –∏–∑ calculator_handlers.py (—Å—Ç—Ä–æ–∫–∞ 257)
# interactive_calculators.py - —É—Å—Ç–∞—Ä–µ–≤—à–∏–π –º–æ–¥—É–ª—å, –Ω–µ –∏—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è
INTERACTIVE_CALCS_AVAILABLE = True  # –ö–∞–ª—å–∫—É–ª—è—Ç–æ—Ä—ã –¥–æ—Å—Ç—É–ø–Ω—ã —á–µ—Ä–µ–∑ calculator_handlers

# –ö–∞—Ç–µ–≥–æ—Ä–∏–∑–∞—Ü–∏—è –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤ v1.0
try:
    from regulations_categories import (
        get_categories_keyboard,
        get_regulations_by_category,
        search_regulations_by_keyword,
        get_all_regulations_text
    )
    REGULATIONS_CATEGORIES_AVAILABLE = True
    logger.info("‚úÖ –ö–∞—Ç–µ–≥–æ—Ä–∏–∑–∞—Ü–∏—è –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤ v1.0 –∑–∞–≥—Ä—É–∂–µ–Ω–∞")
except ImportError:
    REGULATIONS_CATEGORIES_AVAILABLE = False
    logger.warning("‚ö†Ô∏è –ú–æ–¥—É–ª—å regulations_categories.py –Ω–µ –Ω–∞–π–¥–µ–Ω")

# –ö–æ–Ω—Ç–µ–∫—Å—Ç–Ω—ã–µ –ø–æ–¥—Å–∫–∞–∑–∫–∏ v1.0
try:
    from context_hints import (
        is_short_question,
        get_context_hints,
        format_hint_response
    )
    CONTEXT_HINTS_AVAILABLE = True
    logger.info("‚úÖ –ö–æ–Ω—Ç–µ–∫—Å—Ç–Ω—ã–µ –ø–æ–¥—Å–∫–∞–∑–∫–∏ v1.0 –∑–∞–≥—Ä—É–∂–µ–Ω—ã")
except ImportError:
    CONTEXT_HINTS_AVAILABLE = False
    logger.warning("‚ö†Ô∏è –ú–æ–¥—É–ª—å context_hints.py –Ω–µ –Ω–∞–π–¥–µ–Ω")

# –ò–º–ø–æ—Ä—Ç xAI –∫–ª–∏–µ–Ω—Ç–∞
from xai_client import XAIClient, call_xai_with_retry

# –ò–º–ø–æ—Ä—Ç Gemini Live API (–≥–æ–ª–æ—Å–æ–≤–æ–π –∞—Å—Å–∏—Å—Ç–µ–Ω—Ç)
try:
    from gemini_live_bot_integration import start_voice_chat_command
    VOICE_ASSISTANT_AVAILABLE = True
    logger.info("‚úÖ Gemini Live API (–≥–æ–ª–æ—Å–æ–≤–æ–π –∞—Å—Å–∏—Å—Ç–µ–Ω—Ç) –∑–∞–≥—Ä—É–∂–µ–Ω")
except ImportError as e:
    VOICE_ASSISTANT_AVAILABLE = False
    logger.warning(f"‚ö†Ô∏è Gemini Live API –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω: {e}")

# –ò–º–ø–æ—Ä—Ç OpenAI Realtime API (–∞–ª—å—Ç–µ—Ä–Ω–∞—Ç–∏–≤–Ω—ã–π –≥–æ–ª–æ—Å–æ–≤–æ–π –∞—Å—Å–∏—Å—Ç–µ–Ω—Ç)
try:
    from openai_realtime_bot_integration import start_realtime_chat_command
    OPENAI_REALTIME_AVAILABLE = True
    logger.info("‚úÖ OpenAI Realtime API (–≥–æ–ª–æ—Å–æ–≤–æ–π –∞—Å—Å–∏—Å—Ç–µ–Ω—Ç) –∑–∞–≥—Ä—É–∂–µ–Ω")
except ImportError as e:
    OPENAI_REALTIME_AVAILABLE = False
    logger.warning(f"‚ö†Ô∏è OpenAI Realtime API –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω: {e}")

# LLM Council - –°–æ–≤–µ—Ç AI –º–æ–¥–µ–ª–µ–π –¥–ª—è —Å–ª–æ–∂–Ω—ã—Ö –≤–æ–ø—Ä–æ—Å–æ–≤ (Karpathy's approach)
try:
    from llm_council import (
        LLMCouncil,
        get_llm_council,
        is_council_available,
        is_complex_question
    )
    LLM_COUNCIL_AVAILABLE = is_council_available()
    if LLM_COUNCIL_AVAILABLE:
        logger.info("‚úÖ LLM Council –∑–∞–≥—Ä—É–∂–µ–Ω (Grok + Claude + Gemini)")
    else:
        logger.warning("‚ö†Ô∏è LLM Council –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω (–Ω—É–∂–Ω–æ –º–∏–Ω–∏–º—É–º 2 AI –º–æ–¥–µ–ª–∏)")
except ImportError as e:
    LLM_COUNCIL_AVAILABLE = False
    logger.warning(f"‚ö†Ô∏è –ú–æ–¥—É–ª—å llm_council.py –Ω–µ –Ω–∞–π–¥–µ–Ω: {e}")

# –¢–æ–∫–µ–Ω—ã (–∑–∞–≥—Ä—É–∂–∞—é—Ç—Å—è –∏–∑ .env —Ñ–∞–π–ª–∞)
TELEGRAM_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
XAI_API_KEY = os.getenv("XAI_API_KEY")
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")

# –ü—Ä–æ–≤–µ—Ä–∫–∞ –Ω–∞–ª–∏—á–∏—è —Ç–æ–∫–µ–Ω–æ–≤
if not TELEGRAM_TOKEN:
    raise ValueError("‚ùå TELEGRAM_BOT_TOKEN –Ω–µ –Ω–∞–π–¥–µ–Ω –≤ .env —Ñ–∞–π–ª–µ!")
if not XAI_API_KEY:
    raise ValueError("‚ùå XAI_API_KEY –Ω–µ –Ω–∞–π–¥–µ–Ω –≤ .env —Ñ–∞–π–ª–µ!")
if not ANTHROPIC_API_KEY:
    logger.warning("‚ö†Ô∏è ANTHROPIC_API_KEY –Ω–µ –Ω–∞–π–¥–µ–Ω - —Ñ—É–Ω–∫—Ü–∏–∏ Claude –±—É–¥—É—Ç –Ω–µ–¥–æ—Å—Ç—É–ø–Ω—ã")

# –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è xAI –∫–ª–∏–µ–Ω—Ç–∞
grok_client = None

def get_grok_client():
    """–ü–æ–ª—É—á–∏—Ç—å xAI Grok –∫–ª–∏–µ–Ω—Ç (–ª–µ–Ω–∏–≤–∞—è –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è)"""
    global grok_client
    if grok_client is None:
        grok_client = XAIClient(api_key=XAI_API_KEY)
    return grok_client

# –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è Gemini –≥–µ–Ω–µ—Ä–∞—Ç–æ—Ä–∞
gemini_generator = None

def get_gemini_generator():
    """–ü–æ–ª—É—á–∏—Ç—å Gemini –≥–µ–Ω–µ—Ä–∞—Ç–æ—Ä (–ª–µ–Ω–∏–≤–∞—è –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è)"""
    global gemini_generator
    if gemini_generator is None and GEMINI_AVAILABLE:
        gemini_generator = initialize_gemini_generator()
    return gemini_generator


# === RATE LIMITING –°–ò–°–¢–ï–ú–ê ===

# –•—Ä–∞–Ω–∏–ª–∏—â–µ –≤—Ä–µ–º–µ–Ω–∏ –∑–∞–ø—Ä–æ—Å–æ–≤ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª–µ–π
user_request_times = defaultdict(list)

# –ù–∞—Å—Ç—Ä–æ–π–∫–∏ rate limiting
RATE_LIMIT_MAX_REQUESTS = 10  # –ú–∞–∫—Å–∏–º—É–º –∑–∞–ø—Ä–æ—Å–æ–≤
RATE_LIMIT_WINDOW_SECONDS = 60  # –ó–∞ 60 —Å–µ–∫—É–Ω–¥

# üéØ –ù–ê–°–¢–†–û–ô–ö–ê STREAMING –†–ï–ñ–ò–ú–ê
# True = –æ—Ç–≤–µ—Ç—ã –ø–æ—è–≤–ª—è—é—Ç—Å—è –ø–æ—Å—Ç–µ–ø–µ–Ω–Ω–æ (–∫–∞–∫ –≤ ChatGPT)
# False = –æ—Ç–≤–µ—Ç—ã –ø—Ä–∏—Ö–æ–¥—è—Ç —Å—Ä–∞–∑—É —Ü–µ–ª–∏–∫–æ–º (–∫–ª–∞—Å—Å–∏—á–µ—Å–∫–∏–π —Ä–µ–∂–∏–º)
STREAMING_ENABLED = False  # –ü–æ —É–º–æ–ª—á–∞–Ω–∏—é –í–´–ö–õ–Æ–ß–ï–ù

# ü§ñ –ö–û–ù–§–ò–ì–£–†–ê–¶–ò–Ø AI –ú–û–î–ï–õ–ï–ô (xAI Grok)
# –û—Å–Ω–æ–≤–Ω–∞—è –º–æ–¥–µ–ª—å: grok-4-1-fast
#   - –£–ª—É—á—à–µ–Ω–Ω–∞—è reasoning —Å–ø–æ—Å–æ–±–Ω–æ—Å—Ç—å –¥–ª—è –≥–ª—É–±–æ–∫–æ–≥–æ –∞–Ω–∞–ª–∏–∑–∞
#   - –ò—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è –¥–ª—è —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏—Ö –≤–æ–ø—Ä–æ—Å–æ–≤, –∞–Ω–∞–ª–∏–∑–∞ —Ñ–æ—Ç–æ –∏ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤
#   - –ë–æ–ª–µ–µ —Ç–æ—á–Ω—ã–µ –∏ –ø—Ä–æ—Ñ–µ—Å—Å–∏–æ–Ω–∞–ª—å–Ω—ã–µ –æ—Ç–≤–µ—Ç—ã
# –ë—ã—Å—Ç—Ä–∞—è –º–æ–¥–µ–ª—å: grok-4-1-fast
#   - –î–ª—è –∫–ª–∞—Å—Å–∏—Ñ–∏–∫–∞—Ü–∏–∏ –∑–∞–ø—Ä–æ—Å–æ–≤ –∏ –ø—Ä–æ—Å—Ç—ã—Ö –≤–æ–ø—Ä–æ—Å–æ–≤
#   - –ë—ã—Å—Ç—Ä–∞—è –≥–µ–Ω–µ—Ä–∞—Ü–∏—è –æ—Ç–≤–µ—Ç–æ–≤
# Fallback: Claude Sonnet 4.5 (–ø—Ä–∏ –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–æ—Å—Ç–∏ Grok)

def check_rate_limit(user_id: int) -> bool:
    """
    –ü—Ä–æ–≤–µ—Ä–∫–∞ rate limit –¥–ª—è –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è
    Returns: True –µ—Å–ª–∏ –∑–∞–ø—Ä–æ—Å —Ä–∞–∑—Ä–µ—à–µ–Ω, False –µ—Å–ª–∏ –ø—Ä–µ–≤—ã—à–µ–Ω –ª–∏–º–∏—Ç
    """
    now = datetime.now()
    cutoff = now - timedelta(seconds=RATE_LIMIT_WINDOW_SECONDS)

    # –û—á–∏—Å—Ç–∫–∞ —Å—Ç–∞—Ä—ã—Ö –∑–∞–ø—Ä–æ—Å–æ–≤
    user_request_times[user_id] = [
        t for t in user_request_times[user_id] if t > cutoff
    ]

    # –ü—Ä–æ–≤–µ—Ä–∫–∞ –ª–∏–º–∏—Ç–∞
    if len(user_request_times[user_id]) >= RATE_LIMIT_MAX_REQUESTS:
        logger.warning(f"Rate limit exceeded for user {user_id}")
        return False

    # –î–æ–±–∞–≤–ª—è–µ–º –Ω–æ–≤—ã–π –∑–∞–ø—Ä–æ—Å
    user_request_times[user_id].append(now)
    return True


# === –£–õ–£–ß–®–ï–ù–ù–ê–Ø –û–ë–†–ê–ë–û–¢–ö–ê AI API –° FALLBACK –ù–ê CLAUDE ===

import time
from anthropic import Anthropic

# –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è Claude –∫–ª–∏–µ–Ω—Ç–∞ (—Ä–µ–∑–µ—Ä–≤–Ω—ã–π)
claude_client = None

def get_claude_client():
    """–ü–æ–ª—É—á–∏—Ç—å Claude –∫–ª–∏–µ–Ω—Ç (–ª–µ–Ω–∏–≤–∞—è –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è)"""
    global claude_client
    if claude_client is None and ANTHROPIC_API_KEY:
        claude_client = Anthropic(api_key=ANTHROPIC_API_KEY)
    return claude_client

def call_grok_with_retry(client, model, messages, max_tokens, temperature, search_parameters=None):
    """
    –í—ã–∑–æ–≤ xAI Grok API —Å –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏–º fallback –Ω–∞ Claude –ø—Ä–∏ —Å–±–æ–µ

    –õ–æ–≥–∏–∫–∞:
    1. –ü—ã—Ç–∞–µ—Ç—Å—è –∏—Å–ø–æ–ª—å–∑–æ–≤–∞—Ç—å xAI Grok (–æ—Å–Ω–æ–≤–Ω–æ–π)
    2. –ï—Å–ª–∏ –æ—à–∏–±–∫–∞ - –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏ –ø–µ—Ä–µ–∫–ª—é—á–∞–µ—Ç—Å—è –Ω–∞ Claude (—Ä–µ–∑–µ—Ä–≤)
    3. –õ–æ–≥–∏—Ä—É–µ—Ç –∫–∞–∫–æ–π API –±—ã–ª –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω

    Args:
        search_parameters: –ü–∞—Ä–∞–º–µ—Ç—Ä—ã –ø–æ–∏—Å–∫–∞ {"mode": "auto", "return_citations": True, "sources": [{"type": "web"}, {"type": "news"}, {"type": "x"}]}]
    """
    # –°–Ω–∞—á–∞–ª–∞ –ø—ã—Ç–∞–µ–º—Å—è Grok
    try:
        response = call_xai_with_retry(
            client=client,
            model=model,
            messages=messages,
            max_tokens=max_tokens,
            temperature=temperature,
            search_parameters=search_parameters
        )
        logger.info("‚úÖ –û—Ç–≤–µ—Ç –ø–æ–ª—É—á–µ–Ω –æ—Ç xAI Grok")
        return response
    except Exception as grok_error:
        logger.warning(f"‚ö†Ô∏è xAI Grok –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω: {str(grok_error)}")

        # Fallback –Ω–∞ Claude
        if not ANTHROPIC_API_KEY:
            logger.error("‚ùå Claude API —Ç–∞–∫–∂–µ –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω (–Ω–µ—Ç –∫–ª—é—á–∞)")
            raise Exception("‚ö†Ô∏è AI —Å–µ—Ä–≤–∏—Å—ã –≤—Ä–µ–º–µ–Ω–Ω–æ –Ω–µ–¥–æ—Å—Ç—É–ø–Ω—ã. –ü–æ–ø—Ä–æ–±—É–π—Ç–µ –ø–æ–∑–∂–µ.")

        try:
            logger.info("üîÑ –ü–µ—Ä–µ–∫–ª—é—á–µ–Ω–∏–µ –Ω–∞ Claude Sonnet 4.5 (—Ä–µ–∑–µ—Ä–≤)...")
            claude = get_claude_client()

            # –ü—Ä–µ–æ–±—Ä–∞–∑—É–µ–º —Ñ–æ—Ä–º–∞—Ç messages –¥–ª—è Claude
            claude_messages = []
            system_prompt = None

            for msg in messages:
                if msg.get("role") == "system":
                    system_prompt = msg.get("content")
                else:
                    claude_messages.append(msg)

            # –í—ã–∑—ã–≤–∞–µ–º Claude
            claude_response = claude.messages.create(
                model="claude-sonnet-4-5-20250929",
                max_tokens=max_tokens,
                temperature=temperature,
                system=system_prompt if system_prompt else "–í—ã ‚Äî —ç–∫—Å–ø–µ—Ä—Ç –ø–æ —Å—Ç—Ä–æ–∏—Ç–µ–ª—å–Ω—ã–º –Ω–æ—Ä–º–∞—Ç–∏–≤–∞–º –†–§.",
                messages=claude_messages
            )

            # –ü—Ä–µ–æ–±—Ä–∞–∑—É–µ–º –æ—Ç–≤–µ—Ç Claude –≤ —Ñ–æ—Ä–º–∞—Ç —Å–æ–≤–º–µ—Å—Ç–∏–º—ã–π —Å Grok
            response = {
                "choices": [{
                    "message": {
                        "content": claude_response.content[0].text
                    }
                }]
            }

            logger.info("‚úÖ –û—Ç–≤–µ—Ç –ø–æ–ª—É—á–µ–Ω –æ—Ç Claude Sonnet 4.5 (—Ä–µ–∑–µ—Ä–≤)")
            return response

        except Exception as claude_error:
            logger.error(f"‚ùå Claude —Ç–∞–∫–∂–µ –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω: {str(claude_error)}")
            raise Exception("‚ö†Ô∏è –û–±–∞ AI —Å–µ—Ä–≤–∏—Å–∞ (Grok –∏ Claude) –≤—Ä–µ–º–µ–Ω–Ω–æ –Ω–µ–¥–æ—Å—Ç—É–ø–Ω—ã. –ü–æ–ø—Ä–æ–±—É–π—Ç–µ –ø–æ–∑–∂–µ.")


async def call_grok_with_streaming(client, model, messages, max_tokens, temperature, search_parameters=None):
    """
    –í—ã–∑–æ–≤ xAI Grok API —Å streaming —Ä–µ–∂–∏–º–æ–º (–ø–æ—Å—Ç–µ–ø–µ–Ω–Ω–∞—è –æ—Ç–¥–∞—á–∞ –æ—Ç–≤–µ—Ç–∞)

    Args:
        client: XAIClient instance
        model: –ú–æ–¥–µ–ª—å –¥–ª—è –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏—è
        messages: –°–ø–∏—Å–æ–∫ —Å–æ–æ–±—â–µ–Ω–∏–π
        max_tokens: –ú–∞–∫—Å–∏–º—É–º —Ç–æ–∫–µ–Ω–æ–≤
        temperature: –¢–µ–º–ø–µ—Ä–∞—Ç—É—Ä–∞
        search_parameters: –ü–∞—Ä–∞–º–µ—Ç—Ä—ã –ø–æ–∏—Å–∫–∞ {"mode": "auto", "return_citations": True, "sources": [{"type": "web"}, {"type": "news"}, {"type": "x"}]}]

    Yields:
        str - —á–∞—Å—Ç–∏ —Ç–µ–∫—Å—Ç–∞ –ø–æ –º–µ—Ä–µ –ø–æ–ª—É—á–µ–Ω–∏—è –æ—Ç API
    """
    try:
        # –ò—Å–ø–æ–ª—å–∑—É–µ–º streaming –º–µ—Ç–æ–¥ xAI
        async for chunk in client.chat_completions_create_stream(
            model=model,
            messages=messages,
            max_tokens=max_tokens,
            temperature=temperature,
            search_parameters=search_parameters
        ):
            yield chunk

    except Exception as grok_error:
        logger.warning(f"‚ö†Ô∏è xAI Grok streaming –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω: {str(grok_error)}")

        # Fallback –Ω–∞ –æ–±—ã—á–Ω—ã–π —Ä–µ–∂–∏–º –±–µ–∑ streaming
        logger.info("üîÑ –ü–µ—Ä–µ–∫–ª—é—á–µ–Ω–∏–µ –Ω–∞ –æ–±—ã—á–Ω—ã–π —Ä–µ–∂–∏–º –±–µ–∑ streaming...")
        response = call_grok_with_retry(
            client=client,
            model=model,
            messages=messages,
            max_tokens=max_tokens,
            temperature=temperature,
            search_parameters=search_parameters
        )
        # –û—Ç–¥–∞—ë–º –≤–µ—Å—å –æ—Ç–≤–µ—Ç —Ü–µ–ª–∏–∫–æ–º
        yield response["choices"][0]["message"]["content"]


# === –°–ò–°–¢–ï–ú–ê –ö–õ–ê–°–°–ò–§–ò–ö–ê–¶–ò–ò –ù–ê–ú–ï–†–ï–ù–ò–ô (INTENT CLASSIFICATION) ===

def classify_user_intent(user_message: str) -> dict:
    """
    –ë—ã—Å—Ç—Ä–∞—è –∫–ª–∞—Å—Å–∏—Ñ–∏–∫–∞—Ü–∏—è –Ω–∞–º–µ—Ä–µ–Ω–∏—è –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è —Å –ø–æ–º–æ—â—å—é Haiku.
    –í–æ–∑–≤—Ä–∞—â–∞–µ—Ç —Ç–∏–ø –∑–∞–ø—Ä–æ—Å–∞ –¥–ª—è –≤—ã–±–æ—Ä–∞ –æ–ø—Ç–∏–º–∞–ª—å–Ω–æ–π –º–æ–¥–µ–ª–∏.

    –¢–∏–ø—ã:
    - simple_save: —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ, –ø–æ–¥—Ç–≤–µ—Ä–∂–¥–µ–Ω–∏–µ, –ø—Ä–æ—Å—Ç–∞—è —Ñ–∏–∫—Å–∞—Ü–∏—è
    - simple_question: –ø—Ä–æ—Å—Ç–æ–π –≤–æ–ø—Ä–æ—Å, —Ç—Ä–µ–±—É—é—â–∏–π –∫—Ä–∞—Ç–∫–æ–≥–æ –æ—Ç–≤–µ—Ç–∞
    - technical_question: —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏–π –≤–æ–ø—Ä–æ—Å, —Ç—Ä–µ–±—É—é—â–∏–π —ç–∫—Å–ø–µ—Ä—Ç–∏–∑—ã
    - complex_analysis: —Å–ª–æ–∂–Ω—ã–π –∞–Ω–∞–ª–∏–∑, —Ä–∞—Å—á–µ—Ç—ã, –¥–µ—Ç–∞–ª—å–Ω–∞—è —ç–∫—Å–ø–µ—Ä—Ç–∏–∑–∞
    """
    try:
        client = get_grok_client()

        classification_prompt = f"""–û–ø—Ä–µ–¥–µ–ª–∏ —Ç–∏–ø –∑–∞–ø—Ä–æ—Å–∞ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è. –û—Ç–≤–µ—Ç—å –¢–û–õ–¨–ö–û –æ–¥–Ω–∏–º —Å–ª–æ–≤–æ–º:

–¢–ò–ü–´ –ó–ê–ü–†–û–°–û–í:
- simple_save: –µ—Å–ª–∏ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—å –ø—Ä–æ—Å–∏—Ç —Å–æ—Ö—Ä–∞–Ω–∏—Ç—å, –∑–∞—Ñ–∏–∫—Å–∏—Ä–æ–≤–∞—Ç—å, –∑–∞–ø–∏—Å–∞—Ç—å –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é, –∏–ª–∏ –ø—Ä–æ—Å—Ç–æ –ø–æ–¥—Ç–≤–µ—Ä–∂–¥–∞–µ—Ç ("–¥–∞", "–æ–∫", "—Å–æ—Ö—Ä–∞–Ω–∏")
- simple_question: –ø—Ä–æ—Å—Ç–æ–π –≤–æ–ø—Ä–æ—Å, —Ç—Ä–µ–±—É—é—â–∏–π –∫—Ä–∞—Ç–∫–æ–≥–æ –æ—Ç–≤–µ—Ç–∞ –±–µ–∑ –≥–ª—É–±–æ–∫–æ–π —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–æ–π —ç–∫—Å–ø–µ—Ä—Ç–∏–∑—ã
- technical_question: —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏–π –≤–æ–ø—Ä–æ—Å –ø–æ –Ω–æ—Ä–º–∞—Ç–∏–≤–∞–º, –°–ù–∏–ü, —Ä–∞—Å—á–µ—Ç–∞–º, —Ç—Ä–µ–±—É—é—â–∏–π —Å—Å—ã–ª–æ–∫ –Ω–∞ –¥–æ–∫—É–º–µ–Ω—Ç—ã
- complex_analysis: —Å–ª–æ–∂–Ω–∞—è –∑–∞–¥–∞—á–∞ —Å —Ä–∞—Å—á–µ—Ç–∞–º–∏, –º–Ω–æ–≥–æ—Ñ–∞–∫—Ç–æ—Ä–Ω—ã–º –∞–Ω–∞–ª–∏–∑–æ–º, –¥–µ—Ç–∞–ª—å–Ω–æ–π —ç–∫—Å–ø–µ—Ä—Ç–∏–∑–æ–π

–ó–∞–ø—Ä–æ—Å –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è: "{user_message}"

–û—Ç–≤–µ—Ç—å –¢–û–õ–¨–ö–û –æ–¥–Ω–∏–º —Å–ª–æ–≤–æ–º –∏–∑ —Å–ø–∏—Å–∫–∞ –≤—ã—à–µ:"""

        response = call_grok_with_retry(
            client,
            model="grok-4-1-fast",  # –ë—ã—Å—Ç—Ä–∞—è –º–æ–¥–µ–ª—å –¥–ª—è –∫–ª–∞—Å—Å–∏—Ñ–∏–∫–∞—Ü–∏–∏
            max_tokens=50,
            temperature=0.1,
            messages=[{"role": "user", "content": classification_prompt}]
        )

        intent_type = response["choices"][0]["message"]["content"].strip().lower()

        # –í–∞–ª–∏–¥–∞—Ü–∏—è –æ—Ç–≤–µ—Ç–∞
        valid_types = ["simple_save", "simple_question", "technical_question", "complex_analysis"]
        if intent_type not in valid_types:
            # –ü–æ —É–º–æ–ª—á–∞–Ω–∏—é —Å—á–∏—Ç–∞–µ–º —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏–π –≤–æ–ø—Ä–æ—Å
            intent_type = "technical_question"

        # –í—ã–±–æ—Ä –º–æ–¥–µ–ª–∏ –Ω–∞ –æ—Å–Ω–æ–≤–µ —Ç–∏–ø–∞ –∑–∞–ø—Ä–æ—Å–∞
        if intent_type == "simple_save" or intent_type == "simple_question":
            model = "grok-4-1-fast"  # –ë—ã—Å—Ç—Ä–∞—è –º–æ–¥–µ–ª—å –¥–ª—è –ø—Ä–æ—Å—Ç—ã—Ö –∑–∞–ø—Ä–æ—Å–æ–≤
            max_tokens = 1000
        elif intent_type == "technical_question":
            model = "grok-4-1-fast"  # Reasoning –º–æ–¥–µ–ª—å –¥–ª—è —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏—Ö –≤–æ–ø—Ä–æ—Å–æ–≤
            max_tokens = 5000
        else:  # complex_analysis
            model = "grok-4-1-fast"  # Reasoning –º–æ–¥–µ–ª—å –¥–ª—è —Å–ª–æ–∂–Ω–æ–≥–æ –∞–Ω–∞–ª–∏–∑–∞
            max_tokens = 8000

        logger.info(f"üìä Intent: {intent_type} ‚Üí Model: {model}")

        return {
            "intent": intent_type,
            "model": model,
            "max_tokens": max_tokens
        }

    except Exception as e:
        logger.error(f"Error in intent classification: {e}")
        # –ü—Ä–∏ –æ—à–∏–±–∫–µ –∏—Å–ø–æ–ª—å–∑—É–µ–º Grok Reasoning –¥–ª—è –Ω–∞–¥–µ–∂–Ω–æ—Å—Ç–∏
        return {
            "intent": "technical_question",
            "model": "grok-4-1-fast",
            "max_tokens": 5000
        }


# === –°–ò–°–¢–ï–ú–ê –•–†–ê–ù–ï–ù–ò–Ø –ò–°–¢–û–†–ò–ò –î–ò–ê–õ–û–ì–û–í ===

# –î–∏—Ä–µ–∫—Ç–æ—Ä–∏—è –¥–ª—è —Ö—Ä–∞–Ω–µ–Ω–∏—è –∏—Å—Ç–æ—Ä–∏–∏
HISTORY_DIR = Path("user_conversations")
HISTORY_DIR.mkdir(exist_ok=True)

# In-memory —Ö—Ä–∞–Ω–∏–ª–∏—â–µ –∏—Å—Ç–æ—Ä–∏–∏ (–¥–ª—è –±—ã—Å—Ç—Ä–æ–≥–æ –¥–æ—Å—Ç—É–ø–∞)
user_conversations = defaultdict(list)

# –ú–∞–∫—Å–∏–º–∞–ª—å–Ω–æ–µ –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ —Å–æ–æ–±—â–µ–Ω–∏–π –≤ –∏—Å—Ç–æ—Ä–∏–∏ –¥–ª—è –∫–æ–Ω—Ç–µ–∫—Å—Ç–∞
MAX_CONTEXT_MESSAGES = 10

def load_user_history(user_id: int):
    """–ó–∞–≥—Ä—É–∑–∏—Ç—å –∏—Å—Ç–æ—Ä–∏—é –¥–∏–∞–ª–æ–≥–∞ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è –∏–∑ —Ñ–∞–π–ª–∞"""
    history_file = HISTORY_DIR / f"user_{user_id}.json"
    if history_file.exists():
        try:
            with open(history_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
                user_conversations[user_id] = data.get('messages', [])
                logger.info(f"Loaded {len(user_conversations[user_id])} messages for user {user_id}")
        except Exception as e:
            logger.error(f"Error loading history for user {user_id}: {e}")
            user_conversations[user_id] = []
    else:
        user_conversations[user_id] = []

def save_user_history(user_id: int):
    """–°–æ—Ö—Ä–∞–Ω–∏—Ç—å –∏—Å—Ç–æ—Ä–∏—é –¥–∏–∞–ª–æ–≥–∞ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è –≤ —Ñ–∞–π–ª"""
    history_file = HISTORY_DIR / f"user_{user_id}.json"
    try:
        data = {
            'user_id': user_id,
            'last_updated': datetime.now().isoformat(),
            'messages': user_conversations[user_id]
        }
        with open(history_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        logger.info(f"Saved history for user {user_id}")
    except Exception as e:
        logger.error(f"Error saving history for user {user_id}: {e}")

async def add_message_to_history_async(user_id: int, role: str, content: str, image_analyzed: bool = False):
    """–î–æ–±–∞–≤–∏—Ç—å —Å–æ–æ–±—â–µ–Ω–∏–µ –≤ –∏—Å—Ç–æ—Ä–∏—é (PostgreSQL —Å fallback –Ω–∞ JSON)"""
    # –ò–∑–≤–ª–µ–∫–∞–µ–º —Ç–µ–≥–∏
    tags = extract_tags_from_message(content)

    # –°–æ—Ö—Ä–∞–Ω—è–µ–º –≤ PostgreSQL –µ—Å–ª–∏ –¥–æ—Å—Ç—É–ø–µ–Ω
    if DATABASE_AVAILABLE:
        try:
            await save_message(user_id, role, content, image_analyzed, tags)
            # –û–±–Ω–æ–≤–ª—è–µ–º in-memory –∫–µ—à
            load_user_history(user_id)
            message = {
                'role': role,
                'content': content,
                'timestamp': datetime.now().isoformat(),
                'image_analyzed': image_analyzed,
                'tags': tags
            }
            user_conversations[user_id].append(message)
            if len(user_conversations[user_id]) > 50:
                user_conversations[user_id] = user_conversations[user_id][-50:]
            return
        except Exception as e:
            logger.error(f"PostgreSQL save failed, falling back to JSON: {e}")

    # Fallback –Ω–∞ JSON
    load_user_history(user_id)
    message = {
        'role': role,
        'content': content,
        'timestamp': datetime.now().isoformat(),
        'image_analyzed': image_analyzed,
        'tags': tags
    }
    user_conversations[user_id].append(message)
    if len(user_conversations[user_id]) > 50:
        user_conversations[user_id] = user_conversations[user_id][-50:]
    save_user_history(user_id)

def add_message_to_history(user_id: int, role: str, content: str, image_analyzed: bool = False):
    """–°–∏–Ω—Ö—Ä–æ–Ω–Ω–∞—è –æ–±–µ—Ä—Ç–∫–∞ –¥–ª—è —Å–æ–≤–º–µ—Å—Ç–∏–º–æ—Å—Ç–∏"""
    load_user_history(user_id)
    tags = extract_tags_from_message(content)
    message = {
        'role': role,
        'content': content,
        'timestamp': datetime.now().isoformat(),
        'image_analyzed': image_analyzed,
        'tags': tags
    }
    user_conversations[user_id].append(message)
    if len(user_conversations[user_id]) > 50:
        user_conversations[user_id] = user_conversations[user_id][-50:]
    save_user_history(user_id)

def get_conversation_context(user_id: int) -> list:
    """–ü–æ–ª—É—á–∏—Ç—å –∫–æ–Ω—Ç–µ–∫—Å—Ç –¥–∏–∞–ª–æ–≥–∞ –¥–ª—è Claude API (–ø–æ—Å–ª–µ–¥–Ω–∏–µ N —Å–æ–æ–±—â–µ–Ω–∏–π)"""
    load_user_history(user_id)

    # –ë–µ—Ä—ë–º –ø–æ—Å–ª–µ–¥–Ω–∏–µ MAX_CONTEXT_MESSAGES —Å–æ–æ–±—â–µ–Ω–∏–π
    recent_messages = user_conversations[user_id][-MAX_CONTEXT_MESSAGES:]

    # –ü—Ä–µ–æ–±—Ä–∞–∑—É–µ–º –≤ —Ñ–æ—Ä–º–∞—Ç Claude API
    grok_messages = []
    for msg in recent_messages:
        # –ü—Ä–æ–ø—É—Å–∫–∞–µ–º —Å–æ–æ–±—â–µ–Ω–∏—è —Å –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è–º–∏ (–æ–Ω–∏ —É–∂–µ –æ–±—Ä–∞–±–æ—Ç–∞–Ω—ã)
        if not msg.get('image_analyzed', False):
            grok_messages.append({
                'role': msg['role'],
                'content': msg['content']
            })

    return grok_messages

def clear_user_history(user_id: int):
    """–û—á–∏—Å—Ç–∏—Ç—å –∏—Å—Ç–æ—Ä–∏—é –¥–∏–∞–ª–æ–≥–∞ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è"""
    user_conversations[user_id] = []
    save_user_history(user_id)

def get_user_stats(user_id: int) -> dict:
    """–ü–æ–ª—É—á–∏—Ç—å —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫—É –¥–∏–∞–ª–æ–≥–∞ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è"""
    load_user_history(user_id)
    messages = user_conversations[user_id]

    stats = {
        'total_messages': len(messages),
        'user_messages': len([m for m in messages if m['role'] == 'user']),
        'assistant_messages': len([m for m in messages if m['role'] == 'assistant']),
        'images_analyzed': len([m for m in messages if m.get('image_analyzed', False)]),
        'first_message': messages[0]['timestamp'] if messages else None,
        'last_message': messages[-1]['timestamp'] if messages else None
    }

    return stats


# === –°–ò–°–¢–ï–ú–ê –£–ú–ù–´–• –¢–ï–ì–û–í ===

def extract_tags_from_message(content: str) -> list:
    """–ò–∑–≤–ª–µ—á—å —Ç–µ–≥–∏ –∏–∑ —Å–æ–æ–±—â–µ–Ω–∏—è (—É–ø–æ–º–∏–Ω–∞–Ω–∏—è –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤, —Ç–∏–ø—ã –¥–µ—Ñ–µ–∫—Ç–æ–≤)"""
    tags = []

    # –ò–∑–≤–ª–µ–∫–∞–µ–º —É–ø–æ–º–∏–Ω–∞–Ω–∏—è –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤
    for reg_code in REGULATIONS.keys():
        if reg_code in content:
            tags.append(f"–Ω–æ—Ä–º–∞—Ç–∏–≤:{reg_code}")

    # –ò–∑–≤–ª–µ–∫–∞–µ–º –∫–ª—é—á–µ–≤—ã–µ —Å–ª–æ–≤–∞
    keywords = {
        '–±–µ—Ç–æ–Ω': '–º–∞—Ç–µ—Ä–∏–∞–ª:–±–µ—Ç–æ–Ω',
        '–∞—Ä–º–∞—Ç—É—Ä–∞': '–º–∞—Ç–µ—Ä–∏–∞–ª:–∞—Ä–º–∞—Ç—É—Ä–∞',
        '—Ñ—É–Ω–¥–∞–º–µ–Ω—Ç': '–∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏—è:—Ñ—É–Ω–¥–∞–º–µ–Ω—Ç',
        '–∫—Ä–æ–≤–ª—è': '–∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏—è:–∫—Ä–æ–≤–ª—è',
        '—Å—Ç–µ–Ω–∞': '–∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏—è:—Å—Ç–µ–Ω–∞',
        '–ø–µ—Ä–µ–∫—Ä—ã—Ç–∏–µ': '–∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏—è:–ø–µ—Ä–µ–∫—Ä—ã—Ç–∏–µ'
    }

    content_lower = content.lower()
    for keyword, tag in keywords.items():
        if keyword in content_lower:
            tags.append(tag)

    return list(set(tags))  # –£–±–∏—Ä–∞–µ–º –¥—É–±–ª–∏–∫–∞—Ç—ã

def add_message_to_history_with_tags(user_id: int, role: str, content: str, image_analyzed: bool = False):
    """–î–æ–±–∞–≤–∏—Ç—å —Å–æ–æ–±—â–µ–Ω–∏–µ –≤ –∏—Å—Ç–æ—Ä–∏—é —Å –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏–º —Ç–µ–≥–∏—Ä–æ–≤–∞–Ω–∏–µ–º"""
    load_user_history(user_id)

    # –ò–∑–≤–ª–µ–∫–∞–µ–º —Ç–µ–≥–∏
    tags = extract_tags_from_message(content)

    message = {
        'role': role,
        'content': content,
        'timestamp': datetime.now().isoformat(),
        'image_analyzed': image_analyzed,
        'tags': tags
    }

    user_conversations[user_id].append(message)

    # –û–≥—Ä–∞–Ω–∏—á–∏–≤–∞–µ–º —Ä–∞–∑–º–µ—Ä –∏—Å—Ç–æ—Ä–∏–∏
    if len(user_conversations[user_id]) > 50:
        user_conversations[user_id] = user_conversations[user_id][-50:]

    save_user_history(user_id)


# === –°–ò–°–¢–ï–ú–ê –ü–û–ò–°–ö–ê –ü–û –ò–°–¢–û–†–ò–ò ===

def search_in_history(user_id: int, query: str, limit: int = 10) -> list:
    """–ü–æ–∏—Å–∫ –ø–æ –∏—Å—Ç–æ—Ä–∏–∏ –¥–∏–∞–ª–æ–≥–æ–≤"""
    load_user_history(user_id)
    messages = user_conversations[user_id]

    if not messages:
        return []

    query_lower = query.lower()
    results = []

    for msg in messages:
        # –ü–æ–∏—Å–∫ –ø–æ —Å–æ–¥–µ—Ä–∂–∏–º–æ–º—É
        if query_lower in msg['content'].lower():
            results.append(msg)
        # –ü–æ–∏—Å–∫ –ø–æ —Ç–µ–≥–∞–º
        elif 'tags' in msg and any(query_lower in tag.lower() for tag in msg['tags']):
            results.append(msg)

    # –í–æ–∑–≤—Ä–∞—â–∞–µ–º –ø–æ—Å–ª–µ–¥–Ω–∏–µ N —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤
    return results[-limit:]

def search_by_tags(user_id: int, tags: list, limit: int = 10) -> list:
    """–ü–æ–∏—Å–∫ –ø–æ —Ç–µ–≥–∞–º"""
    load_user_history(user_id)
    messages = user_conversations[user_id]

    if not messages:
        return []

    results = []
    for msg in messages:
        if 'tags' in msg:
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º –ø–µ—Ä–µ—Å–µ—á–µ–Ω–∏–µ —Ç–µ–≥–æ–≤
            msg_tags_lower = [t.lower() for t in msg['tags']]
            tags_lower = [t.lower() for t in tags]
            if any(tag in msg_tags_lower for tag in tags_lower):
                results.append(msg)

    return results[-limit:]


# === –°–ò–°–¢–ï–ú–ê –†–ï–ö–û–ú–ï–ù–î–ê–¶–ò–ô ===

def get_recommendations(user_id: int) -> dict:
    """–ü–æ–ª—É—á–∏—Ç—å —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏ –Ω–∞ –æ—Å–Ω–æ–≤–µ –∏—Å—Ç–æ—Ä–∏–∏ –¥–∏–∞–ª–æ–≥–æ–≤"""
    load_user_history(user_id)
    messages = user_conversations[user_id]

    if not messages:
        return {'recommendations': [], 'popular_topics': []}

    # –°–æ–±–∏—Ä–∞–µ–º –≤—Å–µ —Ç–µ–≥–∏
    all_tags = []
    for msg in messages:
        if 'tags' in msg:
            all_tags.extend(msg['tags'])

    # –ü–æ–¥—Å—á–∏—Ç—ã–≤–∞–µ–º —á–∞—Å—Ç–æ—Ç—É —Ç–µ–≥–æ–≤
    tag_counter = Counter(all_tags)
    popular_tags = tag_counter.most_common(5)

    # –§–æ—Ä–º–∏—Ä—É–µ–º —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏ –Ω–∞ –æ—Å–Ω–æ–≤–µ –ø–æ–ø—É–ª—è—Ä–Ω—ã—Ö —Ç–µ–º
    recommendations = []
    for tag, count in popular_tags:
        if tag.startswith('–Ω–æ—Ä–º–∞—Ç–∏–≤:'):
            reg_code = tag.split(':')[1]
            if reg_code in REGULATIONS:
                recommendations.append({
                    'type': 'related_regulation',
                    'code': reg_code,
                    'title': REGULATIONS[reg_code]['title'],
                    'reason': f'–í—ã —á–∞—Å—Ç–æ –æ–±—Ä–∞—â–∞–ª–∏—Å—å –∫ —ç—Ç–æ–º—É –Ω–æ—Ä–º–∞—Ç–∏–≤—É ({count} —Ä–∞–∑)'
                })

    # –ü–æ–ø—É–ª—è—Ä–Ω—ã–µ —Ç–µ–º—ã
    popular_topics = []
    for tag, count in popular_tags:
        category = tag.split(':')[0] if ':' in tag else '–æ–±—â–µ–µ'
        topic = tag.split(':')[1] if ':' in tag else tag
        popular_topics.append({
            'category': category,
            'topic': topic,
            'mentions': count
        })

    return {
        'recommendations': recommendations[:3],
        'popular_topics': popular_topics[:5]
    }


# === –≠–ö–°–ü–û–†–¢ –í PDF ===

def export_history_to_pdf(user_id: int) -> BytesIO:
    """–≠–∫—Å–ø–æ—Ä—Ç–∏—Ä–æ–≤–∞—Ç—å –∏—Å—Ç–æ—Ä–∏—é –≤ PDF"""
    if not PDF_AVAILABLE:
        raise ImportError("ReportLab –Ω–µ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω")

    load_user_history(user_id)
    messages = user_conversations[user_id]

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                           rightMargin=2*cm, leftMargin=2*cm,
                           topMargin=2*cm, bottomMargin=2*cm)

    story = []
    styles = getSampleStyleSheet()

    # –ó–∞–≥–æ–ª–æ–≤–æ–∫
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        alignment=TA_CENTER,
        spaceAfter=20
    )

    story.append(Paragraph("–ò—Å—Ç–æ—Ä–∏—è –¥–∏–∞–ª–æ–≥–æ–≤ –°—Ç—Ä–æ–π–ù–∞–¥–∑–æ—ÄAI", title_style))
    story.append(Spacer(1, 0.5*cm))

    # –ò–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è
    info_text = f"""
    –ü–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—å ID: {user_id}<br/>
    –î–∞—Ç–∞ —ç–∫—Å–ø–æ—Ä—Ç–∞: {datetime.now().strftime('%d.%m.%Y %H:%M')}<br/>
    –í—Å–µ–≥–æ —Å–æ–æ–±—â–µ–Ω–∏–π: {len(messages)}<br/>
    """
    story.append(Paragraph(info_text, styles['Normal']))
    story.append(Spacer(1, 1*cm))

    # –°–æ–æ–±—â–µ–Ω–∏—è
    for msg in messages:
        role = "–ü–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—å" if msg['role'] == 'user' else "–ë–æ—Ç"
        timestamp = datetime.fromisoformat(msg['timestamp']).strftime('%d.%m.%Y %H:%M')

        # –ó–∞–≥–æ–ª–æ–≤–æ–∫ —Å–æ–æ–±—â–µ–Ω–∏—è
        msg_header = f"<b>{role}</b> - {timestamp}"
        story.append(Paragraph(msg_header, styles['Heading3']))

        # –°–æ–¥–µ—Ä–∂–∏–º–æ–µ
        content = msg['content'][:500] + "..." if len(msg['content']) > 500 else msg['content']
        content = content.replace('<', '&lt;').replace('>', '&gt;')
        story.append(Paragraph(content, styles['Normal']))

        # –¢–µ–≥–∏ (–µ—Å–ª–∏ –µ—Å—Ç—å)
        if 'tags' in msg and msg['tags']:
            tags_text = f"<i>–¢–µ–≥–∏: {', '.join(msg['tags'])}</i>"
            story.append(Paragraph(tags_text, styles['Italic']))

        story.append(Spacer(1, 0.5*cm))

    doc.build(story)
    buffer.seek(0)
    return buffer


# === –≠–ö–°–ü–û–†–¢ –í WORD ===

def export_history_to_docx(user_id: int) -> BytesIO:
    """–≠–∫—Å–ø–æ—Ä—Ç–∏—Ä–æ–≤–∞—Ç—å –∏—Å—Ç–æ—Ä–∏—é –≤ Word"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx –Ω–µ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω")

    load_user_history(user_id)
    messages = user_conversations[user_id]

    doc = Document()

    # –ó–∞–≥–æ–ª–æ–≤–æ–∫
    title = doc.add_heading('–ò—Å—Ç–æ—Ä–∏—è –¥–∏–∞–ª–æ–≥–æ–≤ –°—Ç—Ä–æ–π–ù–∞–¥–∑–æ—ÄAI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # –ò–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è
    doc.add_paragraph(f"–ü–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—å ID: {user_id}")
    doc.add_paragraph(f"–î–∞—Ç–∞ —ç–∫—Å–ø–æ—Ä—Ç–∞: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph(f"–í—Å–µ–≥–æ —Å–æ–æ–±—â–µ–Ω–∏–π: {len(messages)}")
    doc.add_paragraph()

    # –°–æ–æ–±—â–µ–Ω–∏—è
    for msg in messages:
        role = "üë§ –ü–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—å" if msg['role'] == 'user' else "ü§ñ –ë–æ—Ç"
        timestamp = datetime.fromisoformat(msg['timestamp']).strftime('%d.%m.%Y %H:%M')

        # –ó–∞–≥–æ–ª–æ–≤–æ–∫ —Å–æ–æ–±—â–µ–Ω–∏—è
        heading = doc.add_heading(f"{role} - {timestamp}", level=2)

        # –°–æ–¥–µ—Ä–∂–∏–º–æ–µ
        content = msg['content']
        p = doc.add_paragraph(content)

        # –¢–µ–≥–∏
        if 'tags' in msg and msg['tags']:
            tags_p = doc.add_paragraph(f"–¢–µ–≥–∏: {', '.join(msg['tags'])}")
            tags_p.italic = True

        doc.add_paragraph()

    # –°–æ—Ö—Ä–∞–Ω—è–µ–º –≤ –±—É—Ñ–µ—Ä
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer



# === –°–ò–°–¢–ï–ú–ê –£–í–ï–î–û–ú–õ–ï–ù–ò–ô –û –ù–û–†–ú–ê–¢–ò–í–ê–• ===

REGULATIONS_UPDATES = {
    'recent': [
        {
            'code': '–°–ü 24.13330.2021',
            'title': '–°–≤–∞–π–Ω—ã–µ —Ñ—É–Ω–¥–∞–º–µ–Ω—Ç—ã',
            'date': '2021-12-01',
            'type': '–Ω–æ–≤–∞—è_—Ä–µ–¥–∞–∫—Ü–∏—è',
            'changes': '–ê–∫—Ç—É–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω—ã —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è –∫ –∏—Å–ø—ã—Ç–∞–Ω–∏—è–º —Å–≤–∞–π'
        },
        {
            'code': '–°–ü 2.13130.2020',
            'title': '–û–±–µ—Å–ø–µ—á–µ–Ω–∏–µ –æ–≥–Ω–µ—Å—Ç–æ–π–∫–æ—Å—Ç–∏',
            'date': '2020-09-01',
            'type': '–Ω–æ–≤—ã–π',
            'changes': '–ù–æ–≤—ã–µ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è –∫ –æ–≥–Ω–µ–∑–∞—â–∏—Ç–µ'
        }
    ],
    'upcoming': []
}

def check_for_regulation_updates() -> list:
    """–ü—Ä–æ–≤–µ—Ä–∏—Ç—å –Ω–∞–ª–∏—á–∏–µ –æ–±–Ω–æ–≤–ª–µ–Ω–∏–π –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤"""
    recent_updates = REGULATIONS_UPDATES['recent']

    # –§–∏–ª—å—Ç—Ä—É–µ–º –æ–±–Ω–æ–≤–ª–µ–Ω–∏—è –∑–∞ –ø–æ—Å–ª–µ–¥–Ω–∏–µ 30 –¥–Ω–µ–π
    thirty_days_ago = datetime.now() - timedelta(days=30)
    new_updates = []

    for update in recent_updates:
        update_date = datetime.fromisoformat(update['date'])
        if update_date > thirty_days_ago:
            new_updates.append(update)

    return new_updates


# –ë–∞–∑–∞ –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤ —Å URL-—Å—Å—ã–ª–∫–∞–º–∏ –Ω–∞ –ø–µ—Ä–≤–æ–∏—Å—Ç–æ—á–Ω–∏–∫–∏ (–æ–±–Ω–æ–≤–ª–µ–Ω–æ 26.11.2025 - –ø—Ä–æ–≤–µ—Ä–µ–Ω–æ –Ω–∞ docs.cntd.ru)
REGULATIONS = {
    # === –ö–û–ù–°–¢–†–£–ö–¢–ò–í–ù–´–ï –†–ï–®–ï–ù–ò–Ø ===

    "–°–ü 63.13330.2018": {
        "title": "–ë–µ—Ç–æ–Ω–Ω—ã–µ –∏ –∂–µ–ª–µ–∑–æ–±–µ—Ç–æ–Ω–Ω—ã–µ –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏. –û—Å–Ω–æ–≤–Ω—ã–µ –ø–æ–ª–æ–∂–µ–Ω–∏—è",
        "url": "https://docs.cntd.ru/document/554403082",
        "year": "2018",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–ö–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏",
        "replaced": "–°–ù–∏–ü 52-01-2003"
    },

    "–°–ü 16.13330.2017": {
        "title": "–°—Ç–∞–ª—å–Ω—ã–µ –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏",
        "url": "https://docs.cntd.ru/document/456044318",
        "year": "2017",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–ö–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏",
        "replaced": "–°–ù–∏–ü II-23-81*"
    },

    "–°–ü 64.13330.2017": {
        "title": "–î–µ—Ä–µ–≤—è–Ω–Ω—ã–µ –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏",
        "url": "https://docs.cntd.ru/document/456069590",
        "year": "2017",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–ö–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏",
        "replaced": "–°–ù–∏–ü II-25-80"
    },

    "–°–ü 15.13330.2020": {
        "title": "–ö–∞–º–µ–Ω–Ω—ã–µ –∏ –∞—Ä–º–æ–∫–∞–º–µ–Ω–Ω—ã–µ –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏",
        "url": "https://docs.cntd.ru/document/573659385",
        "year": "2020",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–ö–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏",
        "replaced": "–°–ü 15.13330.2012"
    },

    "–°–ü 28.13330.2017": {
        "title": "–ó–∞—â–∏—Ç–∞ —Å—Ç—Ä–æ–∏—Ç–µ–ª—å–Ω—ã—Ö –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–π –æ—Ç –∫–æ—Ä—Ä–æ–∑–∏–∏",
        "url": "https://docs.cntd.ru/document/456054198",
        "year": "2017",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–ó–∞—â–∏—Ç–∞ –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–π"
    },

    "–°–ü 266.1325800.2016": {
        "title": "–ö–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏ —Å—Ç–∞–ª–µ–∂–µ–ª–µ–∑–æ–±–µ—Ç–æ–Ω–Ω—ã–µ",
        "url": "https://docs.cntd.ru/document/456069588",
        "year": "2016",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–ö–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏"
    },

    # === –û–°–ù–û–í–ê–ù–ò–Ø –ò –§–£–ù–î–ê–ú–ï–ù–¢–´ ===

    "–°–ü 22.13330.2016": {
        "title": "–û—Å–Ω–æ–≤–∞–Ω–∏—è –∑–¥–∞–Ω–∏–π –∏ —Å–æ–æ—Ä—É–∂–µ–Ω–∏–π",
        "url": "https://docs.cntd.ru/document/456054206",
        "year": "2016",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–§—É–Ω–¥–∞–º–µ–Ω—Ç—ã",
        "replaced": "–°–ù–∏–ü 2.02.01-83*"
    },

    "–°–ü 24.13330.2021": {
        "title": "–°–≤–∞–π–Ω—ã–µ —Ñ—É–Ω–¥–∞–º–µ–Ω—Ç—ã",
        "url": "https://docs.cntd.ru/document/1200177001",
        "year": "2021",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–§—É–Ω–¥–∞–º–µ–Ω—Ç—ã",
        "replaced": "–°–ü 24.13330.2011"
    },

    "–°–ü 50-101-2004": {
        "title": "–ü—Ä–æ–µ–∫—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ –∏ —É—Å—Ç—Ä–æ–π—Å—Ç–≤–æ –æ—Å–Ω–æ–≤–∞–Ω–∏–π –∏ —Ñ—É–Ω–¥–∞–º–µ–Ω—Ç–æ–≤ –∑–¥–∞–Ω–∏–π –∏ —Å–æ–æ—Ä—É–∂–µ–Ω–∏–π",
        "url": "https://docs.cntd.ru/document/1200035505",
        "year": "2004",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–§—É–Ω–¥–∞–º–µ–Ω—Ç—ã"
    },

    "–°–ü 45.13330.2017": {
        "title": "–ó–µ–º–ª—è–Ω—ã–µ —Å–æ–æ—Ä—É–∂–µ–Ω–∏—è, –æ—Å–Ω–æ–≤–∞–Ω–∏—è –∏ —Ñ—É–Ω–¥–∞–º–µ–Ω—Ç—ã",
        "url": "https://docs.cntd.ru/document/456069576",
        "year": "2017",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–§—É–Ω–¥–∞–º–µ–Ω—Ç—ã"
    },

    # === –û–ë–°–õ–ï–î–û–í–ê–ù–ò–ï –ò –ú–û–ù–ò–¢–û–†–ò–ù–ì ===

    "–°–ü 13-102-2003": {
        "title": "–ü—Ä–∞–≤–∏–ª–∞ –æ–±—Å–ª–µ–¥–æ–≤–∞–Ω–∏—è –Ω–µ—Å—É—â–∏—Ö —Å—Ç—Ä–æ–∏—Ç–µ–ª—å–Ω—ã—Ö –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–π –∑–¥–∞–Ω–∏–π –∏ —Å–æ–æ—Ä—É–∂–µ–Ω–∏–π",
        "url": "https://docs.cntd.ru/document/1200035173",
        "year": "2003",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–û–±—Å–ª–µ–¥–æ–≤–∞–Ω–∏–µ"
    },

    "–ì–û–°–¢ 31937-2011": {
        "title": "–ó–¥–∞–Ω–∏—è –∏ —Å–æ–æ—Ä—É–∂–µ–Ω–∏—è. –ü—Ä–∞–≤–∏–ª–∞ –æ–±—Å–ª–µ–¥–æ–≤–∞–Ω–∏—è –∏ –º–æ–Ω–∏—Ç–æ—Ä–∏–Ω–≥–∞ —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–æ–≥–æ —Å–æ—Å—Ç–æ—è–Ω–∏—è",
        "url": "https://docs.cntd.ru/document/1200100941",
        "year": "2011",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–û–±—Å–ª–µ–¥–æ–≤–∞–Ω–∏–µ"
    },

    "–°–ü 255.1325800.2016": {
        "title": "–ó–¥–∞–Ω–∏—è –∏ —Å–æ–æ—Ä—É–∂–µ–Ω–∏—è. –ü—Ä–∞–≤–∏–ª–∞ —ç–∫—Å–ø–ª—É–∞—Ç–∞—Ü–∏–∏. –û—Å–Ω–æ–≤–Ω—ã–µ –ø–æ–ª–æ–∂–µ–Ω–∏—è",
        "url": "https://docs.cntd.ru/document/456050595",
        "year": "2016",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–≠–∫—Å–ø–ª—É–∞—Ç–∞—Ü–∏—è"
    },

    # === –¢–ï–ü–õ–û–¢–ï–•–ù–ò–ö–ê ===

    "–°–ü 50.13330.2012": {
        "title": "–¢–µ–ø–ª–æ–≤–∞—è –∑–∞—â–∏—Ç–∞ –∑–¥–∞–Ω–∏–π",
        "url": "https://docs.cntd.ru/document/1200095525",
        "year": "2012",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç (—Å –∏–∑–º. 2015)",
        "category": "–¢–µ–ø–ª–æ—Ç–µ—Ö–Ω–∏–∫–∞",
        "replaced": "–°–ù–∏–ü 23-02-2003"
    },

    "–°–ü 23-101-2004": {
        "title": "–ü—Ä–æ–µ–∫—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ —Ç–µ–ø–ª–æ–≤–æ–π –∑–∞—â–∏—Ç—ã –∑–¥–∞–Ω–∏–π",
        "url": "https://docs.cntd.ru/document/1200035109",
        "year": "2004",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–¢–µ–ø–ª–æ—Ç–µ—Ö–Ω–∏–∫–∞"
    },

    # === –ö–†–û–í–õ–ò –ò –ì–ò–î–†–û–ò–ó–û–õ–Ø–¶–ò–Ø ===

    "–°–ü 17.13330.2017": {
        "title": "–ö—Ä–æ–≤–ª–∏",
        "url": "https://docs.cntd.ru/document/456044318",
        "year": "2017",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–ö—Ä–æ–≤–ª–∏",
        "replaced": "–°–ù–∏–ü II-26-76"
    },

    "–°–ü 71.13330.2017": {
        "title": "–ò–∑–æ–ª—è—Ü–∏–æ–Ω–Ω—ã–µ –∏ –æ—Ç–¥–µ–ª–æ—á–Ω—ã–µ –ø–æ–∫—Ä—ã—Ç–∏—è",
        "url": "https://docs.cntd.ru/document/456054235",
        "year": "2017",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–ò–∑–æ–ª—è—Ü–∏—è"
    },

    # === –ò–ù–ñ–ï–ù–ï–†–ù–´–ï –°–ò–°–¢–ï–ú–´ ===

    "–°–ü 60.13330.2020": {
        "title": "–û—Ç–æ–ø–ª–µ–Ω–∏–µ, –≤–µ–Ω—Ç–∏–ª—è—Ü–∏—è –∏ –∫–æ–Ω–¥–∏—Ü–∏–æ–Ω–∏—Ä–æ–≤–∞–Ω–∏–µ –≤–æ–∑–¥—É—Ö–∞",
        "url": "https://docs.cntd.ru/document/573659347",
        "year": "2020",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–ò–Ω–∂–µ–Ω–µ—Ä–∏—è",
        "replaced": "–°–ü 60.13330.2016"
    },

    "–°–ü 30.13330.2020": {
        "title": "–í–Ω—É—Ç—Ä–µ–Ω–Ω–∏–π –≤–æ–¥–æ–ø—Ä–æ–≤–æ–¥ –∏ –∫–∞–Ω–∞–ª–∏–∑–∞—Ü–∏—è –∑–¥–∞–Ω–∏–π",
        "url": "https://docs.cntd.ru/document/573659385",
        "year": "2020",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–ò–Ω–∂–µ–Ω–µ—Ä–∏—è",
        "replaced": "–°–ü 30.13330.2016"
    },

    "–°–ü 52.13330.2016": {
        "title": "–ï—Å—Ç–µ—Å—Ç–≤–µ–Ω–Ω–æ–µ –∏ –∏—Å–∫—É—Å—Å—Ç–≤–µ–Ω–Ω–æ–µ –æ—Å–≤–µ—â–µ–Ω–∏–µ",
        "url": "https://docs.cntd.ru/document/456054197",
        "year": "2016",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–ò–Ω–∂–µ–Ω–µ—Ä–∏—è",
        "replaced": "–°–ù–∏–ü 23-05-95*"
    },

    "–°–ü 73.13330.2016": {
        "title": "–í–Ω—É—Ç—Ä–µ–Ω–Ω–∏–µ —Å–∞–Ω–∏—Ç–∞—Ä–Ω–æ-—Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏–µ —Å–∏—Å—Ç–µ–º—ã –∑–¥–∞–Ω–∏–π",
        "url": "https://docs.cntd.ru/document/456069601",
        "year": "2016",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–ò–Ω–∂–µ–Ω–µ—Ä–∏—è"
    },

    # === –ü–û–ñ–ê–†–ù–ê–Ø –ë–ï–ó–û–ü–ê–°–ù–û–°–¢–¨ ===

    "–°–ü 2.13130.2020": {
        "title": "–°–∏—Å—Ç–µ–º—ã –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–∂–∞—Ä–Ω–æ–π –∑–∞—â–∏—Ç—ã. –û–±–µ—Å–ø–µ—á–µ–Ω–∏–µ –æ–≥–Ω–µ—Å—Ç–æ–π–∫–æ—Å—Ç–∏ –æ–±—ä–µ–∫—Ç–æ–≤ –∑–∞—â–∏—Ç—ã",
        "url": "https://docs.cntd.ru/document/565837297",
        "year": "2020",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–ü–æ–∂–∞—Ä–Ω–∞—è –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç—å",
        "replaced": "–°–ü 2.13130.2012"
    },

    "–°–ü 4.13130.2013": {
        "title": "–°–∏—Å—Ç–µ–º—ã –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–∂–∞—Ä–Ω–æ–π –∑–∞—â–∏—Ç—ã. –û–≥—Ä–∞–Ω–∏—á–µ–Ω–∏–µ —Ä–∞—Å–ø—Ä–æ—Å—Ç—Ä–∞–Ω–µ–Ω–∏—è –ø–æ–∂–∞—Ä–∞ –Ω–∞ –æ–±—ä–µ–∫—Ç–∞—Ö –∑–∞—â–∏—Ç—ã. –¢—Ä–µ–±–æ–≤–∞–Ω–∏—è –∫ –æ–±—ä–µ–º–Ω–æ-–ø–ª–∞–Ω–∏—Ä–æ–≤–æ—á–Ω—ã–º –∏ –∫–æ–Ω—Å—Ç—Ä—É–∫—Ç–∏–≤–Ω—ã–º —Ä–µ—à–µ–Ω–∏—è–º",
        "url": "https://docs.cntd.ru/document/1200096437",
        "year": "2013",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç (—Å –∏–∑–º. 2021)",
        "category": "–ü–æ–∂–∞—Ä–Ω–∞—è –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç—å"
    },

    "–°–ü 5.13130.2009": {
        "title": "–°–∏—Å—Ç–µ–º—ã –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–∂–∞—Ä–Ω–æ–π –∑–∞—â–∏—Ç—ã. –£—Å—Ç–∞–Ω–æ–≤–∫–∏ –ø–æ–∂–∞—Ä–Ω–æ–π —Å–∏–≥–Ω–∞–ª–∏–∑–∞—Ü–∏–∏ –∏ –ø–æ–∂–∞—Ä–æ—Ç—É—à–µ–Ω–∏—è –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏–µ",
        "url": "https://docs.cntd.ru/document/1200071872",
        "year": "2009",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç (—Å –∏–∑–º. 2017)",
        "category": "–ü–æ–∂–∞—Ä–Ω–∞—è –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç—å"
    },

    # === –î–û–°–¢–£–ü–ù–û–°–¢–¨ –î–õ–Ø –ú–ì–ù ===

    "–°–ü 59.13330.2020": {
        "title": "–î–æ—Å—Ç—É–ø–Ω–æ—Å—Ç—å –∑–¥–∞–Ω–∏–π –∏ —Å–æ–æ—Ä—É–∂–µ–Ω–∏–π –¥–ª—è –º–∞–ª–æ–º–æ–±–∏–ª—å–Ω—ã—Ö –≥—Ä—É–ø–ø –Ω–∞—Å–µ–ª–µ–Ω–∏—è",
        "url": "https://docs.cntd.ru/document/573659347",
        "year": "2020",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–î–æ—Å—Ç—É–ø–Ω–æ—Å—Ç—å",
        "replaced": "–°–ü 59.13330.2016"
    },

    # === –ö–û–ù–¢–†–û–õ–¨ –ö–ê–ß–ï–°–¢–í–ê ===

    "–ì–û–°–¢ 10180-2012": {
        "title": "–ë–µ—Ç–æ–Ω—ã. –ú–µ—Ç–æ–¥—ã –æ–ø—Ä–µ–¥–µ–ª–µ–Ω–∏—è –ø—Ä–æ—á–Ω–æ—Å—Ç–∏ –ø–æ –∫–æ–Ω—Ç—Ä–æ–ª—å–Ω—ã–º –æ–±—Ä–∞–∑—Ü–∞–º",
        "url": "https://docs.cntd.ru/document/1200100908",
        "year": "2012",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–ö–æ–Ω—Ç—Ä–æ–ª—å –∫–∞—á–µ—Å—Ç–≤–∞"
    },

    "–ì–û–°–¢ 22690-2015": {
        "title": "–ë–µ—Ç–æ–Ω—ã. –û–ø—Ä–µ–¥–µ–ª–µ–Ω–∏–µ –ø—Ä–æ—á–Ω–æ—Å—Ç–∏ –º–µ—Ö–∞–Ω–∏—á–µ—Å–∫–∏–º–∏ –º–µ—Ç–æ–¥–∞–º–∏ –Ω–µ—Ä–∞–∑—Ä—É—à–∞—é—â–µ–≥–æ –∫–æ–Ω—Ç—Ä–æ–ª—è",
        "url": "https://docs.cntd.ru/document/1200121930",
        "year": "2015",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–ö–æ–Ω—Ç—Ä–æ–ª—å –∫–∞—á–µ—Å—Ç–≤–∞"
    },

    "–ì–û–°–¢ 18105-2018": {
        "title": "–ë–µ—Ç–æ–Ω—ã. –ü—Ä–∞–≤–∏–ª–∞ –∫–æ–Ω—Ç—Ä–æ–ª—è –∏ –æ—Ü–µ–Ω–∫–∏ –ø—Ä–æ—á–Ω–æ—Å—Ç–∏",
        "url": "https://docs.cntd.ru/document/1200161950",
        "year": "2018",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–ö–æ–Ω—Ç—Ä–æ–ª—å –∫–∞—á–µ—Å—Ç–≤–∞",
        "replaced": "–ì–û–°–¢ 18105-2010"
    },

    "–ì–û–°–¢ 23055-78": {
        "title": "–ö–æ–Ω—Ç—Ä–æ–ª—å –Ω–µ—Ä–∞–∑—Ä—É—à–∞—é—â–∏–π. –°–≤–∞—Ä–∫–∞ –º–µ—Ç–∞–ª–ª–æ–≤ –ø–ª–∞–≤–ª–µ–Ω–∏–µ–º. –ö–ª–∞—Å—Å–∏—Ñ–∏–∫–∞—Ü–∏—è —Å–≤–∞—Ä–Ω—ã—Ö —Å–æ–µ–¥–∏–Ω–µ–Ω–∏–π –ø–æ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–∞–º —Ä–∞–¥–∏–æ–≥—Ä–∞—Ñ–∏—á–µ—Å–∫–æ–≥–æ –∫–æ–Ω—Ç—Ä–æ–ª—è",
        "url": "https://docs.cntd.ru/document/1200012783",
        "year": "1978",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–ö–æ–Ω—Ç—Ä–æ–ª—å –∫–∞—á–µ—Å—Ç–≤–∞"
    },

    "–ì–û–°–¢ 5781-82": {
        "title": "–°—Ç–∞–ª—å –≥–æ—Ä—è—á–µ–∫–∞—Ç–∞–Ω–∞—è –¥–ª—è –∞—Ä–º–∏—Ä–æ–≤–∞–Ω–∏—è –∂–µ–ª–µ–∑–æ–±–µ—Ç–æ–Ω–Ω—ã—Ö –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–π. –¢–µ—Ö–Ω–∏—á–µ—Å–∫–∏–µ —É—Å–ª–æ–≤–∏—è",
        "url": "https://docs.cntd.ru/document/1200005050",
        "year": "1982",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç (—Å –∏–∑–º. 2021)",
        "category": "–ú–∞—Ç–µ—Ä–∏–∞–ª—ã"
    },

    # === –ù–ê–ì–†–£–ó–ö–ò –ò –í–û–ó–î–ï–ô–°–¢–í–ò–Ø ===

    "–°–ü 20.13330.2016": {
        "title": "–ù–∞–≥—Ä—É–∑–∫–∏ –∏ –≤–æ–∑–¥–µ–π—Å—Ç–≤–∏—è",
        "url": "https://docs.cntd.ru/document/456028291",
        "year": "2016",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–†–∞—Å—á—ë—Ç—ã",
        "replaced": "–°–ù–∏–ü 2.01.07-85*"
    },

    # === –û–•–†–ê–ù–ê –¢–†–£–î–ê –ò –ë–ï–ó–û–ü–ê–°–ù–û–°–¢–¨ ===

    "–°–ü 12-136-2002": {
        "title": "–ë–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç—å —Ç—Ä—É–¥–∞ –≤ —Å—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤–µ. –†–µ—à–µ–Ω–∏—è –ø–æ –æ—Ö—Ä–∞–Ω–µ —Ç—Ä—É–¥–∞ –∏ –ø—Ä–æ–º—ã—à–ª–µ–Ω–Ω–æ–π –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç–∏ –≤ –ø—Ä–æ–µ–∫—Ç–∞—Ö –æ—Ä–≥–∞–Ω–∏–∑–∞—Ü–∏–∏ —Å—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤–∞ –∏ –ø—Ä–æ–µ–∫—Ç–∞—Ö –ø—Ä–æ–∏–∑–≤–æ–¥—Å—Ç–≤–∞ —Ä–∞–±–æ—Ç",
        "url": "https://docs.cntd.ru/document/1200000308",
        "year": "2002",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–û—Ö—Ä–∞–Ω–∞ —Ç—Ä—É–¥–∞"
    },

    # === –°–ï–ô–°–ú–û–°–¢–û–ô–ö–û–°–¢–¨ ===

    "–°–ü 14.13330.2018": {
        "title": "–°—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤–æ –≤ —Å–µ–π—Å–º–∏—á–µ—Å–∫–∏—Ö —Ä–∞–π–æ–Ω–∞—Ö",
        "url": "https://docs.cntd.ru/document/550565571",
        "year": "2018",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–°–µ–π—Å–º–∏–∫–∞",
        "replaced": "–°–ü 14.13330.2014"
    },

    # === –ö–õ–ò–ú–ê–¢–û–õ–û–ì–ò–Ø ===

    "–°–ü 131.13330.2020": {
        "title": "–°—Ç—Ä–æ–∏—Ç–µ–ª—å–Ω–∞—è –∫–ª–∏–º–∞—Ç–æ–ª–æ–≥–∏—è",
        "url": "https://docs.cntd.ru/document/573659347",
        "year": "2020",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–ö–ª–∏–º–∞—Ç",
        "replaced": "–°–ü 131.13330.2018"
    },

    # === –¢–ï–•–ù–û–õ–û–ì–ò–Ø –ë–ï–¢–û–ù–ê ===

    "–ì–û–°–¢ 7473-2010": {
        "title": "–°–º–µ—Å–∏ –±–µ—Ç–æ–Ω–Ω—ã–µ. –¢–µ—Ö–Ω–∏—á–µ—Å–∫–∏–µ —É—Å–ª–æ–≤–∏—è",
        "url": "https://docs.cntd.ru/document/1200084776",
        "year": "2010",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–ú–∞—Ç–µ—Ä–∏–∞–ª—ã"
    },

    # === –û–ì–ù–ï–ó–ê–©–ò–¢–ê ===

    "–°–ü 293.1325800.2017": {
        "title": "–°–∏—Å—Ç–µ–º—ã –æ–≥–Ω–µ–∑–∞—â–∏—Ç—ã. –£–∑–ª—ã —Å—Ç—Ä–æ–∏—Ç–µ–ª—å–Ω—ã—Ö –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–π. –ü—Ä–∞–≤–∏–ª–∞ –ø—Ä–æ–µ–∫—Ç–∏—Ä–æ–≤–∞–Ω–∏—è",
        "url": "https://docs.cntd.ru/document/456082449",
        "year": "2017",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–ü–æ–∂–∞—Ä–Ω–∞—è –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç—å"
    },

    # === –ó–ï–õ–ï–ù–´–ï –°–¢–ê–ù–î–ê–†–¢–´ ===

    "–ì–û–°–¢ –† 57345-2016": {
        "title": "–ó–µ–ª–µ–Ω—ã–µ —Å—Ç–∞–Ω–¥–∞—Ä—Ç—ã. –ó–¥–∞–Ω–∏—è –∂–∏–ª—ã–µ –∏ –æ–±—â–µ—Å—Ç–≤–µ–Ω–Ω—ã–µ. –†–µ–π—Ç–∏–Ω–≥–æ–≤–∞—è —Å–∏—Å—Ç–µ–º–∞ –æ—Ü–µ–Ω–∫–∏ —É—Å—Ç–æ–π—á–∏–≤–æ—Å—Ç–∏ —Å—Ä–µ–¥—ã –æ–±–∏—Ç–∞–Ω–∏—è",
        "url": "https://docs.cntd.ru/document/1200145289",
        "year": "2016",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–≠–∫–æ–ª–æ–≥–∏—è"
    },

    # === –≠–ù–ï–†–ì–û–≠–§–§–ï–ö–¢–ò–í–ù–û–°–¢–¨ ===

    "–°–ü 345.1325800.2017": {
        "title": "–ó–¥–∞–Ω–∏—è –∂–∏–ª—ã–µ –∏ –æ–±—â–µ—Å—Ç–≤–µ–Ω–Ω—ã–µ. –ü—Ä–∞–≤–∏–ª–∞ –ø—Ä–æ–µ–∫—Ç–∏—Ä–æ–≤–∞–Ω–∏—è —Ç–µ–ø–ª–æ–≤–æ–π –∑–∞—â–∏—Ç—ã",
        "url": "https://docs.cntd.ru/document/456054206",
        "year": "2017",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–≠–Ω–µ—Ä–≥–æ—ç—Ñ—Ñ–µ–∫—Ç–∏–≤–Ω–æ—Å—Ç—å"
    },

    # === –û–†–ì–ê–ù–ò–ó–ê–¶–ò–Ø –°–¢–†–û–ò–¢–ï–õ–¨–°–¢–í–ê ===

    "–°–ü 48.13330.2019": {
        "title": "–û—Ä–≥–∞–Ω–∏–∑–∞—Ü–∏—è —Å—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤–∞",
        "url": "https://docs.cntd.ru/document/564477582",
        "year": "2019",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–û—Ä–≥–∞–Ω–∏–∑–∞—Ü–∏—è",
        "replaced": "–°–ü 48.13330.2011"
    },

    "–°–ü 70.13330.2012": {
        "title": "–ù–µ—Å—É—â–∏–µ –∏ –æ–≥—Ä–∞–∂–¥–∞—é—â–∏–µ –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏",
        "url": "https://docs.cntd.ru/document/1200092705",
        "year": "2012",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç (—Å –∏–∑–º. 2020)",
        "category": "–ü—Ä–æ–µ–∫—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ"
    },

    # === –ü–†–û–ï–ö–¢–ò–†–û–í–ê–ù–ò–ï –ó–î–ê–ù–ò–ô (–ê–ö–¢–£–ê–õ–¨–ù–´–ï 2024-2025) ===

    "–°–ü 54.13330.2022": {
        "title": "–ó–¥–∞–Ω–∏—è –∂–∏–ª—ã–µ –º–Ω–æ–≥–æ–∫–≤–∞—Ä—Ç–∏—Ä–Ω—ã–µ",
        "url": "https://docs.cntd.ru/document/1300540989",
        "year": "2022",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç —Å 01.09.2023",
        "category": "–ñ–∏–ª—ã–µ –∑–¥–∞–Ω–∏—è",
        "replaced": "–°–ü 54.13330.2016"
    },

    "–°–ü 118.13330.2022": {
        "title": "–û–±—â–µ—Å—Ç–≤–µ–Ω–Ω—ã–µ –∑–¥–∞–Ω–∏—è –∏ —Å–æ–æ—Ä—É–∂–µ–Ω–∏—è",
        "url": "https://docs.cntd.ru/document/1300540990",
        "year": "2022",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç —Å 01.09.2023",
        "category": "–û–±—â–µ—Å—Ç–≤–µ–Ω–Ω—ã–µ –∑–¥–∞–Ω–∏—è",
        "replaced": "–°–ü 118.13330.2012*"
    },

    "–°–ü 56.13330.2021": {
        "title": "–ü—Ä–æ–∏–∑–≤–æ–¥—Å—Ç–≤–µ–Ω–Ω—ã–µ –∑–¥–∞–Ω–∏—è",
        "url": "https://docs.cntd.ru/document/1200177025",
        "year": "2021",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–ü—Ä–æ–º—ã—à–ª–µ–Ω–Ω—ã–µ –∑–¥–∞–Ω–∏—è",
        "replaced": "–°–ü 56.13330.2011"
    },

    # === –°–ï–ô–°–ú–û–°–¢–û–ô–ö–û–°–¢–¨ ===

    "–°–ü 14.13330.2018": {
        "title": "–°—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤–æ –≤ —Å–µ–π—Å–º–∏—á–µ—Å–∫–∏—Ö —Ä–∞–π–æ–Ω–∞—Ö",
        "url": "https://docs.cntd.ru/document/554402860",
        "year": "2018",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–°–µ–π—Å–º–æ—Å—Ç–æ–π–∫–æ—Å—Ç—å",
        "replaced": "–°–ü 14.13330.2014"
    },

    # === –ì–ï–û–¢–ï–•–ù–ò–ö–ê ===

    "–°–ü 22.13330.2016": {
        "title": "–û—Å–Ω–æ–≤–∞–Ω–∏—è –∑–¥–∞–Ω–∏–π –∏ —Å–æ–æ—Ä—É–∂–µ–Ω–∏–π",
        "url": "https://docs.cntd.ru/document/456054206",
        "year": "2016",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–ì–µ–æ—Ç–µ—Ö–Ω–∏–∫–∞"
    },

    "–°–ü 47.13330.2016": {
        "title": "–ò–Ω–∂–µ–Ω–µ—Ä–Ω—ã–µ –∏–∑—ã—Å–∫–∞–Ω–∏—è –¥–ª—è —Å—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤–∞. –û—Å–Ω–æ–≤–Ω—ã–µ –ø–æ–ª–æ–∂–µ–Ω–∏—è",
        "url": "https://docs.cntd.ru/document/456050589",
        "year": "2016",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç (—Å –∏–∑–º. 2020)",
        "category": "–ò–∑—ã—Å–∫–∞–Ω–∏—è"
    },

    # === –≠–ù–ï–†–ì–û–≠–§–§–ï–ö–¢–ò–í–ù–û–°–¢–¨ ===

    "–°–ü 230.1325800.2015": {
        "title": "–ö–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏ –æ–≥—Ä–∞–∂–¥–∞—é—â–∏–µ –∑–¥–∞–Ω–∏–π. –•–∞—Ä–∞–∫—Ç–µ—Ä–∏—Å—Ç–∏–∫–∏ —Ç–µ–ø–ª–æ—Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏—Ö –Ω–µ–æ–¥–Ω–æ—Ä–æ–¥–Ω–æ—Å—Ç–µ–π",
        "url": "https://docs.cntd.ru/document/1200123147",
        "year": "2015",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–≠–Ω–µ—Ä–≥–æ—ç—Ñ—Ñ–µ–∫—Ç–∏–≤–Ω–æ—Å—Ç—å"
    },

    "–°–ü 345.1325800.2017": {
        "title": "–ó–¥–∞–Ω–∏—è –∂–∏–ª—ã–µ –∏ –æ–±—â–µ—Å—Ç–≤–µ–Ω–Ω—ã–µ. –ü—Ä–∞–≤–∏–ª–∞ –ø—Ä–æ–µ–∫—Ç–∏—Ä–æ–≤–∞–Ω–∏—è —Ç–µ–ø–ª–æ–≤–æ–π –∑–∞—â–∏—Ç—ã",
        "url": "https://docs.cntd.ru/document/456069587",
        "year": "2017",
        "status": "–î–µ–π—Å—Ç–≤—É–µ—Ç",
        "category": "–≠–Ω–µ—Ä–≥–æ—ç—Ñ—Ñ–µ–∫—Ç–∏–≤–Ω–æ—Å—Ç—å"
    }
}


# === –ü–û–°–¢–û–Ø–ù–ù–ê–Ø –ö–õ–ê–í–ò–ê–¢–£–†–ê ===

def get_main_keyboard():
    """–°–æ–∑–¥–∞—Ç—å –ø–æ—Å—Ç–æ—è–Ω–Ω—É—é –∫–ª–∞–≤–∏–∞—Ç—É—Ä—É —Å –∫–Ω–æ–ø–∫–∞–º–∏"""
    keyboard = [
        [KeyboardButton("üé§ Real-time —á–∞—Ç")],
    ]
    return ReplyKeyboardMarkup(
        keyboard,
        resize_keyboard=True,
        one_time_keyboard=False,
        input_field_placeholder="–ù–∞–ø–∏—à–∏—Ç–µ –≤–æ–ø—Ä–æ—Å –∏–ª–∏ –Ω–∞–∂–º–∏—Ç–µ –∫–Ω–æ–ø–∫—É..."
    )


# === –ö–û–ú–ê–ù–î–´ ===

async def start_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /start"""
    user = update.effective_user
    user_id = update.effective_user.id

    # –ó–∞–≥—Ä—É–∂–∞–µ–º —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫—É –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è
    stats = get_user_stats(user_id)

    welcome_message = f"""üëã –ó–¥—Ä–∞–≤—Å—Ç–≤—É–π—Ç–µ, {user.first_name}!

–Ø ‚Äî *–°—Ç—Ä–æ–π–ù–∞–¥–∑–æ—ÄAI v3.0* ‚Äî –ø—Ä–æ—Ñ–µ—Å—Å–∏–æ–Ω–∞–ª—å–Ω—ã–π –∏–Ω–∂–µ–Ω–µ—Ä–Ω–æ-–Ω–æ—Ä–º–∞—Ç–∏–≤–Ω—ã–π AI-–∞—Å—Å–∏—Å—Ç–µ–Ω—Ç –ø–æ —Å—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤—É –†–§.

üéØ *–£—Ä–æ–≤–µ–Ω—å —ç–∫—Å–ø–µ—Ä—Ç–∏–∑—ã:* 20+ –ª–µ—Ç –∏–Ω–∂–µ–Ω–µ—Ä–Ω–æ–≥–æ –æ–ø—ã—Ç–∞

‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ
üîπ *–í–û–ó–ú–û–ñ–ù–û–°–¢–ò*
‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ

üé§ *–ì–æ–ª–æ—Å–æ–≤–æ–π –∞—Å—Å–∏—Å—Ç–µ–Ω—Ç*
   ‚Ä¢ –û—Ç–ø—Ä–∞–≤—å—Ç–µ –≥–æ–ª–æ—Å–æ–≤–æ–µ ‚Äî –æ—Ç–≤–µ—á—É –≥–æ–ª–æ—Å–æ–º
   ‚Ä¢ Whisper + GPT-4 + TTS

üì∏ *–ê–Ω–∞–ª–∏–∑ —Ñ–æ—Ç–æ –æ–±—ä–µ–∫—Ç–∞*
   ‚Ä¢ –¢–µ—Ö–Ω–∏—á–µ—Å–∫–∞—è —ç–∫—Å–ø–µ—Ä—Ç–∏–∑–∞
   ‚Ä¢ –í—ã—è–≤–ª–µ–Ω–∏–µ –¥–µ—Ñ–µ–∫—Ç–æ–≤ –ø–æ –ì–û–°–¢
   ‚Ä¢ –†–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏ –ø–æ —É—Å—Ç—Ä–∞–Ω–µ–Ω–∏—é

üèõÔ∏è *–°–æ–≤–µ—Ç AI* (–¥–ª—è —Å–ª–æ–∂–Ω—ã—Ö –≤–æ–ø—Ä–æ—Å–æ–≤)
   ‚Ä¢ Grok ‚Äî —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏–π –∞–Ω–∞–ª–∏–∑
   ‚Ä¢ Claude ‚Äî –¥–µ—Ç–∞–ª—å–Ω–∞—è —ç–∫—Å–ø–µ—Ä—Ç–∏–∑–∞
   ‚Ä¢ Gemini ‚Äî –ø—Ä–∞–∫—Ç–∏—á–µ—Å–∫–∏–µ —Ä–µ—à–µ–Ω–∏—è

üßÆ *21 —Å—Ç—Ä–æ–∏—Ç–µ–ª—å–Ω—ã–π –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä*
   ‚Ä¢ –ë–µ—Ç–æ–Ω, –∞—Ä–º–∞—Ç—É—Ä–∞, –æ–ø–∞–ª—É–±–∫–∞
   ‚Ä¢ –ö–∏—Ä–ø–∏—á, –ø–ª–∏—Ç–∫–∞, –∫—Ä–∞—Å–∫–∞
   ‚Ä¢ –ö—Ä–æ–≤–ª—è, —Ñ—É–Ω–¥–∞–º–µ–Ω—Ç, –ª–µ—Å—Ç–Ω–∏—Ü—ã

‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ
üîπ *–ë–ê–ó–ê –ó–ù–ê–ù–ò–ô*
‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ

üìö *–ù–æ—Ä–º–∞—Ç–∏–≤—ã –†–§:*
   –°–ü ‚Ä¢ –ì–û–°–¢ ‚Ä¢ –°–ù–∏–ü ‚Ä¢ –§–ó ‚Ä¢ –ü–£–≠ ‚Ä¢ –°–∞–Ω–ü–∏–ù

üìñ *–°–ø—Ä–∞–≤–æ—á–Ω–∏–∫–∏:*
   –ë–∞–π–∫–æ–≤, –¶—ã—Ç–æ–≤–∏—á, –ö—É–¥–∏—à–∏–Ω –∏ –¥—Ä.

‚öñÔ∏è *–ü—Ä–∞–∫—Ç–∏–∫–∞:*
   –¢–µ—Ö–Ω–∞–¥–∑–æ—Ä ‚Ä¢ –≠–∫—Å–ø–µ—Ä—Ç–∏–∑–∞ ‚Ä¢ –°—É–¥—ã

‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ
üîπ *–ö–û–ú–ê–ù–î–´*
‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ

/realtime\\_chat ‚Äî –ì–æ–ª–æ—Å–æ–≤–æ–π —á–∞—Ç
/calculators ‚Äî –ö–∞–ª—å–∫—É–ª—è—Ç–æ—Ä—ã
/regulations ‚Äî –ù–æ—Ä–º–∞—Ç–∏–≤—ã
/templates ‚Äî –®–∞–±–ª–æ–Ω—ã –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤
/help ‚Äî –ü–æ–ª–Ω–∞—è —Å–ø—Ä–∞–≤–∫–∞

"""

    if stats['total_messages'] > 0:
        welcome_message += f"""üìä *–í–∞—à–∞ —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫–∞:*
‚Ä¢ –°–æ–æ–±—â–µ–Ω–∏–π: {stats['total_messages']}
‚Ä¢ –ü—Ä–æ–∞–Ω–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–æ —Ñ–æ—Ç–æ: {stats['images_analyzed']}

"""

    # –ü–æ–∫–∞–∑—ã–≤–∞–µ–º –∞–∫—Ç–∏–≤–Ω—ã–π –ø—Ä–æ–µ–∫—Ç (–µ—Å–ª–∏ –µ—Å—Ç—å)
    if PROJECTS_AVAILABLE:
        current_project = context.user_data.get("current_project")
        if current_project:
            welcome_message += f"üìÅ *–ê–∫—Ç–∏–≤–Ω—ã–π –ø—Ä–æ–µ–∫—Ç:* {current_project}\n"
            welcome_message += "_(–≤—Å–µ –¥–∏–∞–ª–æ–≥–∏ —Å–æ—Ö—Ä–∞–Ω—è—é—Ç—Å—è –≤ –ø—Ä–æ–µ–∫—Ç)_\n\n"

    welcome_message += "–ü–æ–ø—Ä–æ–±—É–π—Ç–µ –æ—Ç–ø—Ä–∞–≤–∏—Ç—å —Ñ–æ—Ç–æ –æ–±—ä–µ–∫—Ç–∞ –∏–ª–∏ –∑–∞–¥–∞—Ç—å –≤–æ–ø—Ä–æ—Å! üëá"

    # Inline –º–µ–Ω—é –ø–æ–¥ —Å–æ–æ–±—â–µ–Ω–∏–µ–º
    inline_keyboard = [
        [InlineKeyboardButton("üé§ Real-time —á–∞—Ç", callback_data="realtime_chat_start")],
        [InlineKeyboardButton("üìÅ –ü—Ä–æ–µ–∫—Ç", callback_data="project_menu"),
         InlineKeyboardButton("üßÆ –ö–∞–ª—å–∫—É–ª—è—Ç–æ—Ä—ã", callback_data="calculators_menu")],
        [InlineKeyboardButton("üìö –ù–æ—Ä–º–∞—Ç–∏–≤—ã", callback_data="regulations"),
         InlineKeyboardButton("‚ùì –ß–∞—Å—Ç—ã–µ –≤–æ–ø—Ä–æ—Å—ã", callback_data="faq_menu")],
        [InlineKeyboardButton("üìã –®–∞–±–ª–æ–Ω—ã", callback_data="templates"),
         InlineKeyboardButton("üëî –í—ã–±—Ä–∞—Ç—å —Ä–æ–ª—å", callback_data="role")],
        [InlineKeyboardButton("üí° –ü—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è", callback_data="suggestions"),
         InlineKeyboardButton("üîß –†–∞–∑—Ä–∞–±–æ—Ç—á–∏–∫", callback_data="dev_mode")],
        [InlineKeyboardButton("üìù –ü—Ä–∏–º–µ—Ä—ã –≤–æ–ø—Ä–æ—Å–æ–≤", callback_data="examples"),
         InlineKeyboardButton("‚ÑπÔ∏è –°–ø—Ä–∞–≤–∫–∞", callback_data="help")]
    ]
    inline_markup = InlineKeyboardMarkup(inline_keyboard)

    # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –ø—Ä–∏–≤–µ—Ç—Å—Ç–≤–∏–µ —Å inline –º–µ–Ω—é –∏ –ø–æ—Å—Ç–æ—è–Ω–Ω–æ–π –∫–ª–∞–≤–∏–∞—Ç—É—Ä–æ–π
    await update.message.reply_text(
        welcome_message,
        parse_mode='Markdown',
        reply_markup=inline_markup
    )


async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /help"""
    help_text = """üìñ *–°–ü–†–ê–í–ö–ê ‚Äî –°—Ç—Ä–æ–π–ù–∞–¥–∑–æ—ÄAI v3.0*

*üé§ –ì–û–õ–û–°–û–í–û–ô –ê–°–°–ò–°–¢–ï–ù–¢:*
   /realtime\\_chat ‚Äî –ù–∞—á–∞—Ç—å –≥–æ–ª–æ—Å–æ–≤–æ–π —á–∞—Ç
   ‚Ä¢ –û—Ç–ø—Ä–∞–≤—å—Ç–µ –≥–æ–ª–æ—Å–æ–≤–æ–µ ‚Üí –ø–æ–ª—É—á–∏—Ç–µ –≥–æ–ª–æ—Å–æ–≤–æ–π –æ—Ç–≤–µ—Ç
   ‚Ä¢ Whisper (—Ä–∞—Å–ø–æ–∑–Ω–∞–≤–∞–Ω–∏–µ) + GPT-4 + TTS (–æ–∑–≤—É—á–∫–∞)

*üì∏ –ê–ù–ê–õ–ò–ó –§–û–¢–û:*
   ‚Ä¢ –û—Ç–ø—Ä–∞–≤—å—Ç–µ —Ñ–æ—Ç–æ –æ–±—ä–µ–∫—Ç–∞/–¥–µ—Ñ–µ–∫—Ç–∞
   ‚Ä¢ –î–æ–±–∞–≤—å—Ç–µ –ø–æ–¥–ø–∏—Å—å —Å –≤–æ–ø—Ä–æ—Å–æ–º
   ‚Ä¢ –ü–æ–ª—É—á–∏—Ç–µ —ç–∫—Å–ø–µ—Ä—Ç–∏–∑—É –ø–æ –ì–û–°–¢

*üèõÔ∏è –°–û–í–ï–¢ AI* (–¥–ª—è —Å–ª–æ–∂–Ω—ã—Ö –≤–æ–ø—Ä–æ—Å–æ–≤):
   /council –≤–æ–ø—Ä–æ—Å ‚Äî –û–ø—Ä–æ—Å 3-—Ö AI –º–æ–¥–µ–ª–µ–π
   ‚Ä¢ Grok ‚Äî —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏–π –∞–Ω–∞–ª–∏–∑
   ‚Ä¢ Claude ‚Äî –¥–µ—Ç–∞–ª—å–Ω–∞—è —ç–∫—Å–ø–µ—Ä—Ç–∏–∑–∞
   ‚Ä¢ Gemini ‚Äî –ø—Ä–∞–∫—Ç–∏—á–µ—Å–∫–∏–µ —Ä–µ—à–µ–Ω–∏—è
   ‚Ä¢ –ê–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏ –¥–ª—è —Å–ª–æ–∂–Ω—ã—Ö –≤–æ–ø—Ä–æ—Å–æ–≤

*üìö –ö–û–ú–ê–ù–î–´ - –ù–û–†–ú–ê–¢–ò–í–´:*
   /regulations - 27 –∞–∫—Ç—É–∞–ª—å–Ω—ã—Ö –°–ü, –ì–û–°–¢, –°–ù–∏–ü
   /regulations_menu - –ö–∞—Ç–µ–≥–æ—Ä–∏–∏ –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤ (—É–¥–æ–±–Ω–∞—è –Ω–∞–≤–∏–≥–∞—Ü–∏—è!)
   /examples - –ü—Ä–∏–º–µ—Ä—ã –≤–æ–ø—Ä–æ—Å–æ–≤

*üìã –ö–û–ú–ê–ù–î–´ - –¢–†–ï–ë–û–í–ê–ù–ò–Ø 2025:*
   /requirements2025 - –ò–Ω—Ç–µ—Ä–∞–∫—Ç–∏–≤–Ω–∞—è –±–∞–∑–∞ –≤—Å–µ—Ö —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π
   /laws - 8 –æ—Å–Ω–æ–≤–Ω—ã—Ö —Ñ–µ–¥–µ—Ä–∞–ª—å–Ω—ã—Ö –∑–∞–∫–æ–Ω–æ–≤
   /checklist - –ß–µ–∫-–ª–∏—Å—Ç –µ–∂–µ–¥–Ω–µ–≤–Ω—ã—Ö –ø—Ä–æ–≤–µ—Ä–æ–∫

*üõ†Ô∏è –ö–û–ú–ê–ù–î–´ - –ü–†–ê–ö–¢–ò–ö–ê –ü–õ–û–©–ê–î–ö–ò:*
   /hse - –û—Ö—Ä–∞–Ω–∞ —Ç—Ä—É–¥–∞ –∏ —Ç–µ—Ö–Ω–∏–∫–∞ –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç–∏
   /technology - –¢–µ—Ö–Ω–æ–ª–æ–≥–∏—è —Å—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤–∞ (–±–µ—Ç–æ–Ω, –∞—Ä–º–∞—Ç—É—Ä–∞)
   /estimating - –°–º–µ—Ç–Ω–æ–µ –¥–µ–ª–æ –∏ —Ñ–∏–Ω–∞–Ω—Å—ã (–ö–°-2/–ö–°-3)
   /legal - –Æ—Ä–∏–¥–∏—á–µ—Å–∫–∏–µ –≤–æ–ø—Ä–æ—Å—ã –∏ –ø—Ä–µ—Ç–µ–Ω–∑–∏–∏
   /management - –£–ø—Ä–∞–≤–ª–µ–Ω–∏–µ –ø—Ä–æ–µ–∫—Ç–∞–º–∏

*üîç –†–ê–ë–û–¢–ê –° –ò–°–¢–û–†–ò–ï–ô v3.5:*
   /history - –ü–æ—Å–ª–µ–¥–Ω–∏–µ 5 –¥–∏–∞–ª–æ–≥–æ–≤ + –∫–Ω–æ–ø–∫–∏ –¥–µ–π—Å—Ç–≤–∏–π
   /stats - –ü–æ–¥—Ä–æ–±–Ω–∞—è —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫–∞ –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏—è
   /search –∑–∞–ø—Ä–æ—Å - –ü–æ–∏—Å–∫ –ø–æ –∏—Å—Ç–æ—Ä–∏–∏ –¥–∏–∞–ª–æ–≥–æ–≤
   /export - –≠–∫—Å–ø–æ—Ä—Ç –∏—Å—Ç–æ—Ä–∏–∏ (TXT/Markdown)
   /clear - –û—á–∏—Å—Ç–∏—Ç—å –∏—Å—Ç–æ—Ä–∏—é

*üí° –£–ú–ù–´–ï –§–£–ù–ö–¶–ò–ò v5.0:*
   /calculators - –ú–µ–Ω—é –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä–æ–≤
   /realtime_chat - –ì–æ–ª–æ—Å–æ–≤–æ–π –∞—Å—Å–∏—Å—Ç–µ–Ω—Ç (OpenAI)
   /saved - –°–æ—Ö—Ä–∞–Ω—ë–Ω–Ω—ã–µ —Ä–∞—Å—á—ë—Ç—ã –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä–æ–≤
   /templates - –®–∞–±–ª–æ–Ω—ã –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤
   /role - –í—ã–±–æ—Ä —Ä–µ–∂–∏–º–∞ —Ä–∞–±–æ—Ç—ã (–ø—Ä–æ—Ä–∞–±/–ì–ò–ü/–û–¢–ö)

*‚ùì –ë–´–°–¢–†–´–ï –û–¢–í–ï–¢–´ v3.7:*
   /faq - –ë–∞–∑–∞ —á–∞—Å—Ç–æ –∑–∞–¥–∞–≤–∞–µ–º—ã—Ö –≤–æ–ø—Ä–æ—Å–æ–≤ (20+ –æ—Ç–≤–µ—Ç–æ–≤)
   /faq_search –∑–∞–ø—Ä–æ—Å - –ü–æ–∏—Å–∫ –≥–æ—Ç–æ–≤—ã—Ö –æ—Ç–≤–µ—Ç–æ–≤

*–ü—Ä–∏–º–µ—Ä—ã –≤–æ–ø—Ä–æ—Å–æ–≤:*
üìå –ö–∞–∫–∏–µ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è –∫ –ø—Ä–æ—á–Ω–æ—Å—Ç–∏ –±–µ—Ç–æ–Ω–∞ –∫–ª–∞—Å—Å–∞ B25?
üìå –°–∫–æ–ª—å–∫–æ –¥–Ω–µ–π –Ω–∞–±–∏—Ä–∞–µ—Ç –ø—Ä–æ—á–Ω–æ—Å—Ç—å –±–µ—Ç–æ–Ω –ø—Ä–∏ -10¬∞C?
üìå –ö–∞–∫ —Ä–∞—Å—Å—á–∏—Ç–∞—Ç—å –Ω–∞—Ö–ª–µ—Å—Ç –∞—Ä–º–∞—Ç—É—Ä—ã –ê400?
üìå –ö—Ç–æ –¥–æ–ª–∂–µ–Ω –ø–æ–¥–ø–∏—Å—ã–≤–∞—Ç—å –∞–∫—Ç –ö–°-2?
üìå –ß—Ç–æ –¥–µ–ª–∞—Ç—å –µ—Å–ª–∏ –∑–∞–∫–∞–∑—á–∏–∫ –Ω–µ –ø—Ä–∏—Ö–æ–¥–∏—Ç –Ω–∞ —Å–∫—Ä—ã—Ç—ã–µ —Ä–∞–±–æ—Ç—ã?

–ï—Å—Ç—å –≤–æ–ø—Ä–æ—Å—ã? –ü—Ä–æ—Å—Ç–æ –Ω–∞–ø–∏—à–∏—Ç–µ! üí¨"""

    await update.message.reply_text(help_text, parse_mode='Markdown')


async def regulations_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /regulations —Å –∫–ª–∏–∫–∞–±–µ–ª—å–Ω—ã–º–∏ —Å—Å—ã–ª–∫–∞–º–∏ –Ω–∞ –Ω–æ—Ä–º–∞—Ç–∏–≤—ã"""
    text = "üìö **–î–æ—Å—Ç—É–ø–Ω—ã–µ –Ω–æ—Ä–º–∞—Ç–∏–≤—ã:**\n\n"
    text += "_–ù–∞–∂–º–∏—Ç–µ –Ω–∞ –Ω–∞–∑–≤–∞–Ω–∏–µ, —á—Ç–æ–±—ã –æ—Ç–∫—Ä—ã—Ç—å –ø–æ–ª–Ω—ã–π —Ç–µ–∫—Å—Ç –¥–æ–∫—É–º–µ–Ω—Ç–∞_\n\n"

    for code, data in REGULATIONS.items():
        title = data['title']
        url = data['url']
        text += f"üìÑ [{code}]({url})\n   _{title}_\n\n"

    text += "\nüí° –ó–∞–¥–∞–π—Ç–µ –≤–æ–ø—Ä–æ—Å –ø–æ –ª—é–±–æ–º—É –Ω–æ—Ä–º–∞—Ç–∏–≤—É!"

    await update.message.reply_text(text, parse_mode='Markdown', disable_web_page_preview=True)


async def council_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    –ö–æ–º–∞–Ω–¥–∞ /council - –ø—Ä–∏–Ω—É–¥–∏—Ç–µ–ª—å–Ω–æ–µ –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏–µ LLM Council
    
    –ü–æ–∑–≤–æ–ª—è–µ—Ç –∑–∞–¥–∞—Ç—å –≤–æ–ø—Ä–æ—Å –°–æ–≤–µ—Ç—É AI –º–æ–¥–µ–ª–µ–π (Grok + Claude + Gemini),
    –∫–æ—Ç–æ—Ä—ã–µ –¥–∞–¥—É—Ç —Å–≤–æ–∏ –æ—Ç–≤–µ—Ç—ã, –æ—Ü–µ–Ω—è—Ç –¥—Ä—É–≥ –¥—Ä—É–≥–∞, –∏ —Å—Ñ–æ—Ä–º–∏—Ä—É—é—Ç
    –∫–æ–Ω—Å–µ–Ω—Å—É—Å–Ω—ã–π –æ—Ç–≤–µ—Ç –≤—ã—Å–æ–∫–æ–≥–æ –∫–∞—á–µ—Å—Ç–≤–∞.
    """
    if not LLM_COUNCIL_AVAILABLE:
        await update.message.reply_text(
            "‚ùå **LLM Council –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω**\n\n"
            "–î–ª—è —Ä–∞–±–æ—Ç—ã –°–æ–≤–µ—Ç–∞ AI –Ω—É–∂–Ω–æ –º–∏–Ω–∏–º—É–º 2 –º–æ–¥–µ–ª–∏:\n"
            "‚Ä¢ Grok (XAI_API_KEY)\n"
            "‚Ä¢ Claude (ANTHROPIC_API_KEY)\n"
            "‚Ä¢ Gemini (GEMINI_API_KEY)\n\n"
            "–ü—Ä–æ–≤–µ—Ä—å—Ç–µ –Ω–∞—Å—Ç—Ä–æ–π–∫–∏ API –∫–ª—é—á–µ–π.",
            parse_mode="Markdown"
        )
        return
    
    # –ü–æ–ª—É—á–∞–µ–º –≤–æ–ø—Ä–æ—Å –ø–æ—Å–ª–µ –∫–æ–º–∞–Ω–¥—ã /council
    args = context.args
    if not args:
        # –ü–æ–∫–∞–∑—ã–≤–∞–µ–º —Å–ø—Ä–∞–≤–∫—É
        await update.message.reply_text(
            "üèõÔ∏è **LLM Council ‚Äî –°–æ–≤–µ—Ç AI –º–æ–¥–µ–ª–µ–π**\n\n"
            "–û—Ç–ø—Ä–∞–≤–ª—è–µ—Ç –≤–∞—à –≤–æ–ø—Ä–æ—Å –Ω–µ—Å–∫–æ–ª—å–∫–∏–º AI –º–æ–¥–µ–ª—è–º, –∫–æ—Ç–æ—Ä—ã–µ:\n"
            "1Ô∏è‚É£ –î–∞—é—Ç –Ω–µ–∑–∞–≤–∏—Å–∏–º—ã–µ –æ—Ç–≤–µ—Ç—ã (Grok, Claude, Gemini)\n"
            "2Ô∏è‚É£ –û—Ü–µ–Ω–∏–≤–∞—é—Ç –æ—Ç–≤–µ—Ç—ã –¥—Ä—É–≥ –¥—Ä—É–≥–∞\n"
            "3Ô∏è‚É£ –§–æ—Ä–º–∏—Ä—É—é—Ç –∫–æ–Ω—Å–µ–Ω—Å—É—Å–Ω—ã–π –æ—Ç–≤–µ—Ç\n\n"
            "**–ò—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏–µ:**\n"
            "`/council –í–∞—à —Å–ª–æ–∂–Ω—ã–π —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏–π –≤–æ–ø—Ä–æ—Å`\n\n"
            "**–ü—Ä–∏–º–µ—Ä:**\n"
            "`/council –ö–∞–∫ –ø—Ä–∞–≤–∏–ª—å–Ω–æ –∞—Ä–º–∏—Ä–æ–≤–∞—Ç—å –ø–ª–∏—Ç—É –ø–µ—Ä–µ–∫—Ä—ã—Ç–∏—è –ø–æ –°–ü 63.13330?`\n\n"
            "‚è±Ô∏è –í—Ä–µ–º—è –æ—Ç–≤–µ—Ç–∞: 30-60 —Å–µ–∫ (–æ–ø—Ä–æ—Å 3 –º–æ–¥–µ–ª–µ–π)\n"
            "üéØ –ö–∞—á–µ—Å—Ç–≤–æ: –º–∞–∫—Å–∏–º–∞–ª—å–Ω–æ–µ (–∫–æ–Ω—Å–µ–Ω—Å—É—Å —ç–∫—Å–ø–µ—Ä—Ç–æ–≤)",
            parse_mode="Markdown"
        )
        return
    
    question = " ".join(args)
    user_id = update.effective_user.id
    
    # –î–æ–±–∞–≤–ª—è–µ–º –≤–æ–ø—Ä–æ—Å –≤ –∏—Å—Ç–æ—Ä–∏—é
    await add_message_to_history_async(user_id, 'user', f"[COUNCIL] {question}")
    
    # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º —Å–æ–æ–±—â–µ–Ω–∏–µ –æ –Ω–∞—á–∞–ª–µ —Ä–∞–±–æ—Ç—ã –°–æ–≤–µ—Ç–∞
    council_msg = await update.message.reply_text(
        "üèõÔ∏è **–°–æ–≤–µ—Ç AI —Å–æ–±–∏—Ä–∞–µ—Ç—Å—è...**\n\n"
        "‚è≥ –≠—Ç–∞–ø 1/3: –ü–æ–ª—É—á–∞—é –º–Ω–µ–Ω–∏—è —ç–∫—Å–ø–µ—Ä—Ç–æ–≤\n"
        "‚Ä¢ Grok ‚Äî —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏–π –∞–Ω–∞–ª–∏–∑\n"
        "‚Ä¢ Claude ‚Äî –¥–µ—Ç–∞–ª—å–Ω–∞—è —ç–∫—Å–ø–µ—Ä—Ç–∏–∑–∞\n"
        "‚Ä¢ Gemini ‚Äî –ø—Ä–∞–∫—Ç–∏—á–µ—Å–∫–∏–µ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏\n\n"
        "_–≠—Ç–æ –∑–∞–π–º—ë—Ç 30-60 —Å–µ–∫—É–Ω–¥..._",
        parse_mode="Markdown"
    )
    
    try:
        council = get_llm_council()
        if not council:
            await council_msg.edit_text(
                "‚ùå –ù–µ —É–¥–∞–ª–æ—Å—å –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä–æ–≤–∞—Ç—å –°–æ–≤–µ—Ç AI",
                parse_mode="Markdown"
            )
            return
        
        # –ü–æ–ª—É—á–∞–µ–º –∫–æ–Ω—Ç–µ–∫—Å—Ç –¥–∏–∞–ª–æ–≥–∞
        conversation_history = get_conversation_context(user_id)
        context_text = ""
        if conversation_history:
            # –ë–µ—Ä—ë–º –ø–æ—Å–ª–µ–¥–Ω–∏–µ 3 —Å–æ–æ–±—â–µ–Ω–∏—è –∫–∞–∫ –∫–æ–Ω—Ç–µ–∫—Å—Ç
            recent = conversation_history[-3:]
            context_text = "\n".join([f"{m['role']}: {m['content'][:200]}" for m in recent])
        
        # –ó–∞–ø—É—Å–∫–∞–µ–º –ø–æ–ª–Ω—É—é –∫–æ–Ω—Å—É–ª—å—Ç–∞—Ü–∏—é
        result = await council.consult(question, context=context_text, skip_review=False)
        
        if result["success"]:
            # –û–±–Ω–æ–≤–ª—è–µ–º —Å–æ–æ–±—â–µ–Ω–∏–µ —Å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–º
            final_answer = result["final_answer"]
            duration = result["duration_seconds"]
            models = result["models_used"]
            
            # –î–æ–±–∞–≤–ª—è–µ–º –º–µ—Ç–∞–∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é
            footer = f"\n\n---\n‚è±Ô∏è _–í—Ä–µ–º—è: {duration:.1f} —Å–µ–∫ | –ú–æ–¥–µ–ª–∏: {', '.join(models)}_"
            
            # –û–≥—Ä–∞–Ω–∏—á–∏–≤–∞–µ–º –¥–ª–∏–Ω—É —Å–æ–æ–±—â–µ–Ω–∏—è –¥–ª—è Telegram (4096 —Å–∏–º–≤–æ–ª–æ–≤)
            max_len = 4000 - len(footer)
            if len(final_answer) > max_len:
                final_answer = final_answer[:max_len] + "..."
            
            await council_msg.edit_text(
                final_answer + footer,
                parse_mode="Markdown"
            )
            
            # –°–æ—Ö—Ä–∞–Ω—è–µ–º –æ—Ç–≤–µ—Ç –≤ –∏—Å—Ç–æ—Ä–∏—é
            await add_message_to_history_async(user_id, 'assistant', final_answer)
            
            logger.info(f"‚úÖ LLM Council: –æ—Ç–≤–µ—Ç –∑–∞ {duration:.1f} —Å–µ–∫ –¥–ª—è user {user_id}")
        else:
            await council_msg.edit_text(
                f"‚ùå **–û—à–∏–±–∫–∞ –°–æ–≤–µ—Ç–∞ AI**\n\n{result.get('final_answer', '–ù–µ–∏–∑–≤–µ—Å—Ç–Ω–∞—è –æ—à–∏–±–∫–∞')}",
                parse_mode="Markdown"
            )
    
    except Exception as e:
        logger.error(f"Council command error: {e}")
        await council_msg.edit_text(
            f"‚ùå **–û—à–∏–±–∫–∞ –ø—Ä–∏ —Ä–∞–±–æ—Ç–µ –°–æ–≤–µ—Ç–∞ AI**\n\n`{str(e)}`",
            parse_mode="Markdown"
        )


async def examples_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /examples"""
    examples_text = """üí° **–ü—Ä–∏–º–µ—Ä—ã –≤–æ–ø—Ä–æ—Å–æ–≤:**

**–û –±–µ—Ç–æ–Ω–µ:**
‚Ä¢ –ö–∞–∫–∏–µ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è –∫ –ø—Ä–æ—á–Ω–æ—Å—Ç–∏ –±–µ—Ç–æ–Ω–∞ –∫–ª–∞—Å—Å–∞ B25?
‚Ä¢ –î–æ–ø—É—Å—Ç–∏–º–∞—è —à–∏—Ä–∏–Ω–∞ —Ç—Ä–µ—â–∏–Ω—ã –≤ –∂–µ–ª–µ–∑–æ–±–µ—Ç–æ–Ω–µ?
‚Ä¢ –ú–µ—Ç–æ–¥—ã –∫–æ–Ω—Ç—Ä–æ–ª—è –∫–∞—á–µ—Å—Ç–≤–∞ –±–µ—Ç–æ–Ω–∞ –ø–æ –ì–û–°–¢ 10180-2012

**–û –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏—è—Ö:**
‚Ä¢ –¢—Ä–µ–±–æ–≤–∞–Ω–∏—è –∫ –Ω–µ—Å—É—â–∏–º —Å—Ç–µ–Ω–∞–º –∂–∏–ª—ã—Ö –¥–æ–º–æ–≤
‚Ä¢ –î–æ–ø—É—Å—Ç–∏–º—ã–µ –¥–µ—Ñ–æ—Ä–º–∞—Ü–∏–∏ –ø–µ—Ä–µ–∫—Ä—ã—Ç–∏–π
‚Ä¢ –ö–∞–∫ –ø—Ä–æ–≤–µ—Ä–∏—Ç—å –∫–∞—á–µ—Å—Ç–≤–æ –∫–∏—Ä–ø–∏—á–Ω–æ–π –∫–ª–∞–¥–∫–∏?

**–ê–Ω–∞–ª–∏–∑ —Ñ–æ—Ç–æ–≥—Ä–∞—Ñ–∏–π:**
‚Ä¢ –¢—Ä–µ—â–∏–Ω–∞ —à–∏—Ä–∏–Ω–æ–π 0.3 –º–º - –∫—Ä–∏—Ç–∏—á–Ω–∞ –ª–∏ –æ–Ω–∞?
‚Ä¢ –ö–∞–∫ –æ—Ü–µ–Ω–∏—Ç—å —Å—Ç–µ–ø–µ–Ω—å –∫–æ—Ä—Ä–æ–∑–∏–∏ –∞—Ä–º–∞—Ç—É—Ä—ã?
‚Ä¢ –ß—Ç–æ –¥–µ–ª–∞—Ç—å –ø—Ä–∏ –æ–±–Ω–∞—Ä—É–∂–µ–Ω–∏–∏ –æ—Ç—Å–ª–æ–µ–Ω–∏—è —à—Ç—É–∫–∞—Ç—É—Ä–∫–∏?

**–û –∫–æ–Ω—Ç—Ä–æ–ª–µ:**
‚Ä¢ –ö–∞–∫ –ø—Ä–æ–≤–µ—Ä–∏—Ç—å –∫–∞—á–µ—Å—Ç–≤–æ —Å–≤–∞—Ä–Ω—ã—Ö —Å–æ–µ–¥–∏–Ω–µ–Ω–∏–π?
‚Ä¢ –ú–µ—Ç–æ–¥—ã –∫–æ–Ω—Ç—Ä–æ–ª—è –≥–∏–¥—Ä–æ–∏–∑–æ–ª—è—Ü–∏–∏ –ø–æ–¥–≤–∞–ª–∞
‚Ä¢ –¢—Ä–µ–±–æ–≤–∞–Ω–∏—è –∫ –ø—Ä–∏–µ–º–∫–µ —Å–∫—Ä—ã—Ç—ã—Ö —Ä–∞–±–æ—Ç

**üÜï –û –ö–†–û–í–õ–ï (–ù–û–í–û–ï):**
‚Ä¢ –ö–∞–∫–æ–π —É–∫–ª–æ–Ω –Ω—É–∂–µ–Ω –¥–ª—è –º–µ—Ç–∞–ª–ª–æ—á–µ—Ä–µ–ø–∏—Ü—ã?
‚Ä¢ –ö–∞–∫ —Ä–∞—Å—Å—á–∏—Ç–∞—Ç—å –ø–ª–æ—â–∞–¥—å –∫—Ä–æ–≤–ª–∏?
‚Ä¢ –ö–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏—è –∫—Ä–æ–≤–µ–ª—å–Ω–æ–≥–æ –ø–∏—Ä–æ–≥–∞ –¥–ª—è –º–∞–Ω—Å–∞—Ä–¥—ã

**üÜï –û –¢–ï–ü–õ–û–ò–ó–û–õ–Ø–¶–ò–ò (–ù–û–í–û–ï):**
‚Ä¢ –ö–∞–∫—É—é —Ç–æ–ª—â–∏–Ω—É —É—Ç–µ–ø–ª–∏—Ç–µ–ª—è –≤—ã–±—Ä–∞—Ç—å –¥–ª—è –ú–æ—Å–∫–≤—ã?
‚Ä¢ –ö–∞–∫ —Ä–∞—Å—Å—á–∏—Ç–∞—Ç—å —Ç–æ—á–∫—É —Ä–æ—Å—ã –≤ —Å—Ç–µ–Ω–µ?
‚Ä¢ –ö–∞–∫–æ–π —É—Ç–µ–ø–ª–∏—Ç–µ–ª—å –ª—É—á—à–µ –¥–ª—è —Ñ–∞—Å–∞–¥–∞?

**üÜï –û –í–ï–ù–¢–ò–õ–Ø–¶–ò–ò (–ù–û–í–û–ï):**
‚Ä¢ –ö–∞–∫–∞—è –≤–µ–Ω—Ç–∏–ª—è—Ü–∏—è –Ω—É–∂–Ω–∞ –¥–ª—è –∫–≤–∞—Ä—Ç–∏—Ä—ã 75 –º¬≤?
‚Ä¢ –†–∞—Å—á–µ—Ç –≤–æ–∑–¥—É—Ö–æ–æ–±–º–µ–Ω–∞ –¥–ª—è –æ—Ñ–∏—Å–∞
‚Ä¢ –ß—Ç–æ —Ç–∞–∫–æ–µ —Ä–µ–∫—É–ø–µ—Ä–∞—Ü–∏—è —Ç–µ–ø–ª–∞?

**üÜï –û –ü–û–ñ–ê–†–ù–û–ô –ë–ï–ó–û–ü–ê–°–ù–û–°–¢–ò (–ù–û–í–û–ï):**
‚Ä¢ –ö–∞–∫–æ–π –∫–ª–∞—Å—Å –æ–≥–Ω–µ—Å—Ç–æ–π–∫–æ—Å—Ç–∏ –Ω—É–∂–µ–Ω –¥–ª—è –ø–µ—Ä–µ–∫—Ä—ã—Ç–∏–π?
‚Ä¢ –¢—Ä–µ–±–æ–≤–∞–Ω–∏—è –∫ —ç–≤–∞–∫—É–∞—Ü–∏–æ–Ω–Ω—ã–º –≤—ã—Ö–æ–¥–∞–º
‚Ä¢ –ù–æ—Ä–º—ã –æ–≥–Ω–µ–∑–∞—â–∏—Ç—ã —Å—Ç–∞–ª—å–Ω—ã—Ö –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–π

üí° **–° –ø–∞–º—è—Ç—å—é –¥–∏–∞–ª–æ–≥–æ–≤ –≤—ã –º–æ–∂–µ—Ç–µ:**
‚Ä¢ –ó–∞–¥–∞–≤–∞—Ç—å —É—Ç–æ—á–Ω—è—é—â–∏–µ –≤–æ–ø—Ä–æ—Å—ã
‚Ä¢ –°—Å—ã–ª–∞—Ç—å—Å—è –Ω–∞ –ø—Ä–µ–¥—ã–¥—É—â–∏–µ –æ–±—Å—É–∂–¥–µ–Ω–∏—è
‚Ä¢ –†–∞–∑–≤–∏–≤–∞—Ç—å –æ–¥–Ω—É —Ç–µ–º—É –≤ –Ω–µ—Å–∫–æ–ª—å–∫–∏—Ö —Å–æ–æ–±—â–µ–Ω–∏—è—Ö

–ü—Ä–æ—Å—Ç–æ –Ω–∞–ø–∏—à–∏—Ç–µ —Å–≤–æ–π –≤–æ–ø—Ä–æ—Å –∏–ª–∏ –æ—Ç–ø—Ä–∞–≤—å—Ç–µ —Ñ–æ—Ç–æ –æ–±—ä–µ–∫—Ç–∞! üì∏"""

    await update.message.reply_text(examples_text, parse_mode='Markdown')


async def visualize_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /visualize - –≤–∏–∑—É–∞–ª–∏–∑–∞—Ü–∏—è –¥–µ—Ñ–µ–∫—Ç–æ–≤ —Å –ø–æ–º–æ—â—å—é Gemini AI"""

    if not GEMINI_AVAILABLE:
        await update.message.reply_text(
            "‚ö†Ô∏è –§—É–Ω–∫—Ü–∏—è –≤–∏–∑—É–∞–ª–∏–∑–∞—Ü–∏–∏ –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∞. –ú–æ–¥—É–ª—å Gemini –Ω–µ –∑–∞–≥—Ä—É–∂–µ–Ω."
        )
        return

    generator = initialize_gemini_generator()
    if not generator:
        await update.message.reply_text(
            "‚ö†Ô∏è Gemini API –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω. –ü—Ä–æ–≤–µ—Ä—å—Ç–µ –Ω–∞—Å—Ç—Ä–æ–π–∫–∏ GEMINI_API_KEY."
        )
        return

    # –ï—Å–ª–∏ –µ—Å—Ç—å –∞—Ä–≥—É–º–µ–Ω—Ç—ã - –≥–µ–Ω–µ—Ä–∏—Ä—É–µ–º –≤–∏–∑—É–∞–ª–∏–∑–∞—Ü–∏—é –¥–µ—Ñ–µ–∫—Ç–∞
    if context.args:
        defect_description = " ".join(context.args)

        generating_msg = await update.message.reply_text(
            "üé® –ì–µ–Ω–µ—Ä–∏—Ä—É—é –≤–∏–∑—É–∞–ª–∏–∑–∞—Ü–∏—é –¥–µ—Ñ–µ–∫—Ç–∞...\n"
            "–≠—Ç–æ –∑–∞–π–º–µ—Ç 15-30 —Å–µ–∫—É–Ω–¥"
        )

        try:
            # –ì–µ–Ω–µ—Ä–∏—Ä—É–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ –¥–µ—Ñ–µ–∫—Ç–∞
            result = await generator.visualize_defect(
                defect_description=defect_description
            )

            if result and result.get("image_data"):
                # –£–¥–∞–ª—è–µ–º —Å–æ–æ–±—â–µ–Ω–∏–µ –æ –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏
                try:
                    await generating_msg.delete()
                except:
                    pass

                # –§–æ—Ä–º–∏—Ä—É–µ–º –ø–æ–¥–ø–∏—Å—å
                caption = f"""üé® **–í–ò–ó–£–ê–õ–ò–ó–ê–¶–ò–Ø –î–ï–§–ï–ö–¢–ê**

**–û–ø–∏—Å–∞–Ω–∏–µ:** {defect_description}

{result.get('text', '')}

---
ü§ñ _–°–≥–µ–Ω–µ—Ä–∏—Ä–æ–≤–∞–Ω–æ —Å –ø–æ–º–æ—â—å—é Gemini AI_"""

                # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ
                result["image_data"].seek(0)
                await update.message.reply_photo(
                    photo=result["image_data"],
                    caption=caption[:1024],  # –û–≥—Ä–∞–Ω–∏—á–µ–Ω–∏–µ Telegram
                    parse_mode='Markdown'
                )
            else:
                await generating_msg.edit_text(
                    "‚ùå –ù–µ —É–¥–∞–ª–æ—Å—å —Å–æ–∑–¥–∞—Ç—å –≤–∏–∑—É–∞–ª–∏–∑–∞—Ü–∏—é. –ü–æ–ø—Ä–æ–±—É–π—Ç–µ –ø–µ—Ä–µ—Ñ–æ—Ä–º—É–ª–∏—Ä–æ–≤–∞—Ç—å –∑–∞–ø—Ä–æ—Å."
                )

        except Exception as e:
            logger.error(f"–û—à–∏–±–∫–∞ –≤–∏–∑—É–∞–ª–∏–∑–∞—Ü–∏–∏ –¥–µ—Ñ–µ–∫—Ç–∞: {e}")
            await generating_msg.edit_text(
                f"‚ùå –ü—Ä–æ–∏–∑–æ—à–ª–∞ –æ—à–∏–±–∫–∞: {str(e)}"
            )

        return

    # –ï—Å–ª–∏ –∞—Ä–≥—É–º–µ–Ω—Ç–æ–≤ –Ω–µ—Ç - –ø–æ–∫–∞–∑—ã–≤–∞–µ–º —Å–ø—Ä–∞–≤–∫—É
    help_text = """üé® **–í–ò–ó–£–ê–õ–ò–ó–ê–¶–ò–Ø –î–ï–§–ï–ö–¢–û–í - Gemini AI**

–≠—Ç–∞ –∫–æ–º–∞–Ω–¥–∞ —Å–æ–∑–¥–∞–µ—Ç –¥–µ—Ç–∞–ª—å–Ω—ã–µ —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏–µ –æ–ø–∏—Å–∞–Ω–∏—è –¥–ª—è –≤–∏–∑—É–∞–ª–∏–∑–∞—Ü–∏–∏ —Å—Ç—Ä–æ–∏—Ç–µ–ª—å–Ω—ã—Ö –¥–µ—Ñ–µ–∫—Ç–æ–≤.

**–ö–∞–∫ –∏—Å–ø–æ–ª—å–∑–æ–≤–∞—Ç—å:**

1Ô∏è‚É£ **–ê–Ω–∞–ª–∏–∑ —Ñ–æ—Ç–æ–≥—Ä–∞—Ñ–∏–∏ –¥–µ—Ñ–µ–∫—Ç–∞:**
   ‚Ä¢ –û—Ç–ø—Ä–∞–≤—å—Ç–µ —Ñ–æ—Ç–æ —Å –ø–æ–¥–ø–∏—Å—å—é /visualize
   ‚Ä¢ –ü–æ–ª—É—á–∏—Ç–µ –¥–µ—Ç–∞–ª—å–Ω–æ–µ –æ–ø–∏—Å–∞–Ω–∏–µ –¥–ª—è –≤–∏–∑—É–∞–ª–∏–∑–∞—Ü–∏–∏
   ‚Ä¢ –ò–ò –≤—ã–¥–µ–ª–∏—Ç –∫–ª—é—á–µ–≤—ã–µ –∑–æ–Ω—ã –∏ –ø—Ä–æ–±–ª–µ–º–Ω—ã–µ —É—á–∞—Å—Ç–∫–∏

2Ô∏è‚É£ **–í–∏–∑—É–∞–ª–∏–∑–∞—Ü–∏—è –ø–æ –æ–ø–∏—Å–∞–Ω–∏—é:**
   ‚Ä¢ /visualize [–æ–ø–∏—Å–∞–Ω–∏–µ –¥–µ—Ñ–µ–∫—Ç–∞]
   ‚Ä¢ –ù–∞–ø—Ä–∏–º–µ—Ä: `/visualize —Ç—Ä–µ—â–∏–Ω–∞ –≤ –±–µ—Ç–æ–Ω–Ω–æ–π –∫–æ–ª–æ–Ω–Ω–µ —à–∏—Ä–∏–Ω–æ–π 2–º–º`
   ‚Ä¢ –ü–æ–ª—É—á–∏—Ç–µ —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–æ–µ –æ–ø–∏—Å–∞–Ω–∏–µ –¥–ª—è —Å—Ö–µ–º—ã

3Ô∏è‚É£ **–°—Ä–∞–≤–Ω–µ–Ω–∏–µ –¥–æ/–ø–æ—Å–ª–µ:**
   ‚Ä¢ –û—Ç–ø—Ä–∞–≤—å—Ç–µ —Ñ–æ—Ç–æ —Å –ø–æ–¥–ø–∏—Å—å—é `/visualize compare`
   ‚Ä¢ –ò–¥–µ–∞–ª—å–Ω–æ –¥–ª—è –æ—Ç—á–µ—Ç–æ–≤ –æ —Ä–µ–º–æ–Ω—Ç–µ

**–ß—Ç–æ –≤—ã –ø–æ–ª—É—á–∏—Ç–µ:**
‚úÖ –¢–µ—Ö–Ω–∏—á–µ—Å–∫–æ–µ –æ–ø–∏—Å–∞–Ω–∏–µ –¥–µ—Ñ–µ–∫—Ç–∞
‚úÖ –†–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏ –ø–æ —Ü–≤–µ—Ç–æ–≤–æ–º—É –≤—ã–¥–µ–ª–µ–Ω–∏—é –∑–æ–Ω
‚úÖ –†–∞–∑–º–µ—Ä—ã –∏ –º–∞—Å—à—Ç–∞–± –ø—Ä–æ–±–ª–µ–º–Ω—ã—Ö —É—á–∞—Å—Ç–∫–æ–≤
‚úÖ –°—Ç—Ä—É–∫—Ç—É—Ä–∏—Ä–æ–≤–∞–Ω–Ω—ã–µ –¥–∞–Ω–Ω—ã–µ –¥–ª—è —Å–æ–∑–¥–∞–Ω–∏—è —Å—Ö–µ–º

**–ü—Ä–∏–º–µ—Ä—ã:**
üìå `/visualize —Ç—Ä–µ—â–∏–Ω–∞ –≤ –Ω–µ—Å—É—â–µ–π —Å—Ç–µ–Ω–µ`
üìå –û—Ç–ø—Ä–∞–≤—å—Ç–µ —Ñ–æ—Ç–æ —Å –ø–æ–¥–ø–∏—Å—å—é `/visualize`
üìå `/visualize –æ—Ç—Å–ª–æ–µ–Ω–∏–µ —à—Ç—É–∫–∞—Ç—É—Ä–∫–∏ 50x30—Å–º`

*–û—Ç–ø—Ä–∞–≤—å—Ç–µ —Ñ–æ—Ç–æ –∏–ª–∏ –æ–ø–∏—Å–∞–Ω–∏–µ –¥–µ—Ñ–µ–∫—Ç–∞ –¥–ª—è –≤–∏–∑—É–∞–ª–∏–∑–∞—Ü–∏–∏!* üéØ"""

    await update.message.reply_text(help_text, parse_mode='Markdown')


async def handle_photo_with_visualization(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–û–±—Ä–∞–±–æ—Ç–∫–∞ —Ñ–æ—Ç–æ–≥—Ä–∞—Ñ–∏–π –¥–ª—è –≤–∏–∑—É–∞–ª–∏–∑–∞—Ü–∏–∏ —á–µ—Ä–µ–∑ Gemini"""

    if not GEMINI_AVAILABLE:
        return False

    caption = update.message.caption or ""

    # –ü—Ä–æ–≤–µ—Ä—è–µ–º, –∑–∞–ø—Ä–æ—à–µ–Ω–∞ –ª–∏ –≤–∏–∑—É–∞–ª–∏–∑–∞—Ü–∏—è
    if "/visualize" not in caption.lower():
        return False

    generator = get_gemini_generator()
    if not generator:
        await update.message.reply_text(
            "‚ö†Ô∏è Gemini API –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω –≤ –¥–∞–Ω–Ω—ã–π –º–æ–º–µ–Ω—Ç."
        )
        return True

    # –ü–æ–ª—É—á–∞–µ–º —Ñ–æ—Ç–æ
    photo = update.message.photo[-1]
    photo_file = await photo.get_file()
    photo_bytes = await photo_file.download_as_bytearray()

    await update.message.reply_text("üé® –ê–Ω–∞–ª–∏–∑–∏—Ä—É—é –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ –¥–ª—è –≤–∏–∑—É–∞–ª–∏–∑–∞—Ü–∏–∏...")

    try:
        # –û–ø—Ä–µ–¥–µ–ª—è–µ–º —Ç–∏–ø –∑–∞–ø—Ä–æ—Å–∞
        is_comparison = "compare" in caption.lower() or "—Å—Ä–∞–≤–Ω–µ–Ω" in caption.lower()

        if is_comparison:
            # –°—Ä–∞–≤–Ω–∏—Ç–µ–ª—å–Ω–∞—è –≤–∏–∑—É–∞–ª–∏–∑–∞—Ü–∏—è
            description = await generator.create_comparison_description(
                before_image=bytes(photo_bytes),
                defect_info=caption.replace("/visualize", "").replace("compare", "").strip()
            )
        else:
            # –û–±—ã—á–Ω—ã–π –∞–Ω–∞–ª–∏–∑ –¥–ª—è –≤–∏–∑—É–∞–ª–∏–∑–∞—Ü–∏–∏
            description = await generator.analyze_and_visualize_defect(
                image_bytes=bytes(photo_bytes),
                analysis_text=caption.replace("/visualize", "").strip()
            )

        if description:
            response_text = f"""üé® **–¢–ï–•–ù–ò–ß–ï–°–ö–û–ï –û–ü–ò–°–ê–ù–ò–ï –î–õ–Ø –í–ò–ó–£–ê–õ–ò–ó–ê–¶–ò–ò**

{description}

---
üí° *–≠—Ç–æ –æ–ø–∏—Å–∞–Ω–∏–µ –º–æ–∂–Ω–æ –∏—Å–ø–æ–ª—å–∑–æ–≤–∞—Ç—å –¥–ª—è —Å–æ–∑–¥–∞–Ω–∏—è —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏—Ö —Å—Ö–µ–º, –¥–∏–∞–≥—Ä–∞–º–º –∏ –∞–Ω–Ω–æ—Ç–∏—Ä–æ–≤–∞–Ω–Ω—ã—Ö –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π –¥–µ—Ñ–µ–∫—Ç–æ–≤.*
"""
            await update.message.reply_text(response_text, parse_mode='Markdown')
        else:
            await update.message.reply_text(
                "‚ùå –ù–µ —É–¥–∞–ª–æ—Å—å —Å–æ–∑–¥–∞—Ç—å –æ–ø–∏—Å–∞–Ω–∏–µ –≤–∏–∑—É–∞–ª–∏–∑–∞—Ü–∏–∏. –ü–æ–ø—Ä–æ–±—É–π—Ç–µ —Å–Ω–æ–≤–∞."
            )

    except Exception as e:
        logger.error(f"–û—à–∏–±–∫–∞ –≤–∏–∑—É–∞–ª–∏–∑–∞—Ü–∏–∏: {e}")
        await update.message.reply_text(
            f"‚ùå –ü—Ä–æ–∏–∑–æ—à–ª–∞ –æ—à–∏–±–∫–∞ –ø—Ä–∏ —Å–æ–∑–¥–∞–Ω–∏–∏ –≤–∏–∑—É–∞–ª–∏–∑–∞—Ü–∏–∏: {str(e)}"
        )

    return True


async def history_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /history - –ø–æ–∫–∞–∑–∞—Ç—å –ø–æ—Å–ª–µ–¥–Ω–∏–µ —Å–æ–æ–±—â–µ–Ω–∏—è –∏–∑ –∏—Å—Ç–æ—Ä–∏–∏"""
    user_id = update.effective_user.id
    load_user_history(user_id)

    messages = user_conversations[user_id]

    if not messages:
        await update.message.reply_text("üì≠ –ò—Å—Ç–æ—Ä–∏—è –¥–∏–∞–ª–æ–≥–æ–≤ –ø—É—Å—Ç–∞. –ù–∞—á–Ω–∏—Ç–µ –æ–±—â–µ–Ω–∏–µ!")
        return

    # –ü–æ–∫–∞–∑—ã–≤–∞–µ–º –ø–æ—Å–ª–µ–¥–Ω–∏–µ 5 —Å–æ–æ–±—â–µ–Ω–∏–π
    recent = messages[-5:]
    history_text = "üìú **–ü–æ—Å–ª–µ–¥–Ω–∏–µ —Å–æ–æ–±—â–µ–Ω–∏—è:**\n\n"

    for msg in recent:
        role_emoji = "üë§" if msg['role'] == 'user' else "ü§ñ"
        timestamp = datetime.fromisoformat(msg['timestamp']).strftime('%d.%m %H:%M')
        content_preview = msg['content'][:100] + "..." if len(msg['content']) > 100 else msg['content']

        if msg.get('image_analyzed', False):
            content_preview = "üì∏ [–ê–Ω–∞–ª–∏–∑ —Ñ–æ—Ç–æ–≥—Ä–∞—Ñ–∏–∏]"

        history_text += f"{role_emoji} **{timestamp}**\n{content_preview}\n\n"

    history_text += f"\n–í—Å–µ–≥–æ —Å–æ–æ–±—â–µ–Ω–∏–π: {len(messages)}\n–ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ /clear –¥–ª—è –æ—á–∏—Å—Ç–∫–∏ –∏—Å—Ç–æ—Ä–∏–∏"

    await update.message.reply_text(history_text, parse_mode='Markdown')


async def stats_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /stats - –ø–æ–∫–∞–∑–∞—Ç—å —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫—É –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏—è"""
    user_id = update.effective_user.id
    stats = get_user_stats(user_id)

    if stats['total_messages'] == 0:
        await update.message.reply_text("üìä –°—Ç–∞—Ç–∏—Å—Ç–∏–∫–∞ –ø–æ–∫–∞ –ø—É—Å—Ç–∞. –ù–∞—á–Ω–∏—Ç–µ –æ–±—â–µ–Ω–∏–µ!")
        return

    stats_text = f"""üìä **–í–∞—à–∞ —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫–∞ –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏—è:**

üìù **–°–æ–æ–±—â–µ–Ω–∏—è:**
   ‚Ä¢ –í—Å–µ–≥–æ: {stats['total_messages']}
   ‚Ä¢ –û—Ç –≤–∞—Å: {stats['user_messages']}
   ‚Ä¢ –û—Ç –±–æ—Ç–∞: {stats['assistant_messages']}

üì∏ **–ê–Ω–∞–ª–∏–∑ —Ñ–æ—Ç–æ:**
   ‚Ä¢ –ü—Ä–æ–∞–Ω–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–æ: {stats['images_analyzed']}

üìÖ **–ü–µ—Ä–∏–æ–¥ –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏—è:**
   ‚Ä¢ –ü–µ—Ä–≤–æ–µ —Å–æ–æ–±—â–µ–Ω–∏–µ: {datetime.fromisoformat(stats['first_message']).strftime('%d.%m.%Y %H:%M') if stats['first_message'] else 'N/A'}
   ‚Ä¢ –ü–æ—Å–ª–µ–¥–Ω–µ–µ: {datetime.fromisoformat(stats['last_message']).strftime('%d.%m.%Y %H:%M') if stats['last_message'] else 'N/A'}

üí° –ë–æ—Ç –ø–æ–º–Ω–∏—Ç –ø–æ—Å–ª–µ–¥–Ω–∏–µ {MAX_CONTEXT_MESSAGES} —Å–æ–æ–±—â–µ–Ω–∏–π –¥–ª—è –∫–æ–Ω—Ç–µ–∫—Å—Ç–∞ –¥–∏–∞–ª–æ–≥–∞."""

    await update.message.reply_text(stats_text, parse_mode='Markdown')


async def clear_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /clear - –æ—á–∏—Å—Ç–∏—Ç—å –∏—Å—Ç–æ—Ä–∏—é –¥–∏–∞–ª–æ–≥–æ–≤"""
    user_id = update.effective_user.id

    # –°–æ–∑–¥–∞–µ–º –∫–Ω–æ–ø–∫–∏ –ø–æ–¥—Ç–≤–µ—Ä–∂–¥–µ–Ω–∏—è
    keyboard = [
        [InlineKeyboardButton("‚úÖ –î–∞, –æ—á–∏—Å—Ç–∏—Ç—å", callback_data="clear_confirm")],
        [InlineKeyboardButton("‚ùå –û—Ç–º–µ–Ω–∞", callback_data="clear_cancel")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)

    await update.message.reply_text(
        "‚ö†Ô∏è **–ü–æ–¥—Ç–≤–µ—Ä–∂–¥–µ–Ω–∏–µ –æ—á–∏—Å—Ç–∫–∏ –∏—Å—Ç–æ—Ä–∏–∏**\n\n"
        "–í—ã —É–≤–µ—Ä–µ–Ω—ã, —á—Ç–æ —Ö–æ—Ç–∏—Ç–µ —É–¥–∞–ª–∏—Ç—å –≤—Å—é –∏—Å—Ç–æ—Ä–∏—é –¥–∏–∞–ª–æ–≥–æ–≤?\n"
        "–≠—Ç–æ –¥–µ–π—Å—Ç–≤–∏–µ –Ω–µ–ª—å–∑—è –æ—Ç–º–µ–Ω–∏—Ç—å.",
        parse_mode='Markdown',
        reply_markup=reply_markup
    )


async def export_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /export - —ç–∫—Å–ø–æ—Ä—Ç–∏—Ä–æ–≤–∞—Ç—å –∏—Å—Ç–æ—Ä–∏—é"""
    user_id = update.effective_user.id

    keyboard = [
        [InlineKeyboardButton("üìÑ PDF", callback_data="export_pdf")],
        [InlineKeyboardButton("üìù Word", callback_data="export_docx")],
        [InlineKeyboardButton("‚ùå –û—Ç–º–µ–Ω–∞", callback_data="export_cancel")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)

    await update.message.reply_text(
        "üì§ **–≠–∫—Å–ø–æ—Ä—Ç –∏—Å—Ç–æ—Ä–∏–∏ –¥–∏–∞–ª–æ–≥–æ–≤**\n\n"
        "–í—ã–±–µ—Ä–∏—Ç–µ —Ñ–æ—Ä–º–∞—Ç –¥–ª—è —ç–∫—Å–ø–æ—Ä—Ç–∞:",
        parse_mode='Markdown',
        reply_markup=reply_markup
    )


async def search_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /search - –ø–æ–∏—Å–∫ –ø–æ –∏—Å—Ç–æ—Ä–∏–∏"""
    user_id = update.effective_user.id

    if not context.args:
        await update.message.reply_text(
            "üîç **–ü–æ–∏—Å–∫ –ø–æ –∏—Å—Ç–æ—Ä–∏–∏**\n\n"
            "–ò—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏–µ: `/search <–∑–∞–ø—Ä–æ—Å>`\n\n"
            "–ü—Ä–∏–º–µ—Ä—ã:\n"
            "‚Ä¢ `/search —Ç—Ä–µ—â–∏–Ω–∞` - –Ω–∞–π—Ç–∏ –≤—Å–µ —Å–æ–æ–±—â–µ–Ω–∏—è –ø—Ä–æ —Ç—Ä–µ—â–∏–Ω—ã\n"
            "‚Ä¢ `/search –°–ü 63` - –Ω–∞–π—Ç–∏ –≤—Å–µ —É–ø–æ–º–∏–Ω–∞–Ω–∏—è –°–ü 63.13330.2018\n"
            "‚Ä¢ `/search –±–µ—Ç–æ–Ω B25` - –Ω–∞–π—Ç–∏ —Å–æ–æ–±—â–µ–Ω–∏—è –ø—Ä–æ –±–µ—Ç–æ–Ω B25",
            parse_mode='Markdown'
        )
        return

    query = " ".join(context.args)
    results = search_in_history(user_id, query, limit=5)

    if not results:
        await update.message.reply_text(
            f"‚ùå –ü–æ –∑–∞–ø—Ä–æ—Å—É ¬´{query}¬ª –Ω–∏—á–µ–≥–æ –Ω–µ –Ω–∞–π–¥–µ–Ω–æ.\n\n"
            "–ü–æ–ø—Ä–æ–±—É–π—Ç–µ –∏–∑–º–µ–Ω–∏—Ç—å –∑–∞–ø—Ä–æ—Å –∏–ª–∏ –ø—Ä–æ–≤–µ—Ä—å—Ç–µ –∏—Å—Ç–æ—Ä–∏—é —á–µ—Ä–µ–∑ /history"
        )
        return

    response = f"üîç **–†–µ–∑—É–ª—å—Ç–∞—Ç—ã –ø–æ–∏—Å–∫–∞ –ø–æ –∑–∞–ø—Ä–æ—Å—É ¬´{query}¬ª**\n\n"
    response += f"–ù–∞–π–¥–µ–Ω–æ: {len(results)} —Å–æ–æ–±—â–µ–Ω–∏–π\n\n"

    for i, msg in enumerate(results, 1):
        role_emoji = "üë§" if msg['role'] == 'user' else "ü§ñ"
        timestamp = datetime.fromisoformat(msg['timestamp']).strftime('%d.%m %H:%M')
        content = msg['content'][:150] + "..." if len(msg['content']) > 150 else msg['content']

        response += f"{i}. {role_emoji} **{timestamp}**\n{content}\n\n"

    response += f"\n–ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ /history –¥–ª—è –ø—Ä–æ—Å–º–æ—Ç—Ä–∞ –ø–æ–ª–Ω–æ–π –∏—Å—Ç–æ—Ä–∏–∏"

    await update.message.reply_text(response, parse_mode='Markdown')


async def recommendations_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /recommendations - –ø–µ—Ä—Å–æ–Ω–∞–ª—å–Ω—ã–µ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏"""
    user_id = update.effective_user.id

    recs = get_recommendations(user_id)

    if not recs['recommendations'] and not recs['popular_topics']:
        await update.message.reply_text(
            "üí° **–†–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏**\n\n"
            "–ü–æ–∫–∞ –Ω–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –¥–∞–Ω–Ω—ã—Ö –¥–ª—è –ø–µ—Ä—Å–æ–Ω–∞–ª—å–Ω—ã—Ö —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–π.\n"
            "–ü—Ä–æ–¥–æ–ª–∂–∞–π—Ç–µ –æ–±—â–µ–Ω–∏–µ —Å –±–æ—Ç–æ–º!"
        )
        return

    response = "üí° **–ü–µ—Ä—Å–æ–Ω–∞–ª—å–Ω—ã–µ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏**\n\n"

    if recs['recommendations']:
        response += "**–ù–∞ –æ—Å–Ω–æ–≤–µ –≤–∞—à–∏—Ö –∏–Ω—Ç–µ—Ä–µ—Å–æ–≤:**\n\n"
        for rec in recs['recommendations']:
            if rec['type'] == 'related_regulation':
                response += f"üìö [{rec['code']}]({REGULATIONS[rec['code']]['url']}) - {rec['title']}\n"
                response += f"_{rec['reason']}_\n\n"

    if recs['popular_topics']:
        response += "\n**–í–∞—à–∏ –ø–æ–ø—É–ª—è—Ä–Ω—ã–µ —Ç–µ–º—ã:**\n\n"
        for topic in recs['popular_topics']:
            emoji_map = {
                '–Ω–æ—Ä–º–∞—Ç–∏–≤': 'üìÑ',
                '–º–∞—Ç–µ—Ä–∏–∞–ª': 'üß±',
                '–∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏—è': 'üèóÔ∏è'
            }
            emoji = emoji_map.get(topic['category'], 'üìå')
            response += f"{emoji} {topic['topic'].capitalize()} - {topic['mentions']} —É–ø–æ–º–∏–Ω–∞–Ω–∏–π\n"

    await update.message.reply_text(response, parse_mode='Markdown', disable_web_page_preview=True)


async def updates_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /updates - –ø—Ä–æ–≤–µ—Ä–∏—Ç—å –æ–±–Ω–æ–≤–ª–µ–Ω–∏—è –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤"""
    recent_updates = REGULATIONS_UPDATES['recent']

    if not recent_updates:
        await update.message.reply_text(
            "‚úÖ –í—Å–µ –Ω–æ—Ä–º–∞—Ç–∏–≤—ã –∞–∫—Ç—É–∞–ª—å–Ω—ã.\n"
            "–ù–æ–≤—ã—Ö –æ–±–Ω–æ–≤–ª–µ–Ω–∏–π –Ω–µ –æ–±–Ω–∞—Ä—É–∂–µ–Ω–æ."
        )
        return

    response = "üÜï **–ù–µ–¥–∞–≤–Ω–∏–µ –æ–±–Ω–æ–≤–ª–µ–Ω–∏—è –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤**\n\n"

    for upd in recent_updates:
        type_emoji = "üÜï" if upd['type'] == '–Ω–æ–≤—ã–π' else "‚ôªÔ∏è"
        update_date = datetime.fromisoformat(upd['date']).strftime('%d.%m.%Y')

        response += f"{type_emoji} **{upd['code']}** - {upd['title']}\n"
        response += f"–î–∞—Ç–∞: {update_date}\n"
        response += f"–ò–∑–º–µ–Ω–µ–Ω–∏—è: {upd['changes']}\n\n"

    response += "\nüí° –ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ /regulations –¥–ª—è –ø—Ä–æ—Å–º–æ—Ç—Ä–∞ –≤—Å–µ—Ö –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤"

    await update.message.reply_text(response, parse_mode='Markdown')


async def requirements2025_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /requirements2025 - –∞–∫—Ç—É–∞–ª—å–Ω—ã–µ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è 2025"""
    if not REGULATIONS_2025_AVAILABLE:
        await update.message.reply_text(
            "‚ö†Ô∏è –ë–∞–∑–∞ –∞–∫—Ç—É–∞–ª—å–Ω—ã—Ö —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π 2025 –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∞.\n"
            "–û–±—Ä–∞—Ç–∏—Ç–µ—Å—å –∫ –∞–¥–º–∏–Ω–∏—Å—Ç—Ä–∞—Ç–æ—Ä—É."
        )
        return

    keyboard = [
        [InlineKeyboardButton("üìã –û—Å–Ω–æ–≤–Ω—ã–µ –∑–∞–∫–æ–Ω—ã", callback_data="req2025_laws")],
        [InlineKeyboardButton("üîÑ –ü—Ä–æ—Ü–µ–¥—É—Ä—ã 2025", callback_data="req2025_procedures")],
        [InlineKeyboardButton("üë∑ –°–†–û –∏ –∫–≤–∞–ª–∏—Ñ–∏–∫–∞—Ü–∏—è", callback_data="req2025_sro")],
        [InlineKeyboardButton("üíª –¢–ò–ú/BIM", callback_data="req2025_bim")],
        [InlineKeyboardButton("üè≠ –ü—Ä–æ–º—ã—à–ª–µ–Ω–Ω–æ–µ", callback_data="req2025_industrial")],
        [InlineKeyboardButton("üèòÔ∏è –ì—Ä–∞–∂–¥–∞–Ω—Å–∫–æ–µ", callback_data="req2025_civil")],
        [InlineKeyboardButton("üè¢ –ö–æ–º–º–µ—Ä—á–µ—Å–∫–æ–µ", callback_data="req2025_commercial")],
        [InlineKeyboardButton("‚úÖ –ß–µ–∫-–ª–∏—Å—Ç –ø—Ä–æ–≤–µ—Ä–æ–∫", callback_data="req2025_checklist")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)

    text = """üìö **–ê–ö–¢–£–ê–õ–¨–ù–´–ï –¢–†–ï–ë–û–í–ê–ù–ò–Ø 2025**

–ë–∞–∑–∞ –ø–æ–ª–Ω–æ—Å—Ç—å—é –æ–±–Ω–æ–≤–ª–µ–Ω–∞ –Ω–∞ 2025-2026 –≥–æ–¥!

–í—ã–±–µ—Ä–∏—Ç–µ —Ä–∞–∑–¥–µ–ª –¥–ª—è –ø–æ–¥—Ä–æ–±–Ω–æ–π –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏–∏:

üìã –û—Å–Ω–æ–≤–Ω—ã–µ –∑–∞–∫–æ–Ω—ã –†–§ (8 –§–ó)
üîÑ –û–±—è–∑–∞—Ç–µ–ª—å–Ω—ã–µ –ø—Ä–æ—Ü–µ–¥—É—Ä—ã 2025
üë∑ –°–†–û –∏ –∫–≤–∞–ª–∏—Ñ–∏–∫–∞—Ü–∏—è
üíª –¢–ò–ú/BIM —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è
üè≠ –ü—Ä–æ–º—ã—à–ª–µ–Ω–Ω–æ–µ —Å—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤–æ
üèòÔ∏è –ì—Ä–∞–∂–¥–∞–Ω—Å–∫–æ–µ —Å—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤–æ
üè¢ –ö–æ–º–º–µ—Ä—á–µ—Å–∫–æ–µ —Å—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤–æ
‚úÖ –ß–µ–∫-–ª–∏—Å—Ç –µ–∂–µ–¥–Ω–µ–≤–Ω—ã—Ö –ø—Ä–æ–≤–µ—Ä–æ–∫"""

    await update.message.reply_text(text, parse_mode='Markdown', reply_markup=reply_markup)


async def laws_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /laws - –æ—Å–Ω–æ–≤–Ω—ã–µ –∑–∞–∫–æ–Ω—ã"""
    if not REGULATIONS_2025_AVAILABLE:
        await update.message.reply_text("‚ö†Ô∏è –ë–∞–∑–∞ –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∞")
        return

    text = "üìã **–û–°–ù–û–í–ù–´–ï –ó–ê–ö–û–ù–´ –†–§ 2025**\n\n"

    for code, data in FEDERAL_LAWS.items():
        text += f"**{code}** - [{data['title']}]({data['url']})\n"
        text += f"_{data['scope']}_\n\n"

    await update.message.reply_text(text, parse_mode='Markdown', disable_web_page_preview=True)


async def checklist_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /checklist - –µ–∂–µ–¥–Ω–µ–≤–Ω—ã–π —á–µ–∫-–ª–∏—Å—Ç"""
    text = """‚úÖ **–ß–ï–ö-–õ–ò–°–¢ –ï–ñ–ï–î–ù–ï–í–ù–´–• –ü–†–û–í–ï–†–û–ö –ù–ê –û–ë–™–ï–ö–¢–ï**

–ß—Ç–æ –ø—Ä–æ–≤–µ—Ä—è—Ç—å –∫–∞–∂–¥—ã–π –¥–µ–Ω—å:

1Ô∏è‚É£ –ï—Å—Ç—å –ª–∏ –¥–µ–π—Å—Ç–≤—É—é—â–µ–µ —Ä–∞–∑—Ä–µ—à–µ–Ω–∏–µ –Ω–∞ —Å—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤–æ (–†–ù–°)?
2Ô∏è‚É£ –ü—Ä–æ—à–ª–∞ –ª–∏ –ü–î —ç–∫—Å–ø–µ—Ä—Ç–∏–∑—É?
3Ô∏è‚É£ –ï—Å—Ç—å –ª–∏ —á–ª–µ–Ω—Å—Ç–≤–æ –≤ –°–†–û —É –≥–µ–Ω–ø–æ–¥—Ä—è–¥—á–∏–∫–∞?
4Ô∏è‚É£ –í–µ–¥—ë—Ç—Å—è –ª–∏ –∏—Å–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω–∞—è –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ü–∏—è (–∞–∫—Ç—ã –æ—Å–≤–∏–¥–µ—Ç–µ–ª—å—Å—Ç–≤–æ–≤–∞–Ω–∏—è —Å–∫—Ä—ã—Ç—ã—Ö —Ä–∞–±–æ—Ç, –∂—É—Ä–Ω–∞–ª—ã)?
5Ô∏è‚É£ –ù–∞–∑–Ω–∞—á–µ–Ω—ã –ª–∏ –ª–∏—Ü–∞ –ø–æ —Ç–µ—Ö–Ω–∞–¥–∑–æ—Ä—É –∏ —Å—Ç—Ä–æ–π–∫–æ–Ω—Ç—Ä–æ–ª—é?
6Ô∏è‚É£ –î–ª—è –ø—Ä–æ–º–∫–∏ - –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω–æ –ª–∏ –û–ü–û? –ï—Å—Ç—å –ª–∏ –°–ó–ó?
7Ô∏è‚É£ –í–µ–¥—ë—Ç—Å—è –ª–∏ –∂—É—Ä–Ω–∞–ª –≤—Ö–æ–¥–Ω–æ–≥–æ –∫–æ–Ω—Ç—Ä–æ–ª—è –º–∞—Ç–µ—Ä–∏–∞–ª–æ–≤?
8Ô∏è‚É£ –ï—Å—Ç—å –ª–∏ –¥–æ–ø—É—Å–∫–∏ —É —Å–≤–∞—Ä—â–∏–∫–æ–≤, —Å—Ç—Ä–æ–ø–∞–ª—å—â–∏–∫–æ–≤, –∫—Ä–∞–Ω–æ–≤—â–∏–∫–æ–≤ (—É–¥–æ—Å—Ç–æ–≤–µ—Ä–µ–Ω–∏—è + –ø—Ä–æ—Ç–æ–∫–æ–ª—ã –ù–û–ö)?
9Ô∏è‚É£ –ü–æ–¥–ø–∏—Å–∞–Ω—ã –ª–∏ –∞–∫—Ç—ã –ö–°-2, –ö–°-3 –µ–∂–µ–º–µ—Å—è—á–Ω–æ?

**–ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ —ç—Ç–æ—Ç —Å–ø–∏—Å–æ–∫ –¥–ª—è –∫–æ–Ω—Ç—Ä–æ–ª—è –æ–±—ä–µ–∫—Ç–∞!**

üí° –°–æ—Ö—Ä–∞–Ω–∏—Ç–µ –≤ –∑–∞–∫–ª–∞–¥–∫–∏ –¥–ª—è –±—ã—Å—Ç—Ä–æ–≥–æ –¥–æ—Å—Ç—É–ø–∞"""

    await update.message.reply_text(text, parse_mode='Markdown')


# === –ü–†–ê–ö–¢–ò–ß–ï–°–ö–ò–ï –ó–ù–ê–ù–ò–Ø 2025 ===

async def hse_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /hse - –æ—Ö—Ä–∞–Ω–∞ —Ç—Ä—É–¥–∞ –∏ —Ç–µ—Ö–Ω–∏–∫–∞ –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç–∏"""
    if not PRACTICAL_KNOWLEDGE_AVAILABLE:
        await update.message.reply_text("‚ö†Ô∏è –ë–∞–∑–∞ –ø—Ä–∞–∫—Ç–∏—á–µ—Å–∫–∏—Ö –∑–Ω–∞–Ω–∏–π –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∞")
        return

    text = """ü¶∫ **–û–•–†–ê–ù–ê –¢–†–£–î–ê –ò –¢–ï–•–ù–ò–ö–ê –ë–ï–ó–û–ü–ê–°–ù–û–°–¢–ò (HSE)**

üìå **–î–æ—Å—Ç—É–ø–Ω—ã–µ —Ä–∞–∑–¥–µ–ª—ã:**

1Ô∏è‚É£ **–†–∞–±–æ—Ç–∞ –Ω–∞ –≤—ã—Å–æ—Ç–µ** (–ü—Ä–∏–∫–∞–∑ –ú–∏–Ω—Ç—Ä—É–¥–∞ ‚Ññ 782–Ω)
   ‚Ä¢ –ì—Ä—É–ø–ø—ã –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç–∏ (1, 2, 3)
   ‚Ä¢ –û–±—è–∑–∞—Ç–µ–ª—å–Ω—ã–µ –°–ò–ó
   ‚Ä¢ –ù–∞—Ä—è–¥-–¥–æ–ø—É—Å–∫

2Ô∏è‚É£ **–≠–ª–µ–∫—Ç—Ä–æ–±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç—å**
   ‚Ä¢ –ì—Ä—É–ø–ø—ã –¥–æ 1000–í / –≤—ã—à–µ 1000–í
   ‚Ä¢ –í—Ä–µ–º–µ–Ω–Ω—ã–µ —ç–ª–µ–∫—Ç—Ä–æ—Å–µ—Ç–∏ –Ω–∞ —Å—Ç—Ä–æ–π–∫–µ

3Ô∏è‚É£ **–ü–æ–≥—Ä—É–∑–æ—á–Ω–æ-—Ä–∞–∑–≥—Ä—É–∑–æ—á–Ω—ã–µ —Ä–∞–±–æ—Ç—ã**
   ‚Ä¢ –°—Ö–µ–º—ã —Å—Ç—Ä–æ–ø–æ–≤–∫–∏
   ‚Ä¢ –°–∏–≥–Ω–∞–ª—ã –∫—Ä–∞–Ω–æ–≤—â–∏–∫—É

üí° **–ü—Ä–∏–º–µ—Ä—ã –≤–æ–ø—Ä–æ—Å–æ–≤:**
‚Ä¢ –ö–∞–∫–∞—è –≥—Ä—É–ø–ø–∞ –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç–∏ –Ω—É–∂–Ω–∞ –¥–ª—è —Ä–∞–±–æ—Ç—ã –Ω–∞ –≤—ã—Å–æ—Ç–µ 5 –º?
‚Ä¢ –ö–∞–∫–∏–µ –°–ò–ó –æ–±—è–∑–∞—Ç–µ–ª—å–Ω—ã –ø—Ä–∏ —Ä–∞–±–æ—Ç–µ –Ω–∞ –≤—ã—Å–æ—Ç–µ?
‚Ä¢ –ö—Ç–æ –º–æ–∂–µ—Ç —Ä–∞–±–æ—Ç–∞—Ç—å —Å —ç–ª–µ–∫—Ç—Ä–æ–∏–Ω—Å—Ç—Ä—É–º–µ–Ω—Ç–æ–º?
‚Ä¢ –ö–∞–∫ –ø—Ä–∞–≤–∏–ª—å–Ω–æ –∑–∞—Å—Ç—Ä–æ–ø–∏—Ç—å –±–∞–ª–∫—É?

üìñ –ü—Ä–æ—Å—Ç–æ –∑–∞–¥–∞–π—Ç–µ –≤–æ–ø—Ä–æ—Å, –∏ —è –¥–∞–º –ø–æ–¥—Ä–æ–±–Ω—ã–π –æ—Ç–≤–µ—Ç —Å–æ —Å—Å—ã–ª–∫–∞–º–∏ –Ω–∞ –Ω–æ—Ä–º–∞—Ç–∏–≤—ã!"""

    await update.message.reply_text(text, parse_mode='Markdown')


async def technology_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /technology - —Ç–µ—Ö–Ω–æ–ª–æ–≥–∏—è —Å—Ç—Ä–æ–∏—Ç–µ–ª—å–Ω–æ–≥–æ –ø—Ä–æ–∏–∑–≤–æ–¥—Å—Ç–≤–∞"""
    if not PRACTICAL_KNOWLEDGE_AVAILABLE:
        await update.message.reply_text("‚ö†Ô∏è –ë–∞–∑–∞ –ø—Ä–∞–∫—Ç–∏—á–µ—Å–∫–∏—Ö –∑–Ω–∞–Ω–∏–π –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∞")
        return

    text = """üèóÔ∏è **–¢–ï–•–ù–û–õ–û–ì–ò–Ø –°–¢–†–û–ò–¢–ï–õ–¨–ù–û–ì–û –ü–†–û–ò–ó–í–û–î–°–¢–í–ê**

üìå **–î–æ—Å—Ç—É–ø–Ω—ã–µ —Ä–∞–∑–¥–µ–ª—ã:**

1Ô∏è‚É£ **–ë–µ—Ç–æ–Ω–Ω—ã–µ —Ä–∞–±–æ—Ç—ã**
   ‚Ä¢ –ù–∞–±–æ—Ä –ø—Ä–æ—á–Ω–æ—Å—Ç–∏ (–≥—Ä–∞—Ñ–∏–∫)
   ‚Ä¢ –ó–∏–º–Ω–µ–µ –±–µ—Ç–æ–Ω–∏—Ä–æ–≤–∞–Ω–∏–µ (–ø—Ä–∏ -10¬∞C, -20¬∞C)
   ‚Ä¢ –ö–æ–Ω—Ç—Ä–æ–ª—å –ø—Ä–æ—á–Ω–æ—Å—Ç–∏ (–∫—É–±–∏–∫–∏, –º–æ–ª–æ—Ç–æ–∫ –§–∏–∑–¥–µ–ª—è)

2Ô∏è‚É£ **–ê—Ä–º–∞—Ç—É—Ä–Ω—ã–µ —Ä–∞–±–æ—Ç—ã**
   ‚Ä¢ –†–∞—Å—á–µ—Ç –Ω–∞—Ö–ª–µ—Å—Ç–æ–≤ (–¥–ª—è –ê400, –ê500)
   ‚Ä¢ –ó–∞—â–∏—Ç–Ω—ã–π —Å–ª–æ–π –±–µ—Ç–æ–Ω–∞
   ‚Ä¢ –°—Ö–µ–º—ã –≤—è–∑–∫–∏

3Ô∏è‚É£ **–í—Ö–æ–¥–Ω–æ–π –∫–æ–Ω—Ç—Ä–æ–ª—å –º–∞—Ç–µ—Ä–∏–∞–ª–æ–≤**
   ‚Ä¢ –ê—Ä–º–∞—Ç—É—Ä–∞ (–¥–æ–ø—É—Å–∫–∏ –ø–æ –¥–∏–∞–º–µ—Ç—Ä—É)
   ‚Ä¢ –ë–µ—Ç–æ–Ω (–ø—Ä–æ—á–Ω–æ—Å—Ç—å, –ø–æ–¥–≤–∏–∂–Ω–æ—Å—Ç—å)
   ‚Ä¢ –ö–∏—Ä–ø–∏—á (–≥–µ–æ–º–µ—Ç—Ä–∏—è, –ø—Ä–æ—á–Ω–æ—Å—Ç—å)

üí° **–ü—Ä–∏–º–µ—Ä—ã –≤–æ–ø—Ä–æ—Å–æ–≤:**
‚Ä¢ –°–∫–æ–ª—å–∫–æ –¥–Ω–µ–π –Ω–∞–±–∏—Ä–∞–µ—Ç –ø—Ä–æ—á–Ω–æ—Å—Ç—å –±–µ—Ç–æ–Ω –∑–∏–º–æ–π?
‚Ä¢ –ö–∞–∫–æ–π –Ω–∞—Ö–ª–µ—Å—Ç –∞—Ä–º–∞—Ç—É—Ä—ã –ê400 d=16 –º–º?
‚Ä¢ –ö–∞–∫–∏–µ –¥–æ–±–∞–≤–∫–∏ –∏—Å–ø–æ–ª—å–∑–æ–≤–∞—Ç—å –ø—Ä–∏ -15¬∞C?
‚Ä¢ –ö–∞–∫ –∫–æ–Ω—Ç—Ä–æ–ª–∏—Ä–æ–≤–∞—Ç—å –ø—Ä–æ—á–Ω–æ—Å—Ç—å –±–µ—Ç–æ–Ω–∞ –Ω–∞ –æ–±—ä–µ–∫—Ç–µ?

üìñ –ü—Ä–æ—Å—Ç–æ –∑–∞–¥–∞–π—Ç–µ –≤–æ–ø—Ä–æ—Å, –∏ —è –¥–∞–º –ø—Ä–∞–∫—Ç–∏—á–µ—Å–∫–∏–π –æ—Ç–≤–µ—Ç —Å —Ä–∞—Å—á–µ—Ç–∞–º–∏!"""

    await update.message.reply_text(text, parse_mode='Markdown')


async def estimating_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /estimating - —Å–º–µ—Ç–Ω–æ–µ –¥–µ–ª–æ –∏ —Ñ–∏–Ω–∞–Ω—Å—ã"""
    if not PRACTICAL_KNOWLEDGE_AVAILABLE:
        await update.message.reply_text("‚ö†Ô∏è –ë–∞–∑–∞ –ø—Ä–∞–∫—Ç–∏—á–µ—Å–∫–∏—Ö –∑–Ω–∞–Ω–∏–π –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∞")
        return

    text = """üí∞ **–°–ú–ï–¢–ù–û–ï –î–ï–õ–û –ò –§–ò–ù–ê–ù–°–´**

üìå **–î–æ—Å—Ç—É–ø–Ω—ã–µ —Ä–∞–∑–¥–µ–ª—ã:**

1Ô∏è‚É£ **–ê–∫—Ç—ã –ö–°-2 –∏ –ö–°-3**
   ‚Ä¢ –¢–∏–ø–∏—á–Ω—ã–µ –æ—à–∏–±–∫–∏
   ‚Ä¢ –ü–æ—Ä—è–¥–æ–∫ –ø–æ–¥–ø–∏—Å–∞–Ω–∏—è
   ‚Ä¢ –°—Ä–æ–∫–∏ —Å–æ–≥–ª–∞—Å–æ–≤–∞–Ω–∏—è

2Ô∏è‚É£ **–î–∞–≤–∞–ª—å—á–µ—Å–∫–∏–µ –º–∞—Ç–µ—Ä–∏–∞–ª—ã**
   ‚Ä¢ –§–æ—Ä–º–∞ –ú-29
   ‚Ä¢ –£—á–µ—Ç –≤ —Å–º–µ—Ç–µ
   ‚Ä¢ –ö—Ç–æ –ø–ª–∞—Ç–∏—Ç –∑–∞ –¥–æ—Å—Ç–∞–≤–∫—É?

3Ô∏è‚É£ **–ù–µ–ø—Ä–µ–¥–≤–∏–¥–µ–Ω–Ω—ã–µ —Ä–∞—Å—Ö–æ–¥—ã**
   ‚Ä¢ –ü—Ä–æ—Ü–µ–Ω—Ç –æ—Ç —Å–º–µ—Ç—ã
   ‚Ä¢ –î–æ—Å—É–¥–µ–±–Ω–æ–µ —É—Ä–µ–≥—É–ª–∏—Ä–æ–≤–∞–Ω–∏–µ
   ‚Ä¢ –î–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω—ã–µ —Å–æ–≥–ª–∞—à–µ–Ω–∏—è

üí° **–ü—Ä–∏–º–µ—Ä—ã –≤–æ–ø—Ä–æ—Å–æ–≤:**
‚Ä¢ –ö–∞–∫–∏–µ –æ—à–∏–±–∫–∏ –±—ã–≤–∞—é—Ç –≤ –ö–°-2?
‚Ä¢ –ö—Ç–æ –ø–ª–∞—Ç–∏—Ç –∑–∞ –¥–æ—Å—Ç–∞–≤–∫—É –¥–∞–≤–∞–ª—å—á–µ—Å–∫–∏—Ö –º–∞—Ç–µ—Ä–∏–∞–ª–æ–≤?
‚Ä¢ –°–∫–æ–ª—å–∫–æ –ø—Ä–æ—Ü–µ–Ω—Ç–æ–≤ –Ω–µ–ø—Ä–µ–¥–≤–∏–¥–µ–Ω–Ω—ã—Ö —Ä–∞—Å—Ö–æ–¥–æ–≤?
‚Ä¢ –ö–∞–∫ –æ—Ñ–æ—Ä–º–∏—Ç—å –¥–æ–ø.—Å–æ–≥–ª–∞—à–µ–Ω–∏–µ –Ω–∞ –∏–∑–º–µ–Ω–µ–Ω–∏–µ –æ–±—ä–µ–º–æ–≤?

üìñ –ü—Ä–æ—Å—Ç–æ –∑–∞–¥–∞–π—Ç–µ –≤–æ–ø—Ä–æ—Å, –∏ —è –ø–æ–º–æ–≥—É —Ä–∞–∑–æ–±—Ä–∞—Ç—å—Å—è —Å —Ñ–∏–Ω–∞–Ω—Å–∞–º–∏!"""

    await update.message.reply_text(text, parse_mode='Markdown')


async def legal_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /legal - —é—Ä–∏–¥–∏—á–µ—Å–∫–∏–µ –≤–æ–ø—Ä–æ—Å—ã –∏ –ø—Ä–µ—Ç–µ–Ω–∑–∏–æ–Ω–Ω–∞—è —Ä–∞–±–æ—Ç–∞"""
    if not PRACTICAL_KNOWLEDGE_AVAILABLE:
        await update.message.reply_text("‚ö†Ô∏è –ë–∞–∑–∞ –ø—Ä–∞–∫—Ç–∏—á–µ—Å–∫–∏—Ö –∑–Ω–∞–Ω–∏–π –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∞")
        return

    text = """‚öñÔ∏è **–Æ–†–ò–î–ò–ß–ï–°–ö–ò–ï –í–û–ü–†–û–°–´ –ò –ü–†–ï–¢–ï–ù–ó–ò–û–ù–ù–ê–Ø –†–ê–ë–û–¢–ê**

üìå **–î–æ—Å—Ç—É–ø–Ω—ã–µ —Ä–∞–∑–¥–µ–ª—ã:**

1Ô∏è‚É£ **–°—Ä–æ–∫–∏ –∏ –Ω–µ—É—Å—Ç–æ–π–∫–∏**
   ‚Ä¢ –ì–ö –†–§ —Å—Ç. 330
   ‚Ä¢ –†–∞—Å—á–µ—Ç –ø–µ–Ω–µ–π
   ‚Ä¢ –°–ø–æ—Å–æ–±—ã —Å–Ω–∏–∂–µ–Ω–∏—è

2Ô∏è‚É£ **–û—Å–≤–∏–¥–µ—Ç–µ–ª—å—Å—Ç–≤–æ–≤–∞–Ω–∏–µ —Å–∫—Ä—ã—Ç—ã—Ö —Ä–∞–±–æ—Ç**
   ‚Ä¢ –ö—Ç–æ –ø–æ–¥–ø–∏—Å—ã–≤–∞–µ—Ç –∞–∫—Ç—ã?
   ‚Ä¢ –ß—Ç–æ –µ—Å–ª–∏ –∑–∞–∫–∞–∑—á–∏–∫ –Ω–µ —è–≤–ª—è–µ—Ç—Å—è?
   ‚Ä¢ 3 –¥–Ω—è –Ω–∞ –æ—Å–≤–∏–¥–µ—Ç–µ–ª—å—Å—Ç–≤–æ–≤–∞–Ω–∏–µ

3Ô∏è‚É£ **–ü—Ä–µ—Ç–µ–Ω–∑–∏–æ–Ω–Ω–∞—è —Ä–∞–±–æ—Ç–∞**
   ‚Ä¢ –î–æ—Å—É–¥–µ–±–Ω–æ–µ —É—Ä–µ–≥—É–ª–∏—Ä–æ–≤–∞–Ω–∏–µ
   ‚Ä¢ –û–±—Ä–∞–∑—Ü—ã –ø–∏—Å–µ–º
   ‚Ä¢ –°—Ä–æ–∫–∏ –æ—Ç–≤–µ—Ç–æ–≤

üí° **–ü—Ä–∏–º–µ—Ä—ã –≤–æ–ø—Ä–æ—Å–æ–≤:**
‚Ä¢ –ö–∞–∫ —Ä–∞—Å—Å—á–∏—Ç–∞—Ç—å –Ω–µ—É—Å—Ç–æ–π–∫—É –∑–∞ –ø—Ä–æ—Å—Ä–æ—á–∫—É?
‚Ä¢ –ß—Ç–æ –¥–µ–ª–∞—Ç—å, –µ—Å–ª–∏ –∑–∞–∫–∞–∑—á–∏–∫ –Ω–µ –ø–æ–¥–ø–∏—Å—ã–≤–∞–µ—Ç –∞–∫—Ç —Å–∫—Ä—ã—Ç—ã—Ö —Ä–∞–±–æ—Ç?
‚Ä¢ –ö–∞–∫ –Ω–∞–ø–∏—Å–∞—Ç—å –ø—Ä–µ—Ç–µ–Ω–∑–∏—é –Ω–∞ –Ω–µ–≤—ã–ø–ª–∞—Ç—É?
‚Ä¢ –ú–æ–∂–Ω–æ –ª–∏ —Å–Ω–∏–∑–∏—Ç—å –ø–µ–Ω–∏?

üìñ –ü—Ä–æ—Å—Ç–æ –∑–∞–¥–∞–π—Ç–µ –≤–æ–ø—Ä–æ—Å, –∏ —è –ø–æ–º–æ–≥—É —Å —é—Ä–∏–¥–∏—á–µ—Å–∫–∏–º–∏ —Ç–æ–Ω–∫–æ—Å—Ç—è–º–∏!"""

    await update.message.reply_text(text, parse_mode='Markdown')


async def management_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /management - —É–ø—Ä–∞–≤–ª–µ–Ω–∏–µ –ø—Ä–æ–µ–∫—Ç–∞–º–∏ –∏ soft skills"""
    if not PRACTICAL_KNOWLEDGE_AVAILABLE:
        await update.message.reply_text("‚ö†Ô∏è –ë–∞–∑–∞ –ø—Ä–∞–∫—Ç–∏—á–µ—Å–∫–∏—Ö –∑–Ω–∞–Ω–∏–π –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∞")
        return

    text = """üìä **–£–ü–†–ê–í–õ–ï–ù–ò–ï –ü–†–û–ï–ö–¢–ê–ú–ò –ò SOFT SKILLS**

üìå **–î–æ—Å—Ç—É–ø–Ω—ã–µ —Ä–∞–∑–¥–µ–ª—ã:**

1Ô∏è‚É£ **–ü–ª–∞–Ω–∏—Ä–æ–≤–∞–Ω–∏–µ —Ä–∞–±–æ—Ç**
   ‚Ä¢ –î–∏–∞–≥—Ä–∞–º–º–∞ –ì–∞–Ω—Ç–∞
   ‚Ä¢ –°–µ—Ç–µ–≤—ã–µ –≥—Ä–∞—Ñ–∏–∫–∏
   ‚Ä¢ –ö—Ä–∏—Ç–∏—á–µ—Å–∫–∏–π –ø—É—Ç—å

2Ô∏è‚É£ **–ü—Ä–æ—Ç–æ–∫–æ–ª—ã —Å–æ–≤–µ—â–∞–Ω–∏–π**
   ‚Ä¢ –û–±—è–∑–∞—Ç–µ–ª—å–Ω—ã–µ –ø—É–Ω–∫—Ç—ã
   ‚Ä¢ –§–æ—Ä–º–∞—Ç –æ—Ñ–æ—Ä–º–ª–µ–Ω–∏—è
   ‚Ä¢ –Æ—Ä–∏–¥–∏—á–µ—Å–∫–∞—è —Å–∏–ª–∞

3Ô∏è‚É£ **–†–∞—Å—á–µ—Ç —á–∏—Å–ª–µ–Ω–Ω–æ—Å—Ç–∏**
   ‚Ä¢ –§–æ—Ä–º—É–ª–∞ —Ä–∞—Å—á–µ—Ç–∞ —Ä–∞–±–æ—á–∏—Ö
   ‚Ä¢ –ö–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç —Å–æ–≤–º–µ—â–µ–Ω–∏—è
   ‚Ä¢ –£—á–µ—Ç —Å–º–µ–Ω–Ω–æ—Å—Ç–∏

4Ô∏è‚É£ **–†–∞–∑—Ä–µ—à–µ–Ω–∏–µ –∫–æ–Ω—Ñ–ª–∏–∫—Ç–æ–≤**
   ‚Ä¢ –° –∑–∞–∫–∞–∑—á–∏–∫–æ–º
   ‚Ä¢ –° —Å—É–±–ø–æ–¥—Ä—è–¥—á–∏–∫–æ–º
   ‚Ä¢ –í–Ω—É—Ç—Ä–∏ –∫–æ–º–∞–Ω–¥—ã

üí° **–ü—Ä–∏–º–µ—Ä—ã –≤–æ–ø—Ä–æ—Å–æ–≤:**
‚Ä¢ –ö–∞–∫ –ø–æ—Å—Ç—Ä–æ–∏—Ç—å –¥–∏–∞–≥—Ä–∞–º–º—É –ì–∞–Ω—Ç–∞ –¥–ª—è —Å—Ç—Ä–æ–π–∫–∏?
‚Ä¢ –ß—Ç–æ –æ–±—è–∑–∞—Ç–µ–ª—å–Ω–æ –¥–æ–ª–∂–Ω–æ –±—ã—Ç—å –≤ –ø—Ä–æ—Ç–æ–∫–æ–ª–µ —Å–æ–≤–µ—â–∞–Ω–∏—è?
‚Ä¢ –ö–∞–∫ —Ä–∞—Å—Å—á–∏—Ç–∞—Ç—å —Å–∫–æ–ª—å–∫–æ –Ω—É–∂–Ω–æ –∫–∞–º–µ–Ω—â–∏–∫–æ–≤?
‚Ä¢ –ß—Ç–æ –¥–µ–ª–∞—Ç—å –µ—Å–ª–∏ —Å—É–±–ø–æ–¥—Ä—è–¥—á–∏–∫ —Å—Ä—ã–≤–∞–µ—Ç —Å—Ä–æ–∫–∏?

üìñ –ü—Ä–æ—Å—Ç–æ –∑–∞–¥–∞–π—Ç–µ –≤–æ–ø—Ä–æ—Å, –∏ —è –ø–æ–º–æ–≥—É —Å —É–ø—Ä–∞–≤–ª–µ–Ω–∏–µ–º –ø—Ä–æ–µ–∫—Ç–æ–º!"""

    await update.message.reply_text(text, parse_mode='Markdown')


# === –ù–û–í–´–ï –ö–û–ú–ê–ù–î–´ v3.9 ===

async def templates_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /templates - –ø–æ–∫–∞–∑–∞—Ç—å —à–∞–±–ª–æ–Ω—ã –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤"""
    if not TEMPLATES_AVAILABLE:
        await update.message.reply_text("‚ö†Ô∏è –®–∞–±–ª–æ–Ω—ã –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ –Ω–µ–¥–æ—Å—Ç—É–ø–Ω—ã")
        return

    keyboard = []
    for template_id, info in DOCUMENT_TEMPLATES.items():
        keyboard.append([
            InlineKeyboardButton(
                text=info["name"],
                callback_data=f"template_{template_id}"
            )
        ])

    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text(
        "üìÑ **–®–ê–ë–õ–û–ù–´ –î–û–ö–£–ú–ï–ù–¢–û–í**\n\n"
        "–í—ã–±–µ—Ä–∏—Ç–µ —Ç–∏–ø –¥–æ–∫—É–º–µ–Ω—Ç–∞ –¥–ª—è –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏:",
        reply_markup=reply_markup,
        parse_mode="Markdown"
    )


async def handle_template_selection(update: Update, context: ContextTypes.DEFAULT_TYPE, template_id: str):
    """–û–±—Ä–∞–±–æ—Ç—á–∏–∫ –≤—ã–±–æ—Ä–∞ —à–∞–±–ª–æ–Ω–∞ –¥–æ–∫—É–º–µ–Ω—Ç–∞"""
    query = update.callback_query
    await query.answer()

    if template_id not in DOCUMENT_TEMPLATES:
        await query.edit_message_text("‚ùå –®–∞–±–ª–æ–Ω –Ω–µ –Ω–∞–π–¥–µ–Ω")
        return

    template_info = DOCUMENT_TEMPLATES[template_id]

    # –§–æ—Ä–º–∏—Ä—É–µ–º —Å–ø–∏—Å–æ–∫ –ø–∞—Ä–∞–º–µ—Ç—Ä–æ–≤ –¥–ª—è –æ—Ç–æ–±—Ä–∞–∂–µ–Ω–∏—è –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—é
    params_display = template_info.get("params_display", template_info["params"])
    params_list = "\n".join([f"‚Ä¢ {param}" for param in params_display])

    # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é –æ —à–∞–±–ª–æ–Ω–µ —Å –≤—ã–±–æ—Ä–æ–º –¥–µ–π—Å—Ç–≤–∏—è
    message_text = (
        f"üìÑ **{template_info['name']}**\n\n"
        f"{template_info['description']}\n\n"
        f"**–ù–µ–æ–±—Ö–æ–¥–∏–º—ã–µ –¥–∞–Ω–Ω—ã–µ:**\n{params_list}\n\n"
        "–í—ã–±–µ—Ä–∏—Ç–µ –¥–µ–π—Å—Ç–≤–∏–µ:"
    )

    # –ö–Ω–æ–ø–∫–∏ –≤—ã–±–æ—Ä–∞: —Å–∫–∞—á–∞—Ç—å –ø—É—Å—Ç–æ–π –∏–ª–∏ –∑–∞–ø–æ–ª–Ω–∏—Ç—å —Å –±–æ—Ç–æ–º
    keyboard = [
        [InlineKeyboardButton("üì• –°–∫–∞—á–∞—Ç—å –ø—É—Å—Ç–æ–π —à–∞–±–ª–æ–Ω", callback_data=f"download_empty_{template_id}")],
        [InlineKeyboardButton("‚úèÔ∏è –ó–∞–ø–æ–ª–Ω–∏—Ç—å —Å –±–æ—Ç–æ–º", callback_data=f"fill_{template_id}")],
        [InlineKeyboardButton("¬´ –ù–∞–∑–∞–¥ –∫ —à–∞–±–ª–æ–Ω–∞–º", callback_data="templates")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)

    await query.edit_message_text(
        message_text,
        reply_markup=reply_markup,
        parse_mode="Markdown"
    )


async def handle_download_empty_template(update: Update, context: ContextTypes.DEFAULT_TYPE, template_id: str):
    """–û–±—Ä–∞–±–æ—Ç—á–∏–∫ —Å–∫–∞—á–∏–≤–∞–Ω–∏—è –ø—É—Å—Ç–æ–≥–æ —à–∞–±–ª–æ–Ω–∞"""
    query = update.callback_query
    await query.answer()

    if template_id not in DOCUMENT_TEMPLATES:
        await query.edit_message_text("‚ùå –®–∞–±–ª–æ–Ω –Ω–µ –Ω–∞–π–¥–µ–Ω")
        return

    template_info = DOCUMENT_TEMPLATES[template_id]

    # –°–æ–∑–¥–∞–µ–º –ø—É—Å—Ç–æ–π –¥–æ–∫—É–º–µ–Ω—Ç —Å –∑–∞–ø–æ–ª–Ω–∏—Ç–µ–ª—è–º–∏
    params = {param: "___________" for param in template_info["params"]}

    # –ì–µ–Ω–µ—Ä–∏—Ä—É–µ–º –¥–æ–∫—É–º–µ–Ω—Ç
    result = generate_document(template_id, params)

    if not result["success"]:
        await query.edit_message_text(f"‚ùå –û—à–∏–±–∫–∞ –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ –¥–æ–∫—É–º–µ–Ω—Ç–∞: {result['error']}")
        return

    # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –ø—É—Å—Ç–æ–π —à–∞–±–ª–æ–Ω
    with open(result["filepath"], 'rb') as doc_file:
        await update.effective_chat.send_document(
            document=doc_file,
            filename=os.path.basename(result["filepath"]),
            caption=f"üìÑ **–ü—É—Å—Ç–æ–π —à–∞–±–ª–æ–Ω**: {template_info['name']}\n\n"
                    f"–ó–∞–ø–æ–ª–Ω–∏—Ç–µ –¥–æ–∫—É–º–µ–Ω—Ç –≤—Ä—É—á–Ω—É—é, –∑–∞–º–µ–Ω–∏–≤ –ø–æ–¥—á—ë—Ä–∫–∏–≤–∞–Ω–∏—è –Ω–∞ –≤–∞—à–∏ –¥–∞–Ω–Ω—ã–µ.",
            parse_mode="Markdown"
        )

    # –ö–Ω–æ–ø–∫–∞ –Ω–∞–∑–∞–¥ –∫ —à–∞–±–ª–æ–Ω–∞–º
    keyboard = [[InlineKeyboardButton("¬´ –ù–∞–∑–∞–¥ –∫ —à–∞–±–ª–æ–Ω–∞–º", callback_data="templates")]]
    reply_markup = InlineKeyboardMarkup(keyboard)

    await query.edit_message_text(
        f"‚úÖ –ü—É—Å—Ç–æ–π —à–∞–±–ª–æ–Ω –æ—Ç–ø—Ä–∞–≤–ª–µ–Ω!\n\n–ß—Ç–æ –¥–∞–ª—å—à–µ?",
        reply_markup=reply_markup
    )


async def projects_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /projects - –ø–æ–∫–∞–∑–∞—Ç—å –ø—Ä–æ–µ–∫—Ç—ã –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è"""
    if not PROJECTS_AVAILABLE:
        await update.message.reply_text("‚ö†Ô∏è –£–ø—Ä–∞–≤–ª–µ–Ω–∏–µ –ø—Ä–æ–µ–∫—Ç–∞–º–∏ –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–æ")
        return

    user_id = update.effective_user.id
    projects = get_user_projects(user_id)

    if not projects:
        await update.message.reply_text(
            "üìÅ –£ –≤–∞—Å –ø–æ–∫–∞ –Ω–µ—Ç –ø—Ä–æ–µ–∫—Ç–æ–≤.\n\n"
            "–°–æ–∑–¥–∞–π—Ç–µ –Ω–æ–≤—ã–π –ø—Ä–æ–µ–∫—Ç:\n"
            "`/new_project –ù–∞–∑–≤–∞–Ω–∏–µ –ø—Ä–æ–µ–∫—Ç–∞`",
            parse_mode="Markdown"
        )
        return

    keyboard = []
    for project_name in projects:
        keyboard.append([
            InlineKeyboardButton(
                text=f"üìÅ {project_name}",
                callback_data=f"proj_open_{project_name}"
            )
        ])

    keyboard.append([
        InlineKeyboardButton(
            text="‚ûï –°–æ–∑–¥–∞—Ç—å –Ω–æ–≤—ã–π –ø—Ä–æ–µ–∫—Ç",
            callback_data="proj_new"
        )
    ])

    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text(
        f"üìÅ **–í–ê–®–ò –ü–†–û–ï–ö–¢–´** ({len(projects)})\n\n"
        "–í—ã–±–µ—Ä–∏—Ç–µ –ø—Ä–æ–µ–∫—Ç –¥–ª—è —Ä–∞–±–æ—Ç—ã:",
        reply_markup=reply_markup,
        parse_mode="Markdown"
    )


async def new_project_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /new_project - —Å–æ–∑–¥–∞—Ç—å –Ω–æ–≤—ã–π –ø—Ä–æ–µ–∫—Ç"""
    if not PROJECTS_AVAILABLE:
        await update.message.reply_text("‚ö†Ô∏è –£–ø—Ä–∞–≤–ª–µ–Ω–∏–µ –ø—Ä–æ–µ–∫—Ç–∞–º–∏ –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–æ")
        return

    user_id = update.effective_user.id

    if not context.args:
        await update.message.reply_text(
            "‚ùå –£–∫–∞–∂–∏—Ç–µ –Ω–∞–∑–≤–∞–Ω–∏–µ –ø—Ä–æ–µ–∫—Ç–∞:\n"
            "`/new_project –ù–∞–∑–≤–∞–Ω–∏–µ –ø—Ä–æ–µ–∫—Ç–∞`",
            parse_mode="Markdown"
        )
        return

    project_name = " ".join(context.args)
    result = create_project(user_id, project_name)

    if result["success"]:
        context.user_data["current_project"] = project_name
        await update.message.reply_text(
            f"‚úÖ –ü—Ä–æ–µ–∫—Ç **{project_name}** —Å–æ–∑–¥–∞–Ω –∏ –∞–∫—Ç–∏–≤–∏—Ä–æ–≤–∞–Ω!\n\n"
            "üìå –í—Å–µ –≤–∞—à–∏ –≤–æ–ø—Ä–æ—Å—ã –∏ –æ—Ç–≤–µ—Ç—ã —Ç–µ–ø–µ—Ä—å —Å–æ—Ö—Ä–∞–Ω—è—é—Ç—Å—è –≤ –ø—Ä–æ–µ–∫—Ç.\n\n"
            "–î–æ—Å—Ç—É–ø–Ω—ã–µ –∫–æ–º–∞–Ω–¥—ã:\n"
            "‚Ä¢ `/project_info` - –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è –æ –ø—Ä–æ–µ–∫—Ç–µ\n"
            "‚Ä¢ `/project_log` - –∂—É—Ä–Ω–∞–ª —Ä–∞–±–æ—Ç—ã\n"
            "‚Ä¢ `/set_project` - –ø–µ—Ä–µ–∫–ª—é—á–∏—Ç—å –ø—Ä–æ–µ–∫—Ç\n"
            "‚Ä¢ `/projects` - —Å–ø–∏—Å–æ–∫ –ø—Ä–æ–µ–∫—Ç–æ–≤",
            parse_mode="Markdown"
        )
    else:
        await update.message.reply_text(
            f"‚ùå –û—à–∏–±–∫–∞ —Å–æ–∑–¥–∞–Ω–∏—è –ø—Ä–æ–µ–∫—Ç–∞:\n{result.get('error', '')}"
        )


async def set_project_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /set_project - —É—Å—Ç–∞–Ω–æ–≤–∏—Ç—å –∞–∫—Ç–∏–≤–Ω—ã–π –ø—Ä–æ–µ–∫—Ç"""
    if not PROJECTS_AVAILABLE:
        await update.message.reply_text("‚ö†Ô∏è –£–ø—Ä–∞–≤–ª–µ–Ω–∏–µ –ø—Ä–æ–µ–∫—Ç–∞–º–∏ –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–æ")
        return

    user_id = update.effective_user.id
    projects = get_user_projects(user_id)

    if not projects:
        await update.message.reply_text(
            "üìÅ –£ –≤–∞—Å –Ω–µ—Ç –ø—Ä–æ–µ–∫—Ç–æ–≤.\n\n"
            "–°–æ–∑–¥–∞–π—Ç–µ –Ω–æ–≤—ã–π: `/new_project –ù–∞–∑–≤–∞–Ω–∏–µ`",
            parse_mode="Markdown"
        )
        return

    if context.args:
        # –£—Å—Ç–∞–Ω–æ–≤–∏—Ç—å –ø—Ä–æ–µ–∫—Ç –ø–æ –Ω–∞–∑–≤–∞–Ω–∏—é
        project_name = " ".join(context.args)
        if project_name in projects:
            context.user_data["current_project"] = project_name
            project = load_project(user_id, project_name)
            await update.message.reply_text(
                f"‚úÖ –ê–∫—Ç–∏–≤–Ω—ã–π –ø—Ä–æ–µ–∫—Ç: **{project_name}**\n\n"
                f"{project.get_log_summary()}\n\n"
                "–í—Å–µ –≤–æ–ø—Ä–æ—Å—ã –∏ –æ—Ç–≤–µ—Ç—ã —Å–æ—Ö—Ä–∞–Ω—è—é—Ç—Å—è –≤ —ç—Ç–æ—Ç –ø—Ä–æ–µ–∫—Ç.",
                parse_mode="Markdown"
            )
        else:
            await update.message.reply_text(f"‚ùå –ü—Ä–æ–µ–∫—Ç '{project_name}' –Ω–µ –Ω–∞–π–¥–µ–Ω.")
    else:
        # –ü–æ–∫–∞–∑–∞—Ç—å –º–µ–Ω—é –≤—ã–±–æ—Ä–∞
        keyboard = []
        for proj in projects:
            keyboard.append([
                InlineKeyboardButton(
                    text=f"üìÅ {proj}",
                    callback_data=f"setproj_{proj}"
                )
            ])
        reply_markup = InlineKeyboardMarkup(keyboard)
        current = context.user_data.get("current_project", "–ù–µ –≤—ã–±—Ä–∞–Ω")
        await update.message.reply_text(
            f"**–í–´–ë–û–† –ê–ö–¢–ò–í–ù–û–ì–û –ü–†–û–ï–ö–¢–ê**\n\n"
            f"–¢–µ–∫—É—â–∏–π –ø—Ä–æ–µ–∫—Ç: {current}\n\n"
            "–í—ã–±–µ—Ä–∏—Ç–µ –ø—Ä–æ–µ–∫—Ç –¥–ª—è —Ä–∞–±–æ—Ç—ã:",
            reply_markup=reply_markup,
            parse_mode="Markdown"
        )


async def project_info_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /project_info - –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è –æ–± –∞–∫—Ç–∏–≤–Ω–æ–º –ø—Ä–æ–µ–∫—Ç–µ"""
    if not PROJECTS_AVAILABLE:
        await update.message.reply_text("‚ö†Ô∏è –£–ø—Ä–∞–≤–ª–µ–Ω–∏–µ –ø—Ä–æ–µ–∫—Ç–∞–º–∏ –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–æ")
        return

    current_project_name = context.user_data.get("current_project")
    if not current_project_name:
        await update.message.reply_text(
            "‚ùå –ù–µ—Ç –∞–∫—Ç–∏–≤–Ω–æ–≥–æ –ø—Ä–æ–µ–∫—Ç–∞.\n\n"
            "–í—ã–±–µ—Ä–∏—Ç–µ –ø—Ä–æ–µ–∫—Ç: `/set_project`",
            parse_mode="Markdown"
        )
        return

    user_id = update.effective_user.id
    project = load_project(user_id, current_project_name)

    if not project:
        await update.message.reply_text("‚ùå –û—à–∏–±–∫–∞ –∑–∞–≥—Ä—É–∑–∫–∏ –ø—Ä–æ–µ–∫—Ç–∞")
        return

    info = project.get_project_summary()
    log_summary = project.get_log_summary()

    await update.message.reply_text(
        f"{info}\n\n"
        f"üìä **–ñ—É—Ä–Ω–∞–ª —Ä–∞–±–æ—Ç—ã:** {log_summary}",
        parse_mode="Markdown"
    )


async def project_log_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /project_log - –∂—É—Ä–Ω–∞–ª —Ä–∞–±–æ—Ç—ã –Ω–∞–¥ –ø—Ä–æ–µ–∫—Ç–æ–º"""
    if not PROJECTS_AVAILABLE:
        await update.message.reply_text("‚ö†Ô∏è –£–ø—Ä–∞–≤–ª–µ–Ω–∏–µ –ø—Ä–æ–µ–∫—Ç–∞–º–∏ –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–æ")
        return

    current_project_name = context.user_data.get("current_project")
    if not current_project_name:
        await update.message.reply_text(
            "‚ùå –ù–µ—Ç –∞–∫—Ç–∏–≤–Ω–æ–≥–æ –ø—Ä–æ–µ–∫—Ç–∞.\n\n"
            "–í—ã–±–µ—Ä–∏—Ç–µ –ø—Ä–æ–µ–∫—Ç: `/set_project`",
            parse_mode="Markdown"
        )
        return

    user_id = update.effective_user.id
    project = load_project(user_id, current_project_name)

    if not project:
        await update.message.reply_text("‚ùå –û—à–∏–±–∫–∞ –∑–∞–≥—Ä—É–∑–∫–∏ –ø—Ä–æ–µ–∫—Ç–∞")
        return

    log = project.get_conversation_log()

    if not log:
        await update.message.reply_text("üìã –ñ—É—Ä–Ω–∞–ª —Ä–∞–±–æ—Ç—ã –ø—É—Å—Ç")
        return

    # –ü–æ–∫–∞–∑—ã–≤–∞–µ–º –ø–æ—Å–ª–µ–¥–Ω–∏–µ 10 –∑–∞–ø–∏—Å–µ–π
    recent_log = log[-10:]
    response = f"üìã **–ñ–£–†–ù–ê–õ: {current_project_name}**\n\n"
    response += f"–í—Å–µ–≥–æ –∑–∞–ø–∏—Å–µ–π: {len(log)}\n"
    response += f"–ü–æ–∫–∞–∑–∞–Ω—ã –ø–æ—Å–ª–µ–¥–Ω–∏–µ {len(recent_log)} –∑–∞–ø–∏—Å–µ–π:\n\n"

    for i, entry in enumerate(recent_log, 1):
        timestamp = entry["timestamp"][:16].replace("T", " ")
        question = entry.get("question", "")[:50]
        response += f"{i}. {timestamp}\n   Q: {question}...\n\n"

    response += "\nüí° –î–ª—è —ç–∫—Å–ø–æ—Ä—Ç–∞ –ø–æ–ª–Ω–æ–≥–æ –∂—É—Ä–Ω–∞–ª–∞ –∏—Å–ø–æ–ª—å–∑—É–π—Ç–µ `/export_project`"

    await update.message.reply_text(response, parse_mode="Markdown")


# === –û–ë–†–ê–ë–û–¢–ö–ê –°–û–û–ë–©–ï–ù–ò–ô ===

async def handle_photo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–û–±—Ä–∞–±–æ—Ç–∫–∞ —Ñ–æ—Ç–æ–≥—Ä–∞—Ñ–∏–π"""
    user_id = update.effective_user.id

    # –ü—Ä–æ–≤–µ—Ä–∫–∞ rate limit
    if not check_rate_limit(user_id):
        await update.message.reply_text(
            "‚è±Ô∏è –°–ª–∏—à–∫–æ–º –º–Ω–æ–≥–æ –∑–∞–ø—Ä–æ—Å–æ–≤!\n\n"
            f"–í—ã –º–æ–∂–µ—Ç–µ –æ—Ç–ø—Ä–∞–≤–ª—è—Ç—å –¥–æ {RATE_LIMIT_MAX_REQUESTS} –∑–∞–ø—Ä–æ—Å–æ–≤ –≤ –º–∏–Ω—É—Ç—É.\n"
            "–ü–æ–∂–∞–ª—É–π—Å—Ç–∞, –ø–æ–¥–æ–∂–¥–∏—Ç–µ –Ω–µ–º–Ω–æ–≥–æ –∏ –ø–æ–ø—Ä–æ–±—É–π—Ç–µ —Å–Ω–æ–≤–∞."
        )
        return

    # –ü—Ä–æ–≤–µ—Ä—è–µ–º, –Ω—É–∂–Ω–∞ –ª–∏ –≤–∏–∑—É–∞–ª–∏–∑–∞—Ü–∏—è —á–µ—Ä–µ–∑ Gemini
    if await handle_photo_with_visualization(update, context):
        return

    # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º —Å–æ–æ–±—â–µ–Ω–∏–µ –æ –ø—Ä–æ—Ü–µ—Å—Å–µ –∏ —Å–æ—Ö—Ä–∞–Ω—è–µ–º –µ–≥–æ –¥–ª—è –ø–æ—Å–ª–µ–¥—É—é—â–µ–≥–æ —É–¥–∞–ª–µ–Ω–∏—è
    thinking_message = await update.message.reply_text("üì∏ –ê–Ω–∞–ª–∏–∑–∏—Ä—É—é —Ñ–æ—Ç–æ–≥—Ä–∞—Ñ–∏—é...\n\n–í—ã –º–æ–∂–µ—Ç–µ –Ω–µ –∂–¥–∞—Ç—å, —è –ø—Ä–∏—à–ª—é —É–≤–µ–¥–æ–º–ª–µ–Ω–∏–µ üòâ")

    try:
        # –ü–æ–ª—É—á–∞–µ–º —Ñ–æ—Ç–æ (—Å–∞–º–æ–µ –±–æ–ª—å—à–æ–µ —Ä–∞–∑—Ä–µ—à–µ–Ω–∏–µ)
        photo = update.message.photo[-1]
        caption_text = update.message.caption or "–ü—Ä–æ–∞–Ω–∞–ª–∏–∑–∏—Ä—É–π —ç—Ç–æ —Ñ–æ—Ç–æ"

        # ============================================================================
        # –û–ü–¢–ò–ú–ò–ó–ê–¶–ò–Ø: –£–º–Ω—ã–π –≤—ã–±–æ—Ä –º–æ–¥–µ–ª–∏ –¥–ª—è —Ñ–æ—Ç–æ (Gemini –¥–ª—è –¥–µ—Ñ–µ–∫—Ç–æ–≤)
        # ============================================================================
        if SMART_WRAPPER_AVAILABLE:
            smart_result = await smart_model_selection_photo(
                question=caption_text,
                photo_file_id=photo.file_id,
                update=update,
                context=context
            )

            # –ï—Å–ª–∏ —É–º–Ω—ã–π –≤—ã–±–æ—Ä –æ–±—Ä–∞–±–æ—Ç–∞–ª —Ñ–æ—Ç–æ - –≤—ã—Ö–æ–¥–∏–º
            if smart_result and smart_result.get("success"):
                try:
                    await thinking_message.delete()
                except:
                    pass
                return  # –û—Ç–≤–µ—Ç —É–∂–µ –æ—Ç–ø—Ä–∞–≤–ª–µ–Ω

        # –ï—Å–ª–∏ —É–º–Ω—ã–π –≤—ã–±–æ—Ä –Ω–µ —Å—Ä–∞–±–æ—Ç–∞–ª - –ø—Ä–æ–¥–æ–ª–∂–∞–µ–º —Å Grok/Gemini –ø–æ —Å—Ç–∞–Ω–¥–∞—Ä—Ç–Ω–æ–π –ª–æ–≥–∏–∫–µ

        # –ü—Ä–æ–≤–µ—Ä—è–µ–º —Ä–∞–∑–º–µ—Ä —Ñ–æ—Ç–æ
        if photo.file_size and photo.file_size > 20 * 1024 * 1024:  # 20 –ú–ë
            await thinking_message.edit_text(
                f"‚ùå **–§–æ—Ç–æ —Å–ª–∏—à–∫–æ–º –±–æ–ª—å—à–æ–µ**\n\n"
                f"üìä –†–∞–∑–º–µ—Ä: {photo.file_size / (1024 * 1024):.1f} –ú–ë\n"
                f"üìè –ú–∞–∫—Å–∏–º—É–º: 20 –ú–ë\n\n"
                f"üí° –°–æ–∂–º–∏—Ç–µ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ –ø–µ—Ä–µ–¥ –æ—Ç–ø—Ä–∞–≤–∫–æ–π"
            )
            return

        photo_file = await photo.get_file()

        # –°–∫–∞—á–∏–≤–∞–µ–º —Ñ–æ—Ç–æ
        photo_bytes = await photo_file.download_as_bytearray()

        # –ö–æ–¥–∏—Ä—É–µ–º –≤ base64
        photo_base64 = base64.b64encode(photo_bytes).decode('utf-8')

        # –ü–æ–ª—É—á–∞–µ–º –ø–æ–¥–ø–∏—Å—å (–µ—Å–ª–∏ –µ—Å—Ç—å)
        caption = update.message.caption or ""

        # ============================================
        # –í–´–ë–û–† AI –î–í–ò–ñ–ö–ê –î–õ–Ø –ê–ù–ê–õ–ò–ó–ê –§–û–¢–û
        # ============================================
        # –ü—Ä–∏–æ—Ä–∏—Ç–µ—Ç: 1) Gemini 2.5 Flash (–±—ã—Å—Ç—Ä–µ–µ, –¥–µ—à–µ–≤–ª–µ)
        #            2) xAI Grok (fallback)

        # –ü—Ä–æ–±—É–µ–º Gemini Vision —Å–Ω–∞—á–∞–ª–∞
        if GEMINI_VISION_AVAILABLE and gemini_vision_analyzer:
            try:
                logger.info("üì∏ –ò—Å–ø–æ–ª—å–∑—É–µ–º Gemini 2.5 Flash –¥–ª—è –∞–Ω–∞–ª–∏–∑–∞ —Ñ–æ—Ç–æ")

                # –ê–Ω–∞–ª–∏–∑–∏—Ä—É–µ–º —á–µ—Ä–µ–∑ Gemini
                analysis_result = await gemini_vision_analyzer.analyze_defect_photo(
                    image_data=bytes(photo_bytes),
                    user_prompt=caption if caption else None
                )

                if analysis_result:
                    # –£–¥–∞–ª—è–µ–º —Å–æ–æ–±—â–µ–Ω–∏–µ "–∞–Ω–∞–ª–∏–∑–∏—Ä—É—é —Ñ–æ—Ç–æ–≥—Ä–∞—Ñ–∏—é"
                    try:
                        await thinking_message.delete()
                    except Exception as e:
                        logger.warning(f"Could not delete thinking message: {e}")

                    # –§–æ—Ä–º–∏—Ä—É–µ–º –æ—Ç–≤–µ—Ç
                    result = f"üîç **–ê–Ω–∞–ª–∏–∑ —Ñ–æ—Ç–æ–≥—Ä–∞—Ñ–∏–∏ (Gemini 2.5 Flash):**\n\n{analysis_result}\n\n"
                    result += f"‚è∞ {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}"

                    # –†–∞–∑–±–∏–≤–∞–µ–º –¥–ª–∏–Ω–Ω—ã–µ —Å–æ–æ–±—â–µ–Ω–∏—è –Ω–∞ —á–∞—Å—Ç–∏
                    max_length = 4000
                    if len(result) > max_length:
                        parts = []
                        current_part = ""
                        for line in result.split('\n'):
                            if len(current_part) + len(line) + 1 > max_length:
                                parts.append(current_part)
                                current_part = line + '\n'
                            else:
                                current_part += line + '\n'
                        if current_part:
                            parts.append(current_part)

                        for part in parts:
                            await update.message.reply_text(
                                part,
                                parse_mode='Markdown'
                            )
                    else:
                        await update.message.reply_text(
                            result,
                            parse_mode='Markdown'
                        )

                    # –°–æ—Ö—Ä–∞–Ω—è–µ–º –≤ –∏—Å—Ç–æ—Ä–∏—é
                    if HISTORY_MANAGER_AVAILABLE:
                        await history_manager.add_to_history(
                            user_id=user_id,
                            message_type="photo",
                            content=caption if caption else "–§–æ—Ç–æ –±–µ–∑ –ø–æ–¥–ø–∏—Å–∏",
                            response=analysis_result,
                            metadata={"ai_model": "gemini-2.5-flash"}
                        )

                    logger.info(f"‚úÖ –§–æ—Ç–æ –ø—Ä–æ–∞–Ω–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–æ —á–µ—Ä–µ–∑ Gemini –¥–ª—è –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è {user_id}")
                    return  # –í—ã—Ö–æ–¥–∏–º, –∞–Ω–∞–ª–∏–∑ –≤—ã–ø–æ–ª–Ω–µ–Ω

            except Exception as e:
                logger.warning(f"‚ö†Ô∏è Gemini Vision –Ω–µ —É–¥–∞–ª–æ—Å—å: {e}. –ü–µ—Ä–µ–∫–ª—é—á–∞–µ–º—Å—è –Ω–∞ Grok")
                # –ü—Ä–æ–¥–æ–ª–∂–∞–µ–º –∫ Grok fallback

        # –ï—Å–ª–∏ Gemini –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω –∏–ª–∏ –Ω–µ —Å—Ä–∞–±–æ—Ç–∞–ª - –∏—Å–ø–æ–ª—å–∑—É–µ–º Grok
        logger.info("üì∏ –ò—Å–ø–æ–ª—å–∑—É–µ–º xAI Grok –¥–ª—è –∞–Ω–∞–ª–∏–∑–∞ —Ñ–æ—Ç–æ (fallback)")

        # –§–æ—Ä–º–∏—Ä—É–µ–º –ø—Ä–æ—Ñ–µ—Å—Å–∏–æ–Ω–∞–ª—å–Ω—ã–π –ø—Ä–æ–º–ø—Ç –¥–ª—è Grok
        system_prompt = """–í—ã ‚Äî –°—Ç—Ä–æ–π–ù–∞–¥–∑–æ—ÄAI v2.3, AI-—ç–∫—Å–ø–µ—Ä—Ç –ø–æ —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–æ–º—É –Ω–∞–¥–∑–æ—Ä—É –≤ —Å—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤–µ –†–§.

üì± –§–û–†–ú–ê–¢–ò–†–û–í–ê–ù–ò–ï –î–õ–Ø –¢–ï–õ–ï–§–û–ù–ê:
‚Ä¢ –ü–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—å –Ω–∞ –º–æ–±–∏–ª—å–Ω–æ–º —É—Å—Ç—Ä–æ–π—Å—Ç–≤–µ!
‚Ä¢ –ú–∞–∫—Å–∏–º—É–º 30-35 —Å–∏–º–≤–æ–ª–æ–≤ –≤ —Å—Ç—Ä–æ–∫–µ
‚Ä¢ –ù–ï –∏—Å–ø–æ–ª—å–∑—É–π—Ç–µ —à–∏—Ä–æ–∫–∏–µ —Ç–∞–±–ª–∏—Ü—ã –∏ ASCII-—Å—Ö–µ–º—ã
‚Ä¢ –ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ –≤–µ—Ä—Ç–∏–∫–∞–ª—å–Ω—ã–µ —Å–ø–∏—Å–∫–∏ –∏ –ø—Ä–æ—Å—Ç—ã–µ —Å–∏–º–≤–æ–ª—ã ‚Üí ‚Üì ‚Ä¢

üì∏ –ê–ù–ê–õ–ò–ó –ò–ó–û–ë–†–ê–ñ–ï–ù–ò–ô - –ö–†–ò–¢–ò–ß–ï–°–ö–ò –í–ê–ñ–ù–û:

üéØ –ü–†–ò–ù–¶–ò–ü –û–ë–™–ï–ö–¢–ò–í–ù–û–°–¢–ò:
‚Ä¢ –ü—Ä–æ–≤–µ–¥–∏—Ç–µ –ë–ï–°–ü–†–ò–°–¢–†–ê–°–¢–ù–´–ô –∞–Ω–∞–ª–∏–∑ —Ñ–æ—Ç–æ–≥—Ä–∞—Ñ–∏–∏
‚Ä¢ –ù–ï –ø—Ä–µ–¥–ø–æ–ª–∞–≥–∞–π—Ç–µ –Ω–∞–ª–∏—á–∏–µ –¥–µ—Ñ–µ–∫—Ç–æ–≤ –∑–∞—Ä–∞–Ω–µ–µ
‚Ä¢ –ï—Å–ª–∏ –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏—è –≤ –Ω–æ—Ä–º–µ ‚Äî –ø—Ä—è–º–æ —É–∫–∞–∂–∏—Ç–µ —ç—Ç–æ
‚Ä¢ –î–µ—Ñ–µ–∫—Ç–∞–º–∏ —Å—á–∏—Ç–∞–π—Ç–µ –¢–û–õ–¨–ö–û —è–≤–Ω—ã–µ –æ—Ç–∫–ª–æ–Ω–µ–Ω–∏—è –æ—Ç –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤

üìã –°–¢–†–£–ö–¢–£–†–ê –û–¢–í–ï–¢–ê (–∫—Ä–∞—Ç–∫–æ, –¥–æ 300 —Å–ª–æ–≤):

**–ê–Ω–∞–ª–∏–∑:**
‚Ä¢ –¢–∏–ø —ç–ª–µ–º–µ–Ω—Ç–∞ (—Ñ—É–Ω–¥–∞–º–µ–Ω—Ç/—Å—Ç–µ–Ω–∞/–ø–µ—Ä–µ–∫—Ä—ã—Ç–∏–µ/–∫—Ä–æ–≤–ª—è)
‚Ä¢ –ú–∞—Ç–µ—Ä–∏–∞–ª (–±–µ—Ç–æ–Ω/–∫–∏—Ä–ø–∏—á/–º–µ—Ç–∞–ª–ª/–¥–µ—Ä–µ–≤–æ)
‚Ä¢ –°–æ—Å—Ç–æ—è–Ω–∏–µ: ‚úÖ –Ω–æ—Ä–º–∞ / ‚ö†Ô∏è –¥–µ—Ñ–µ–∫—Ç –≤—ã—è–≤–ª–µ–Ω

**–ï—Å–ª–∏ –¥–µ—Ñ–µ–∫—Ç –æ–±–Ω–∞—Ä—É–∂–µ–Ω:**
‚Ä¢ –¢–∏–ø –¥–µ—Ñ–µ–∫—Ç–∞ (—Ç—Ä–µ—â–∏–Ω–∞/–∫–æ—Ä—Ä–æ–∑–∏—è/–æ—Ç—Å–ª–æ–µ–Ω–∏–µ/–¥–µ—Ñ–æ—Ä–º–∞—Ü–∏—è)
‚Ä¢ –ö—Ä–∏—Ç–∏—á–Ω–æ—Å—Ç—å: –Ω–∏–∑–∫–∞—è/—Å—Ä–µ–¥–Ω—è—è/–≤—ã—Å–æ–∫–∞—è
‚Ä¢ –†–∞—Å–ø–æ–ª–æ–∂–µ–Ω–∏–µ (–≥–¥–µ –∏–º–µ–Ω–Ω–æ –Ω–∞ —ç–ª–µ–º–µ–Ω—Ç–µ)
‚Ä¢ –ü–∞—Ä–∞–º–µ—Ç—Ä—ã (—à–∏—Ä–∏–Ω–∞, –¥–ª–∏–Ω–∞, –≥–ª—É–±–∏–Ω–∞ - –µ—Å–ª–∏ –≤–∏–¥–Ω–æ)

**–ù–æ—Ä–º–∞—Ç–∏–≤—ã** (—Ç–æ–ª—å–∫–æ –ø—Ä–∏ –¥–µ—Ñ–µ–∫—Ç–∞—Ö):
‚Ä¢ –°–ü 63.13330.2018 ‚Äî –¥–ª—è –±–µ—Ç–æ–Ω–∞/–∂.–± (—Ç—Ä–µ—â–∏–Ω—ã, –ø—Ä–æ—á–Ω–æ—Å—Ç—å)
‚Ä¢ –°–ü 13-102-2003 ‚Äî –ø—Ä–∞–≤–∏–ª–∞ –æ–±—Å–ª–µ–¥–æ–≤–∞–Ω–∏—è
‚Ä¢ –°–ü 28.13330.2017 ‚Äî –∑–∞—â–∏—Ç–∞ –æ—Ç –∫–æ—Ä—Ä–æ–∑–∏–∏
‚Ä¢ –ì–û–°–¢ –† 31937-2011 ‚Äî –∫–ª–∞—Å—Å—ã —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–æ–≥–æ —Å–æ—Å—Ç–æ—è–Ω–∏—è

**–†–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏:**
‚Ä¢ –ï—Å–ª–∏ –Ω–æ—Ä–º–∞: –∫—Ä–∞—Ç–∫–∏–µ —Å–æ–≤–µ—Ç—ã –ø–æ —ç–∫—Å–ø–ª—É–∞—Ç–∞—Ü–∏–∏
‚Ä¢ –ï—Å–ª–∏ –¥–µ—Ñ–µ–∫—Ç: –º–µ—Ç–æ–¥ —É—Å—Ç—Ä–∞–Ω–µ–Ω–∏—è (–∏–Ω—ä–µ–∫—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ/—É—Å–∏–ª–µ–Ω–∏–µ/–∑–∞–º–µ–Ω–∞)

üîß –£–ß–Å–¢ –ù–ê–í–´–ö–û–í –°–¢–†–û–ò–¢–ï–õ–Ø:
‚Ä¢ –ê–¥–∞–ø—Ç–∏—Ä—É–π—Ç–µ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏ –ø–æ–¥ —É—Ä–æ–≤–µ–Ω—å –∫–≤–∞–ª–∏—Ñ–∏–∫–∞—Ü–∏–∏ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è
‚Ä¢ –î–ª—è –Ω–∞—á–∏–Ω–∞—é—â–∏—Ö: –¥–∞–≤–∞–π—Ç–µ –ø–æ—à–∞–≥–æ–≤—ã–µ –∏–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏ —Å —É–∫–∞–∑–∞–Ω–∏–µ–º –∏–Ω—Å—Ç—Ä—É–º–µ–Ω—Ç–æ–≤
‚Ä¢ –î–ª—è –æ–ø—ã—Ç–Ω—ã—Ö: —Å–æ—Å—Ä–µ–¥–æ—Ç–æ—á—å—Ç–µ—Å—å –Ω–∞ —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏—Ö –¥–µ—Ç–∞–ª—è—Ö –∏ –Ω–æ—Ä–º–∞—Ç–∏–≤–∞—Ö

üí° –ü–†–ê–ö–¢–ò–ß–ï–°–ö–ò–ô –ü–û–î–•–û–î:
‚Ä¢ –ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ —Å—Ç—Ä–æ–∏—Ç–µ–ª—å–Ω—É—é —Ç–µ—Ä–º–∏–Ω–æ–ª–æ–≥–∏—é
‚Ä¢ –ü—Ä–∏ –¥–µ—Ñ–µ–∫—Ç–∞—Ö —É–∫–∞–∑—ã–≤–∞–π—Ç–µ —á–∏—Å–ª–∞ —Å –¥–æ–ø—É—Å–∫–∞–º–∏
‚Ä¢ –ï—Å–ª–∏ –Ω–µ—è—Å–Ω–æ ‚Äî –∑–∞–ø—Ä–æ—Å–∏—Ç–µ –¥–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω—ã–µ —Ñ–æ—Ç–æ
‚Ä¢ –î–µ—Ä–∂–∏—Ç–µ –¥—Ä—É–∂–µ–ª—é–±–Ω—ã–π —Ç–æ–Ω, —Å—Ç—Ä—É–∫—Ç—É—Ä–∏—Ä—É–π—Ç–µ (bullets, emojis)

–í–ê–ñ–ù–û: –û–±—ä–µ–∫—Ç–∏–≤–Ω–æ—Å—Ç—å –ø—Ä–µ–≤—ã—à–µ –≤—Å–µ–≥–æ. –ï—Å–ª–∏ –¥–µ—Ñ–µ–∫—Ç–æ–≤ –Ω–µ—Ç ‚Äî —Ç–∞–∫ –∏ –Ω–∞–ø–∏—à–∏—Ç–µ. –ù–µ –∏—â–∏—Ç–µ –ø—Ä–æ–±–ª–µ–º—ã —Ç–∞–º, –≥–¥–µ –∏—Ö –Ω–µ—Ç!"""

        user_message = "–ü—Ä–æ–≤–µ–¥–∏—Ç–µ –æ–±—ä–µ–∫—Ç–∏–≤–Ω—ã–π —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏–π –æ—Å–º–æ—Ç—Ä –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è. –û–ø–∏—à–∏—Ç–µ —Å–æ—Å—Ç–æ—è–Ω–∏–µ –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏ –∏ —É–∫–∞–∂–∏—Ç–µ, –µ—Å—Ç—å –ª–∏ –¥–µ—Ñ–µ–∫—Ç—ã –∏–ª–∏ –Ω–∞—Ä—É—à–µ–Ω–∏—è —Å—Ç—Ä–æ–∏—Ç–µ–ª—å–Ω—ã—Ö –Ω–æ—Ä–º."
        if caption:
            user_message += f"\n\n–î–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω–∞—è –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è –æ—Ç –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è: {caption}"

        # –í—ã–∑—ã–≤–∞–µ–º xAI Grok API –¥–ª—è –∞–Ω–∞–ª–∏–∑–∞ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è —Å retry logic
        client = get_grok_client()
        loop = asyncio.get_event_loop()

        # –í–∫–ª—é—á–∞–µ–º web_search –¥–ª—è –∞–Ω–∞–ª–∏–∑–∞ —Ñ–æ—Ç–æ (–ø–æ–∏—Å–∫ –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏–∏ –æ –¥–µ—Ñ–µ–∫—Ç–∞—Ö)
        search_params = {
            "mode": "auto", "return_citations": True, "sources": [{"type": "web"}, {"type": "news"}, {"type": "x"}]}

        response = await loop.run_in_executor(
            None,
            lambda: call_grok_with_retry(
                client,
                model="grok-4-1-fast",  # Reasoning –º–æ–¥–µ–ª—å –¥–ª—è –∞–Ω–∞–ª–∏–∑–∞ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π
                max_tokens=6000,
                temperature=0.7,
                messages=[
                    {
                        "role": "system",
                        "content": system_prompt
                    },
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": "image/jpeg",
                                    "data": photo_base64
                                }
                            },
                            {
                                "type": "text",
                                "text": user_message
                            }
                        ]
                    }
                ],
                search_parameters=search_params
            )
        )
        analysis = response["choices"][0]["message"]["content"]

        # –£–¥–∞–ª—è–µ–º —Å–æ–æ–±—â–µ–Ω–∏–µ "–∞–Ω–∞–ª–∏–∑–∏—Ä—É—é —Ñ–æ—Ç–æ–≥—Ä–∞—Ñ–∏—é"
        try:
            await thinking_message.delete()
        except Exception as e:
            logger.warning(f"Could not delete thinking message: {e}")

        # –§–æ—Ä–º–∏—Ä—É–µ–º –æ—Ç–≤–µ—Ç
        result = f"üîç **–ê–Ω–∞–ª–∏–∑ —Ñ–æ—Ç–æ–≥—Ä–∞—Ñ–∏–∏:**\n\n{analysis}\n\n"
        result += f"‚è∞ {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}"

        # –†–∞–∑–±–∏–≤–∞–µ–º –¥–ª–∏–Ω–Ω—ã–µ —Å–æ–æ–±—â–µ–Ω–∏—è –Ω–∞ —á–∞—Å—Ç–∏ (–ª–∏–º–∏—Ç Telegram: 4096 —Å–∏–º–≤–æ–ª–æ–≤)
        max_length = 4000  # –û—Å—Ç–∞–≤–ª—è–µ–º –∑–∞–ø–∞—Å
        if len(result) > max_length:
            parts = []
            current_part = ""
            for line in result.split('\n'):
                if len(current_part) + len(line) + 1 > max_length:
                    parts.append(current_part)
                    current_part = line + '\n'
                else:
                    current_part += line + '\n'
            if current_part:
                parts.append(current_part)

            # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –ø–æ —á–∞—Å—Ç—è–º –ë–ï–ó parse_mode (–∏–∑–±–µ–≥–∞–µ–º –æ—à–∏–±–æ–∫ –ø–∞—Ä—Å–∏–Ω–≥–∞)
            for i, part in enumerate(parts):
                # –î–æ–±–∞–≤–ª—è–µ–º –∫–Ω–æ–ø–∫—É "–ü—Ä–∏–º–µ–Ω–∏—Ç—å –∏–∑–º–µ–Ω–µ–Ω–∏—è" –∫ –ø–æ—Å–ª–µ–¥–Ω–µ–π —á–∞—Å—Ç–∏ (—Ç–æ–ª—å–∫–æ –¥–ª—è —Ä–∞–∑—Ä–∞–±–æ—Ç—á–∏–∫–∞)
                part_reply_markup = None
                user_id = update.effective_user.id
                if i == len(parts) - 1 and AUTO_APPLY_AVAILABLE and should_show_apply_button(analysis) and is_developer(user_id):
                    part_reply_markup = add_apply_button()

                if i == 0:
                    await update.message.reply_text(part, reply_markup=part_reply_markup)
                else:
                    await update.message.reply_text(
                        f"(–ø—Ä–æ–¥–æ–ª–∂–µ–Ω–∏–µ {i+1}/{len(parts)})\n\n{part}",
                        reply_markup=part_reply_markup
                    )
        else:
            # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –ë–ï–ó parse_mode –¥–ª—è –∏–∑–±–µ–∂–∞–Ω–∏—è –æ—à–∏–±–æ–∫ "can't parse entities"
            await update.message.reply_text(result)

        logger.info(f"Photo analyzed for user {update.effective_user.id} by Claude")

    except Exception as e:
        logger.error(f"Error analyzing photo: {e}")
        # –£–¥–∞–ª—è–µ–º —Å–æ–æ–±—â–µ–Ω–∏–µ "–∞–Ω–∞–ª–∏–∑–∏—Ä—É—é —Ñ–æ—Ç–æ–≥—Ä–∞—Ñ–∏—é" –¥–∞–∂–µ –≤ —Å–ª—É—á–∞–µ –æ—à–∏–±–∫–∏
        try:
            await thinking_message.delete()
        except:
            pass
        await update.message.reply_text(
            f"‚ùå –û—à–∏–±–∫–∞ –ø—Ä–∏ –∞–Ω–∞–ª–∏–∑–µ —Ñ–æ—Ç–æ–≥—Ä–∞—Ñ–∏–∏: {str(e)}\n\n–ü–æ–ø—Ä–æ–±—É–π—Ç–µ –µ—â–µ —Ä–∞–∑ –∏–ª–∏ –æ–±—Ä–∞—Ç–∏—Ç–µ—Å—å –∫ –∞–¥–º–∏–Ω–∏—Å—Ç—Ä–∞—Ç–æ—Ä—É."
        )


async def handle_voice(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–û–±—Ä–∞–±–æ—Ç–∫–∞ –≥–æ–ª–æ—Å–æ–≤—ã—Ö —Å–æ–æ–±—â–µ–Ω–∏–π"""
    if not VOICE_HANDLER_AVAILABLE:
        await update.message.reply_text("‚ö†Ô∏è –ì–æ–ª–æ—Å–æ–≤—ã–µ —Å–æ–æ–±—â–µ–Ω–∏—è –Ω–µ–¥–æ—Å—Ç—É–ø–Ω—ã")
        return

    user_id = update.effective_user.id
    thinking_msg = await update.message.reply_text("üé§ –†–∞—Å–ø–æ–∑–Ω–∞—é –≥–æ–ª–æ—Å–æ–≤–æ–µ —Å–æ–æ–±—â–µ–Ω–∏–µ...")

    try:
        voice_file_id = update.message.voice.file_id
        result = await process_voice_message(
            bot=context.bot,
            voice_file_id=voice_file_id,
            user_id=user_id
        )

        if result["success"]:
            recognized_text = result["text"]

            # –ü–æ–∫–∞–∑—ã–≤–∞–µ–º —Ä–∞—Å–ø–æ–∑–Ω–∞–Ω–Ω—ã–π —Ç–µ–∫—Å—Ç
            await thinking_msg.edit_text(
                f"‚úÖ –†–∞—Å–ø–æ–∑–Ω–∞–Ω–æ:\n\n{recognized_text}\n\n‚è≥ –û–±—Ä–∞–±–∞—Ç—ã–≤–∞—é –∑–∞–ø—Ä–æ—Å..."
            )

            # –°–æ—Ö—Ä–∞–Ω—è–µ–º —Ä–∞—Å–ø–æ–∑–Ω–∞–Ω–Ω—ã–π —Ç–µ–∫—Å—Ç –≤ context –¥–ª—è handle_text
            context.user_data['_voice_recognized_text'] = recognized_text

            # –û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º –∫–∞–∫ —Ç–µ–∫—Å—Ç–æ–≤—ã–π –≤–æ–ø—Ä–æ—Å
            await handle_text(update, context)

        else:
            error_msg = result.get("error", "–ù–µ–∏–∑–≤–µ—Å—Ç–Ω–∞—è –æ—à–∏–±–∫–∞")
            await thinking_msg.edit_text(
                f"‚ùå –û—à–∏–±–∫–∞ —Ä–∞—Å–ø–æ–∑–Ω–∞–≤–∞–Ω–∏—è:\n{error_msg}\n\n"
                "–ü–æ–ø—Ä–æ–±—É–π—Ç–µ –∑–∞–ø–∏—Å–∞—Ç—å –≥–æ–ª–æ—Å–æ–≤–æ–µ —Å–æ–æ–±—â–µ–Ω–∏–µ –µ—â—ë —Ä–∞–∑ –∏–ª–∏ –Ω–∞–ø–∏—à–∏—Ç–µ —Ç–µ–∫—Å—Ç–æ–º."
            )

    except Exception as e:
        logger.error(f"–û—à–∏–±–∫–∞ –æ–±—Ä–∞–±–æ—Ç–∫–∏ –≥–æ–ª–æ—Å–∞: {e}")
        try:
            await thinking_msg.delete()
        except:
            pass
        await update.message.reply_text(
            f"‚ùå –û—à–∏–±–∫–∞: {str(e)}\n\n–ü–æ–∂–∞–ª—É–π—Å—Ç–∞, –Ω–∞–ø–∏—à–∏—Ç–µ –≤–∞—à –≤–æ–ø—Ä–æ—Å —Ç–µ–∫—Å—Ç–æ–º."
        )


async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–û–±—Ä–∞–±–æ—Ç–∫–∞ –∑–∞–≥—Ä—É–∑–∫–∏ —Ñ–∞–π–ª–æ–≤ —Å –∞–Ω–∞–ª–∏–∑–æ–º –∏ —ç–∫—Å–ø–µ—Ä—Ç–Ω—ã–º –∑–∞–∫–ª—é—á–µ–Ω–∏–µ–º"""
    user_id = update.effective_user.id
    current_project_name = context.user_data.get("current_project")

    # –ï—Å–ª–∏ –Ω–µ—Ç –∞–∫—Ç–∏–≤–Ω–æ–≥–æ –ø—Ä–æ–µ–∫—Ç–∞ - –ø—Ä–µ–¥–ª–∞–≥–∞–µ–º —Å–æ–∑–¥–∞—Ç—å
    if not PROJECTS_AVAILABLE or not current_project_name:
        await update.message.reply_text(
            "üìÅ –î–ª—è –∞–Ω–∞–ª–∏–∑–∞ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ —Å–æ–∑–¥–∞–π—Ç–µ –∏–ª–∏ –≤—ã–±–µ—Ä–∏—Ç–µ –ø—Ä–æ–µ–∫—Ç:\n"
            "–ù–∞–∂–º–∏—Ç–µ –∫–Ω–æ–ø–∫—É **üìÅ –ü—Ä–æ–µ–∫—Ç** –≤ –≥–ª–∞–≤–Ω–æ–º –º–µ–Ω—é",
            parse_mode="Markdown"
        )
        return

    project = load_project(user_id, current_project_name)
    if not project:
        await update.message.reply_text("‚ùå –ü—Ä–æ–µ–∫—Ç –Ω–µ –Ω–∞–π–¥–µ–Ω")
        return

    try:
        # –ü—Ä–æ–≤–µ—Ä—è–µ–º —Ä–∞–∑–º–µ—Ä —Ñ–∞–π–ª–∞
        file_size = update.message.document.file_size
        max_size = 20 * 1024 * 1024  # 20 –ú–ë - –ª–∏–º–∏—Ç Telegram Bot API

        if file_size > max_size:
            await update.message.reply_text(
                f"‚ùå **–§–∞–π–ª —Å–ª–∏—à–∫–æ–º –±–æ–ª—å—à–æ–π**\n\n"
                f"üìä –†–∞–∑–º–µ—Ä —Ñ–∞–π–ª–∞: {file_size / (1024 * 1024):.1f} –ú–ë\n"
                f"üìè –ú–∞–∫—Å–∏–º–∞–ª—å–Ω—ã–π —Ä–∞–∑–º–µ—Ä: 20 –ú–ë\n\n"
                f"üí° **–†–µ—à–µ–Ω–∏—è:**\n"
                f"‚Ä¢ –°–æ–∂–º–∏—Ç–µ PDF (–æ–Ω–ª–∞–π–Ω-—Å–µ—Ä–≤–∏—Å—ã: ilovepdf.com, smallpdf.com)\n"
                f"‚Ä¢ –†–∞–∑–±–µ–π—Ç–µ –¥–æ–∫—É–º–µ–Ω—Ç –Ω–∞ —á–∞—Å—Ç–∏\n"
                f"‚Ä¢ –£–¥–∞–ª–∏—Ç–µ –Ω–µ–Ω—É–∂–Ω—ã–µ —Å—Ç—Ä–∞–Ω–∏—Ü—ã"
            )
            return

        # –°–∫–∞—á–∏–≤–∞–µ–º —Ñ–∞–π–ª
        file = await update.message.document.get_file()
        file_name = update.message.document.file_name
        file_path = f"temp_{user_id}_{file_name}"
        await file.download_to_drive(file_path)

        description = update.message.caption or ""
        file_type = update.message.document.mime_type or "unknown"

        # –û–ø—Ä–µ–¥–µ–ª—è–µ–º —Ç–∏–ø –¥–æ–∫—É–º–µ–Ω—Ç–∞
        is_pdf = file_name.lower().endswith('.pdf') or 'pdf' in file_type.lower()

        # –°–æ–æ–±—â–µ–Ω–∏–µ –æ –Ω–∞—á–∞–ª–µ –∞–Ω–∞–ª–∏–∑–∞
        thinking_msg = await update.message.reply_text(
            f"üìÑ –ü–æ–ª—É—á–µ–Ω –¥–æ–∫—É–º–µ–Ω—Ç: **{file_name}**\n\n"
            f"‚è≥ –ê–Ω–∞–ª–∏–∑–∏—Ä—É—é –¥–æ–∫—É–º–µ–Ω—Ç...",
            parse_mode="Markdown"
        )

        # –ê–Ω–∞–ª–∏–∑–∏—Ä—É–µ–º PDF
        expert_opinion = None
        if is_pdf:
            try:
                # –ò–∑–≤–ª–µ–∫–∞–µ–º —Ç–µ–∫—Å—Ç –∏–∑ PDF
                import PyPDF2
                pdf_text = ""
                with open(file_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    num_pages = len(pdf_reader.pages)

                    # –ß–∏—Ç–∞–µ–º –º–∞–∫—Å–∏–º—É–º –ø–µ—Ä–≤—ã–µ 10 —Å—Ç—Ä–∞–Ω–∏—Ü
                    max_pages = min(num_pages, 10)
                    for page_num in range(max_pages):
                        page = pdf_reader.pages[page_num]
                        pdf_text += page.extract_text() + "\n"

                # –û–≥—Ä–∞–Ω–∏—á–∏–≤–∞–µ–º —Ä–∞–∑–º–µ—Ä —Ç–µ–∫—Å—Ç–∞
                pdf_text = pdf_text[:15000]  # ~3000 —Ç–æ–∫–µ–Ω–æ–≤

                if pdf_text.strip():
                    # –§–æ—Ä–º–∏—Ä—É–µ–º –ø—Ä–æ–º–ø—Ç –¥–ª—è –∞–Ω–∞–ª–∏–∑–∞
                    analysis_prompt = f"""–í—ã ‚Äî –≤–µ–¥—É—â–∏–π –∏–Ω–∂–µ–Ω–µ—Ä-—ç–∫—Å–ø–µ—Ä—Ç –ø–æ —Å—Ç—Ä–æ–∏—Ç–µ–ª—å–Ω—ã–º –Ω–æ—Ä–º–∞—Ç–∏–≤–∞–º –†–§ —Å 20-–ª–µ—Ç–Ω–∏–º –æ–ø—ã—Ç–æ–º.

üìã **–ó–ê–î–ê–ß–ê:** –ü—Ä–æ–∞–Ω–∞–ª–∏–∑–∏—Ä—É–π—Ç–µ –ø—Ä–µ–¥–æ—Å—Ç–∞–≤–ª–µ–Ω–Ω—ã–π —Å—Ç—Ä–æ–∏—Ç–µ–ª—å–Ω—ã–π –¥–æ–∫—É–º–µ–Ω—Ç –∏ –¥–∞–π—Ç–µ —ç–∫—Å–ø–µ—Ä—Ç–Ω–æ–µ –∑–∞–∫–ª—é—á–µ–Ω–∏–µ.

{'üìù **–ó–ê–ü–†–û–° –ü–û–õ–¨–ó–û–í–ê–¢–ï–õ–Ø:** ' + description if description else ''}

üéØ **–¢–†–ï–ë–û–í–ê–ù–ò–Ø –ö –ê–ù–ê–õ–ò–ó–£:**

1. **–ò–î–ï–ù–¢–ò–§–ò–ö–ê–¶–ò–Ø –î–û–ö–£–ú–ï–ù–¢–ê:**
   ‚Ä¢ –û–ø—Ä–µ–¥–µ–ª–∏—Ç–µ —Ç–∏–ø –¥–æ–∫—É–º–µ–Ω—Ç–∞ (–ø—Ä–æ–µ–∫—Ç, —Å–º–µ—Ç–∞, –∞–∫—Ç, –∑–∞–∫–ª—é—á–µ–Ω–∏–µ, —ç–∫—Å–ø–µ—Ä—Ç–∏–∑–∞, –∏ —Ç.–¥.)
   ‚Ä¢ –£–∫–∞–∂–∏—Ç–µ –æ—Å–Ω–æ–≤–Ω—ã–µ —Ä–µ–∫–≤–∏–∑–∏—Ç—ã (–µ—Å–ª–∏ –µ—Å—Ç—å)
   ‚Ä¢ –û–ø—Ä–µ–¥–µ–ª–∏—Ç–µ –æ–±—ä–µ–∫—Ç —Å—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤–∞

2. **–°–û–î–ï–†–ñ–ê–ù–ò–ï:**
   ‚Ä¢ –ö—Ä–∞—Ç–∫–æ –æ–ø–∏—à–∏—Ç–µ –æ—Å–Ω–æ–≤–Ω–æ–µ —Å–æ–¥–µ—Ä–∂–∞–Ω–∏–µ (3-5 –ø—É–Ω–∫—Ç–æ–≤)
   ‚Ä¢ –í—ã–¥–µ–ª–∏—Ç–µ –∫–ª—é—á–µ–≤—ã–µ —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏–µ —Ä–µ—à–µ–Ω–∏—è
   ‚Ä¢ –£–∫–∞–∂–∏—Ç–µ –ø—Ä–∏–º–µ–Ω—ë–Ω–Ω—ã–µ –Ω–æ—Ä–º–∞—Ç–∏–≤—ã

3. **–≠–ö–°–ü–ï–†–¢–ù–ê–Ø –û–¶–ï–ù–ö–ê:**
   ‚Ä¢ –°–æ–æ—Ç–≤–µ—Ç—Å—Ç–≤–∏–µ –∞–∫—Ç—É–∞–ª—å–Ω—ã–º –Ω–æ—Ä–º–∞—Ç–∏–≤–∞–º 2024-2025
   ‚Ä¢ –í—ã—è–≤–ª–µ–Ω–Ω—ã–µ –ø—Ä–æ–±–ª–µ–º—ã –∏–ª–∏ –Ω–µ—Å–æ–æ—Ç–≤–µ—Ç—Å—Ç–≤–∏—è (–µ—Å–ª–∏ –µ—Å—Ç—å)
   ‚Ä¢ –ß—Ç–æ —Ç—Ä–µ–±—É–µ—Ç –æ—Å–æ–±–æ–≥–æ –≤–Ω–∏–º–∞–Ω–∏—è

4. **–†–ï–ö–û–ú–ï–ù–î–ê–¶–ò–ò:**
   ‚Ä¢ –ß—Ç–æ –Ω—É–∂–Ω–æ –ø—Ä–æ–≤–µ—Ä–∏—Ç—å –¥–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω–æ
   ‚Ä¢ –ö–∞–∫–∏–µ –¥–æ–∫—É–º–µ–Ω—Ç—ã –º–æ–≥—É—Ç –ø–æ—Ç—Ä–µ–±–æ–≤–∞—Ç—å—Å—è
   ‚Ä¢ –ü—Ä–∞–∫—Ç–∏—á–µ—Å–∫–∏–µ —Å–æ–≤–µ—Ç—ã –ø–æ –ø—Ä–∏–º–µ–Ω–µ–Ω–∏—é

**–í–ê–ñ–ù–û:**
- –ï—Å–ª–∏ –¥–æ–∫—É–º–µ–Ω—Ç —á–∞—Å—Ç–∏—á–Ω–æ –Ω–µ—á–∏—Ç–∞–µ–º - —É–∫–∞–∂–∏—Ç–µ —ç—Ç–æ
- –°—Å—ã–ª–∞–π—Ç–µ—Å—å –Ω–∞ –∫–æ–Ω–∫—Ä–µ—Ç–Ω—ã–µ –°–ü/–ì–û–°–¢ —Å –ø—É–Ω–∫—Ç–∞–º–∏
- –ë—É–¥—å—Ç–µ –æ–±—ä–µ–∫—Ç–∏–≤–Ω—ã –∏ –∫–æ–Ω–∫—Ä–µ—Ç–Ω—ã

---

üìÑ **–¢–ï–ö–°–¢ –î–û–ö–£–ú–ï–ù–¢–ê:**

{pdf_text}"""

                    # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –Ω–∞ –∞–Ω–∞–ª–∏–∑ Grok
                    client = get_grok_client()
                    loop = asyncio.get_event_loop()

                    # –í–∫–ª—é—á–∞–µ–º web_search –¥–ª—è –∞–Ω–∞–ª–∏–∑–∞ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ (–ø–æ–∏—Å–∫ –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤)
                    search_params = {
                        "mode": "auto", "return_citations": True, "sources": [{"type": "web"}, {"type": "news"}, {"type": "x"}]}

                    response = await loop.run_in_executor(
                        None,
                        lambda: call_grok_with_retry(
                            client,
                            model="grok-4-1-fast",  # Reasoning –º–æ–¥–µ–ª—å –¥–ª—è –∞–Ω–∞–ª–∏–∑–∞ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤
                            max_tokens=6000,
                            temperature=0.3,
                            messages=[
                                {"role": "system", "content": "–í—ã ‚Äî —ç–∫—Å–ø–µ—Ä—Ç –ø–æ —Å—Ç—Ä–æ–∏—Ç–µ–ª—å–Ω—ã–º –Ω–æ—Ä–º–∞—Ç–∏–≤–∞–º –†–§. –î–∞—ë—Ç–µ –ø—Ä–æ—Ñ–µ—Å—Å–∏–æ–Ω–∞–ª—å–Ω—ã–µ –∑–∞–∫–ª—é—á–µ–Ω–∏—è –ø–æ –¥–æ–∫—É–º–µ–Ω—Ç–∞–º."},
                                {"role": "user", "content": analysis_prompt}
                            ],
                            search_parameters=search_params
                        )
                    )
                    expert_opinion = response["choices"][0]["message"]["content"]

                    # –°–æ—Ö—Ä–∞–Ω—è–µ–º –∞–Ω–∞–ª–∏–∑ –≤ –ø—Ä–æ–µ–∫—Ç
                    if expert_opinion:
                        project.add_conversation_entry(
                            f"[–î–û–ö–£–ú–ï–ù–¢] {file_name}" + (f": {description}" if description else ""),
                            expert_opinion,
                            "document_analysis"
                        )

            except ImportError:
                expert_opinion = "‚ö†Ô∏è –î–ª—è –∞–Ω–∞–ª–∏–∑–∞ PDF —É—Å—Ç–∞–Ω–æ–≤–∏—Ç–µ –±–∏–±–ª–∏–æ—Ç–µ–∫—É PyPDF2:\n`pip install PyPDF2`"
            except Exception as e:
                logger.error(f"–û—à–∏–±–∫–∞ –∞–Ω–∞–ª–∏–∑–∞ PDF: {e}")
                expert_opinion = f"‚ö†Ô∏è –ù–µ —É–¥–∞–ª–æ—Å—å –ø—Ä–æ–∞–Ω–∞–ª–∏–∑–∏—Ä–æ–≤–∞—Ç—å PDF: {str(e)}"

        # –°–æ—Ö—Ä–∞–Ω—è–µ–º —Ñ–∞–π–ª –≤ –ø—Ä–æ–µ–∫—Ç
        result = project.add_file(file_path, file_type, description)

        # –£–¥–∞–ª—è–µ–º –≤—Ä–µ–º–µ–Ω–Ω—ã–π —Ñ–∞–π–ª
        import os
        os.remove(file_path)

        # –£–¥–∞–ª—è–µ–º —Å–æ–æ–±—â–µ–Ω–∏–µ "–∞–Ω–∞–ª–∏–∑–∏—Ä—É—é"
        try:
            await thinking_msg.delete()
        except:
            pass

        # –§–æ—Ä–º–∏—Ä—É–µ–º –æ—Ç–≤–µ—Ç
        response_text = f"‚úÖ **–î–æ–∫—É–º–µ–Ω—Ç –¥–æ–±–∞–≤–ª–µ–Ω –≤ –ø—Ä–æ–µ–∫—Ç:** {current_project_name}\n\n"
        response_text += f"üìÑ **–§–∞–π–ª:** {file_name}\n"

        if result["success"]:
            file_info = result["file_info"]
            response_text += f"üíæ **–†–∞–∑–º–µ—Ä:** {file_info['size_bytes'] / 1024:.1f} –ö–ë\n"
            if description:
                response_text += f"üìù **–û–ø–∏—Å–∞–Ω–∏–µ:** {description}\n"

        # –î–æ–±–∞–≤–ª—è–µ–º —ç–∫—Å–ø–µ—Ä—Ç–Ω–æ–µ –∑–∞–∫–ª—é—á–µ–Ω–∏–µ
        if expert_opinion and is_pdf:
            response_text += f"\n{'='*40}\n\n"
            response_text += f"üéì **–≠–ö–°–ü–ï–†–¢–ù–û–ï –ó–ê–ö–õ–Æ–ß–ï–ù–ò–ï:**\n\n{expert_opinion}"

        # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –æ—Ç–≤–µ—Ç —á–∞—Å—Ç—è–º–∏ –µ—Å–ª–∏ –Ω—É–∂–Ω–æ (–ë–ï–ó parse_mode –¥–ª—è –∏–∑–±–µ–∂–∞–Ω–∏—è –æ—à–∏–±–æ–∫ –ø–∞—Ä—Å–∏–Ω–≥–∞)
        max_length = 4000
        if len(response_text) > max_length:
            parts = [response_text[i:i+max_length] for i in range(0, len(response_text), max_length)]
            for part in parts:
                await update.message.reply_text(part)
        else:
            await update.message.reply_text(response_text)

    except Exception as e:
        logger.error(f"–û—à–∏–±–∫–∞ –æ–±—Ä–∞–±–æ—Ç–∫–∏ –¥–æ–∫—É–º–µ–Ω—Ç–∞: {e}")
        await update.message.reply_text(f"‚ùå –û—à–∏–±–∫–∞ –æ–±—Ä–∞–±–æ—Ç–∫–∏: {str(e)}")


async def handle_project_creation(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–û–±—Ä–∞–±–æ—Ç–∫–∞ —Å–æ–∑–¥–∞–Ω–∏—è –Ω–æ–≤–æ–≥–æ –ø—Ä–æ–µ–∫—Ç–∞"""
    user_id = update.effective_user.id
    project_name = update.message.text.strip()

    context.user_data["waiting_for_project_name"] = False

    if PROJECTS_AVAILABLE:
        result = create_project(user_id, project_name)
        if result["success"]:
            context.user_data["current_project"] = project_name

            keyboard = [[InlineKeyboardButton("¬´ –ö –ø—Ä–æ–µ–∫—Ç–∞–º", callback_data="project_menu")]]
            reply_markup = InlineKeyboardMarkup(keyboard)

            await update.message.reply_text(
                f"‚úÖ **–ü—Ä–æ–µ–∫—Ç —Å–æ–∑–¥–∞–Ω:** {project_name}\n\n"
                "üìå –ü—Ä–æ–µ–∫—Ç –∞–∫—Ç–∏–≤–∏—Ä–æ–≤–∞–Ω!\n"
                "–í—Å–µ –≤–∞—à–∏ –≤–æ–ø—Ä–æ—Å—ã –∏ –æ—Ç–≤–µ—Ç—ã —Ç–µ–ø–µ—Ä—å —Å–æ—Ö—Ä–∞–Ω—è—é—Ç—Å—è –≤ —ç—Ç–æ—Ç –ø—Ä–æ–µ–∫—Ç.\n\n"
                "–ú–æ–∂–µ—Ç–µ –Ω–∞—á–∏–Ω–∞—Ç—å —Ä–∞–±–æ—Ç—É!",
                reply_markup=reply_markup,
                parse_mode="Markdown"
            )
        else:
            await update.message.reply_text(f"‚ùå –û—à–∏–±–∫–∞ —Å–æ–∑–¥–∞–Ω–∏—è –ø—Ä–æ–µ–∫—Ç–∞: {result.get('error', '')}")


async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–û–±—Ä–∞–±–æ—Ç–∫–∞ —Ç–µ–∫—Å—Ç–æ–≤—ã—Ö —Å–æ–æ–±—â–µ–Ω–∏–π —Å –∫–æ–Ω—Ç–µ–∫—Å—Ç–æ–º –∏—Å—Ç–æ—Ä–∏–∏"""
    user_id = update.effective_user.id
    # –ü—Ä–æ–≤–µ—Ä—è–µ–º, –µ—Å—Ç—å –ª–∏ —Ä–∞—Å–ø–æ–∑–Ω–∞–Ω–Ω—ã–π —Ç–µ–∫—Å—Ç –∏–∑ –≥–æ–ª–æ—Å–æ–≤–æ–≥–æ —Å–æ–æ–±—â–µ–Ω–∏—è
    question = context.user_data.pop('_voice_recognized_text', None) or update.message.text

    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –∫–Ω–æ–ø–∫–∏ "üé§ Real-time —á–∞—Ç"
    if question and question.strip() == "üé§ Real-time —á–∞—Ç":
        if OPENAI_REALTIME_AVAILABLE:
            await start_realtime_chat_command(update, context)
        else:
            await update.message.reply_text(
                "‚ùå **Real-time —á–∞—Ç –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω**\n\n"
                "–¢—Ä–µ–±—É–µ—Ç—Å—è OPENAI_API_KEY.\n"
                "–ü–æ–ø—Ä–æ–±—É–π—Ç–µ –∫–æ–º–∞–Ω–¥—É: /realtime_chat",
                parse_mode="Markdown"
            )
        return

    # –ü—Ä–æ–≤–µ—Ä–∫–∞ –æ–∂–∏–¥–∞–Ω–∏—è –Ω–∞–∑–≤–∞–Ω–∏—è –ø—Ä–æ–µ–∫—Ç–∞
    if context.user_data.get("waiting_for_project_name"):
        await handle_project_creation(update, context)
        return

    # –ü—Ä–æ–≤–µ—Ä–∫–∞ –æ–∂–∏–¥–∞–Ω–∏—è –∑–∞–º–µ—Ç–∫–∏
    if context.user_data.get("waiting_for_note"):
        project_name = context.user_data["waiting_for_note"]
        context.user_data["waiting_for_note"] = None
        note_text = question.strip()

        if PROJECTS_AVAILABLE:
            project = load_project(user_id, project_name)
            if project:
                # –°–æ—Ö—Ä–∞–Ω—è–µ–º –∑–∞–º–µ—Ç–∫—É –∫–∞–∫ –æ—Ç–¥–µ–ª—å–Ω—É—é –∑–∞–ø–∏—Å—å
                project.add_conversation_entry(
                    f"[–ó–ê–ú–ï–¢–ö–ê] {note_text[:50]}...",
                    note_text,
                    "note"
                )

                keyboard = [[InlineKeyboardButton("¬´ –ö –ø—Ä–æ–µ–∫—Ç—É", callback_data=f"proj_open_{project_name}")]]
                reply_markup = InlineKeyboardMarkup(keyboard)

                await update.message.reply_text(
                    f"‚úÖ –ó–∞–º–µ—Ç–∫–∞ –¥–æ–±–∞–≤–ª–µ–Ω–∞ –≤ –ø—Ä–æ–µ–∫—Ç **{project_name}**",
                    reply_markup=reply_markup,
                    parse_mode="Markdown"
                )
            else:
                await update.message.reply_text("‚ùå –û—à–∏–±–∫–∞ –∑–∞–≥—Ä—É–∑–∫–∏ –ø—Ä–æ–µ–∫—Ç–∞")


    # === –°–¢–ê–†–ê–Ø –û–ë–†–ê–ë–û–¢–ö–ê –®–ê–ë–õ–û–ù–û–í (–ó–ê–ú–ï–ù–ï–ù–ê –ù–ê –ò–ù–¢–ï–†–ê–ö–¢–ò–í–ù–´–ï –û–ë–†–ê–ë–û–¢–ß–ò–ö–ò v1.0) ===
    # –¢–µ–ø–µ—Ä—å –∏—Å–ø–æ–ª—å–∑—É—é—Ç—Å—è ConversationHandler –∏–∑ document_handlers.py
    # –≠—Ç–∞ —Å–µ–∫—Ü–∏—è –æ—Å—Ç–∞–≤–ª–µ–Ω–∞ –¥–ª—è —Å–æ–≤–º–µ—Å—Ç–∏–º–æ—Å—Ç–∏, –Ω–æ –±–æ–ª—å—à–µ –Ω–µ –∏—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è
    #
    # if context.user_data.get("waiting_for_template_params"):
    #     ... (—Å—Ç–∞—Ä—ã–π –∫–æ–¥ —É–¥–∞–ª—ë–Ω)

    # –ü—Ä–æ–≤–µ—Ä–∫–∞ rate limit
    if not check_rate_limit(user_id):
        await update.message.reply_text(
            "‚è±Ô∏è –°–ª–∏—à–∫–æ–º –º–Ω–æ–≥–æ –∑–∞–ø—Ä–æ—Å–æ–≤!\n\n"
            f"–í—ã –º–æ–∂–µ—Ç–µ –æ—Ç–ø—Ä–∞–≤–ª—è—Ç—å –¥–æ {RATE_LIMIT_MAX_REQUESTS} –∑–∞–ø—Ä–æ—Å–æ–≤ –≤ –º–∏–Ω—É—Ç—É.\n"
            "–ü–æ–∂–∞–ª—É–π—Å—Ç–∞, –ø–æ–¥–æ–∂–¥–∏—Ç–µ –Ω–µ–º–Ω–æ–≥–æ –∏ –ø–æ–ø—Ä–æ–±—É–π—Ç–µ —Å–Ω–æ–≤–∞."
        )
        return

    # –î–æ–±–∞–≤–ª—è–µ–º –≤–æ–ø—Ä–æ—Å –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è –≤ –∏—Å—Ç–æ—Ä–∏—é
    await add_message_to_history_async(user_id, 'user', question)

    # ============================================================================
    # LLM COUNCIL: –ê–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–æ–µ –æ–ø—Ä–µ–¥–µ–ª–µ–Ω–∏–µ —Å–ª–æ–∂–Ω—ã—Ö –≤–æ–ø—Ä–æ—Å–æ–≤
    # ============================================================================
    if LLM_COUNCIL_AVAILABLE:
        is_complex, complexity_reason = is_complex_question(question)
        
        if is_complex:
            logger.info(f"üèõÔ∏è LLM Council: –°–ª–æ–∂–Ω—ã–π –≤–æ–ø—Ä–æ—Å –æ–±–Ω–∞—Ä—É–∂–µ–Ω - {complexity_reason}")
            
            # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º —Å–æ–æ–±—â–µ–Ω–∏–µ –æ —Ä–∞–±–æ—Ç–µ –°–æ–≤–µ—Ç–∞
            council_thinking = await update.message.reply_text(
                "üèõÔ∏è **–°–ª–æ–∂–Ω—ã–π –≤–æ–ø—Ä–æ—Å ‚Äî —Å–æ–±–∏—Ä–∞—é –°–æ–≤–µ—Ç AI...**\n\n"
                f"üìä –ü—Ä–∏—á–∏–Ω–∞: _{complexity_reason}_\n\n"
                "‚è≥ –≠—Ç–∞–ø 1/3: –ü–æ–ª—É—á–∞—é –º–Ω–µ–Ω–∏—è —ç–∫—Å–ø–µ—Ä—Ç–æ–≤\n"
                "‚Ä¢ Grok ‚Äî —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏–π –∞–Ω–∞–ª–∏–∑\n"
                "‚Ä¢ Claude ‚Äî –¥–µ—Ç–∞–ª—å–Ω–∞—è —ç–∫—Å–ø–µ—Ä—Ç–∏–∑–∞\n"
                "‚Ä¢ Gemini ‚Äî –ø—Ä–∞–∫—Ç–∏—á–µ—Å–∫–∏–µ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏\n\n"
                "_–≠—Ç–æ –∑–∞–π–º—ë—Ç 30-60 —Å–µ–∫—É–Ω–¥ –¥–ª—è –º–∞–∫—Å–∏–º–∞–ª—å–Ω–æ–≥–æ –∫–∞—á–µ—Å—Ç–≤–∞..._",
                parse_mode="Markdown"
            )
            
            try:
                council = get_llm_council()
                if council:
                    # –ü–æ–ª—É—á–∞–µ–º –∫–æ–Ω—Ç–µ–∫—Å—Ç –¥–∏–∞–ª–æ–≥–∞
                    conversation_history = get_conversation_context(user_id)
                    context_text = ""
                    if conversation_history:
                        recent = conversation_history[-3:]
                        context_text = "\n".join([f"{m['role']}: {m['content'][:200]}" for m in recent])
                    
                    # –ó–∞–ø—É—Å–∫–∞–µ–º –∫–æ–Ω—Å—É–ª—å—Ç–∞—Ü–∏—é (skip_review=True –¥–ª—è —É—Å–∫–æ—Ä–µ–Ω–∏—è)
                    result = await council.consult(question, context=context_text, skip_review=True)
                    
                    if result["success"]:
                        final_answer = result["final_answer"]
                        duration = result["duration_seconds"]
                        models = result["models_used"]
                        
                        footer = f"\n\n---\nüèõÔ∏è _–°–æ–≤–µ—Ç AI: {', '.join(models)} | {duration:.1f} —Å–µ–∫_"
                        
                        max_len = 4000 - len(footer)
                        if len(final_answer) > max_len:
                            final_answer = final_answer[:max_len] + "..."
                        
                        await council_thinking.edit_text(
                            final_answer + footer,
                            parse_mode="Markdown"
                        )
                        
                        await add_message_to_history_async(user_id, 'assistant', final_answer)
                        logger.info(f"‚úÖ LLM Council auto: –æ—Ç–≤–µ—Ç –∑–∞ {duration:.1f} —Å–µ–∫ –¥–ª—è user {user_id}")
                        return  # –û—Ç–≤–µ—Ç –æ—Ç –°–æ–≤–µ—Ç–∞ –æ—Ç–ø—Ä–∞–≤–ª–µ–Ω
                    else:
                        # –°–æ–≤–µ—Ç –Ω–µ —Å–º–æ–≥ –æ—Ç–≤–µ—Ç–∏—Ç—å - –ø—Ä–æ–¥–æ–ª–∂–∞–µ–º –æ–±—ã—á–Ω—É—é –æ–±—Ä–∞–±–æ—Ç–∫—É
                        await council_thinking.delete()
                        logger.warning("LLM Council: –Ω–µ —É–¥–∞–ª–æ—Å—å –ø–æ–ª—É—á–∏—Ç—å –æ—Ç–≤–µ—Ç, fallback to single model")
                else:
                    await council_thinking.delete()
            except Exception as e:
                logger.error(f"LLM Council auto error: {e}")
                try:
                    await council_thinking.delete()
                except:
                    pass
                # –ü—Ä–æ–¥–æ–ª–∂–∞–µ–º –æ–±—ã—á–Ω—É—é –æ–±—Ä–∞–±–æ—Ç–∫—É

    # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º —Å–æ–æ–±—â–µ–Ω–∏–µ –æ –ø—Ä–æ—Ü–µ—Å—Å–µ –∏ —Å–æ—Ö—Ä–∞–Ω—è–µ–º –µ–≥–æ –¥–ª—è –ø–æ—Å–ª–µ–¥—É—é—â–µ–≥–æ —É–¥–∞–ª–µ–Ω–∏—è
    thinking_text = "ü§î –î—É–º–∞—é –Ω–∞–¥ –≤–∞—à–∏–º –≤–æ–ø—Ä–æ—Å–æ–º... \n\n–í—ã –º–æ–∂–µ—Ç–µ –Ω–µ –∂–¥–∞—Ç—å, —è –ø—Ä–∏—à–ª—é —É–≤–µ–¥–æ–º–ª–µ–Ω–∏–µ üòâ"

    # –ï—Å–ª–∏ —Ä–∞–±–æ—Ç–∞–µ–º –≤ –ø—Ä–æ–µ–∫—Ç–µ - —É–∫–∞–∑—ã–≤–∞–µ–º —ç—Ç–æ
    if PROJECTS_AVAILABLE:
        active_project = context.user_data.get("current_project")
        if active_project:
            thinking_text = f"ü§î –ê–Ω–∞–ª–∏–∑–∏—Ä—É—é –≤ –∫–æ–Ω—Ç–µ–∫—Å—Ç–µ –ø—Ä–æ–µ–∫—Ç–∞ **{active_project}**...\n\n–í—ã –º–æ–∂–µ—Ç–µ –Ω–µ –∂–¥–∞—Ç—å, —è –ø—Ä–∏—à–ª—é —É–≤–µ–¥–æ–º–ª–µ–Ω–∏–µ üòâ"

    thinking_message = await update.message.reply_text(thinking_text, parse_mode="Markdown")

    try:
        # ============================================================================
        # –û–ü–¢–ò–ú–ò–ó–ê–¶–ò–Ø: –£–º–Ω—ã–π –≤—ã–±–æ—Ä –º–æ–¥–µ–ª–∏ (Claude/Gemini/Grok)
        # ============================================================================
        if SMART_WRAPPER_AVAILABLE:
            smart_result = await smart_model_selection_text(
                question=question,
                user_id=user_id,
                thinking_message=thinking_message,
                update=update,
                context=context
            )

            # –ï—Å–ª–∏ —É–º–Ω—ã–π –≤—ã–±–æ—Ä –æ–±—Ä–∞–±–æ—Ç–∞–ª –∑–∞–ø—Ä–æ—Å - –≤—ã—Ö–æ–¥–∏–º
            if smart_result and smart_result.get("success"):
                return  # –û—Ç–≤–µ—Ç —É–∂–µ –æ—Ç–ø—Ä–∞–≤–ª–µ–Ω

        # –ï—Å–ª–∏ —É–º–Ω—ã–π –≤—ã–±–æ—Ä –Ω–µ —Å—Ä–∞–±–æ—Ç–∞–ª –∏–ª–∏ –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω - –ø—Ä–æ–¥–æ–ª–∂–∞–µ–º —Å Grok

        # –í–ê–ñ–ù–û: –≥–ª–∞–≤–Ω—ã–π —Ä–µ–∂–∏–º ‚Äî –µ–¥–∏–Ω—ã–π —É–Ω–∏–≤–µ—Ä—Å–∞–ª—å–Ω—ã–π –ø—Ä–æ–º–ø—Ç —Å –∞–≤—Ç–æ-–∞–¥–∞–ø—Ç–∞—Ü–∏–µ–π.
        # –†–æ–ª–∏ (/role) –æ—Å—Ç–∞—é—Ç—Å—è –∫–∞–∫ —Ñ—É–Ω–∫—Ü–∏—è –∏–Ω—Ç–µ—Ä—Ñ–µ–π—Å–∞, –Ω–æ –ù–ï –¥–æ–ª–∂–Ω—ã –ø–µ—Ä–µ–æ–ø—Ä–µ–¥–µ–ª—è—Ç—å system_prompt.
        # –ü–æ—ç—Ç–æ–º—É –∑–¥–µ—Å—å –Ω–µ –ø–æ–¥—Å—Ç–∞–≤–ª—è–µ–º role-based –ø—Ä–æ–º–ø—Ç.
        system_prompt = ""  # –±—É–¥–µ—Ç —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω –Ω–∏–∂–µ –ø–µ—Ä–µ–¥ –≤—ã–∑–æ–≤–æ–º –º–æ–¥–µ–ª–∏

        # --- Legacy –±–ª–æ–∫ –ø—Ä–æ–º–ø—Ç–∞ (–æ—Å—Ç–∞–≤–ª–µ–Ω –∫–∞–∫ ¬´—Ç–µ–∫—Å—Ç-—Å–ø—Ä–∞–≤–∫–∞¬ª, –Ω–µ –≤–ª–∏—è–µ—Ç –Ω–∞ —Ä–∞–±–æ—Ç—É –±–æ—Ç–∞) ---
        """


üìã –û–°–ù–û–í–ù–´–ï –ó–ê–ö–û–ù–´ –†–§:
‚Ä¢ 190-–§–ó - –ì—Ä–∞–¥–æ—Å—Ç—Ä–æ–∏—Ç–µ–ª—å–Ω—ã–π –∫–æ–¥–µ–∫—Å –†–§
‚Ä¢ 384-–§–ó - –¢–µ—Ö–Ω–∏—á–µ—Å–∫–∏–π —Ä–µ–≥–ª–∞–º–µ–Ω—Ç –æ –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç–∏ –∑–¥–∞–Ω–∏–π –∏ —Å–æ–æ—Ä—É–∂–µ–Ω–∏–π
‚Ä¢ 123-–§–ó - –¢–µ—Ö–Ω–∏—á–µ—Å–∫–∏–π —Ä–µ–≥–ª–∞–º–µ–Ω—Ç –æ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è—Ö –ø–æ–∂–∞—Ä–Ω–æ–π –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç–∏
‚Ä¢ 116-–§–ó - –û –ø—Ä–æ–º—ã—à–ª–µ–Ω–Ω–æ–π –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç–∏ –û–ü–û
‚Ä¢ 214-–§–ó - –û–± —É—á–∞—Å—Ç–∏–∏ –≤ –¥–æ–ª–µ–≤–æ–º —Å—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤–µ (–∂–∏–ª—å—ë + —ç—Å–∫—Ä–æ—É-—Å—á–µ—Ç–∞ —Å 2019)
‚Ä¢ 44-–§–ó –∏ 223-–§–ó - –ì–æ—Å–∑–∞–∫—É–ø–∫–∏ –∏ –∑–∞–∫—É–ø–∫–∏ –≥–æ—Å–∫–æ—Ä–ø–æ—Ä–∞—Ü–∏–π
‚Ä¢ 248-–§–ó - –û –≥–æ—Å—É–¥–∞—Ä—Å—Ç–≤–µ–Ω–Ω–æ–º –∫–æ–Ω—Ç—Ä–æ–ª–µ (–Ω–∞–¥–∑–æ—Ä–µ) –∏ –º—É–Ω–∏—Ü–∏–ø–∞–ª—å–Ω–æ–º –∫–æ–Ω—Ç—Ä–æ–ª–µ

üîÑ –û–ë–Ø–ó–ê–¢–ï–õ–¨–ù–´–ï –ü–†–û–¶–ï–î–£–†–´ 2025:
1. –†–ù–° (–†–∞–∑—Ä–µ—à–µ–Ω–∏–µ –Ω–∞ —Å—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤–æ) - –æ–±—è–∑–∞—Ç–µ–ª—å–Ω–æ –¥–ª—è –≤—Å–µ—Ö –∫—Ä–æ–º–µ –ò–ñ–° –¥–æ 3 —ç—Ç–∞–∂–µ–π
2. –≠–∫—Å–ø–µ—Ä—Ç–∏–∑–∞ –ø—Ä–æ–µ–∫—Ç–Ω–æ–π –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ü–∏–∏ (–≥–æ—Å—É–¥–∞—Ä—Å—Ç–≤–µ–Ω–Ω–∞—è –∏–ª–∏ –Ω–µ–≥–æ—Å—É–¥–∞—Ä—Å—Ç–≤–µ–Ω–Ω–∞—è)
3. –ó–û–° (–ó–∞–∫–ª—é—á–µ–Ω–∏–µ –æ —Å–æ–æ—Ç–≤–µ—Ç—Å—Ç–≤–∏–∏) - –≤—ã–¥–∞—ë—Ç —Å—Ç—Ä–æ–π–Ω–∞–¥–∑–æ—Ä (–ü–ü ‚Ññ 1431)
4. –£–≤–µ–¥–æ–º–ª–µ–Ω–∏–µ –æ –Ω–∞—á–∞–ª–µ –∏ –æ–∫–æ–Ω—á–∞–Ω–∏–∏ —Å—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤–∞ —á–µ—Ä–µ–∑ –ì–æ—Å—É—Å–ª—É–≥–∏
5. –í–≤–æ–¥ –≤ —ç–∫—Å–ø–ª—É–∞—Ç–∞—Ü–∏—é —á–µ—Ä–µ–∑ –ì–æ—Å—É—Å–ª—É–≥–∏ (—Ä–µ–µ—Å—Ç—Ä–æ–≤–∞—è –º–æ–¥–µ–ª—å —Å 2022)

üë∑ –°–†–û –ò –ö–í–ê–õ–ò–§–ò–ö–ê–¶–ò–Ø 2025:
‚Ä¢ –ì–µ–Ω–ø–æ–¥—Ä—è–¥—á–∏–∫ - –û–ë–Ø–ó–ê–¢–ï–õ–¨–ù–û —á–ª–µ–Ω—Å—Ç–≤–æ –≤ –°–†–û (–∫–æ–º–ø–µ–Ω—Å–∞—Ü–∏–æ–Ω–Ω—ã–µ —Ñ–æ–Ω–¥—ã)
‚Ä¢ –°—É–±–ø–æ–¥—Ä—è–¥—á–∏–∫ - —Å 2023 –±–µ–∑ –°–†–û –µ—Å–ª–∏ –¥–æ–≥–æ–≤–æ—Ä ‚â§ 3 –º–ª–Ω —Ä—É–±.
‚Ä¢ –ì–ò–ü –∏ –ì–ê–ü - –û–ë–Ø–ó–ê–¢–ï–õ–¨–ù–û –ù–û–ö (–Ω–µ–∑–∞–≤–∏—Å–∏–º–∞—è –æ—Ü–µ–Ω–∫–∞ –∫–≤–∞–ª–∏—Ñ–∏–∫–∞—Ü–∏–∏) + –∑–∞–ø–∏—Å—å –≤ –ù–†–°

üíª –¢–ò–ú/BIM:
‚Ä¢ 2022-2025: –æ–±—è–∑–∞—Ç–µ–ª—å–Ω–æ –¥–ª—è –æ–±—ä–µ–∫—Ç–æ–≤ –≥–æ—Å–∑–∞–∫–∞–∑–∞
‚Ä¢ –° 01.01.2027: –æ–±—è–∑–∞—Ç–µ–ª—å–Ω–æ –¥–ª—è –í–°–ï–• –æ–±—ä–µ–∫—Ç–æ–≤ –∫–∞–ø—Å—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤–∞ (OpenBIM Level 2)
‚Ä¢ –°—Ç–∞–Ω–¥–∞—Ä—Ç—ã: –ì–û–°–¢ –† 57580, –°–ü 301.1325800.2017, –ü—Ä–∏–∫–∞–∑ ‚Ññ 926/–ø—Ä

üí∞ –¶–ï–ù–û–û–ë–†–ê–ó–û–í–ê–ù–ò–ï 2025:
‚Ä¢ –§–ì–ò–° –¶–° - –æ–±—è–∑–∞—Ç–µ–ª—å–Ω–æ
‚Ä¢ –†–µ—Å—É—Ä—Å–Ω—ã–π –º–µ—Ç–æ–¥ - –æ—Å–Ω–æ–≤–Ω–æ–π —Å 2025
‚Ä¢ –°–º–µ—Ç—ã: –ì–≠–°–ù/–§–ï–† + –∏–Ω–¥–µ–∫—Å—ã –ú–∏–Ω—Å—Ç—Ä–æ—è

üè≠ –ü–†–û–ú–´–®–õ–ï–ù–ù–û–ï –°–¢–†–û–ò–¢–ï–õ–¨–°–¢–í–û:
–ö–ª–∞—Å—Å—ã –æ–ø–∞—Å–Ω–æ—Å—Ç–∏ –û–ü–û:
‚Ä¢ I - —á—Ä–µ–∑–≤—ã—á–∞–π–Ω–æ –≤—ã—Å–æ–∫–∏–π (–Ω–µ—Ñ—Ç–µ—Ö–∏–º, –ê–≠–°): –¥–µ–∫–ª–∞—Ä–∞—Ü–∏—è –ø—Ä–æ–º–±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç–∏ –æ–±—è–∑–∞—Ç–µ–ª—å–Ω–∞
‚Ä¢ II - –≤—ã—Å–æ–∫–∏–π (—Ö–∏–º–∑–∞–≤–æ–¥—ã): –ª–∏—Ü–µ–Ω–∑–∏—è –†–æ—Å—Ç–µ—Ö–Ω–∞–¥–∑–æ—Ä–∞ + –¥–µ–∫–ª–∞—Ä–∞—Ü–∏—è
‚Ä¢ III - —Å—Ä–µ–¥–Ω–∏–π (–∑–∞–≤–æ–¥—ã): –ª–∏—Ü–µ–Ω–∑–∏—è –†–æ—Å—Ç–µ—Ö–Ω–∞–¥–∑–æ—Ä–∞ + –ü–ú–õ–ê
‚Ä¢ IV - –Ω–∏–∑–∫–∏–π (–º–µ–ª–∫–∏–µ –ø—Ä–æ–∏–∑–≤–æ–¥—Å—Ç–≤–∞): —Ä–µ–≥–∏—Å—Ç—Ä–∞—Ü–∏—è –≤ —Ä–µ–µ—Å—Ç—Ä–µ

–û–±—è–∑–∞—Ç–µ–ª—å–Ω—ã–µ –¥–æ–∫—É–º–µ–Ω—Ç—ã:
‚Ä¢ –°–ó–ó (—Å–∞–Ω–∏—Ç–∞—Ä–Ω–æ-–∑–∞—â–∏—Ç–Ω–∞—è –∑–æ–Ω–∞) - –°–∞–Ω–ü–∏–ù 2.2.1/2.1.1.1200-03
‚Ä¢ –õ–∏—Ü–µ–Ω–∑–∏—è –†–æ—Å—Ç–µ—Ö–Ω–∞–¥–∑–æ—Ä–∞ (–¥–ª—è II-III –∫–ª–∞—Å—Å–æ–≤)
‚Ä¢ –ü–ú–õ–ê (–ü–ª–∞–Ω –º–µ—Ä–æ–ø—Ä–∏—è—Ç–∏–π –ø–æ –ª–æ–∫–∞–ª–∏–∑–∞—Ü–∏–∏ –∏ –ª–∏–∫–≤–∏–¥–∞—Ü–∏–∏ –∞–≤–∞—Ä–∏–π)
‚Ä¢ –î–µ–∫–ª–∞—Ä–∞—Ü–∏—è –ø—Ä–æ–º–±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç–∏ (–¥–ª—è I-II –∫–ª–∞—Å—Å–æ–≤)
‚Ä¢ –†–∞–∑–¥–µ–ª 9 –ü–î - –ø—Ä–æ–º—ã—à–ª–µ–Ω–Ω–∞—è –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç—å
‚Ä¢ –ò–¢–ú –ì–û–ß–°, –≤–∑—Ä—ã–≤–æ–∑–∞—â–∏—Ç–∞ (–ì–û–°–¢ IEC 60079), –º–æ–ª–Ω–∏–µ–∑–∞—â–∏—Ç–∞

üèòÔ∏è –ì–†–ê–ñ–î–ê–ù–°–ö–û–ï –°–¢–†–û–ò–¢–ï–õ–¨–°–¢–í–û (–∂–∏–ª—å—ë, —à–∫–æ–ª—ã, –±–æ–ª—å–Ω–∏—Ü—ã):
–¢—Ä–µ–±–æ–≤–∞–Ω–∏—è 2025:
‚Ä¢ 214-–§–ó + —ç—Å–∫—Ä–æ—É-—Å—á–µ—Ç–∞ (–æ–±—è–∑–∞—Ç–µ–ª—å–Ω–æ —Å 2019)
‚Ä¢ –°–ü 54.13330.2022 - –∂–∏–ª—ã–µ –∑–¥–∞–Ω–∏—è
‚Ä¢ –°–ü 59.13330.2020 - –¥–æ—Å—Ç—É–ø–Ω–æ—Å—Ç—å –¥–ª—è –ú–ì–ù
‚Ä¢ –ö–ª–∞—Å—Å —ç–Ω–µ—Ä–≥–æ—ç—Ñ—Ñ–µ–∫—Ç–∏–≤–Ω–æ—Å—Ç–∏ –ù–ï –ù–ò–ñ–ï ¬´–°¬ª
‚Ä¢ –£–º–Ω—ã–µ —Å—á—ë—Ç—á–∏–∫–∏ (—ç–ª–µ–∫—Ç—Ä–æ, —Ç–µ–ø–ª–æ, –≤–æ–¥–∞) - –û–ë–Ø–ó–ê–¢–ï–õ–¨–ù–û
‚Ä¢ –£–º–Ω—ã–π –¥–æ–º (–≤–∏–¥–µ–æ–Ω–∞–±–ª—é–¥–µ–Ω–∏–µ, –∫–æ–Ω—Ç—Ä–æ–ª—å –¥–æ—Å—Ç—É–ø–∞, –°–û–£–≠)
‚Ä¢ 5-10% –º–∞—à–∏–Ω–æ–º–µ—Å—Ç –¥–ª—è –∏–Ω–≤–∞–ª–∏–¥–æ–≤
‚Ä¢ –≠–∫—Å–ø–µ—Ä—Ç–∏–∑–∞: –æ–±—è–∑–∞—Ç–µ–ª—å–Ω–∞ –¥–ª—è >3 —ç—Ç–∞–∂–µ–π –∏–ª–∏ >1500 –º¬≤

üè¢ –ö–û–ú–ú–ï–†–ß–ï–°–ö–û–ï –°–¢–†–û–ò–¢–ï–õ–¨–°–¢–í–û:
–¢–†–¶:
‚Ä¢ –°–ü 4.13130.2013 - –∂—ë—Å—Ç–∫–∏–µ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è –ø–æ —ç–≤–∞–∫—É–∞—Ü–∏–∏
‚Ä¢ 2 –Ω–µ–∑–∞–≤–∏—Å–∏–º—ã—Ö —ç–≤–∞–∫—É–∞—Ü–∏–æ–Ω–Ω—ã—Ö –≤—ã—Ö–æ–¥–∞ —Å –∫–∞–∂–¥–æ–≥–æ —ç—Ç–∞–∂–∞
‚Ä¢ –î—ã–º–æ—É–¥–∞–ª–µ–Ω–∏–µ, —Å–ø—Ä–∏–Ω–∫–ª–µ—Ä—ã, –°–û–£–≠ 3-5 —Ç–∏–ø–∞
‚Ä¢ –ü—Ä–æ–≤–µ—Ä–∫–∏ –∫–∞–∂–¥—ã–µ 2 –≥–æ–¥–∞ (–ø–æ—Å–ª–µ "–ó–∏–º–Ω–µ–π –≤–∏—à–Ω–∏")

–ì–æ—Å—Ç–∏–Ω–∏—Ü—ã:
‚Ä¢ –ö–ª–∞—Å—Å–∏—Ñ–∏–∫–∞—Ü–∏—è –ø–æ "–∑–≤—ë–∑–¥–∞–º" (–ü—Ä–∏–∫–∞–∑ –ú–∏–Ω–∫—É–ª—å—Ç—É—Ä—ã ‚Ññ 1215)
‚Ä¢ –°–û–£–≠ –Ω–µ –Ω–∏–∂–µ 3 —Ç–∏–ø–∞

–ê–ø–∞—Ä—Ç–∞–º–µ–Ω—Ç—ã:
‚Ä¢ –Æ—Ä–∏–¥–∏—á–µ—Å–∫–∏ - –Ω–µ–∂–∏–ª—ã–µ –ø–æ–º–µ—â–µ–Ω–∏—è
‚Ä¢ –ë–µ–∑ 214-–§–ó –∏ —ç—Å–∫—Ä–æ—É
‚Ä¢ –ù–µ–ª—å–∑—è –ø—Ä–æ–ø–∏—Å–∞—Ç—å—Å—è, –≤—ã—à–µ –Ω–∞–ª–æ–≥

üìä –ö–õ–Æ–ß–ï–í–´–ï –ù–û–†–ú–ê–¢–ò–í–´ 2025:
‚Ä¢ –ü–ü ‚Ññ 985 + –ü—Ä–∏–ª–æ–∂–µ–Ω–∏–µ - —á—Ç–æ –ø–æ–¥–ª–µ–∂–∏—Ç —ç–∫—Å–ø–µ—Ä—Ç–∏–∑–µ
‚Ä¢ –ü–ü ‚Ññ 1431 - –ø–æ—Ä—è–¥–æ–∫ –≤—ã–¥–∞—á–∏ –ó–û–°
‚Ä¢ –ü—Ä–∏–∫–∞–∑ –ú–∏–Ω—Å—Ç—Ä–æ—è ‚Ññ 783/–ø—Ä - —Å—Ç—Ä–æ–π–∫–æ–Ω—Ç—Ä–æ–ª—å –∏ –Ω–∞–¥–∑–æ—Ä
‚Ä¢ –ü—Ä–∏–∫–∞–∑ ‚Ññ 926/–ø—Ä - –¢–ò–ú/BIM
‚Ä¢ –°–ü 48.13330.2019 - –û—Ä–≥–∞–Ω–∏–∑–∞—Ü–∏—è —Å—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤–∞
‚Ä¢ –†–î-11-02-2006 - –∏—Å–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω–∞—è –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ü–∏—è
‚Ä¢ –ü–ü ‚Ññ 815 - –∫–∞–∫–∏–µ —Ä–∞–∑–¥–µ–ª—ã –ü–î –æ–±—è–∑–∞—Ç–µ–ª—å–Ω—ã
‚Ä¢ –ì–û–°–¢ –† 57580.1-2017 - —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è –∫ BIM-–º–æ–¥–µ–ª—è–º

‚úÖ –ß–ï–ö-–õ–ò–°–¢ –ï–ñ–ï–î–ù–ï–í–ù–´–• –ü–†–û–í–ï–†–û–ö:
1. –î–µ–π—Å—Ç–≤—É—é—â–µ–µ –†–ù–°
2. –ü–î –ø—Ä–æ—à–ª–∞ —ç–∫—Å–ø–µ—Ä—Ç–∏–∑—É
3. –ß–ª–µ–Ω—Å—Ç–≤–æ –≥–µ–Ω–ø–æ–¥—Ä—è–¥—á–∏–∫–∞ –≤ –°–†–û
4. –ò—Å–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω–∞—è –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ü–∏—è (–∞–∫—Ç—ã, –∂—É—Ä–Ω–∞–ª—ã)
5. –õ–∏—Ü–∞ –ø–æ —Ç–µ—Ö–Ω–∞–¥–∑–æ—Ä—É –∏ —Å—Ç—Ä–æ–π–∫–æ–Ω—Ç—Ä–æ–ª—é –Ω–∞–∑–Ω–∞—á–µ–Ω—ã
6. –î–ª—è –ø—Ä–æ–º–∫–∏: –û–ü–û –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω–æ, –°–ó–ó –µ—Å—Ç—å
7. –ñ—É—Ä–Ω–∞–ª –≤—Ö–æ–¥–Ω–æ–≥–æ –∫–æ–Ω—Ç—Ä–æ–ª—è –º–∞—Ç–µ—Ä–∏–∞–ª–æ–≤
8. –î–æ–ø—É—Å–∫–∏ —Å–ø–µ—Ü–∏–∞–ª–∏—Å—Ç–æ–≤ (—Å–≤–∞—Ä—â–∏–∫–∏, —Å—Ç—Ä–æ–ø–∞–ª—å—â–∏–∫–∏, –∫—Ä–∞–Ω–æ–≤—â–∏–∫–∏)
9. –ê–∫—Ç—ã –ö–°-2, –ö–°-3 –ø–æ–¥–ø–∏—Å—ã–≤–∞—é—Ç—Å—è –µ–∂–µ–º–µ—Å—è—á–Ω–æ

üîÆ –¢–†–ï–ù–î–´ 2025-2027:
‚Ä¢ –ü–æ–ª–Ω—ã–π –ø–µ—Ä–µ—Ö–æ–¥ –Ω–∞ —Ä–µ–µ—Å—Ç—Ä–æ–≤—É—é –º–æ–¥–µ–ª—å (—á–µ—Ä–µ–∑ –ì–æ—Å—É—Å–ª—É–≥–∏)
‚Ä¢ –¢–ò–ú Level 2 –æ–±—è–∑–∞—Ç–µ–ª–µ–Ω —Å 2027
‚Ä¢ ESG –∏ "–∑–µ–ª—ë–Ω–æ–µ" —Å—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤–æ (LEED, BREEAM)
‚Ä¢ –†–æ–±–æ—Ç–∏–∑–∞—Ü–∏—è –∏ –¥—Ä–æ–Ω—ã (–¥–æ–ø—É—Å–∫–∞—é—Ç—Å—è –¥–ª—è —Å—Ç—Ä–æ–π–∫–æ–Ω—Ç—Ä–æ–ª—è)
‚Ä¢ –û–±—è–∑–∞—Ç–µ–ª—å–Ω–æ–µ —Å—Ç—Ä–∞—Ö–æ–≤–∞–Ω–∏–µ –æ—Ç–≤–µ—Ç—Å—Ç–≤–µ–Ω–Ω–æ—Å—Ç–∏ –∑–∞—Å—Ç—Ä–æ–π—â–∏–∫–∞ (—Å 2025 –¥–ª—è –≤—Å–µ—Ö)

**–ü–†–ê–ö–¢–ò–ß–ï–°–ö–ò–ï –ó–ù–ê–ù–ò–Ø –ü–õ–û–©–ê–î–ö–ò:**

ü¶∫ –û–•–†–ê–ù–ê –¢–†–£–î–ê (HSE):
‚Ä¢ –†–∞–±–æ—Ç–∞ –Ω–∞ –≤—ã—Å–æ—Ç–µ >1.8–º (–ü—Ä–∏–∫–∞–∑ –ú–∏–Ω—Ç—Ä—É–¥–∞ ‚Ññ 782–Ω): –≥—Ä—É–ø–ø—ã –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç–∏ 1/2/3, –°–ò–ó –æ–±—è–∑–∞—Ç–µ–ª—å–Ω—ã
‚Ä¢ –≠–ª–µ–∫—Ç—Ä–æ–±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç—å: –≥—Ä—É–ø–ø—ã –¥–æ 1000–í / –≤—ã—à–µ 1000–í, –≤—Ä–µ–º–µ–Ω–Ω—ã–µ —Å–µ—Ç–∏ 380/220–í
‚Ä¢ –ü–æ–≥—Ä—É–∑–æ—á–Ω–æ-—Ä–∞–∑–≥—Ä—É–∑–æ—á–Ω—ã–µ: —Å—Ö–µ–º—ã —Å—Ç—Ä–æ–ø–æ–≤–∫–∏, —Å–∏–≥–Ω–∞–ª—ã –∫—Ä–∞–Ω–æ–≤—â–∏–∫—É

üèóÔ∏è –¢–ï–•–ù–û–õ–û–ì–ò–Ø –ü–†–û–ò–ó–í–û–î–°–¢–í–ê:
‚Ä¢ –ë–µ—Ç–æ–Ω: –Ω–∞–±–æ—Ä –ø—Ä–æ—á–Ω–æ—Å—Ç–∏ –ø—Ä–∏ +20¬∞C = 28 —Å—É—Ç–æ–∫ –¥–æ 100%, –ø—Ä–∏ -10¬∞C —Å –ø—Ä–æ–≥—Ä–µ–≤–æ–º = 7-10 —Å—É—Ç–æ–∫ –¥–æ 70%
‚Ä¢ –ó–∏–º–Ω–µ–µ –±–µ—Ç–æ–Ω–∏—Ä–æ–≤–∞–Ω–∏–µ: –¥–æ–±–∞–≤–∫–∏ (–Ω–∏—Ç—Ä–∏—Ç –Ω–∞—Ç—Ä–∏—è 3-5%), –º–µ—Ç–æ–¥—ã –ø—Ä–æ–≥—Ä–µ–≤–∞, —Ç–µ–º–ø–µ—Ä–∞—Ç—É—Ä–Ω—ã–π –∫–æ–Ω—Ç—Ä–æ–ª—å
‚Ä¢ –ê—Ä–º–∞—Ç—É—Ä–∞ –ê400: –Ω–∞—Ö–ª–µ—Å—Ç –≤ —Ä–∞—Å—Ç—è–∂–µ–Ω–∏–∏ = 40d (–¥–ª—è d=16–º–º = 640–º–º), –∑–∞—â–∏—Ç–Ω—ã–π —Å–ª–æ–π –ø–æ –°–ü 63
‚Ä¢ –í—Ö–æ–¥–Ω–æ–π –∫–æ–Ω—Ç—Ä–æ–ª—å: –¥–æ–ø—É—Å–∫–∏ –Ω–∞ –∞—Ä–º–∞—Ç—É—Ä—É ¬±0.3–º–º, –±–µ—Ç–æ–Ω (–ø—Ä–æ—á–Ω–æ—Å—Ç—å, –ø–æ–¥–≤–∏–∂–Ω–æ—Å—Ç—å), –∫–∏—Ä–ø–∏—á

üí∞ –°–ú–ï–¢–ù–û–ï –î–ï–õ–û:
‚Ä¢ –ê–∫—Ç—ã –ö–°-2/–ö–°-3: —Ç–∏–ø–∏—á–Ω—ã–µ –æ—à–∏–±–∫–∏ (–Ω–µ—Å–æ–æ—Ç–≤–µ—Ç—Å—Ç–≤–∏–µ –æ–±—ä–µ–º–æ–≤, –æ—Ç—Å—É—Ç—Å—Ç–≤–∏–µ –ø–æ–¥–ø–∏—Å–µ–π), —Å—Ä–æ–∫–∏ —Å–æ–≥–ª–∞—Å–æ–≤–∞–Ω–∏—è
‚Ä¢ –î–∞–≤–∞–ª—å—á–µ—Å–∫–∏–µ –º–∞—Ç–µ—Ä–∏–∞–ª—ã: —Ñ–æ—Ä–º–∞ –ú-29, —É—á–µ—Ç –≤ —Å–º–µ—Ç–µ, –¥–æ—Å—Ç–∞–≤–∫–∞ - –ø–æ –¥–æ–≥–æ–≤–æ—Ä–µ–Ω–Ω–æ—Å—Ç–∏
‚Ä¢ –ù–µ–ø—Ä–µ–¥–≤–∏–¥–µ–Ω–Ω—ã–µ —Ä–∞—Å—Ö–æ–¥—ã: 2-3% –æ—Ç —Å–º–µ—Ç—ã, –æ—Ñ–æ—Ä–º–ª–µ–Ω–∏–µ –¥–æ–ø.—Å–æ–≥–ª–∞—à–µ–Ω–∏–π

‚öñÔ∏è –Æ–†–ò–î–ò–ß–ï–°–ö–ò–ï –í–û–ü–†–û–°–´:
‚Ä¢ –ù–µ—É—Å—Ç–æ–π–∫–∞: –ì–ö –†–§ —Å—Ç. 330, —Ä–∞—Å—á–µ—Ç = 1/300 –∫–ª—é—á–µ–≤–æ–π —Å—Ç–∞–≤–∫–∏ –¶–ë –æ—Ç —Å—É–º–º—ã –¥–æ–ª–≥–∞ –∑–∞ –∫–∞–∂–¥—ã–π –¥–µ–Ω—å
‚Ä¢ –°–∫—Ä—ã—Ç—ã–µ —Ä–∞–±–æ—Ç—ã: –∞–∫—Ç—ã –æ—Å–≤–∏–¥–µ—Ç–µ–ª—å—Å—Ç–≤–æ–≤–∞–Ω–∏—è, 3 –¥–Ω—è –Ω–∞ –≤—ã–∑–æ–≤ –∑–∞–∫–∞–∑—á–∏–∫–∞, –º–æ–∂–Ω–æ –∑–∞–∫—Ä—ã–≤–∞—Ç—å –±–µ–∑ –Ω–µ–≥–æ –ø—Ä–∏ –Ω–µ—è–≤–∫–µ
‚Ä¢ –ü—Ä–µ—Ç–µ–Ω–∑–∏–æ–Ω–Ω–∞—è —Ä–∞–±–æ—Ç–∞: –¥–æ—Å—É–¥–µ–±–Ω–æ–µ —É—Ä–µ–≥—É–ª–∏—Ä–æ–≤–∞–Ω–∏–µ 30 –¥–Ω–µ–π, –æ–±—Ä–∞–∑—Ü—ã –ø–∏—Å–µ–º, —Å—Ä–æ–∫–∏ –æ—Ç–≤–µ—Ç–æ–≤

üìä –£–ü–†–ê–í–õ–ï–ù–ò–ï –ü–†–û–ï–ö–¢–ê–ú–ò:
‚Ä¢ –ü–ª–∞–Ω–∏—Ä–æ–≤–∞–Ω–∏–µ: –¥–∏–∞–≥—Ä–∞–º–º–∞ –ì–∞–Ω—Ç–∞, —Å–µ—Ç–µ–≤—ã–µ –≥—Ä–∞—Ñ–∏–∫–∏, –∫—Ä–∏—Ç–∏—á–µ—Å–∫–∏–π –ø—É—Ç—å
‚Ä¢ –ü—Ä–æ—Ç–æ–∫–æ–ª—ã —Å–æ–≤–µ—â–∞–Ω–∏–π: –æ–±—è–∑–∞—Ç–µ–ª—å–Ω—ã–µ –ø—É–Ω–∫—Ç—ã (–¥–∞—Ç–∞, —É—á–∞—Å—Ç–Ω–∏–∫–∏, –≤–æ–ø—Ä–æ—Å—ã, —Ä–µ—à–µ–Ω–∏—è, —Å—Ä–æ–∫–∏, –æ—Ç–≤–µ—Ç—Å—Ç–≤–µ–Ω–Ω—ã–µ)
‚Ä¢ –†–∞—Å—á–µ—Ç —á–∏—Å–ª–µ–Ω–Ω–æ—Å—Ç–∏: N = V / (H √ó K √ó T), –≥–¥–µ V-–æ–±—ä–µ–º, H-–Ω–æ—Ä–º–∞ –≤—ã—Ä–∞–±–æ—Ç–∫–∏, K-–∫–æ—ç—Ñ.—Å–æ–≤–º–µ—â–µ–Ω–∏—è, T-–≤—Ä–µ–º—è
‚Ä¢ –ö–æ–Ω—Ñ–ª–∏–∫—Ç—ã: –¥–æ–∫—É–º–µ–Ω—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ –≤—Å–µ–≥–æ, –ø—Ä–µ—Ç–µ–Ω–∑–∏–∏ –≤ –ø–∏—Å—å–º–µ–Ω–Ω–æ–º –≤–∏–¥–µ, —Å—Å—ã–ª–∫–∏ –Ω–∞ –¥–æ–≥–æ–≤–æ—Ä

**–†–ê–°–®–ò–†–ï–ù–ù–´–ï –ü–†–ê–ö–¢–ò–ß–ï–°–ö–ò–ï –ó–ù–ê–ù–ò–Ø:**

üë• –ö–ê–î–†–´ –ò –ú–ò–ì–†–ê–¶–ò–Ø:
‚Ä¢ –ò–Ω–æ—Å—Ç—Ä–∞–Ω–Ω—ã–µ —Ä–∞–±–æ—Ç–Ω–∏–∫–∏: –ø–∞—Ç–µ–Ω—Ç—ã (—Å—Ä–æ–∫ 1-12 –º–µ—Å), —É–≤–µ–¥–æ–º–ª–µ–Ω–∏–µ –ú–í–î (3 –¥–Ω—è), —à—Ç—Ä–∞—Ñ 400-800 —Ç—ã—Å.—Ä—É–±
‚Ä¢ –í–∞—Ö—Ç–æ–≤—ã–π –º–µ—Ç–æ–¥: –Ω–∞–¥–±–∞–≤–∫–∞ = –¥–Ω–µ–≤–Ω–∞—è —Å—Ç–∞–≤–∫–∞ √ó –¥–Ω–∏ –≤–∞—Ö—Ç—ã, –º–µ–∂–¥—É–≤–∞—Ö—Ç–æ–≤—ã–π –æ—Ç–¥—ã—Ö ‚â• –≤—Ä–µ–º—è –≤–∞—Ö—Ç—ã
‚Ä¢ –ö–≤–∞–ª–∏—Ñ–∏–∫–∞—Ü–∏—è: –ù–ê–ö–° –¥–ª—è —Å–≤–∞—Ä—â–∏–∫–æ–≤ (2 –≥–æ–¥–∞), —Å—Ç—Ä–æ–ø–∞–ª—å—â–∏–∫–∏ (72 —á –æ–±—É—á–µ–Ω–∏—è), —ç–ª–µ–∫—Ç—Ä–∏–∫–∏ (–≥—Ä—É–ø–ø—ã 2-5)

üìê –ì–ï–û–î–ï–ó–ò–Ø:
‚Ä¢ –ì–†–û –æ–±—è–∑–∞—Ç–µ–ª—å–Ω–∞! –î–æ–ø—É—Å–∫–∏: —Ñ—É–Ω–¥–∞–º–µ–Ω—Ç—ã ¬±10–º–º, –∫–æ–ª–æ–Ω–Ω—ã ¬±5–º–º/—ç—Ç–∞–∂, –ø–µ—Ä–µ–∫—Ä—ã—Ç–∏—è ¬±10–º–º
‚Ä¢ –ü—Ä–∏–±–æ—Ä—ã: —Ç–∞—Ö–µ–æ–º–µ—Ç—Ä (—É–≥–ª—ã+—Ä–∞—Å—Å—Ç), –Ω–∏–≤–µ–ª–∏—Ä (–≤—ã—Å–æ—Ç—ã), GNSS-RTK (¬±10-20–º–º)
‚Ä¢ –ò—Å–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω–∞—è —Å—ä–µ–º–∫–∞: —Ñ—É–Ω–¥–∞–º–µ–Ω—Ç—ã, –∫–æ–ª–æ–Ω–Ω—ã, –ø–µ—Ä–µ–∫—Ä—ã—Ç–∏—è (–¥–æ –∑–∞—Å—ã–ø–∫–∏/–∑–∞–∫—Ä—ã—Ç–∏—è)

üì¶ –õ–û–ì–ò–°–¢–ò–ö–ê:
‚Ä¢ –¶–µ–º–µ–Ω—Ç: –∑–∞–∫—Ä—ã—Ç—ã–π —Å–∫–ª–∞–¥, —Å—Ä–æ–∫ 2-3 –º–µ—Å –≤ –º–µ—à–∫–∞—Ö, –æ–∫–∞–º–µ–Ω–µ–≤–∞–µ—Ç –æ—Ç –≤–ª–∞–≥–∏
‚Ä¢ –ê—Ä–º–∞—Ç—É—Ä–∞: –ø–æ–≤–µ—Ä—Ö–Ω–æ—Å—Ç–Ω–∞—è —Ä–∂–∞–≤—á–∏–Ω–∞ –¥–æ–ø—É—Å—Ç–∏–º–∞, –≥–ª—É–±–æ–∫–∞—è - –±—Ä–∞–∫
‚Ä¢ –ë–µ—Ç–æ–Ω: –ø—Ä–æ–≤–µ—Ä–∫–∞ –æ—Å–∞–¥–∫–∏ –∫–æ–Ω—É—Å–∞ (–ü1-–ü4), –æ—Ç–±–æ—Ä –∫—É–±–∏–∫–æ–≤ (1 —Å–µ—Ä–∏—è –Ω–∞ 100 –º¬≥)
‚Ä¢ –û–ø–∞–ª—É–±–∫–∞: –æ–±–æ—Ä–∞—á–∏–≤–∞–µ–º–æ—Å—Ç—å = 30 –¥–Ω–µ–π / (1+7+1) = 3.3 —Ä–∞–∑–∞/–º–µ—Å

üåç –≠–ö–û–õ–û–ì–ò–Ø:
‚Ä¢ –û—Ç—Ö–æ–¥—ã: —Ä–∞–∑–¥–µ–ª–µ–Ω–∏–µ (–±–µ—Ç–æ–Ω, –º–µ—Ç–∞–ª–ª, –¥–µ—Ä–µ–≤–æ), —Ç–∞–ª–æ–Ω—ã –û–°–°–∏–ì (–ú–æ—Å–∫–≤–∞), —à—Ç—Ä–∞—Ñ –¥–æ 600 —Ç—ã—Å
‚Ä¢ –ú–æ–π–∫–∞ –∫–æ–ª–µ—Å: –æ–±—è–∑–∞—Ç–µ–ª—å–Ω–∞ –ø—Ä–∏ –≤—ã–µ–∑–¥–µ –Ω–∞ –¥–æ—Ä–æ–≥–∏
‚Ä¢ –®—É–º: 23:00-7:00 —Ç–∏—à–∏–Ω–∞ (—Ä–µ–≥–∏–æ–Ω–∞–ª—å–Ω—ã–π –∑–∞–∫–æ–Ω), –∏—Å–∫–ª—é—á–µ–Ω–∏–µ - —Ä–∞–∑—Ä–µ—à–µ–Ω–∏–µ –∞–¥–º–∏–Ω–∏—Å—Ç—Ä–∞—Ü–∏–∏

‚ùÑÔ∏è –°–ü–ï–¶–£–°–õ–û–í–ò–Ø:
‚Ä¢ –ó–∏–º–Ω–µ–µ –±–µ—Ç–æ–Ω–∏—Ä–æ–≤–∞–Ω–∏–µ: —ç–ª–µ–∫—Ç—Ä–æ–ø—Ä–æ–≥—Ä–µ–≤ (100 –∫–í—Ç¬∑—á/–º¬≥), –Ω–∏—Ç—Ä–∏—Ç –Ω–∞—Ç—Ä–∏—è 3-5%, –∫–æ–Ω—Ç—Ä–æ–ª—å T –∫–∞–∂–¥—ã–µ 4 —á
‚Ä¢ –°–µ–π—Å–º–∏–∫–∞: –∑–∞–º–∫–Ω—É—Ç—ã–µ –∫–∞—Ä–∫–∞—Å—ã, —Ö–æ–º—É—Ç—ã 100–º–º, –Ω–∞—Ö–ª–µ—Å—Ç—ã 50d, –¥–∏–∞—Ñ—Ä–∞–≥–º—ã –∂–µ—Å—Ç–∫–æ—Å—Ç–∏
‚Ä¢ –í–µ—á–Ω–∞—è –º–µ—Ä–∑–ª–æ—Ç–∞: —Ç–µ—Ä–º–æ—Å—Ç–∞–±–∏–ª–∏–∑–∞—Ç–æ—Ä—ã, –ø—Ä–æ–≤–µ—Ç—Ä–∏–≤–∞–µ–º–æ–µ –ø–æ–¥–ø–æ–ª—å–µ

‚ö° –ò–ù–ñ–ï–ù–ï–†–ù–´–ï –°–ï–¢–ò:
‚Ä¢ –í—Ä–µ–º–µ–Ω–Ω–æ–µ —ç–ª–µ–∫—Ç—Ä–æ: –±–∞—à–µ–Ω–Ω—ã–π –∫—Ä–∞–Ω 40-80 –∫–í—Ç, –±–µ—Ç–æ–Ω–æ–Ω–∞—Å–æ—Å 40-50 –∫–í—Ç, —Å–≤–∞—Ä–∫–∞ 5-10 –∫–í—Ç
‚Ä¢ –í–æ–¥–∞: 15-25 –ª/—á–µ–ª –ø–∏—Ç—å–µ, 30-50 –ª –¥—É—à, –ø–æ–∂–∞—Ä–æ—Ç—É—à–µ–Ω–∏–µ 10-20 –ª/—Å
‚Ä¢ –ü–µ—Ä–µ—Å–µ—á–∫–∏: –≥–∏–ª—å–∑—ã –æ–±—è–∑–∞—Ç–µ–ª—å–Ω—ã (d+50-100 –º–º), —É—Å–∏–ª–µ–Ω–∏–µ –µ—Å–ª–∏ –æ—Ç–≤–µ—Ä—Å—Ç–∏–µ >1/3 –≤—ã—Å–æ—Ç—ã –±–∞–ª–∫–∏
‚Ä¢ –ü–ù–† ‚â† –º–æ–Ω—Ç–∞–∂: –ü–ù–† = –Ω–∞—Å—Ç—Ä–æ–π–∫–∞ + –∏—Å–ø—ã—Ç–∞–Ω–∏—è + –∞–∫—Ç—ã

**–Æ–†–ò–î–ò–ß–ï–°–ö–ê–Ø –ó–ê–©–ò–¢–ê –ò –ü–†–ï–¢–ï–ù–ó–ò–û–ù–ù–ê–Ø –†–ê–ë–û–¢–ê:**

‚öñÔ∏è –ì–†–ê–ñ–î–ê–ù–°–ö–ò–ô –ö–û–î–ï–ö–° (–ü–æ–¥—Ä—è–¥, –≥–ª. 37):
‚Ä¢ –°—Ç. 716 –ì–ö –†–§: –æ–±—è–∑–∞–Ω–Ω–æ—Å—Ç—å –ø–æ–¥—Ä—è–¥—á–∏–∫–∞ –ø—Ä–µ–¥—É–ø—Ä–µ–¥–∏—Ç—å –∑–∞–∫–∞–∑—á–∏–∫–∞ –æ –ø—Ä–æ–±–ª–µ–º–∞—Ö (–Ω–µ–≥–æ–¥–Ω–æ—Å—Ç—å/–Ω–µ–ø—Ä–∏–≥–æ–¥–Ω–æ—Å—Ç—å –º–∞—Ç–µ—Ä–∏–∞–ª–æ–≤, —É–∫–∞–∑–∞–Ω–∏–π –∑–∞–∫–∞–∑—á–∏–∫–∞)
‚Ä¢ –°—Ç. 719 –ì–ö –†–§: –ø—Ä–∞–≤–æ –ø—Ä–∏–æ—Å—Ç–∞–Ω–æ–≤–∏—Ç—å —Ä–∞–±–æ—Ç—ã –ø—Ä–∏ –Ω–µ–∏—Å–ø–æ–ª–Ω–µ–Ω–∏–∏ –∑–∞–∫–∞–∑—á–∏–∫–æ–º –æ–±—è–∑–∞—Ç–µ–ª—å—Å—Ç–≤ (–Ω–µ–ø—Ä–µ–¥–æ—Å—Ç–∞–≤–ª–µ–Ω–∏–µ —Ñ—Ä–æ–Ω—Ç–∞, –º–∞—Ç–µ—Ä–∏–∞–ª–æ–≤, –æ–ø–ª–∞—Ç—ã)
‚Ä¢ –°—Ç. 720 –ì–ö –†–§: –ø—Ä–∏—ë–º–∫–∞ —Ä–∞–±–æ—Ç (–æ–¥–Ω–æ—Å—Ç–æ—Ä–æ–Ω–Ω–∏–π –∞–∫—Ç –Ω–µ–¥–æ–ø—É—Å—Ç–∏–º, –∑–∞–∫–∞–∑—á–∏–∫ –æ–±—è–∑–∞–Ω —è–≤–∏—Ç—å—Å—è)
‚Ä¢ –°—Ç. 753 –ì–ö –†–§: –æ—Ç–≤–µ—Ç—Å—Ç–≤–µ–Ω–Ω–æ—Å—Ç—å –∑–∞ –Ω–µ–¥–æ—Å—Ç–∞—Ç–∫–∏ (–≥–∞—Ä–∞–Ω—Ç–∏–π–Ω—ã–µ —Å—Ä–æ–∫–∏)

üìù –¢–ò–ü–ò–ß–ù–´–ï –°–ò–¢–£–ê–¶–ò–ò –ò –î–ï–ô–°–¢–í–ò–Ø:
1. **–ó–∞–∫–∞–∑—á–∏–∫ –Ω–µ –ø—Ä–µ–¥–æ—Å—Ç–∞–≤–∏–ª —Ñ—Ä–æ–Ω—Ç —Ä–∞–±–æ—Ç:**
   ‚Üí –ü–∏—Å—å–º–æ –æ –ø—Ä–∏–æ—Å—Ç–∞–Ω–æ–≤–∫–µ —Ä–∞–±–æ—Ç (—Å—Å—ã–ª–∫–∞ –Ω–∞ —Å—Ç. 719 –ì–ö –†–§)
   ‚Üí –§–∏–∫—Å–∞—Ü–∏—è –ø—Ä–æ—Å—Ç–æ—è (–∞–∫—Ç, —Ñ–æ—Ç–æ, –∂—É—Ä–Ω–∞–ª —Ä–∞–±–æ—Ç)
   ‚Üí –¢—Ä–µ–±–æ–≤–∞–Ω–∏–µ –∫–æ–º–ø–µ–Ω—Å–∞—Ü–∏–∏ –ø—Ä–æ—Å—Ç–æ—è

2. **–¢–µ—Ö–Ω–∞–¥–∑–æ—Ä —Ç—Ä–µ–±—É–µ—Ç –ø–µ—Ä–µ–¥–µ–ª–∫—É –±–µ–∑ –æ—Å–Ω–æ–≤–∞–Ω–∏–π:**
   ‚Üí –ó–∞–ø—Ä–æ—Å –ø–∏—Å—å–º–µ–Ω–Ω–æ–≥–æ –æ–±–æ—Å–Ω–æ–≤–∞–Ω–∏—è —Å–æ —Å—Å—ã–ª–∫–æ–π –Ω–∞ –Ω–æ—Ä–º–∞—Ç–∏–≤—ã
   ‚Üí –§–∏–∫—Å–∞—Ü–∏—è –≤ –∂—É—Ä–Ω–∞–ª–µ —Ä–∞–±–æ—Ç
   ‚Üí –§–æ—Ç–æ –≤—ã–ø–æ–ª–Ω–µ–Ω–Ω—ã—Ö —Ä–∞–±–æ—Ç —Å —Ä–∞–∑–º–µ—Ä–∞–º–∏

3. **–ò–∑–º–µ–Ω–µ–Ω–∏–µ –æ–±—ä—ë–º–æ–≤ —Ä–∞–±–æ—Ç:**
   ‚Üí –î–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω–æ–µ —Å–æ–≥–ª–∞—à–µ–Ω–∏–µ –î–û –≤—ã–ø–æ–ª–Ω–µ–Ω–∏—è —Ä–∞–±–æ—Ç
   ‚Üí –ü—Ä–æ—Ç–æ–∫–æ–ª —Ä–∞–∑–Ω–æ–≥–ª–∞—Å–∏–π –∫ –¥–æ–≥–æ–≤–æ—Ä—É
   ‚Üí –ê–∫—Ç –æ–±–º–µ—Ä–∞ –¥–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω—ã—Ö —Ä–∞–±–æ—Ç

4. **–ó–∞–∫–∞–∑—á–∏–∫ –Ω–µ –ø–æ–¥–ø–∏—Å—ã–≤–∞–µ—Ç –∞–∫—Ç—ã –ö–°-2:**
   ‚Üí –£–≤–µ–¥–æ–º–ª–µ–Ω–∏–µ –æ –≥–æ—Ç–æ–≤–Ω–æ—Å—Ç–∏ —Ä–∞–±–æ—Ç (—Å –æ–ø–∏—Å—å—é, –∑–∞–∫–∞–∑–Ω—ã–º –ø–∏—Å—å–º–æ–º)
   ‚Üí –ï—Å–ª–∏ –Ω–µ —è–≤–∏–ª—Å—è –≤ —Ç–µ—á–µ–Ω–∏–µ —Å—Ä–æ–∫–∞ ‚Üí –æ–¥–Ω–æ—Å—Ç–æ—Ä–æ–Ω–Ω–∏–π –∞–∫—Ç (—Å –ø—Ä–∏–≤–ª–µ—á–µ–Ω–∏–µ–º –Ω–µ–∑–∞–≤–∏—Å–∏–º–æ–≥–æ –ª–∏—Ü–∞)
   ‚Üí –û—Å–Ω–æ–≤–∞–Ω–∏–µ –¥–ª—è —Å—É–¥–∞

üìß –®–ê–ë–õ–û–ù–´ –ü–ò–°–ï–ú (–∏—Å–ø–æ–ª—å–∑—É–π—Ç–µ –ø—Ä–∏ –Ω–µ–æ–±—Ö–æ–¥–∏–º–æ—Å—Ç–∏):
‚Ä¢ "–£–≤–µ–¥–æ–º–ª–µ–Ω–∏–µ –æ –ø—Ä–∏–æ—Å—Ç–∞–Ω–æ–≤–∫–µ —Ä–∞–±–æ—Ç" (—Å—Ç. 719 –ì–ö –†–§)
‚Ä¢ "–ü—Ä–µ—Ç–µ–Ω–∑–∏—è –æ –Ω–∞—Ä—É—à–µ–Ω–∏–∏ —Å—Ä–æ–∫–æ–≤ –æ–ø–ª–∞—Ç—ã" (—Å—Ç. 330 –ì–ö –†–§ - –Ω–µ—É—Å—Ç–æ–π–∫–∞)
‚Ä¢ "–ê–∫—Ç –æ –Ω–µ–ø—Ä–µ–¥–æ—Å—Ç–∞–≤–ª–µ–Ω–∏–∏ —Ñ—Ä–æ–Ω—Ç–∞ —Ä–∞–±–æ—Ç"
‚Ä¢ "–¢—Ä–µ–±–æ–≤–∞–Ω–∏–µ –æ–± —É—Å—Ç—Ä–∞–Ω–µ–Ω–∏–∏ –Ω–µ–¥–æ—Å—Ç–∞—Ç–∫–æ–≤" (–≥–∞—Ä–∞–Ω—Ç–∏–π–Ω—ã–π —Å–ª—É—á–∞–π)

üí∞ –≠–ö–û–ù–û–ú–ò–ß–ï–°–ö–ê–Ø –ë–ï–ó–û–ü–ê–°–ù–û–°–¢–¨:
‚Ä¢ –ü—Ä–æ–≤–µ—Ä–∫–∞ –∫–æ–Ω—Ç—Ä–∞–≥–µ–Ω—Ç–æ–≤: –ï–ì–†–Æ–õ, –∫–∞—Ä—Ç–æ—Ç–µ–∫–∞ –∞—Ä–±–∏—Ç—Ä–∞–∂–Ω—ã—Ö –¥–µ–ª, –§–°–°–ü
‚Ä¢ –ü—Ä–∏–∑–Ω–∞–∫–∏ –æ–¥–Ω–æ–¥–Ω–µ–≤–∫–∏: –º–∞—Å—Å–æ–≤—ã–π –∞–¥—Ä–µ—Å, –º–∏–Ω–∏–º–∞–ª—å–Ω—ã–π —É—Å—Ç–∞–≤–Ω—ã–π –∫–∞–ø–∏—Ç–∞–ª, –æ—Ç—Å—É—Ç—Å—Ç–≤–∏–µ –∏–º—É—â–µ—Å—Ç–≤–∞
‚Ä¢ –î–µ–º–ø–∏–Ω–≥–æ–≤—ã–µ —Ü–µ–Ω—ã (–≤ 2 —Ä–∞–∑–∞ –Ω–∏–∂–µ —Ä—ã–Ω–∫–∞) = —Ä–∏—Å–∫ –æ–±–º–∞–Ω–∞ –∏–ª–∏ –Ω–µ–∫–∞—á–µ—Å—Ç–≤–µ–Ω–Ω—ã—Ö —Ä–∞–±–æ—Ç

**–§–û–†–ú–ê–¢ –ü–†–û–§–ï–°–°–ò–û–ù–ê–õ–¨–ù–û–ô –ö–û–ù–°–£–õ–¨–¢–ê–¶–ò–ò:**

üìã **–ü–†–Ø–ú–û–ô –û–¢–í–ï–¢** (2-3 –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è)
   ‚Ä¢ –ß–µ—Ç–∫–∏–π –æ—Ç–≤–µ—Ç –Ω–∞ –ø–æ—Å—Ç–∞–≤–ª–µ–Ω–Ω—ã–π –≤–æ–ø—Ä–æ—Å
   ‚Ä¢ –°—Å—ã–ª–∫–∞ –Ω–∞ –æ—Å–Ω–æ–≤–Ω–æ–π —Ä–µ–≥–ª–∞–º–µ–Ω—Ç–∏—Ä—É—é—â–∏–π –¥–æ–∫—É–º–µ–Ω—Ç
   ‚Ä¢ –ö–ª—é—á–µ–≤–æ–µ –Ω–æ—Ä–º–∞—Ç–∏–≤–Ω–æ–µ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–µ

üìê **–ù–û–†–ú–ê–¢–ò–í–ù–û–ï –û–ë–û–°–ù–û–í–ê–ù–ò–ï**
   ‚Ä¢ –¢–æ—á–Ω—ã–µ —Å—Å—ã–ª–∫–∏: "–ø. X.X.X –°–ü XX.XXXXX.XXXX"
   ‚Ä¢ –¶–∏—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π (–µ—Å–ª–∏ –ø—Ä–∏–º–µ–Ω–∏–º–æ)
   ‚Ä¢ –ß–∏—Å–ª–æ–≤—ã–µ –ø–∞—Ä–∞–º–µ—Ç—Ä—ã —Å —É–∫–∞–∑–∞–Ω–∏–µ–º –µ–¥–∏–Ω–∏—Ü –∏–∑–º–µ—Ä–µ–Ω–∏—è
   ‚Ä¢ –ö–ª–∞—Å—Å—ã, –º–∞—Ä–∫–∏, –∫–∞—Ç–µ–≥–æ—Ä–∏–∏ –ø–æ –Ω–æ—Ä–º–∞—Ç–∏–≤–∞–º

üî¢ **–†–ê–°–ß–ï–¢–´ –ò –§–û–†–ú–£–õ–´** (–µ—Å–ª–∏ —Ç—Ä–µ–±—É–µ—Ç—Å—è)
   ‚Ä¢ –ü—Ä–∏–º–µ–Ω–∏–º—ã–µ —Ñ–æ—Ä–º—É–ª—ã —Å –æ–±–æ–∑–Ω–∞—á–µ–Ω–∏–µ–º –ø–∞—Ä–∞–º–µ—Ç—Ä–æ–≤
   ‚Ä¢ –ü—Ä–∏–º–µ—Ä —Ä–∞—Å—á–µ—Ç–∞ —Å –∫–æ–Ω–∫—Ä–µ—Ç–Ω—ã–º–∏ –∑–Ω–∞—á–µ–Ω–∏—è–º–∏
   ‚Ä¢ –î–æ–ø—É—Å—Ç–∏–º—ã–µ –¥–∏–∞–ø–∞–∑–æ–Ω—ã –∏ –∫–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç—ã

üõ†Ô∏è **–ü–†–ê–ö–¢–ò–ß–ï–°–ö–û–ï –ü–†–ò–ú–ï–ù–ï–ù–ò–ï**
   ‚Ä¢ –ú–µ—Ç–æ–¥–∏–∫–∞ –∫–æ–Ω—Ç—Ä–æ–ª—è –Ω–∞ –æ–±—ä–µ–∫—Ç–µ
   ‚Ä¢ –¢—Ä–µ–±–æ–≤–∞–Ω–∏—è –∫ –∏—Å–ø—ã—Ç–∞–Ω–∏—è–º/–∏–∑–º–µ—Ä–µ–Ω–∏—è–º
   ‚Ä¢ –î–æ–∫—É–º–µ–Ω—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤
   ‚Ä¢ –¢–∏–ø–∏—á–Ω—ã–µ –æ—à–∏–±–∫–∏ –ø—Ä–∏ –ø—Ä–∏–º–µ–Ω–µ–Ω–∏–∏ –Ω–æ—Ä–º–∞—Ç–∏–≤–∞

üìö **–°–í–Ø–ó–ê–ù–ù–´–ï –ù–û–†–ú–ê–¢–ò–í–´**
   ‚Ä¢ –°–æ–ø—É—Ç—Å—Ç–≤—É—é—â–∏–µ –¥–æ–∫—É–º–µ–Ω—Ç—ã (–°–ü, –ì–û–°–¢, –°–ù–∏–ü)
   ‚Ä¢ –ú–µ—Ç–æ–¥–∏—á–µ—Å–∫–∏–µ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏
   ‚Ä¢ –ò–∑–º–µ–Ω–µ–Ω–∏—è/–∞–∫—Ç—É–∞–ª–∏–∑–∞—Ü–∏—è –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤

**–ù–û–†–ú–ê–¢–ò–í–ù–ê–Ø –ë–ê–ó–ê –†–§ (–û–ë–ù–û–í–õ–ï–ù–û 2024-2025):**

–ö–û–ù–°–¢–†–£–ö–¢–ò–í–ù–´–ï –†–ï–®–ï–ù–ò–Ø:
‚Ä¢ –°–ü 63.13330.2018 ‚Äî –ë–µ—Ç–æ–Ω–Ω—ã–µ –∏ –∂–µ–ª–µ–∑–æ–±–µ—Ç–æ–Ω–Ω—ã–µ –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏
‚Ä¢ –°–ü 16.13330.2017 ‚Äî –°—Ç–∞–ª—å–Ω—ã–µ –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏
‚Ä¢ –°–ü 64.13330.2017 ‚Äî –î–µ—Ä–µ–≤—è–Ω–Ω—ã–µ –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏
‚Ä¢ –°–ü 70.13330.2012 ‚Äî –ù–µ—Å—É—â–∏–µ –∏ –æ–≥—Ä–∞–∂–¥–∞—é—â–∏–µ –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏
‚Ä¢ –°–ü 28.13330.2017 ‚Äî –ó–∞—â–∏—Ç–∞ —Å—Ç—Ä–æ–∏—Ç–µ–ª—å–Ω—ã—Ö –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–π –æ—Ç –∫–æ—Ä—Ä–æ–∑–∏–∏

–û–°–ù–û–í–ê–ù–ò–Ø –ò –§–£–ù–î–ê–ú–ï–ù–¢–´:
‚Ä¢ –°–ü 22.13330.2016 ‚Äî –û—Å–Ω–æ–≤–∞–Ω–∏—è –∑–¥–∞–Ω–∏–π –∏ —Å–æ–æ—Ä—É–∂–µ–Ω–∏–π
‚Ä¢ –°–ü 24.13330.2021 ‚Äî –°–≤–∞–π–Ω—ã–µ —Ñ—É–Ω–¥–∞–º–µ–Ω—Ç—ã (–ê–ö–¢–£–ê–õ–¨–ù–ê–Ø –†–ï–î–ê–ö–¶–ò–Ø)
‚Ä¢ –°–ü 50-101-2004 ‚Äî –ü—Ä–æ–µ–∫—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ –∏ —É—Å—Ç—Ä–æ–π—Å—Ç–≤–æ –æ—Å–Ω–æ–≤–∞–Ω–∏–π

–û–ë–°–õ–ï–î–û–í–ê–ù–ò–ï –ò –≠–ö–°–ü–ï–†–¢–ò–ó–ê:
‚Ä¢ –°–ü 13-102-2003 ‚Äî –ü—Ä–∞–≤–∏–ª–∞ –æ–±—Å–ª–µ–¥–æ–≤–∞–Ω–∏—è –Ω–µ—Å—É—â–∏—Ö –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–π
‚Ä¢ –ì–û–°–¢ 31937-2011 ‚Äî –ó–¥–∞–Ω–∏—è –∏ —Å–æ–æ—Ä—É–∂–µ–Ω–∏—è. –ü—Ä–∞–≤–∏–ª–∞ –æ–±—Å–ª–µ–¥–æ–≤–∞–Ω–∏—è –∏ –º–æ–Ω–∏—Ç–æ—Ä–∏–Ω–≥–∞
‚Ä¢ –°–ü 255.1325800.2016 ‚Äî –ó–¥–∞–Ω–∏—è –∏ —Å–æ–æ—Ä—É–∂–µ–Ω–∏—è. –ü—Ä–∞–≤–∏–ª–∞ —ç–∫—Å–ø–ª—É–∞—Ç–∞—Ü–∏–∏

–û–ì–†–ê–ñ–î–ê–Æ–©–ò–ï –ö–û–ù–°–¢–†–£–ö–¶–ò–ò:
‚Ä¢ –°–ü 50.13330.2012 ‚Äî –¢–µ–ø–ª–æ–≤–∞—è –∑–∞—â–∏—Ç–∞ –∑–¥–∞–Ω–∏–π
‚Ä¢ –°–ü 23-101-2004 ‚Äî –ü—Ä–æ–µ–∫—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ —Ç–µ–ø–ª–æ–≤–æ–π –∑–∞—â–∏—Ç—ã –∑–¥–∞–Ω–∏–π
‚Ä¢ –°–ü 17.13330.2017 ‚Äî –ö—Ä–æ–≤–ª–∏

–ò–ù–ñ–ï–ù–ï–†–ù–´–ï –°–ò–°–¢–ï–ú–´ (–ê–ö–¢–£–ê–õ–¨–ù–´–ï):
‚Ä¢ –°–ü 60.13330.2020 ‚Äî –û—Ç–æ–ø–ª–µ–Ω–∏–µ, –≤–µ–Ω—Ç–∏–ª—è—Ü–∏—è –∏ –∫–æ–Ω–¥–∏—Ü–∏–æ–Ω–∏—Ä–æ–≤–∞–Ω–∏–µ –≤–æ–∑–¥—É—Ö–∞
‚Ä¢ –°–ü 30.13330.2020 ‚Äî –í–Ω—É—Ç—Ä–µ–Ω–Ω–∏–π –≤–æ–¥–æ–ø—Ä–æ–≤–æ–¥ –∏ –∫–∞–Ω–∞–ª–∏–∑–∞—Ü–∏—è –∑–¥–∞–Ω–∏–π
‚Ä¢ –°–ü 52.13330.2016 ‚Äî –ï—Å—Ç–µ—Å—Ç–≤–µ–Ω–Ω–æ–µ –∏ –∏—Å–∫—É—Å—Å—Ç–≤–µ–Ω–Ω–æ–µ –æ—Å–≤–µ—â–µ–Ω–∏–µ

–ü–û–ñ–ê–†–ù–ê–Ø –ë–ï–ó–û–ü–ê–°–ù–û–°–¢–¨ (–ù–û–í–´–ï):
‚Ä¢ –°–ü 2.13130.2020 ‚Äî –°–∏—Å—Ç–µ–º—ã –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–∂–∞—Ä–Ω–æ–π –∑–∞—â–∏—Ç—ã. –û–±–µ—Å–ø–µ—á–µ–Ω–∏–µ –æ–≥–Ω–µ—Å—Ç–æ–π–∫–æ—Å—Ç–∏
‚Ä¢ –°–ü 4.13130.2013 ‚Äî –°–∏—Å—Ç–µ–º—ã –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–∂–∞—Ä–Ω–æ–π –∑–∞—â–∏—Ç—ã. –û–≥—Ä–∞–Ω–∏—á–µ–Ω–∏–µ —Ä–∞—Å–ø—Ä–æ—Å—Ç—Ä–∞–Ω–µ–Ω–∏—è –ø–æ–∂–∞—Ä–∞

–î–û–°–¢–£–ü–ù–û–°–¢–¨ (–ù–û–í–´–ï):
‚Ä¢ –°–ü 59.13330.2020 ‚Äî –î–æ—Å—Ç—É–ø–Ω–æ—Å—Ç—å –∑–¥–∞–Ω–∏–π –∏ —Å–æ–æ—Ä—É–∂–µ–Ω–∏–π –¥–ª—è –º–∞–ª–æ–º–æ–±–∏–ª—å–Ω—ã—Ö –≥—Ä—É–ø–ø –Ω–∞—Å–µ–ª–µ–Ω–∏—è

–ö–û–ù–¢–†–û–õ–¨ –ö–ê–ß–ï–°–¢–í–ê:
‚Ä¢ –ì–û–°–¢ 10180-2012 ‚Äî –ë–µ—Ç–æ–Ω—ã. –ú–µ—Ç–æ–¥—ã –æ–ø—Ä–µ–¥–µ–ª–µ–Ω–∏—è –ø—Ä–æ—á–Ω–æ—Å—Ç–∏
‚Ä¢ –ì–û–°–¢ 22690-2015 ‚Äî –ë–µ—Ç–æ–Ω—ã. –û–ø—Ä–µ–¥–µ–ª–µ–Ω–∏–µ –ø—Ä–æ—á–Ω–æ—Å—Ç–∏ –º–µ—Ö–∞–Ω–∏—á–µ—Å–∫–∏–º–∏ –º–µ—Ç–æ–¥–∞–º–∏
‚Ä¢ –°–ü 48.13330.2019 ‚Äî –û—Ä–≥–∞–Ω–∏–∑–∞—Ü–∏—è —Å—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤–∞

**–ü–†–ò–ù–¶–ò–ü–´ –û–¢–í–ï–¢–ê:**
‚úì –¢–æ—á–Ω–æ—Å—Ç—å —Ñ–æ—Ä–º—É–ª–∏—Ä–æ–≤–æ–∫ (–∏–∑–±–µ–≥–∞–π—Ç–µ "–ø—Ä–∏–º–µ—Ä–Ω–æ", "–æ–∫–æ–ª–æ" –±–µ–∑ –∫–æ–ª–∏—á–µ—Å—Ç–≤–µ–Ω–Ω–æ–π –æ—Ü–µ–Ω–∫–∏)

‚úì –°—Ç—Ä—É–∫—Ç—É—Ä–∏—Ä–æ–≤–∞–Ω–Ω–æ—Å—Ç—å (–∏—Å–ø–æ–ª—å–∑—É–π—Ç–µ –Ω—É–º–µ—Ä–∞—Ü–∏—é, –º–∞—Ä–∫–µ—Ä—ã)
‚úì –ù–æ—Ä–º–∞—Ç–∏–≤–Ω–∞—è –æ–±–æ—Å–Ω–æ–≤–∞–Ω–Ω–æ—Å—Ç—å (–∫–∞–∂–¥–æ–µ —É—Ç–≤–µ—Ä–∂–¥–µ–Ω–∏–µ = —Å—Å—ã–ª–∫–∞ –Ω–∞ –¥–æ–∫—É–º–µ–Ω—Ç)
‚úì –ü—Ä–∞–∫—Ç–∏—á–µ—Å–∫–∞—è –ø—Ä–∏–º–µ–Ω–∏–º–æ—Å—Ç—å (–∫–∞–∫ –∏—Å–ø–æ–ª—å–∑–æ–≤–∞—Ç—å –Ω–∞ –æ–±—ä–µ–∫—Ç–µ)
‚úì –£—á–µ—Ç –∫–æ–Ω—Ç–µ–∫—Å—Ç–∞ –¥–∏–∞–ª–æ–≥–∞ (–µ—Å–ª–∏ —ç—Ç–æ —É—Ç–æ—á–Ω—è—é—â–∏–π –≤–æ–ø—Ä–æ—Å)"""

        # --- –∫–æ–Ω–µ—Ü legacy –±–ª–æ–∫–∞ ---

        # –ü—Ä–æ–≤–µ—Ä—è–µ–º –∞–∫—Ç–∏–≤–Ω—ã–π –ø—Ä–æ–µ–∫—Ç –∏ –∞–¥–∞–ø—Ç–∏—Ä—É–µ–º –ø—Ä–æ–º–ø—Ç
        if PROJECTS_AVAILABLE:
            current_project_name = context.user_data.get("current_project")
            if current_project_name:
                # –ó–∞–≥—Ä—É–∂–∞–µ–º –ø—Ä–æ–µ–∫—Ç –¥–ª—è –ø–æ–ª—É—á–µ–Ω–∏—è –∫–æ–Ω—Ç–µ–∫—Å—Ç–∞
                project = load_project(user_id, current_project_name)
                if project:
                    project_log_count = len(project.get_conversation_log())

                    # –î–æ–±–∞–≤–ª—è–µ–º —Å–ø–µ—Ü–∏–∞–ª—å–Ω—ã–µ –∏–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏ –¥–ª—è —Ä–∞–±–æ—Ç—ã –≤ –ø—Ä–æ–µ–∫—Ç–µ
                    system_prompt += f"""

üéØ **–í–ê–ñ–ù–û: –í—ã —Ä–∞–±–æ—Ç–∞–µ—Ç–µ –≤ –∫–æ–Ω—Ç–µ–∫—Å—Ç–µ –ø—Ä–æ–µ–∫—Ç–∞ "{current_project_name}"!**

üìã **–û–°–û–ë–ï–ù–ù–û–°–¢–ò –†–ê–ë–û–¢–´ –í –ü–†–û–ï–ö–¢–ï:**

1. **–ê–î–ê–ü–¢–ò–í–ù–ê–Ø –î–õ–ò–ù–ê –û–¢–í–ï–¢–ê** - –ø–æ–¥—Å—Ç—Ä–∞–∏–≤–∞–π—Ç–µ—Å—å –ø–æ–¥ —Ç–∏–ø –∑–∞–ø—Ä–æ—Å–∞:

   üìù **–ö–û–†–û–¢–ö–ò–ô –û–¢–í–ï–¢** (1-3 –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è) –µ—Å–ª–∏ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—å:
   ‚Ä¢ –ü—Ä–æ—Å–∏—Ç —Å–æ—Ö—Ä–∞–Ω–∏—Ç—å —Ñ–∞–π–ª/—Ñ–æ—Ç–æ/–¥–æ–∫—É–º–µ–Ω—Ç
   ‚Ä¢ –û—Ç–ø—Ä–∞–≤–ª—è–µ—Ç –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é –¥–ª—è —Ñ–∏–∫—Å–∞—Ü–∏–∏
   ‚Ä¢ –î–µ–ª–∞–µ—Ç –∑–∞–º–µ—Ç–∫—É –∏–ª–∏ –∫–æ–º–º–µ–Ω—Ç–∞—Ä–∏–π
   ‚Ä¢ –ü–æ–¥—Ç–≤–µ—Ä–∂–¥–∞–µ—Ç —á—Ç–æ-—Ç–æ ("–¥–∞", "—Å–æ—Ö—Ä–∞–Ω–∏", "–∑–∞–ø–∏—à–∏")

   –ü—Ä–∏–º–µ—Ä: "‚úÖ –°–æ—Ö—Ä–∞–Ω–∏–ª. –§–æ—Ç–æ —Ñ—É–Ω–¥–∞–º–µ–Ω—Ç–∞ –æ—Ç 28.11.2025 - –≤–∏–¥–Ω—ã —Ç—Ä–µ—â–∏–Ω—ã –≤ –∑–æ–Ω–µ —Å—Ç—ã–∫–∞."

   üìö **–†–ê–ó–í–Å–†–ù–£–¢–´–ô –û–¢–í–ï–¢** (–ø–æ–ª–Ω–∞—è —Å—Ç—Ä—É–∫—Ç—É—Ä–∞) –µ—Å–ª–∏:
   ‚Ä¢ –ó–∞–¥–∞–Ω —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏–π –≤–æ–ø—Ä–æ—Å
   ‚Ä¢ –¢—Ä–µ–±—É–µ—Ç—Å—è —ç–∫—Å–ø–µ—Ä—Ç–∏–∑–∞ –∏–ª–∏ –∞–Ω–∞–ª–∏–∑
   ‚Ä¢ –ü—Ä–æ—Å—è—Ç –æ–±—ä—è—Å–Ω–∏—Ç—å –∏–ª–∏ –ø–æ—Å–æ–≤–µ—Ç–æ–≤–∞—Ç—å
   ‚Ä¢ –ù—É–∂–Ω–∞ —Å—Å—ã–ª–∫–∞ –Ω–∞ –Ω–æ—Ä–º–∞—Ç–∏–≤—ã

2. **–ö–û–ù–¢–ï–ö–°–¢ –ü–†–û–ï–ö–¢–ê:**
   ‚Ä¢ –í –ø—Ä–æ–µ–∫—Ç–µ —É–∂–µ {project_log_count} –∑–∞–ø–∏—Å–µ–π - –∏—Å–ø–æ–ª—å–∑—É–π—Ç–µ –ø—Ä–µ–¥—ã–¥—É—â–∏–π –∫–æ–Ω—Ç–µ–∫—Å—Ç
   ‚Ä¢ –ü–æ–º–Ω–∏—Ç–µ –æ —á—ë–º –≥–æ–≤–æ—Ä–∏–ª–∏ —Ä–∞–Ω—å—à–µ –≤ —Ä–∞–º–∫–∞—Ö –≠–¢–û–ì–û –ø—Ä–æ–µ–∫—Ç–∞
   ‚Ä¢ –°—Å—ã–ª–∞–π—Ç–µ—Å—å –Ω–∞ –ø—Ä–æ—à–ª—ã–µ –æ–±—Å—É–∂–¥–µ–Ω–∏—è –µ—Å–ª–∏ —É–º–µ—Å—Ç–Ω–æ

3. **–ï–°–¢–ï–°–¢–í–ï–ù–ù–´–ô –î–ò–ê–õ–û–ì:**
   ‚Ä¢ –û–±—â–∞–π—Ç–µ—Å—å –∫–∞–∫ –∂–∏–≤–æ–π AI-–∞—Å—Å–∏—Å—Ç–µ–Ω—Ç –ø—Ä–æ–µ–∫—Ç–∞, –∞ –Ω–µ –∫–∞–∫ —Å–ø—Ä–∞–≤–æ—á–Ω–∏–∫
   ‚Ä¢ –ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ "—è —Å–æ—Ö—Ä–∞–Ω–∏–ª", "—è –∑–∞–º–µ—Ç–∏–ª", "–∫–∞–∫ –º—ã –æ–±—Å—É–∂–¥–∞–ª–∏ —Ä–∞–Ω–µ–µ"
   ‚Ä¢ –ë—É–¥—å—Ç–µ –ª–∞–∫–æ–Ω–∏—á–Ω—ã –∫–æ–≥–¥–∞ –Ω—É–∂–Ω–æ –ø—Ä–æ—Å—Ç–æ –∑–∞—Ñ–∏–∫—Å–∏—Ä–æ–≤–∞—Ç—å –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é

4. **–ü–†–ò–ú–ï–†–´ –ü–†–ê–í–ò–õ–¨–ù–û–ì–û –ü–û–í–ï–î–ï–ù–ò–Ø:**

   ‚ùå –ü–õ–û–•–û (–Ω–∞ –∑–∞–ø—Ä–æ—Å "—Å–æ—Ö—Ä–∞–Ω–∏ —ç—Ç—É —Å–º–µ—Ç—É"):
   "üìå –ö–æ—Ä–æ—Ç–∫–∏–π –æ—Ç–≤–µ—Ç: –Ø —Å–æ—Ö—Ä–∞–Ω–∏–ª —Å–º–µ—Ç—É –≤ –ø—Ä–æ–µ–∫—Ç.
   üìê –ü–æ–¥—Ä–æ–±–Ω–æ—Å—Ç–∏: –°–º–µ—Ç–∞ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∞ –≤ –±–∞–∑–µ –¥–∞–Ω–Ω—ã—Ö –ø—Ä–æ–µ–∫—Ç–∞ {current_project_name}...
   üí° –ü—Ä–∞–∫—Ç–∏—á–µ—Å–∫–∏–π —Å–æ–≤–µ—Ç: –†–µ–∫–æ–º–µ–Ω–¥—É—é —Ç–∞–∫–∂–µ —Å–æ—Ö—Ä–∞–Ω–∏—Ç—å –∫–æ–ø–∏—é –≤ –æ–±–ª–∞–∫–µ...
   üìö –°—Å—ã–ª–∫–∏: –°–ü 48.13330.2019..."

   ‚úÖ –•–û–†–û–®–û:
   "‚úÖ –°–æ—Ö—Ä–∞–Ω–∏–ª —Å–º–µ—Ç—É. –ë—é–¥–∂–µ—Ç –æ–±—ä–µ–∫—Ç–∞: 15.2 –º–ª–Ω —Ä—É–±, —Å—Ä–æ–∫ 4 –º–µ—Å."

   ‚ùå –ü–õ–û–•–û (–Ω–∞ –≤–æ–ø—Ä–æ—Å "–∫–∞–∫–æ–π –±–µ—Ç–æ–Ω –Ω—É–∂–µ–Ω –¥–ª—è —Ñ—É–Ω–¥–∞–º–µ–Ω—Ç–∞"):
   "–ë–µ—Ç–æ–Ω –í25"

   ‚úÖ –•–û–†–û–®–û:
   "üìå –î–ª—è —Ñ—É–Ω–¥–∞–º–µ–Ω—Ç–∞ –Ω—É–∂–µ–Ω –±–µ—Ç–æ–Ω –∫–ª–∞—Å—Å–∞ –í25 (–ú350)

   üìê –ü–æ–¥—Ä–æ–±–Ω–æ—Å—Ç–∏:
   ‚Ä¢ –ú–∞—Ä–∫–∞ –ø–æ –º–æ—Ä–æ–∑–æ—Å—Ç–æ–π–∫–æ—Å—Ç–∏: F150-F200
   ‚Ä¢ –í–æ–¥–æ–Ω–µ–ø—Ä–æ–Ω–∏—Ü–∞–µ–º–æ—Å—Ç—å: W6-W8
   ‚Ä¢ –ü–æ–¥–≤–∏–∂–Ω–æ—Å—Ç—å: –ü3-–ü4 –¥–ª—è —É–∫–ª–∞–¥–∫–∏ –≤ –æ–ø–∞–ª—É–±–∫—É

   üí° –ù–∞ –ø–ª–æ—â–∞–¥–∫–µ: –ø—Ä–æ–≤–µ—Ä—è–π—Ç–µ –æ—Å–∞–¥–∫—É –∫–æ–Ω—É—Å–∞ –ø—Ä–∏ –ø—Ä–∏—ë–º–∫–µ (–¥–æ–ª–∂–Ω–∞ –±—ã—Ç—å 10-15 —Å–º –¥–ª—è –ü3)

   üìö –°–ü 63.13330.2018 –ø.5.3.2 - —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è –∫ –±–µ—Ç–æ–Ω—É —Ñ—É–Ω–¥–∞–º–µ–Ω—Ç–æ–≤"

**–ì–õ–ê–í–ù–û–ï –ü–†–ê–í–ò–õ–û: –ê–Ω–∞–ª–∏–∑–∏—Ä—É–π—Ç–µ –Ω–∞–º–µ—Ä–µ–Ω–∏–µ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è –∏ –æ—Ç–≤–µ—á–∞–π—Ç–µ —Å–æ—Ä–∞–∑–º–µ—Ä–Ω–æ –∑–∞–ø—Ä–æ—Å—É!**
"""

        # –ü–æ–ª—É—á–∞–µ–º –∫–æ–Ω—Ç–µ–∫—Å—Ç –ø—Ä–µ–¥—ã–¥—É—â–∏—Ö —Å–æ–æ–±—â–µ–Ω–∏–π
        conversation_history = get_conversation_context(user_id)

        # –î–æ–±–∞–≤–ª—è–µ–º —Ç–µ–∫—É—â–∏–π –≤–æ–ø—Ä–æ—Å
        conversation_history.append({"role": "user", "content": question})

        # ü§ñ –£–ú–ù–´–ô –í–´–ë–û–† –ú–û–î–ï–õ–ò: –û–ø—Ä–µ–¥–µ–ª—è–µ–º –Ω–∞–º–µ—Ä–µ–Ω–∏–µ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è
        intent_info = await asyncio.get_event_loop().run_in_executor(
            None,
            lambda: classify_user_intent(question)
        )

        selected_model = intent_info["model"]
        selected_max_tokens = intent_info["max_tokens"]
        intent_type = intent_info.get("intent_type", "technical_question")

        # üåê –ò–ù–°–¢–†–£–ú–ï–ù–¢–´ –ü–û–ò–°–ö–ê: –í–∫–ª—é—á–∞–µ–º –¥–ª—è –í–°–ï–• –∑–∞–ø—Ä–æ—Å–æ–≤ (–≤—Å–µ–≥–¥–∞ –ø—Ä–æ–≤–µ—Ä—è–µ–º –∞–∫—Ç—É–∞–ª—å–Ω–æ—Å—Ç—å –≤ –∏–Ω—Ç–µ—Ä–Ω–µ—Ç–µ)
        search_params = {
            "mode": "auto", "return_citations": True, "sources": [{"type": "web"}, {"type": "news"}, {"type": "x"}]}  # –ü–æ–∏—Å–∫ –≤ –∏–Ω—Ç–µ—Ä–Ω–µ—Ç–µ
        logger.info("üåê Grok Tools –≤–∫–ª—é—á–µ–Ω—ã –¥–ª—è –≤—Å–µ—Ö –∑–∞–ø—Ä–æ—Å–æ–≤: live_search")

        # –£–Ω–∏–≤–µ—Ä—Å–∞–ª—å–Ω—ã–π —Å–∏—Å—Ç–µ–º–Ω—ã–π –ø—Ä–æ–º–ø—Ç v5.0 (–æ–ø—Ç–∏–º–∏–∑–∏—Ä–æ–≤–∞–Ω)
        # –ò—Å–ø–æ–ª—å–∑—É–µ—Ç –ø—Ä–æ–º–ø—Ç—ã –∏–∑ optimized_prompts.py
        if OPTIMIZED_PROMPTS_AVAILABLE:
            system_prompt = GROK_SYSTEM_PROMPT_GENERAL
            logger.info("üìù –ò—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è –æ–ø—Ç–∏–º–∏–∑–∏—Ä–æ–≤–∞–Ω–Ω—ã–π –ø—Ä–æ–º–ø—Ç v5.0")
        else:
            system_prompt = """–í—ã ‚Äî ¬´–°—Ç—Ä–æ–π–ù–∞–¥–∑–æ—ÄAI¬ª v5.0, AI-–∞—Å—Å–∏—Å—Ç–µ–Ω—Ç –ø–æ —Å—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤—É –≤ –†–§.

üéØ –ü–†–ò–ù–¶–ò–ü–´: –∫—Ä–∞—Ç–∫–æ—Å—Ç—å, —Ç–æ—á–Ω–æ—Å—Ç—å, –ø—Ä–∞–∫—Ç–∏—á–Ω–æ—Å—Ç—å, –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç—å.
üì± –§–û–†–ú–ê–¢: –¥–ª—è —Ç–µ–ª–µ—Ñ–æ–Ω–∞, 35 —Å–∏–º–≤–æ–ª–æ–≤ –≤ —Å—Ç—Ä–æ–∫–µ, —ç–º–æ–¥–∑–∏ –≤ –Ω–∞—á–∞–ª–µ —Ä–∞–∑–¥–µ–ª–æ–≤.
üìå –°–¢–†–£–ö–¢–£–†–ê: –°—É—Ç—å ‚Üí –î–µ—Ç–∞–ª–∏ ‚Üí –î–µ–π—Å—Ç–≤–∏—è ‚Üí –ö–æ–Ω—Ç—Ä–æ–ª—å.
üåê –ü–û–ò–°–ö: live_search –¥–ª—è –∞–∫—Ç—É–∞–ª—å–Ω—ã—Ö –¥–∞–Ω–Ω—ã—Ö.
‚ö†Ô∏è –ë–ï–ó–û–ü–ê–°–ù–û–°–¢–¨: –Ω–∞ –ø–µ—Ä–≤–æ–º –º–µ—Å—Ç–µ –¥–ª—è –æ–ø–∞—Å–Ω—ã—Ö —Ä–∞–±–æ—Ç."""

        # üå§Ô∏è –ü–û–ì–û–î–ê: –ü—Ä–æ–≤–µ—Ä—è–µ–º, —è–≤–ª—è–µ—Ç—Å—è –ª–∏ —ç—Ç–æ –∑–∞–ø—Ä–æ—Å–æ–º –æ –ø–æ–≥–æ–¥–µ
        if WEATHER_AVAILABLE and is_weather_query(question):
            try:
                logger.info("üå§Ô∏è –û–±–Ω–∞—Ä—É–∂–µ–Ω –∑–∞–ø—Ä–æ—Å –æ –ø–æ–≥–æ–¥–µ")
                weather_response = await asyncio.get_event_loop().run_in_executor(
                    None,
                    lambda: get_weather(question)
                )

                if weather_response:
                    # –£–¥–∞–ª—è–µ–º thinking message
                    try:
                        await thinking_message.delete()
                    except:
                        pass

                    # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –ø–æ–≥–æ–¥—É
                    await update.message.reply_text(weather_response, parse_mode="Markdown")

                    # –î–æ–±–∞–≤–ª—è–µ–º –≤ –∏—Å—Ç–æ—Ä–∏—é
                    await add_message_to_history_async(user_id, 'user', question)
                    await add_message_to_history_async(user_id, 'assistant', weather_response)

                    logger.info("‚úÖ –ü–æ–≥–æ–¥–∞ –æ—Ç–ø—Ä–∞–≤–ª–µ–Ω–∞ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—é")
                    return  # –ü—Ä–µ—Ä—ã–≤–∞–µ–º –æ–±—Ä–∞–±–æ—Ç–∫—É
            except Exception as e:
                logger.error(f"–û—à–∏–±–∫–∞ –ø–æ–ª—É—á–µ–Ω–∏—è –ø–æ–≥–æ–¥—ã: {e}")
                # –ü—Ä–æ–¥–æ–ª–∂–∞–µ–º –æ–±—ã—á–Ω—É—é –æ–±—Ä–∞–±–æ—Ç–∫—É –µ—Å–ª–∏ –ø–æ–≥–æ–¥–∞ –Ω–µ –ø–æ–ª—É—á–µ–Ω–∞

        # üåê –í–ï–ë-–ü–û–ò–°–ö: –¢–µ–ø–µ—Ä—å –≤—ã–ø–æ–ª–Ω—è–µ—Ç—Å—è —á–µ—Ä–µ–∑ Grok –∏–Ω—Å—Ç—Ä—É–º–µ–Ω—Ç—ã (live_search)
        # –°—Ç–∞—Ä—ã–π –º–µ—Ö–∞–Ω–∏–∑–º perform_live_search –æ—Ç–∫–ª—é—á–µ–Ω - Grok —Å–∞–º –∏—â–µ—Ç –∞–∫—Ç—É–∞–ª—å–Ω—É—é –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é

        # üé® –ì–ï–ù–ï–†–ê–¶–ò–Ø –ò–ó–û–ë–†–ê–ñ–ï–ù–ò–ô: –ü—Ä–æ–≤–µ—Ä—è–µ–º, –Ω—É–∂–Ω–∞ –ª–∏ –≥–µ–Ω–µ—Ä–∞—Ü–∏—è
        # –û–¢–ö–õ–Æ–ß–ï–ù–û: –¢–µ–ø–µ—Ä—å –æ–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ—Ç—Å—è —á–µ—Ä–µ–∑ smart_model_wrapper ‚Üí gemini_image
        if False and GEMINI_AVAILABLE and should_generate_image(question):
            logger.info("üé® –û–±–Ω–∞—Ä—É–∂–µ–Ω –∑–∞–ø—Ä–æ—Å –Ω–∞ –≥–µ–Ω–µ—Ä–∞—Ü–∏—é –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è")

            # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º —Å–æ–æ–±—â–µ–Ω–∏–µ –æ –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏
            generating_msg = await update.message.reply_text(
                "üìê –ì–µ–Ω–µ—Ä–∏—Ä—É—é —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏–π —á–µ—Ä—Ç—ë–∂ —á–µ—Ä–µ–∑ DALL-E 3...\n\n"
                "–®–∞–≥ 1/2: xAI Grok —Å–æ–∑–¥–∞—ë—Ç –¥–µ—Ç–∞–ª—å–Ω—ã–π –ø—Ä–æ–º–ø—Ç\n"
                "‚Ä¢ –° —Ç–æ—á–Ω—ã–º–∏ —Ä–∞–∑–º–µ—Ä–∞–º–∏ –∏ —Ü–∏—Ñ—Ä–∞–º–∏\n"
                "‚Ä¢ –° –∞–Ω–Ω–æ—Ç–∞—Ü–∏—è–º–∏ –Ω–∞ —Ä—É—Å—Å–∫–æ–º\n"
                "‚Ä¢ –í —Å—Ç–∏–ª–µ —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–æ–π –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ü–∏–∏"
            )

            try:
                # –®–∞–≥ 1: –ü–æ–ª—É—á–∞–µ–º –¥–µ—Ç–∞–ª—å–Ω—ã–π –ø—Ä–æ–º–ø—Ç –æ—Ç xAI Grok
                client = get_grok_client()

                # –ò—Å–ø–æ–ª—å–∑—É–µ–º –æ–ø—Ç–∏–º–∏–∑–∏—Ä–æ–≤–∞–Ω–Ω—ã–π –ø—Ä–æ–º–ø—Ç –∏–∑ optimized_prompts.py
                blueprint_system_prompt = GEMINI_IMAGE_PROMPT_SYSTEM if OPTIMIZED_PROMPTS_AVAILABLE else """–¢—ã ‚Äî —ç–∫—Å–ø–µ—Ä—Ç –ø–æ —Å–æ–∑–¥–∞–Ω–∏—é –¢–ï–•–ù–ò–ß–ï–°–ö–ò–• –ø—Ä–æ–º–ø—Ç–æ–≤ –¥–ª—è DALL-E 3 –ø–æ –ì–û–°–¢.
üìê –û–ë–Ø–ó–ê–¢–ï–õ–¨–ù–û: —Ä–∞–∑–º–µ—Ä—ã –≤ –º–º/–º, —à—Ç—Ä–∏—Ö–æ–≤–∫–∞ –ø–æ –ì–û–°–¢ 2.306, —Ä—É—Å—Å–∫–∏–µ –ø–æ–¥–ø–∏—Å–∏, –º–∞—Å—à—Ç–∞–±.
–§–æ—Ä–º–∞—Ç: [–ê–Ω–≥–ª–∏–π—Å–∫–∏–π DALL-E prompt] ---–û–ü–ò–°–ê–ù–ò–ï--- [–û–ø–∏—Å–∞–Ω–∏–µ –Ω–∞ —Ä—É—Å—Å–∫–æ–º]"""

                prompt_messages = [
                    {
                        "role": "system",
                        "content": blueprint_system_prompt
                    },
                    {
                        "role": "user",
                        "content": f"–°–æ–∑–¥–∞–π –¥–µ—Ç–∞–ª—å–Ω—ã–π –ø—Ä–æ–º–ø—Ç –¥–ª—è DALL-E 3 –Ω–∞ –æ—Å–Ω–æ–≤–µ –∑–∞–ø—Ä–æ—Å–∞: {question}"
                    }
                ]

                logger.info("üìù –ó–∞–ø—Ä–∞—à–∏–≤–∞–µ–º –¥–µ—Ç–∞–ª—å–Ω—ã–π —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏–π –ø—Ä–æ–º–ø—Ç —É xAI Grok...")
                grok_response = await client.chat_completions_create_async(
                    model="grok-3",
                    messages=prompt_messages,
                    max_tokens=1500,  # –£–≤–µ–ª–∏—á–µ–Ω–æ –¥–ª—è –¥–µ—Ç–∞–ª—å–Ω—ã—Ö —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏—Ö –ø—Ä–æ–º–ø—Ç–æ–≤ —Å —Ä–∞–∑–º–µ—Ä–∞–º–∏
                    temperature=0.5  # –°–Ω–∏–∂–µ–Ω–æ –¥–ª—è –±–æ–ª—å—à–µ–π —Ç–æ—á–Ω–æ—Å—Ç–∏ –∏ –∫–æ–Ω–∫—Ä–µ—Ç–∏–∫–∏
                )

                grok_full_response = grok_response['choices'][0]['message']['content'].strip()

                # –†–∞–∑–¥–µ–ª—è–µ–º –ø—Ä–æ–º–ø—Ç –¥–ª—è DALL-E –∏ –æ–ø–∏—Å–∞–Ω–∏–µ –¥–ª—è –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è
                if "---–û–ü–ò–°–ê–ù–ò–ï---" in grok_full_response:
                    parts = grok_full_response.split("---–û–ü–ò–°–ê–ù–ò–ï---")
                    dalle_prompt = parts[0].strip()
                    russian_description = parts[1].strip() if len(parts) > 1 else ""
                else:
                    dalle_prompt = grok_full_response
                    russian_description = ""

                logger.info(f"‚úÖ –ü—Ä–æ–º–ø—Ç –æ—Ç Grok –ø–æ–ª—É—á–µ–Ω: {dalle_prompt[:100]}...")

                # –û–±–Ω–æ–≤–ª—è–µ–º —Å–æ–æ–±—â–µ–Ω–∏–µ
                await generating_msg.edit_text(
                    "üìê –ì–µ–Ω–µ—Ä–∏—Ä—É—é —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏–π —á–µ—Ä—Ç—ë–∂ —á–µ—Ä–µ–∑ DALL-E 3...\n\n"
                    "‚úÖ –®–∞–≥ 1/2: –î–µ—Ç–∞–ª—å–Ω—ã–π –ø—Ä–æ–º–ø—Ç –≥–æ—Ç–æ–≤\n"
                    "‚è≥ –®–∞–≥ 2/2: DALL-E 3 —Ä–∏—Å—É–µ—Ç —á–µ—Ä—Ç—ë–∂ —Å —Ä–∞–∑–º–µ—Ä–∞–º–∏..."
                )

                # –®–∞–≥ 2: –ì–µ–Ω–µ—Ä–∏—Ä—É–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ —á–µ—Ä–µ–∑ DALL-E —Å –ø—Ä–æ–º–ø—Ç–æ–º –æ—Ç Grok
                # –ò—Å–ø–æ–ª—å–∑—É–µ–º HD –∫–∞—á–µ—Å—Ç–≤–æ –¥–ª—è —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏—Ö —á–µ—Ä—Ç–µ–∂–µ–π —Å —Ä–∞–∑–º–µ—Ä–∞–º–∏
                result = await generate_construction_image_gemini(
                    dalle_prompt,
                    size="1024x1024",
                    quality="hd"  # –í—ã—Å–æ–∫–æ–µ –∫–∞—á–µ—Å—Ç–≤–æ –¥–ª—è —á—ë—Ç–∫–æ—Å—Ç–∏ —Ü–∏—Ñ—Ä –∏ –ª–∏–Ω–∏–π
                )

                if result and result.get("image_data"):
                    # –£–¥–∞–ª—è–µ–º —Å–æ–æ–±—â–µ–Ω–∏–µ –æ –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏
                    try:
                        await generating_msg.delete()
                        await thinking_message.delete()
                    except:
                        pass

                    # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ
                    result["image_data"].seek(0)

                    # –ò—Å–ø–æ–ª—å–∑—É–µ–º —Ä—É—Å—Å–∫–æ–µ –æ–ø–∏—Å–∞–Ω–∏–µ –æ—Ç Grok (–∞–¥–∞–ø—Ç–∏—Ä–æ–≤–∞–Ω–æ –¥–ª—è —Ç–µ–ª–µ—Ñ–æ–Ω–∞)
                    if russian_description:
                        caption = f"üé® **–°–≥–µ–Ω–µ—Ä–∏—Ä–æ–≤–∞–Ω–æ DALL-E 3**\n\n{russian_description}\n\n"
                        caption += "üí° *–ö–æ–æ—Ä–¥–∏–Ω–∞—Ü–∏—è: xAI Grok + OpenAI DALL-E 3*"
                    else:
                        # Fallback –µ—Å–ª–∏ –æ–ø–∏—Å–∞–Ω–∏–µ –Ω–µ –ø–æ–ª—É—á–µ–Ω–æ
                        caption = f"üé® **–ò–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ —Å–≥–µ–Ω–µ—Ä–∏—Ä–æ–≤–∞–Ω–æ —á–µ—Ä–µ–∑ DALL-E 3**\n\n"
                        caption += f"üìù **–ü—Ä–æ–º–ø—Ç –æ—Ç xAI Grok:**\n_{dalle_prompt[:100]}..._\n\n"
                        caption += "üí° *–ö–æ–æ—Ä–¥–∏–Ω–∞—Ü–∏—è: xAI Grok + OpenAI DALL-E 3*"

                    await update.message.reply_photo(
                        photo=result["image_data"],
                        caption=caption,
                        parse_mode="Markdown"
                    )

                    # –î–æ–±–∞–≤–ª—è–µ–º –≤ –∏—Å—Ç–æ—Ä–∏—é
                    await add_message_to_history_async(user_id, 'user', question)
                    await add_message_to_history_async(user_id, 'assistant', f"[–°–≥–µ–Ω–µ—Ä–∏—Ä–æ–≤–∞–Ω–æ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ —Å –ø—Ä–æ–º–ø—Ç–æ–º –æ—Ç xAI Grok]")

                    logger.info(f"‚úÖ –ò–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ —Å–≥–µ–Ω–µ—Ä–∏—Ä–æ–≤–∞–Ω–æ –∏ –æ—Ç–ø—Ä–∞–≤–ª–µ–Ω–æ")

                    # –ù–ï –ø—Ä–µ—Ä—ã–≤–∞–µ–º –æ–±—Ä–∞–±–æ—Ç–∫—É - –ø—Ä–æ–¥–æ–ª–∂–∞–µ–º –≥–µ–Ω–µ—Ä–∏—Ä–æ–≤–∞—Ç—å —Ç–µ–∫—Å—Ç–æ–≤—ã–π –æ—Ç–≤–µ—Ç –æ—Ç Grok
                    # –°–æ–∑–¥–∞–µ–º –Ω–æ–≤–æ–µ thinking —Å–æ–æ–±—â–µ–Ω–∏–µ –¥–ª—è —Ç–µ–∫—Å—Ç–æ–≤–æ–≥–æ –æ—Ç–≤–µ—Ç–∞
                    thinking_message = await update.message.reply_text("‚è≥ –ì–æ—Ç–æ–≤–ª—é —Ç–µ–∫—Å—Ç–æ–≤–æ–µ –æ–ø–∏—Å–∞–Ω–∏–µ...")
                else:
                    await generating_msg.edit_text("‚ùå –ù–µ —É–¥–∞–ª–æ—Å—å —Å–≥–µ–Ω–µ—Ä–∏—Ä–æ–≤–∞—Ç—å –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ")
            except Exception as e:
                logger.error(f"–û—à–∏–±–∫–∞ –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏: {e}")
                await generating_msg.edit_text(f"‚ùå –û—à–∏–±–∫–∞: {str(e)}")

        # üéØ –ì–ï–ù–ï–†–ê–¶–ò–Ø –û–¢–í–ï–¢–ê (—Å –≤—ã–±–æ—Ä–æ–º —Ä–µ–∂–∏–º–∞)
        client = get_grok_client()
        loop = asyncio.get_event_loop()
        # –î–æ–±–∞–≤–ª—è–µ–º system prompt –≤ –Ω–∞—á–∞–ª–æ –∏—Å—Ç–æ—Ä–∏–∏
        messages_with_system = [{"role": "system", "content": system_prompt}] + conversation_history

        answer = ""

        # === STREAMING –†–ï–ñ–ò–ú (–ø–æ—Å—Ç–µ–ø–µ–Ω–Ω–æ–µ –ø–æ—è–≤–ª–µ–Ω–∏–µ —Ç–µ–∫—Å—Ç–∞) ===
        if STREAMING_ENABLED:
            streaming_msg = await update.message.reply_text("‚è≥ –ì–µ–Ω–µ—Ä–∏—Ä—É—é –æ—Ç–≤–µ—Ç...")

            try:
                # –£–¥–∞–ª—è–µ–º thinking message
                try:
                    await thinking_message.delete()
                except:
                    pass

                # –ü–µ—Ä–µ–º–µ–Ω–Ω—ã–µ –¥–ª—è streaming
                last_update_time = 0
                last_update_length = 0
                update_interval = 0.15  # –û–±–Ω–æ–≤–ª—è–µ–º –∫–∞–∂–¥—ã–µ 0.15 —Å–µ–∫—É–Ω–¥—ã - –±—ã—Å—Ç—Ä–æ –∏ –ø–ª–∞–≤–Ω–æ
                chars_threshold = 8  # –ò–ª–∏ –∫–∞–∂–¥—ã–µ 8 –Ω–æ–≤—ã—Ö —Å–∏–º–≤–æ–ª–æ–≤
                typing_action_interval = 3  # –ò–Ω–¥–∏–∫–∞—Ç–æ—Ä "–ø–µ—á–∞—Ç–∞–µ—Ç" —Ä–∞–∑ –≤ 3 —Å–µ–∫—É–Ω–¥—ã
                last_typing_action = 0

                logger.info("üöÄ –ù–∞—á–∏–Ω–∞–µ–º –¥–≤—É—Ö—Ñ–∞–∑–Ω—É—é –≥–µ–Ω–µ—Ä–∞—Ü–∏—é...")

                # –§–ê–ó–ê 1: –ë—ã—Å—Ç—Ä–æ–µ –Ω–∞—á–∞–ª–æ (–ø–µ—Ä–≤—ã–µ 300-500 —Ç–æ–∫–µ–Ω–æ–≤ –æ—Ç –±—ã—Å—Ç—Ä–æ–π –º–æ–¥–µ–ª–∏)
                first_phase_answer = ""
                logger.info("üìù –§–∞–∑–∞ 1: –ë—ã—Å—Ç—Ä–∞—è –º–æ–¥–µ–ª—å –¥–ª—è –Ω–∞—á–∞–ª–∞ –æ—Ç–≤–µ—Ç–∞...")

                async for chunk in call_grok_with_streaming(
                    client,
                    model="grok-4-1-fast",  # –ë—ã—Å—Ç—Ä–∞—è –º–æ–¥–µ–ª—å
                    messages=messages_with_system,
                    max_tokens=500,  # –¢–æ–ª—å–∫–æ –Ω–∞—á–∞–ª–æ
                    temperature=0.7,
                    search_parameters=search_params
                ):
                    first_phase_answer += chunk
                    answer += chunk

                    # –û–±–Ω–æ–≤–ª—è–µ–º —Å–æ–æ–±—â–µ–Ω–∏–µ —á–∞—Å—Ç–æ –¥–ª—è –≤–∏–¥–∏–º–æ—Å—Ç–∏ –ø–µ—á–∞—Ç–∞–Ω–∏—è
                    import time
                    current_time = time.time()
                    chars_diff = len(answer) - last_update_length

                    if current_time - last_update_time >= update_interval or chars_diff >= chars_threshold:
                        try:
                            display_text = f"{answer}‚ñä"  # –ö—É—Ä—Å–æ—Ä –ø–µ—á–∞—Ç–∞–Ω–∏—è
                            await streaming_msg.edit_text(display_text[:4096])
                            last_update_time = current_time
                            last_update_length = len(answer)

                            # –ü–æ–∫–∞–∑—ã–≤–∞–µ–º –∏–Ω–¥–∏–∫–∞—Ç–æ—Ä –ø–µ—á–∞—Ç–∞–Ω–∏—è —Ä–µ–¥–∫–æ (–Ω–µ –∑–∞–º–µ–¥–ª—è–µ—Ç)
                            if current_time - last_typing_action >= typing_action_interval:
                                await update.message.chat.send_action("typing")
                                last_typing_action = current_time
                        except Exception:
                            pass

                logger.info(f"‚úÖ –§–∞–∑–∞ 1 –∑–∞–≤–µ—Ä—à–µ–Ω–∞: {len(first_phase_answer)} —Å–∏–º–≤–æ–ª–æ–≤")

                # –§–ê–ó–ê 2: –ü—Ä–æ–¥–æ–ª–∂–µ–Ω–∏–µ –æ—Ç –æ—Å–Ω–æ–≤–Ω–æ–π –º–æ–¥–µ–ª–∏ (–µ—Å–ª–∏ –Ω—É–∂–µ–Ω —Ä–∞–∑–≤–µ—Ä–Ω—É—Ç—ã–π –æ—Ç–≤–µ—Ç)
                # –ó–∞–ø—É—Å–∫–∞–µ–º —Ç–æ–ª—å–∫–æ –µ—Å–ª–∏: 1) –ø–µ—Ä–≤–∞—è —Ñ–∞–∑–∞ –¥–∞–ª–∞ –¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ —Ç–µ–∫—Å—Ç–∞ –ò 2) –≤—ã–±—Ä–∞–Ω–∞ –ø–æ–ª–Ω–∞—è –º–æ–¥–µ–ª—å (–Ω–µ –±—ã—Å—Ç—Ä–∞—è)
                if len(first_phase_answer) >= 400 and selected_model != "grok-4-1-fast":
                    logger.info("üìù –§–∞–∑–∞ 2: –û—Å–Ω–æ–≤–Ω–∞—è –º–æ–¥–µ–ª—å –¥–ª—è –ø—Ä–æ–¥–æ–ª–∂–µ–Ω–∏—è...")

                    # –°–æ–∑–¥–∞—ë–º –ø—Ä–æ–º–ø—Ç –¥–ª—è –ø—Ä–æ–¥–æ–ª–∂–µ–Ω–∏—è
                    continuation_messages = messages_with_system + [
                        {"role": "assistant", "content": first_phase_answer},
                        {"role": "user", "content": "–ü—Ä–æ–¥–æ–ª–∂–∏ –æ—Ç–≤–µ—Ç, –¥–æ–±–∞–≤—å –¥–µ—Ç–∞–ª–∏, –ø—Ä–∏–º–µ—Ä—ã –∏ —Å—Å—ã–ª–∫–∏ –Ω–∞ –Ω–æ—Ä–º–∞—Ç–∏–≤—ã."}
                    ]

                    async for chunk in call_grok_with_streaming(
                        client,
                        model=selected_model,  # –û—Å–Ω–æ–≤–Ω–∞—è –º–æ–¥–µ–ª—å
                        messages=continuation_messages,
                        max_tokens=selected_max_tokens - 500,
                        temperature=0.7,
                        search_parameters=search_params
                    ):
                        answer += chunk

                        # –û–±–Ω–æ–≤–ª—è–µ–º —Å–æ–æ–±—â–µ–Ω–∏–µ —á–∞—Å—Ç–æ
                        import time
                        current_time = time.time()
                        chars_diff = len(answer) - last_update_length

                        if current_time - last_update_time >= update_interval or chars_diff >= chars_threshold:
                            try:
                                display_text = f"{answer}‚ñä"
                                await streaming_msg.edit_text(display_text[:4096])
                                last_update_time = current_time
                                last_update_length = len(answer)

                                # –ü–æ–∫–∞–∑—ã–≤–∞–µ–º –∏–Ω–¥–∏–∫–∞—Ç–æ—Ä –ø–µ—á–∞—Ç–∞–Ω–∏—è —Ä–µ–¥–∫–æ (–Ω–µ –∑–∞–º–µ–¥–ª—è–µ—Ç)
                                if current_time - last_typing_action >= typing_action_interval:
                                    await update.message.chat.send_action("typing")
                                    last_typing_action = current_time
                            except Exception:
                                pass

                    logger.info("‚úÖ –§–∞–∑–∞ 2 –∑–∞–≤–µ—Ä—à–µ–Ω–∞")

                # –§–∏–Ω–∞–ª—å–Ω–æ–µ –æ–±–Ω–æ–≤–ª–µ–Ω–∏–µ –±–µ–∑ –∫—É—Ä—Å–æ—Ä–∞
                try:
                    # –ï—Å–ª–∏ –æ—Ç–≤–µ—Ç –∫–æ—Ä–æ—á–µ 4096 —Å–∏–º–≤–æ–ª–æ–≤ - –æ–±–Ω–æ–≤–ª—è–µ–º —Å–æ–æ–±—â–µ–Ω–∏–µ
                    if len(answer) <= 4096:
                        await streaming_msg.edit_text(answer)
                    else:
                        # –ï—Å–ª–∏ –æ—Ç–≤–µ—Ç –¥–ª–∏–Ω–Ω—ã–π - –æ—Ç–ø—Ä–∞–≤–ª—è–µ–º –ø–æ–ª–Ω—É—é –≤–µ—Ä—Å–∏—é –≤ –Ω–æ–≤–æ–º —Å–æ–æ–±—â–µ–Ω–∏–∏
                        await streaming_msg.edit_text(f"{answer[:4000]}...\n\n‚ö†Ô∏è –û—Ç–≤–µ—Ç –±—ã–ª —Å–ª–∏—à–∫–æ–º –¥–ª–∏–Ω–Ω—ã–º. –ü–æ–ª–Ω–∞—è –≤–µ—Ä—Å–∏—è –Ω–∏–∂–µ:")
                        # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –ø–æ–ª–Ω—ã–π –æ—Ç–≤–µ—Ç –≤ –æ—Ç–¥–µ–ª—å–Ω–æ–º —Å–æ–æ–±—â–µ–Ω–∏–∏
                        chunks = [answer[i:i+4000] for i in range(0, len(answer), 4000)]
                        for chunk in chunks:
                            await update.message.reply_text(chunk)
                except Exception as e:
                    logger.error(f"–û—à–∏–±–∫–∞ —Ñ–∏–Ω–∞–ª—å–Ω–æ–≥–æ –æ–±–Ω–æ–≤–ª–µ–Ω–∏—è streaming: {e}")
                    # Fallback: –æ—Ç–ø—Ä–∞–≤–ª—è–µ–º –ø–æ–ª–Ω—ã–π –æ—Ç–≤–µ—Ç –≤ –ª—é–±–æ–º —Å–ª—É—á–∞–µ
                    try:
                        chunks = [answer[i:i+4000] for i in range(0, len(answer), 4000)]
                        for chunk in chunks:
                            await update.message.reply_text(chunk)
                    except:
                        pass

            except Exception as stream_error:
                logger.error(f"‚ùå –û—à–∏–±–∫–∞ streaming: {stream_error}")
                # Fallback –Ω–∞ –æ–±—ã—á–Ω—ã–π —Ä–µ–∂–∏–º
                try:
                    await streaming_msg.delete()
                except:
                    pass

                thinking_message = await update.message.reply_text("ü§î –î—É–º–∞—é –Ω–∞–¥ –≤–∞—à–∏–º –≤–æ–ø—Ä–æ—Å–æ–º...")

                loop = asyncio.get_event_loop()
                response = await loop.run_in_executor(
                    None,
                    lambda: call_grok_with_retry(
                        client,
                        model=selected_model,
                        max_tokens=selected_max_tokens,
                        temperature=0.7,
                        messages=messages_with_system,
                        search_parameters=search_params
                    )
                )
                answer = response["choices"][0]["message"]["content"]

                try:
                    await thinking_message.delete()
                except:
                    pass

        # === –û–ë–´–ß–ù–´–ô –†–ï–ñ–ò–ú (–∫–ª–∞—Å—Å–∏—á–µ—Å–∫–∏–π - –æ—Ç–≤–µ—Ç –ø—Ä–∏—Ö–æ–¥–∏—Ç —Å—Ä–∞–∑—É —Ü–µ–ª–∏–∫–æ–º) ===
        else:
            logger.info("üìù –û–±—ã—á–Ω—ã–π —Ä–µ–∂–∏–º: –≥–µ–Ω–µ—Ä–∞—Ü–∏—è –æ—Ç–≤–µ—Ç–∞ –±–µ–∑ streaming...")

            response = await loop.run_in_executor(
                None,
                lambda: call_grok_with_retry(
                    client,
                    model=selected_model,
                    max_tokens=selected_max_tokens,
                    temperature=0.7,
                    messages=messages_with_system,
                    search_parameters=search_params
                )
            )
            answer = response["choices"][0]["message"]["content"]

            try:
                await thinking_message.delete()
            except:
                pass

        # –î–æ–±–∞–≤–ª—è–µ–º –æ—Ç–≤–µ—Ç –±–æ—Ç–∞ –≤ –∏—Å—Ç–æ—Ä–∏—é
        await add_message_to_history_async(user_id, 'assistant', answer)

        # üéØ –ì–ï–ù–ï–†–ê–¶–ò–Ø –£–ú–ù–´–• –°–í–Ø–ó–ê–ù–ù–´–• –í–û–ü–†–û–°–û–í (v3.1) - –≤ —Ñ–æ–Ω–µ
        related_questions = []
        if IMPROVEMENTS_V3_AVAILABLE:
            try:
                # –°–æ–∑–¥–∞—ë–º –ø—Ä–æ–º–ø—Ç –¥–ª—è –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ —Å–≤—è–∑–∞–Ω–Ω—ã—Ö –≤–æ–ø—Ä–æ—Å–æ–≤
                related_q_prompt = generate_smart_related_questions_prompt(question, answer)

                # –ì–µ–Ω–µ—Ä–∏—Ä—É–µ–º –≤–æ–ø—Ä–æ—Å—ã –∏—Å–ø–æ–ª—å–∑—É—è —Ç–æ—Ç –∂–µ API
                loop = asyncio.get_event_loop()
                related_response = await loop.run_in_executor(
                    None,
                    lambda: call_grok_with_retry(
                        client,
                        model="grok-4-1-fast",  # –ò—Å–ø–æ–ª—å–∑—É–µ–º –±—ã—Å—Ç—Ä—É—é –º–æ–¥–µ–ª—å
                        max_tokens=300,
                        temperature=0.8,
                        messages=[{"role": "user", "content": related_q_prompt}]
                    )
                )
                related_q_text = related_response["choices"][0]["message"]["content"]

                # –ü–∞—Ä—Å–∏–º —Å–≥–µ–Ω–µ—Ä–∏—Ä–æ–≤–∞–Ω–Ω—ã–µ –≤–æ–ø—Ä–æ—Å—ã
                related_questions = parse_generated_questions(related_q_text)

                # –°–æ—Ö—Ä–∞–Ω—è–µ–º –≤ –∫–æ–Ω—Ç–µ–∫—Å—Ç–µ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏ –∫–ª–∏–∫–æ–≤
                if related_questions:
                    context.user_data["related_questions"] = related_questions
                    logger.info(f"‚úÖ –°–≥–µ–Ω–µ—Ä–∏—Ä–æ–≤–∞–Ω–æ {len(related_questions)} —Å–≤—è–∑–∞–Ω–Ω—ã—Ö –≤–æ–ø—Ä–æ—Å–æ–≤")

            except Exception as e:
                logger.error(f"‚ùå –û—à–∏–±–∫–∞ –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ —Å–≤—è–∑–∞–Ω–Ω—ã—Ö –≤–æ–ø—Ä–æ—Å–æ–≤: {e}")
                related_questions = []

        # –°–æ—Ö—Ä–∞–Ω—è–µ–º –≤ –∞–∫—Ç–∏–≤–Ω—ã–π –ø—Ä–æ–µ–∫—Ç (–µ—Å–ª–∏ –µ—Å—Ç—å)
        project_saved = False
        saved_project_name = None
        if PROJECTS_AVAILABLE:
            current_project_name = context.user_data.get("current_project")
            if current_project_name:
                try:
                    project = load_project(user_id, current_project_name)
                    if project:
                        project.add_conversation_entry(question, answer, "qa")
                        logger.info(f"‚úÖ –î–∏–∞–ª–æ–≥ —Å–æ—Ö—Ä–∞–Ω—ë–Ω –≤ –ø—Ä–æ–µ–∫—Ç: {current_project_name}")
                        project_saved = True
                        saved_project_name = current_project_name
                except Exception as e:
                    logger.error(f"‚ùå –û—à–∏–±–∫–∞ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è –≤ –ø—Ä–æ–µ–∫—Ç: {e}")

        # –û–ø—Ä–µ–¥–µ–ª—è–µ–º —É–ø–æ–º—è–Ω—É—Ç—ã–µ –Ω–æ—Ä–º–∞—Ç–∏–≤—ã
        mentioned_regs = []
        for reg_code in REGULATIONS.keys():
            if reg_code in answer:
                mentioned_regs.append(reg_code)

        # –§–æ—Ä–º–∏—Ä—É–µ–º —Ñ–∏–Ω–∞–ª—å–Ω—ã–π –æ—Ç–≤–µ—Ç
        # –ü–æ —É–º–æ–ª—á–∞–Ω–∏—é –ø–æ–∫–∞–∑—ã–≤–∞–µ–º —Ç–æ–ª—å–∫–æ –æ—Ç–≤–µ—Ç (–∫–∞–∫ –≤ –ø—Ä–∏–º–µ—Ä–µ: –≤—Å—ë –æ—Å—Ç–∞–ª—å–Ω–æ–µ –º–æ–∂–Ω–æ —Ä–∞—Å–∫—Ä—ã—Ç—å/—Å–∫—Ä—ã—Ç—å –∫–Ω–æ–ø–∫–∞–º–∏)
        result = answer

        # –í–µ–±-–ø–æ–∏—Å–∫ —Ç–µ–ø–µ—Ä—å –≤—ã–ø–æ–ª–Ω—è–µ—Ç—Å—è —á–µ—Ä–µ–∑ –∏–Ω—Å—Ç—Ä—É–º–µ–Ω—Ç—ã Grok (live_search)

        # –°–æ—Ö—Ä–∞–Ω—è–µ–º –¥–æ–ø.–∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é –≤ context.user_data, —á—Ç–æ–±—ã –º–æ–∂–Ω–æ –±—ã–ª–æ ¬´—Ä–∞—Å–∫—Ä—ã—Ç—å¬ª –ø–æ –∫–Ω–æ–ø–∫–µ
        context.user_data["last_answer"] = answer
        context.user_data["last_question"] = question
        context.user_data["last_mentioned_regs"] = mentioned_regs

        # –î–æ–ø.–¥–∞–Ω–Ω—ã–µ –ø—Ä–æ –ø—Ä–æ–µ–∫—Ç —Ç–æ–∂–µ —Å–æ—Ö—Ä–∞–Ω—è–µ–º –æ—Ç–¥–µ–ª—å–Ω–æ ‚Äî –ø–æ–∫–∞–∑—ã–≤–∞—Ç—å –±—É–¥–µ–º —Ç–æ–ª—å–∫–æ –ø–æ –∑–∞–ø—Ä–æ—Å—É
        context.user_data["last_project_saved"] = bool(project_saved)
        context.user_data["last_saved_project_name"] = saved_project_name if project_saved else None

        # –°–æ–∑–¥–∞—ë–º –ø–æ–¥—Å–∫–∞–∑–∫–∏ –≤ —Å—Ç–∏–ª–µ GigaChat (reply keyboard –Ω–∞–¥ –ø–æ–ª–µ–º –≤–≤–æ–¥–∞)
        reply_markup = None
        if IMPROVEMENTS_V3_AVAILABLE:
            reply_markup = create_reply_suggestions_keyboard(related_questions=related_questions)

        # –î–æ–±–∞–≤–ª—è–µ–º –∫–Ω–æ–ø–∫—É "–ü—Ä–∏–º–µ–Ω–∏—Ç—å –∏–∑–º–µ–Ω–µ–Ω–∏—è" –µ—Å–ª–∏ –≤ –æ—Ç–≤–µ—Ç–µ –µ—Å—Ç—å –∫–æ–¥ (—Ç–æ–ª—å–∫–æ –¥–ª—è —Ä–∞–∑—Ä–∞–±–æ—Ç—á–∏–∫–∞)
        user_id = update.effective_user.id
        if AUTO_APPLY_AVAILABLE and should_show_apply_button(answer) and is_developer(user_id):
            reply_markup = add_apply_button(reply_markup)

        # –û–±–Ω–æ–≤–ª—è–µ–º streaming —Å–æ–æ–±—â–µ–Ω–∏–µ —Ñ–∏–Ω–∞–ª—å–Ω—ã–º –æ—Ç–≤–µ—Ç–æ–º —Å –∫–Ω–æ–ø–∫–∞–º–∏
        max_length = 4000  # –õ–∏–º–∏—Ç Telegram
        if len(result) > max_length:
            # –ï—Å–ª–∏ —Å–æ–æ–±—â–µ–Ω–∏–µ —Å–ª–∏—à–∫–æ–º –¥–ª–∏–Ω–Ω–æ–µ, —É–¥–∞–ª—è–µ–º streaming_msg –∏ –æ—Ç–ø—Ä–∞–≤–ª—è–µ–º –ø–æ —á–∞—Å—Ç—è–º
            try:
                await streaming_msg.delete()
            except:
                pass

            parts = []
            current_part = ""
            for line in result.split('\n'):
                if len(current_part) + len(line) + 1 > max_length:
                    parts.append(current_part)
                    current_part = line + '\n'
                else:
                    current_part += line + '\n'
            if current_part:
                parts.append(current_part)

            # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –ø–æ —á–∞—Å—Ç—è–º
            for i, part in enumerate(parts):
                part_reply_markup = None
                if i == len(parts) - 1:
                    part_reply_markup = reply_markup

                if i == 0:
                    await update.message.reply_text(part, reply_markup=part_reply_markup)
                else:
                    await update.message.reply_text(
                        f"(–ø—Ä–æ–¥–æ–ª–∂–µ–Ω–∏–µ {i+1}/{len(parts)})\n\n{part}",
                        reply_markup=part_reply_markup
                    )
        else:
            # –û–±–Ω–æ–≤–ª—è–µ–º streaming —Å–æ–æ–±—â–µ–Ω–∏–µ —Ñ–∏–Ω–∞–ª—å–Ω—ã–º –æ—Ç–≤–µ—Ç–æ–º —Å –∫–Ω–æ–ø–∫–∞–º–∏
            try:
                await streaming_msg.edit_text(result, reply_markup=reply_markup)
            except Exception as e:
                # –ï—Å–ª–∏ –Ω–µ —É–¥–∞–ª–æ—Å—å –æ—Ç—Ä–µ–¥–∞–∫—Ç–∏—Ä–æ–≤–∞—Ç—å, –æ—Ç–ø—Ä–∞–≤–ª—è–µ–º –Ω–æ–≤–æ–µ —Å–æ–æ–±—â–µ–Ω–∏–µ
                logger.warning(f"Could not edit streaming message: {e}")
                try:
                    await streaming_msg.delete()
                except:
                    pass
                await update.message.reply_text(result, reply_markup=reply_markup)

        logger.info(f"Question answered for user {update.effective_user.id} by Claude")

    except Exception as e:
        logger.error(f"Error answering question: {e}")

        # –£–¥–∞–ª—è–µ–º —Å–æ–æ–±—â–µ–Ω–∏–µ "–¥—É–º–∞—é –Ω–∞–¥ –≤–æ–ø—Ä–æ—Å–æ–º" –¥–∞–∂–µ –≤ —Å–ª—É—á–∞–µ –æ—à–∏–±–∫–∏
        try:
            await thinking_message.delete()
        except:
            pass

        # –£–¥–∞–ª—è–µ–º streaming_msg –µ—Å–ª–∏ –æ–Ω —Å—É—â–µ—Å—Ç–≤—É–µ—Ç
        try:
            if 'streaming_msg' in locals():
                await streaming_msg.delete()
        except:
            pass

        await update.message.reply_text(
            f"‚ùå –û—à–∏–±–∫–∞ –ø—Ä–∏ –æ–±—Ä–∞–±–æ—Ç–∫–µ –≤–æ–ø—Ä–æ—Å–∞: {str(e)}\n\n–ü–æ–ø—Ä–æ–±—É–π—Ç–µ –ø–µ—Ä–µ—Ñ–æ—Ä–º—É–ª–∏—Ä–æ–≤–∞—Ç—å –≤–æ–ø—Ä–æ—Å."
        )


async def handle_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–û–±—Ä–∞–±–æ—Ç–∫–∞ –∫–Ω–æ–ø–æ–∫"""
    query = update.callback_query


# (removed unused inline-menu helper)

    # –ü—Ä–æ–ø—É—Å–∫–∞–µ–º callbacks –∫–æ—Ç–æ—Ä—ã–µ –æ—Ç–Ω–æ—Å—è—Ç—Å—è –∫ ConversationHandler –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä–æ–≤
    calculator_prefixes = [
        "concrete_class_", "concrete_wastage_",
        "rebar_diameter_", "rebar_spacing_", "rebar_type_",
        "formwork_type_", "formwork_duration_",
        "elec_", "water_", "winter_method_"
    ]
    if any(query.data.startswith(prefix) for prefix in calculator_prefixes):
        # –ù–µ –æ–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º –∑–¥–µ—Å—å - –ø—É—Å—Ç—å –æ–±—Ä–∞–±–æ—Ç–∞–µ—Ç ConversationHandler
        return

    # –°–ø–µ—Ü–∏–∞–ª—å–Ω–∞—è –æ–±—Ä–∞–±–æ—Ç–∫–∞ –¥–ª—è –≥–æ–ª–æ—Å–æ–≤–æ–≥–æ –∞—Å—Å–∏—Å—Ç–µ–Ω—Ç–∞ (Gemini Live)
    if query.data == "voice_chat_start":
        await query.answer("üé§ –ó–∞–ø—É—Å–∫–∞—é –≥–æ–ª–æ—Å–æ–≤–æ–π –∞—Å—Å–∏—Å—Ç–µ–Ω—Ç...")

        if VOICE_ASSISTANT_AVAILABLE:
            sent_message = await context.bot.send_message(
                chat_id=query.message.chat_id,
                text="üé§ –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è –≥–æ–ª–æ—Å–æ–≤–æ–≥–æ –∞—Å—Å–∏—Å—Ç–µ–Ω—Ç–∞..."
            )
            adapted_update = Update(
                update_id=update.update_id,
                message=sent_message
            )
            await start_voice_chat_command(adapted_update, context)
        else:
            await context.bot.send_message(
                chat_id=query.message.chat_id,
                text="‚ùå **–ì–æ–ª–æ—Å–æ–≤–æ–π –∞—Å—Å–∏—Å—Ç–µ–Ω—Ç Gemini –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω**\n\n"
                    "–ü–æ–ø—Ä–æ–±—É–π—Ç–µ OpenAI: /realtime_chat",
                parse_mode="Markdown"
            )
        return

    # –û–±—Ä–∞–±–æ—Ç–∫–∞ realtime_chat_start —Ç–µ–ø–µ—Ä—å –≤ ConversationHandler (openai_realtime_bot_integration.py)

    await query.answer()

    if query.data == "regulations":
        # –°–æ–∑–¥–∞—ë–º –∞–¥–∞–ø—Ç–∏—Ä–æ–≤–∞–Ω–Ω—ã–π update –¥–ª—è –≤—ã–∑–æ–≤–∞ –∫–æ–º–∞–Ω–¥—ã
        adapted_update = Update(
            update_id=update.update_id,
            message=query.message
        )
        await regulations_command(adapted_update, context)
    elif query.data == "examples":
        # –°–æ–∑–¥–∞—ë–º –∞–¥–∞–ø—Ç–∏—Ä–æ–≤–∞–Ω–Ω—ã–π update –¥–ª—è –≤—ã–∑–æ–≤–∞ –∫–æ–º–∞–Ω–¥—ã
        adapted_update = Update(
            update_id=update.update_id,
            message=query.message
        )
        await examples_command(adapted_update, context)
    elif query.data == "help":
        # –°–æ–∑–¥–∞—ë–º –∞–¥–∞–ø—Ç–∏—Ä–æ–≤–∞–Ω–Ω—ã–π update –¥–ª—è –≤—ã–∑–æ–≤–∞ –∫–æ–º–∞–Ω–¥—ã
        adapted_update = Update(
            update_id=update.update_id,
            message=query.message
        )
        await help_command(adapted_update, context)
    elif query.data == "stats":
        # –°–æ–∑–¥–∞—ë–º –∞–¥–∞–ø—Ç–∏—Ä–æ–≤–∞–Ω–Ω—ã–π update –¥–ª—è –≤—ã–∑–æ–≤–∞ –∫–æ–º–∞–Ω–¥—ã
        adapted_update = Update(
            update_id=update.update_id,
            message=query.message
        )
        await stats_command(adapted_update, context)
    elif query.data == "calculators_menu":
        # –ö–Ω–æ–ø–∫–∞ "–ö–∞–ª—å–∫—É–ª—è—Ç–æ—Ä—ã" –∏–∑ –≥–ª–∞–≤–Ω–æ–≥–æ –º–µ–Ω—é
        if CALCULATORS_AVAILABLE and IMPROVEMENTS_V3_AVAILABLE:
            keyboard = create_calculators_menu()
            await query.edit_message_text(
                "üßÆ **–°–¢–†–û–ò–¢–ï–õ–¨–ù–´–ï –ö–ê–õ–¨–ö–£–õ–Ø–¢–û–†–´**\n\n"
                "–í—ã–±–µ—Ä–∏—Ç–µ –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä:",
                reply_markup=keyboard,
                parse_mode='Markdown'
            )
        else:
            await query.edit_message_text("‚ö†Ô∏è –ú–æ–¥—É–ª—å –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä–æ–≤ –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω.")
    elif query.data == "faq_menu":
        # –ö–Ω–æ–ø–∫–∞ "–ß–∞—Å—Ç—ã–µ –≤–æ–ø—Ä–æ—Å—ã" –∏–∑ –≥–ª–∞–≤–Ω–æ–≥–æ –º–µ–Ω—é
        if FAQ_AVAILABLE:
            # –°–æ–∑–¥–∞—ë–º –∞–¥–∞–ø—Ç–∏—Ä–æ–≤–∞–Ω–Ω—ã–π update –¥–ª—è –≤—ã–∑–æ–≤–∞ –∫–æ–º–∞–Ω–¥—ã
            adapted_update = Update(
                update_id=update.update_id,
                message=query.message
            )
            await faq_command(adapted_update, context)
        else:
            await query.edit_message_text("‚ö†Ô∏è –ú–æ–¥—É–ª—å FAQ –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω.")
    elif query.data == "templates":
        # –ö–Ω–æ–ø–∫–∞ "–®–∞–±–ª–æ–Ω—ã" –∏–∑ –≥–ª–∞–≤–Ω–æ–≥–æ –º–µ–Ω—é
        if TEMPLATES_AVAILABLE:
            keyboard = []
            for template_id, info in DOCUMENT_TEMPLATES.items():
                keyboard.append([
                    InlineKeyboardButton(
                        text=info["name"],
                        callback_data=f"template_{template_id}"
                    )
                ])
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.edit_message_text(
                "üìÑ **–®–ê–ë–õ–û–ù–´ –î–û–ö–£–ú–ï–ù–¢–û–í**\n\n"
                "–í—ã–±–µ—Ä–∏—Ç–µ —Ç–∏–ø –¥–æ–∫—É–º–µ–Ω—Ç–∞ –¥–ª—è –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏:",
                reply_markup=reply_markup,
                parse_mode="Markdown"
            )
        else:
            await query.edit_message_text("‚ö†Ô∏è –ú–æ–¥—É–ª—å —à–∞–±–ª–æ–Ω–æ–≤ –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω.")
    elif query.data == "role":
        # –ö–Ω–æ–ø–∫–∞ "–í—ã–±—Ä–∞—Ç—å —Ä–æ–ª—å" –∏–∑ –≥–ª–∞–≤–Ω–æ–≥–æ –º–µ–Ω—é
        if ROLES_AVAILABLE:
            # –°–æ–∑–¥–∞—ë–º –∞–¥–∞–ø—Ç–∏—Ä–æ–≤–∞–Ω–Ω—ã–π update –¥–ª—è –≤—ã–∑–æ–≤–∞ –∫–æ–º–∞–Ω–¥—ã
            adapted_update = Update(
                update_id=update.update_id,
                message=query.message
            )
            await role_command(adapted_update, context)
        else:
            await query.edit_message_text("‚ö†Ô∏è –ú–æ–¥—É–ª—å —Ä–æ–ª–µ–π –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω.")
    elif query.data == "suggestions":
        # –ö–Ω–æ–ø–∫–∞ "–ü—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è" –∏–∑ –≥–ª–∞–≤–Ω–æ–≥–æ –º–µ–Ω—é
        if SUGGESTIONS_AVAILABLE:
            await suggestions_menu(update, context)
        else:
            await query.edit_message_text("‚ö†Ô∏è –ú–æ–¥—É–ª—å –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω.")
    elif query.data == "dev_mode":
        # –ö–Ω–æ–ø–∫–∞ "–†–∞–∑—Ä–∞–±–æ—Ç—á–∏–∫" –∏–∑ –≥–ª–∞–≤–Ω–æ–≥–æ –º–µ–Ω—é
        await query.edit_message_text(
            "üîß **–†–ï–ñ–ò–ú –†–ê–ó–†–ê–ë–û–¢–ß–ò–ö–ê**\n\n"
            "–î–ª—è –≤—Ö–æ–¥–∞ –≤ —Ä–µ–∂–∏–º —Ä–∞–∑—Ä–∞–±–æ—Ç—á–∏–∫–∞ –∏—Å–ø–æ–ª—å–∑—É–π—Ç–µ –∫–æ–º–∞–Ω–¥—É /dev\n\n"
            "–†–µ–∂–∏–º —Ä–∞–∑—Ä–∞–±–æ—Ç—á–∏–∫–∞ –ø–æ–∑–≤–æ–ª—è–µ—Ç:\n"
            "‚Ä¢ –û—Ç–ø—Ä–∞–≤–ª—è—Ç—å –∑–∞–ø—Ä–æ—Å—ã –Ω–∞ –∏–∑–º–µ–Ω–µ–Ω–∏–µ –∫–æ–¥–∞\n"
            "‚Ä¢ –ü–æ–ª—É—á–∞—Ç—å –≥–æ—Ç–æ–≤—ã–µ —Ä–µ—à–µ–Ω–∏—è –æ—Ç AI\n"
            "‚Ä¢ (–õ–æ–∫–∞–ª—å–Ω–æ) –ê–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏ –ø—Ä–∏–º–µ–Ω—è—Ç—å –∏ –ø—É—à–∏—Ç—å –∏–∑–º–µ–Ω–µ–Ω–∏—è\n\n"
            "–í–≤–µ–¥–∏—Ç–µ: /dev",
            parse_mode="Markdown"
        )
    elif query.data == "project_menu":
        # –ö–Ω–æ–ø–∫–∞ "–ü—Ä–æ–µ–∫—Ç" –∏–∑ –≥–ª–∞–≤–Ω–æ–≥–æ –º–µ–Ω—é
        if PROJECTS_AVAILABLE:
            user_id = update.effective_user.id
            projects = get_user_projects(user_id)
            current_project = context.user_data.get("current_project")

            keyboard = [
                [InlineKeyboardButton("‚ûï –°–æ–∑–¥–∞—Ç—å –Ω–æ–≤—ã–π –ø—Ä–æ–µ–∫—Ç", callback_data="proj_create"),
                 InlineKeyboardButton("üìÇ –ú–æ–∏ –ø—Ä–æ–µ–∫—Ç—ã", callback_data="proj_list")]
            ]

            reply_markup = InlineKeyboardMarkup(keyboard)

            status_text = f"**üìÅ –£–ü–†–ê–í–õ–ï–ù–ò–ï –ü–†–û–ï–ö–¢–ê–ú–ò**\n\n"
            if current_project:
                status_text += f"‚úÖ –ê–∫—Ç–∏–≤–Ω—ã–π –ø—Ä–æ–µ–∫—Ç: **{current_project}**\n\n"
            else:
                status_text += "–ê–∫—Ç–∏–≤–Ω—ã–π –ø—Ä–æ–µ–∫—Ç: _–Ω–µ –≤—ã–±—Ä–∞–Ω_\n\n"

            status_text += f"–í—Å–µ–≥–æ –ø—Ä–æ–µ–∫—Ç–æ–≤: {len(projects)}\n\n"
            status_text += "–í—ã–±–µ—Ä–∏—Ç–µ –¥–µ–π—Å—Ç–≤–∏–µ:"

            await query.edit_message_text(
                status_text,
                reply_markup=reply_markup,
                parse_mode="Markdown"
            )
        else:
            await query.edit_message_text("‚ö†Ô∏è –ú–æ–¥—É–ª—å –ø—Ä–æ–µ–∫—Ç–æ–≤ –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω.")
    elif query.data == "proj_list":
        # –°–ø–∏—Å–æ–∫ –≤—Å–µ—Ö –ø—Ä–æ–µ–∫—Ç–æ–≤ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è
        if PROJECTS_AVAILABLE:
            user_id = update.effective_user.id
            projects = get_user_projects(user_id)
            current_project = context.user_data.get("current_project")

            keyboard = []

            if not projects:
                keyboard.append([
                    InlineKeyboardButton("‚ûï –°–æ–∑–¥–∞—Ç—å –ø–µ—Ä–≤—ã–π –ø—Ä–æ–µ–∫—Ç", callback_data="proj_create")
                ])
            else:
                # –°–ø–∏—Å–æ–∫ —Å—É—â–µ—Å—Ç–≤—É—é—â–∏—Ö –ø—Ä–æ–µ–∫—Ç–æ–≤
                for proj in projects:
                    emoji = "‚úÖ " if proj == current_project else "üìÅ "
                    keyboard.append([
                        InlineKeyboardButton(
                            text=f"{emoji}{proj}",
                            callback_data=f"proj_open_{proj}"
                        )
                    ])

            # –ö–Ω–æ–ø–∫–∞ –Ω–∞–∑–∞–¥
            keyboard.append([
                InlineKeyboardButton("¬´ –ù–∞–∑–∞–¥ –∫ –º–µ–Ω—é –ø—Ä–æ–µ–∫—Ç–æ–≤", callback_data="project_menu")
            ])

            reply_markup = InlineKeyboardMarkup(keyboard)

            status_text = f"**üìÇ –ú–û–ò –ü–†–û–ï–ö–¢–´**\n\n"
            if current_project:
                status_text += f"‚úÖ –ê–∫—Ç–∏–≤–Ω—ã–π: **{current_project}**\n\n"

            if projects:
                status_text += f"–í—Å–µ–≥–æ –ø—Ä–æ–µ–∫—Ç–æ–≤: {len(projects)}\n\n"
                status_text += "–í—ã–±–µ—Ä–∏—Ç–µ –ø—Ä–æ–µ–∫—Ç –¥–ª—è —Ä–∞–±–æ—Ç—ã:"
            else:
                status_text += "–£ –≤–∞—Å –ø–æ–∫–∞ –Ω–µ—Ç –ø—Ä–æ–µ–∫—Ç–æ–≤.\n\n"
                status_text += "–°–æ–∑–¥–∞–π—Ç–µ –ø–µ—Ä–≤—ã–π –ø—Ä–æ–µ–∫—Ç –¥–ª—è –Ω–∞—á–∞–ª–∞ —Ä–∞–±–æ—Ç—ã!"

            await query.edit_message_text(
                status_text,
                reply_markup=reply_markup,
                parse_mode="Markdown"
            )
        else:
            await query.edit_message_text("‚ö†Ô∏è –ú–æ–¥—É–ª—å –ø—Ä–æ–µ–∫—Ç–æ–≤ –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω.")
    elif query.data == "clear_confirm":
        # –ü–æ–¥—Ç–≤–µ—Ä–∂–¥–µ–Ω–∏–µ –æ—á–∏—Å—Ç–∫–∏ –∏—Å—Ç–æ—Ä–∏–∏
        user_id = update.effective_user.id
        clear_user_history(user_id)
        await query.edit_message_text(
            "‚úÖ –ò—Å—Ç–æ—Ä–∏—è –¥–∏–∞–ª–æ–≥–æ–≤ —É—Å–ø–µ—à–Ω–æ –æ—á–∏—â–µ–Ω–∞!\n\n"
            "–í—ã –º–æ–∂–µ—Ç–µ –Ω–∞—á–∞—Ç—å –Ω–æ–≤—ã–π –¥–∏–∞–ª–æ–≥ —Å —á–∏—Å—Ç–æ–≥–æ –ª–∏—Å—Ç–∞.",
            parse_mode='Markdown'
        )
    elif query.data == "clear_cancel":
        # –û—Ç–º–µ–Ω–∞ –æ—á–∏—Å—Ç–∫–∏
        await query.edit_message_text(
            "‚ùå –û—á–∏—Å—Ç–∫–∞ –∏—Å—Ç–æ—Ä–∏–∏ –æ—Ç–º–µ–Ω–µ–Ω–∞.\n\n"
            "–í–∞—à–∏ –¥–∞–Ω–Ω—ã–µ —Å–æ—Ö—Ä–∞–Ω–µ–Ω—ã.",
            parse_mode='Markdown'
        )
    elif query.data == "export_pdf":
        # –≠–∫—Å–ø–æ—Ä—Ç –≤ PDF
        user_id = update.effective_user.id
        try:
            await query.edit_message_text("‚è≥ –°–æ–∑–¥–∞—é PDF —Ñ–∞–π–ª...")
            pdf_buffer = export_history_to_pdf(user_id)
            filename = f"history_{user_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            await query.message.reply_document(
                document=pdf_buffer,
                filename=filename,
                caption="üìÑ –ò—Å—Ç–æ—Ä–∏—è –¥–∏–∞–ª–æ–≥–æ–≤ –≤ —Ñ–æ—Ä–º–∞—Ç–µ PDF"
            )
            await query.edit_message_text("‚úÖ PDF —Ñ–∞–π–ª —É—Å–ø–µ—à–Ω–æ —Å–æ–∑–¥–∞–Ω!")
        except Exception as e:
            logger.error(f"Error exporting to PDF: {e}")
            await query.edit_message_text(
                f"‚ùå –û—à–∏–±–∫–∞ –ø—Ä–∏ —Å–æ–∑–¥–∞–Ω–∏–∏ PDF:\n{str(e)}\n\n"
                "–ü–æ–ø—Ä–æ–±—É–π—Ç–µ —ç–∫—Å–ø–æ—Ä—Ç –≤ Word –∏–ª–∏ –æ–±—Ä–∞—Ç–∏—Ç–µ—Å—å –∫ –∞–¥–º–∏–Ω–∏—Å—Ç—Ä–∞—Ç–æ—Ä—É."
            )
    elif query.data == "export_docx":
        # –≠–∫—Å–ø–æ—Ä—Ç –≤ Word
        user_id = update.effective_user.id
        try:
            await query.edit_message_text("‚è≥ –°–æ–∑–¥–∞—é Word —Ñ–∞–π–ª...")
            docx_buffer = export_history_to_docx(user_id)
            filename = f"history_{user_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            await query.message.reply_document(
                document=docx_buffer,
                filename=filename,
                caption="üìù –ò—Å—Ç–æ—Ä–∏—è –¥–∏–∞–ª–æ–≥–æ–≤ –≤ —Ñ–æ—Ä–º–∞—Ç–µ Word"
            )
            await query.edit_message_text("‚úÖ Word —Ñ–∞–π–ª —É—Å–ø–µ—à–Ω–æ —Å–æ–∑–¥–∞–Ω!")
        except Exception as e:
            logger.error(f"Error exporting to Word: {e}")
            await query.edit_message_text(
                f"‚ùå –û—à–∏–±–∫–∞ –ø—Ä–∏ —Å–æ–∑–¥–∞–Ω–∏–∏ Word:\n{str(e)}\n\n"
                "–ü–æ–ø—Ä–æ–±—É–π—Ç–µ —ç–∫—Å–ø–æ—Ä—Ç –≤ PDF –∏–ª–∏ –æ–±—Ä–∞—Ç–∏—Ç–µ—Å—å –∫ –∞–¥–º–∏–Ω–∏—Å—Ç—Ä–∞—Ç–æ—Ä—É."
            )
    elif query.data == "export_cancel":
        # –û—Ç–º–µ–Ω–∞ —ç–∫—Å–ø–æ—Ä—Ç–∞
        await query.edit_message_text(
            "‚ùå –≠–∫—Å–ø–æ—Ä—Ç –æ—Ç–º–µ–Ω–µ–Ω.",
            parse_mode='Markdown'
        )

    # === –û–ë–†–ê–ë–û–¢–ß–ò–ö–ò –ö–ù–û–ü–û–ö v3.0 ===

    elif query.data == "clarify":
        # –ö–Ω–æ–ø–∫–∞ "–£—Ç–æ—á–Ω–∏—Ç—å"
        await query.edit_message_text(
            "üîç **–£—Ç–æ—á–Ω—è—é—â–∏–µ –≤–æ–ø—Ä–æ—Å—ã:**\n\n"
            "–í—ã –º–æ–∂–µ—Ç–µ –∑–∞–¥–∞—Ç—å –¥–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω—ã–π –≤–æ–ø—Ä–æ—Å –ø–æ —Ç–µ–º–µ.\n\n"
            "–ù–∞–ø—Ä–∏–º–µ—Ä:\n"
            "‚Ä¢ –ö–∞–∫–∏–µ –µ—â—ë —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è?\n"
            "‚Ä¢ –ü–æ–∫–∞–∂–∏ —Ä–∞—Å—á—ë—Ç\n"
            "‚Ä¢ –ö–∞–∫–∏–µ –¥–æ–∫—É–º–µ–Ω—Ç—ã –Ω—É–∂–Ω—ã?\n"
            "‚Ä¢ –ß—Ç–æ –µ—Å–ª–∏ –∏–∑–º–µ–Ω–∏—Ç—å –ø–∞—Ä–∞–º–µ—Ç—Ä—ã?\n\n"
            "–ü—Ä–æ—Å—Ç–æ –Ω–∞–ø–∏—à–∏—Ç–µ –≤–∞—à –≤–æ–ø—Ä–æ—Å –≤ —á–∞—Ç!",
            parse_mode='Markdown'
        )

    elif query.data == "example":
        # –ö–Ω–æ–ø–∫–∞ "–ü—Ä–∏–º–µ—Ä"
        await query.edit_message_text(
            "üí° **–ü—Ä–∞–∫—Ç–∏—á–µ—Å–∫–∏–π –ø—Ä–∏–º–µ—Ä:**\n\n"
            "–ß—Ç–æ–±—ã –ø–æ–ª—É—á–∏—Ç—å –ø—Ä–∏–º–µ—Ä –ø–æ –≤–∞—à–µ–π —Ç–µ–º–µ, –Ω–∞–ø–∏—à–∏—Ç–µ:\n\n"
            "‚Ä¢ \"–ü–æ–∫–∞–∂–∏ –ø—Ä–∏–º–µ—Ä —Ä–∞—Å—á—ë—Ç–∞\"\n"
            "‚Ä¢ \"–î–∞–π –ø—Ä–∞–∫—Ç–∏—á–µ—Å–∫–∏–π –ø—Ä–∏–º–µ—Ä\"\n"
            "‚Ä¢ \"–ö–∞–∫ —ç—Ç–æ –ø—Ä–∏–º–µ–Ω–∏—Ç—å –Ω–∞ –ø—Ä–∞–∫—Ç–∏–∫–µ?\"\n\n"
            "–Ø –ø—Ä–µ–¥–æ—Å—Ç–∞–≤–ª—é –∫–æ–Ω–∫—Ä–µ—Ç–Ω—ã–π –ø—Ä–∏–º–µ—Ä —Å —Ü–∏—Ñ—Ä–∞–º–∏ –∏ —Ä–∞—Å—á—ë—Ç–∞–º–∏!",
            parse_mode='Markdown'
        )

    # legacy callbacks (—Ä–∞–Ω—å—à–µ –±—ã–ª –æ—Ç–¥–µ–ª—å–Ω—ã–π —ç–∫—Ä–∞–Ω —Å–≤—è–∑–∞–Ω–Ω—ã—Ö –≤–æ–ø—Ä–æ—Å–æ–≤)
    elif query.data == "show_related_questions":
        await query.answer("‚ÑπÔ∏è –°–≤—è–∑–∞–Ω–Ω—ã–µ –≤–æ–ø—Ä–æ—Å—ã —Ç–µ–ø–µ—Ä—å —Å—Ä–∞–∑—É –ø–æ–¥ –æ—Ç–≤–µ—Ç–æ–º", show_alert=True)

    elif query.data == "hide_related_questions":
        await query.answer("‚ÑπÔ∏è –°–≤—è–∑–∞–Ω–Ω—ã–µ –≤–æ–ø—Ä–æ—Å—ã —Ç–µ–ø–µ—Ä—å —Å—Ä–∞–∑—É –ø–æ–¥ –æ—Ç–≤–µ—Ç–æ–º", show_alert=True)

    elif query.data.startswith("related_q_"):
        # –ö–ª–∏–∫ –Ω–∞ —Å–≤—è–∑–∞–Ω–Ω—ã–π –≤–æ–ø—Ä–æ—Å - –æ—Ç–ø—Ä–∞–≤–ª—è–µ–º –µ–≥–æ –∫–∞–∫ –Ω–æ–≤—ã–π –≤–æ–ø—Ä–æ—Å
        try:
            question_index = int(query.data.split("_")[-1])
            related_questions = context.user_data.get("related_questions", [])

            if question_index < len(related_questions):
                selected_question = related_questions[question_index]

                # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –≤–æ–ø—Ä–æ—Å –æ—Ç –∏–º–µ–Ω–∏ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è
                await query.answer(f"–ó–∞–¥–∞—é –≤–æ–ø—Ä–æ—Å: {selected_question[:50]}...")

                # –°–æ–∑–¥–∞—ë–º —Ñ–µ–π–∫–æ–≤–æ–µ —Å–æ–æ–±—â–µ–Ω–∏–µ —Å –≤–æ–ø—Ä–æ—Å–æ–º –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏
                from telegram import Message, Chat, User as TelegramUser
                fake_message = Message(
                    message_id=0,
                    date=datetime.now(),
                    chat=query.message.chat,
                    from_user=query.from_user,
                    text=selected_question
                )
                fake_update = Update(update_id=0, message=fake_message)

                # –û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º –∫–∞–∫ –æ–±—ã—á–Ω—ã–π —Ç–µ–∫—Å—Ç–æ–≤—ã–π –≤–æ–ø—Ä–æ—Å
                await handle_text(fake_update, context)
                logger.info(f"‚úÖ –û–±—Ä–∞–±–æ—Ç–∞–Ω —Å–≤—è–∑–∞–Ω–Ω—ã–π –≤–æ–ø—Ä–æ—Å #{question_index}")
            else:
                await query.answer("‚ö†Ô∏è –í–æ–ø—Ä–æ—Å –Ω–µ –Ω–∞–π–¥–µ–Ω", show_alert=True)
        except Exception as e:
            logger.error(f"‚ùå –û—à–∏–±–∫–∞ –æ–±—Ä–∞–±–æ—Ç–∫–∏ —Å–≤—è–∑–∞–Ω–Ω–æ–≥–æ –≤–æ–ø—Ä–æ—Å–∞: {e}")
            await query.answer("‚ö†Ô∏è –û—à–∏–±–∫–∞ –æ–±—Ä–∞–±–æ—Ç–∫–∏ –≤–æ–ø—Ä–æ—Å–∞", show_alert=True)


    # ===== New inline "menu in one row" actions =====

    elif query.data == "answer_hide":
        # –°–∫—Ä—ã—Ç—å ¬´—Ñ—É–Ω–∫—Ü–∏–∏¬ª (–∫–Ω–æ–ø–∫–∏ –ø–æ–¥ —Å–æ–æ–±—â–µ–Ω–∏–µ–º) –∏ –æ—Å—Ç–∞–≤–∏—Ç—å —Ç–æ–ª—å–∫–æ —Ç–µ–∫—Å—Ç –æ—Ç–≤–µ—Ç–∞.
        # –≠—Ç–æ –∞–Ω–∞–ª–æ–≥ —Å–∏—Å—Ç–µ–º–Ω–æ–π –∫–Ω–æ–ø–∫–∏ Telegram ¬´—Å–∫—Ä—ã—Ç—å –∫–ª–∞–≤–∏–∞—Ç—É—Ä—É¬ª, –Ω–æ –¥–ª—è InlineKeyboard.
        try:
            await query.edit_message_reply_markup(reply_markup=None)
            await query.answer("ü´• –°–∫—Ä—ã—Ç–æ")
        except Exception as e:
            logger.error(f"‚ùå –û—à–∏–±–∫–∞ answer_hide: {e}")
            await query.answer("‚ö†Ô∏è –ù–µ —É–¥–∞–ª–æ—Å—å —Å–∫—Ä—ã—Ç—å –∫–Ω–æ–ø–∫–∏", show_alert=True)

    elif query.data == "answer_menu":
        # –ü–æ–∫–∞–∑–∞—Ç—å/–≤–µ—Ä–Ω—É—Ç—å –º–µ–Ω—é –ø–æ–¥ –æ—Ç–≤–µ—Ç
        if not IMPROVEMENTS_V3_AVAILABLE:
            await query.answer("‚ö†Ô∏è –§—É–Ω–∫—Ü–∏—è –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∞", show_alert=True)
            return

        related_questions = context.user_data.get("related_questions", [])
        keyboard = create_answer_buttons(related_questions=related_questions)
        try:
            await query.edit_message_reply_markup(reply_markup=keyboard)
            await query.answer("‚ò∞ –ú–µ–Ω—é")
        except Exception as e:
            logger.error(f"‚ùå –û—à–∏–±–∫–∞ answer_menu: {e}")
            await query.answer("‚ö†Ô∏è –ù–µ —É–¥–∞–ª–æ—Å—å –ø–æ–∫–∞–∑–∞—Ç—å –º–µ–Ω—é", show_alert=True)

    elif query.data == "answer_more":
        # –ü–æ–ø—Ä–æ—Å–∏—Ç—å –º–æ–¥–µ–ª—å –¥–∞—Ç—å –µ—â—ë –æ–¥–Ω—É –≤–µ—Ä—Å–∏—é –æ—Ç–≤–µ—Ç–∞
        original_q = context.user_data.get("last_question")
        last_answer = context.user_data.get("last_answer")
        if not original_q:
            await query.answer("‚ö†Ô∏è –ù–µ –Ω–∞–π–¥–µ–Ω –∏—Å—Ö–æ–¥–Ω—ã–π –≤–æ–ø—Ä–æ—Å", show_alert=True)
            return

        await query.answer("üîÅ –ì–µ–Ω–µ—Ä–∏—Ä—É—é –µ—â—ë –≤–∞—Ä–∏–∞–Ω—Ç‚Ä¶")

        try:
            # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –∫–æ—Ä–æ—Ç–∫–∏–π follow-up –ø—Ä–æ–º–ø—Ç –≤ –æ–±—â–∏–π –æ–±—Ä–∞–±–æ—Ç—á–∏–∫ (–∫–∞–∫ –Ω–æ–≤—ã–π –≤–æ–ø—Ä–æ—Å)
            followup = (
                f"–î–∞–π –∞–ª—å—Ç–µ—Ä–Ω–∞—Ç–∏–≤–Ω—É—é –≤–µ—Ä—Å–∏—é –æ—Ç–≤–µ—Ç–∞ –Ω–∞ –≤–æ–ø—Ä–æ—Å: {original_q}\n\n"
                f"–¢–µ–∫—É—â–∏–π –æ—Ç–≤–µ—Ç (–¥–ª—è –∫–æ–Ω—Ç–µ–∫—Å—Ç–∞):\n{last_answer}\n\n"
                "–°–¥–µ–ª–∞–π –ø–æ-–¥—Ä—É–≥–æ–º—É: –¥—Ä—É–≥–∞—è —Å—Ç—Ä—É–∫—Ç—É—Ä–∞/—Ñ–æ—Ä–º—É–ª–∏—Ä–æ–≤–∫–∏, –Ω–æ –±–µ–∑ –≤—ã–¥—É–º–∞–Ω–Ω—ã—Ö —Ñ–∞–∫—Ç–æ–≤."
            )

            from telegram import Message, Update
            fake_message = Message(
                message_id=0,
                date=datetime.now(),
                chat=query.message.chat,
                from_user=query.from_user,
                text=followup,
            )
            fake_update = Update(update_id=0, message=fake_message)
            await handle_text(fake_update, context)
        except Exception as e:
            logger.error(f"‚ùå –û—à–∏–±–∫–∞ answer_more: {e}")
            await query.answer("‚ö†Ô∏è –ù–µ —É–¥–∞–ª–æ—Å—å —Å–≥–µ–Ω–µ—Ä–∏—Ä–æ–≤–∞—Ç—å –≤–∞—Ä–∏–∞–Ω—Ç", show_alert=True)

    elif query.data == "answer_edit":
        # –ü–æ–¥—Å–∫–∞–∑–∫–∞ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—é –∫–∞–∫ —É—Ç–æ—á–Ω–∏—Ç—å/–ø–µ—Ä–µ–æ—Ñ–æ—Ä–º–∏—Ç—å –∑–∞–ø—Ä–æ—Å
        await query.answer("‚úèÔ∏è", show_alert=False)
        try:
            await query.message.reply_text(
                "‚úèÔ∏è –ù–∞–ø–∏—à–∏—Ç–µ —É—Ç–æ—á–Ω–µ–Ω–∏–µ –æ–¥–Ω–∏–º —Å–æ–æ–±—â–µ–Ω–∏–µ–º ‚Äî —è –ø–µ—Ä–µ—Ä–∞–±–æ—Ç–∞—é –æ—Ç–≤–µ—Ç.\n\n"
                "–ü—Ä–∏–º–µ—Ä—ã:\n"
                "‚Ä¢ ¬´–°–¥–µ–ª–∞–π –∫–æ—Ä–æ—á–µ –∏ –ø–æ –ø—É–Ω–∫—Ç–∞–º¬ª\n"
                "‚Ä¢ ¬´–î–∞–π 2 –≤–∞—Ä–∏–∞–Ω—Ç–∞ —Ä–µ—à–µ–Ω–∏—è¬ª\n"
                "‚Ä¢ ¬´–£–∫–∞–∂–∏ —Ä–∏—Å–∫–∏ –∏ –∫–∞–∫ –ø—Ä–æ–≤–µ—Ä–∏—Ç—å –∫–∞—á–µ—Å—Ç–≤–æ¬ª"
            )
        except Exception as e:
            logger.error(f"‚ùå –û—à–∏–±–∫–∞ answer_edit: {e}")

    elif query.data == "show_regulations":
        # –ö–Ω–æ–ø–∫–∞ "–ù–æ—Ä–º–∞—Ç–∏–≤—ã" - –ø–æ–∫–∞–∑–∞—Ç—å –∫–∞—Ç–µ–≥–æ—Ä–∏–∏
        if IMPROVEMENTS_V3_AVAILABLE:
            keyboard = create_regulations_category_menu()
            await query.edit_message_text(
                "üìö **–ù–û–†–ú–ê–¢–ò–í–ù–´–ï –î–û–ö–£–ú–ï–ù–¢–´**\n\n"
                "–í—ã–±–µ—Ä–∏—Ç–µ –∫–∞—Ç–µ–≥–æ—Ä–∏—é –¥–ª—è –ø—Ä–æ—Å–º–æ—Ç—Ä–∞ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤:",
                reply_markup=keyboard,
                parse_mode='Markdown'
            )
        else:
            await query.edit_message_text(
                "üìö –ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ –∫–æ–º–∞–Ω–¥—É /regulations –¥–ª—è –ø—Ä–æ—Å–º–æ—Ç—Ä–∞ –≤—Å–µ—Ö –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤.",
                parse_mode='Markdown'
            )

    elif query.data == "calculator":
        # –ö–Ω–æ–ø–∫–∞ "–ö–∞–ª—å–∫—É–ª—è—Ç–æ—Ä" - –ø–æ–∫–∞–∑–∞—Ç—å –º–µ–Ω—é –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä–æ–≤
        if CALCULATORS_AVAILABLE:
            keyboard = create_calculators_menu()
            await query.edit_message_text(
                "üßÆ **–°–¢–†–û–ò–¢–ï–õ–¨–ù–´–ï –ö–ê–õ–¨–ö–£–õ–Ø–¢–û–†–´**\n\n"
                "–í—ã–±–µ—Ä–∏—Ç–µ –Ω—É–∂–Ω—ã–π –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä:",
                reply_markup=keyboard,
                parse_mode='Markdown'
            )
        else:
            await query.edit_message_text(
                "‚ö†Ô∏è –ú–æ–¥—É–ª—å –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä–æ–≤ –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω.",
                parse_mode='Markdown'
            )

    elif query.data == "save_query":
        # –ö–Ω–æ–ø–∫–∞ "–°–æ—Ö—Ä–∞–Ω–∏—Ç—å"
        await query.edit_message_text(
            "üíæ **–°–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ –∑–∞–ø—Ä–æ—Å–∞**\n\n"
            "–§—É–Ω–∫—Ü–∏—è —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è –±—É–¥–µ—Ç –¥–æ—Å—Ç—É–ø–Ω–∞ –≤ —Å–ª–µ–¥—É—é—â–µ–º –æ–±–Ω–æ–≤–ª–µ–Ω–∏–∏ (v3.1).\n\n"
            "–ó–∞–ø–ª–∞–Ω–∏—Ä–æ–≤–∞–Ω–æ:\n"
            "‚Ä¢ –°–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ –∏–∑–±—Ä–∞–Ω–Ω—ã—Ö –∑–∞–ø—Ä–æ—Å–æ–≤\n"
            "‚Ä¢ –ë—ã—Å—Ç—Ä—ã–π –¥–æ—Å—Ç—É–ø –∫ —Å–æ—Ö—Ä–∞–Ω—ë–Ω–Ω—ã–º\n"
            "‚Ä¢ –≠–∫—Å–ø–æ—Ä—Ç –≤ PDF/Word\n\n"
            "–°–ª–µ–¥–∏—Ç–µ –∑–∞ –æ–±–Ω–æ–≤–ª–µ–Ω–∏—è–º–∏!",
            parse_mode='Markdown'
        )

    # === –ò–ù–¢–ï–†–ê–ö–¢–ò–í–ù–´–ï –ö–ê–õ–¨–ö–£–õ–Ø–¢–û–†–´ v5.0 ===
    # –í–ê–ñ–ù–û: calc_concrete –∏ calc_reinforcement –æ–±—Ä–∞–±–∞—Ç—ã–≤–∞—é—Ç—Å—è ConversationHandler
    # –∏–∑ calculator_handlers.py - –ù–ï –¥—É–±–ª–∏—Ä–æ–≤–∞—Ç—å –∑–¥–µ—Å—å!

    elif query.data == "calc_brick":
        await query.edit_message_text(
            "üß± **–ö–ê–õ–¨–ö–£–õ–Ø–¢–û–† –ö–ò–†–ü–ò–ß–ê/–ë–õ–û–ö–û–í**\n\n"
            "–ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ –ø–æ—à–∞–≥–æ–≤—ã–π —Ä–µ–∂–∏–º —á–µ—Ä–µ–∑ –º–µ–Ω—é –∏–ª–∏:\n"
            "`/calc_brick –¥–ª–∏–Ω–∞ –≤—ã—Å–æ—Ç–∞ —Ç–æ–ª—â–∏–Ω–∞ —Ç–∏–ø –ø—Ä–æ—ë–º—ã`\n\n"
            "**–ü—Ä–∏–º–µ—Ä:**\n"
            "`/calc_brick 10 3 0.25 single 5`\n\n"
            "–ü–∞—Ä–∞–º–µ—Ç—Ä—ã:\n"
            "‚Ä¢ –î–ª–∏–Ω–∞ –∏ –≤—ã—Å–æ—Ç–∞ —Å—Ç–µ–Ω—ã (–º)\n"
            "‚Ä¢ –¢–æ–ª—â–∏–Ω–∞: 0.12, 0.25, 0.38, 0.51 –º\n"
            "‚Ä¢ –¢–∏–ø: single, one_and_half, double\n"
            "‚Ä¢ –ü–ª–æ—â–∞–¥—å –ø—Ä–æ—ë–º–æ–≤ (–º¬≤)",
            parse_mode='Markdown'
        )

    elif query.data == "calc_tile":
        await query.edit_message_text(
            "üî≤ **–ö–ê–õ–¨–ö–£–õ–Ø–¢–û–† –ü–õ–ò–¢–ö–ò**\n\n"
            "–ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ –ø–æ—à–∞–≥–æ–≤—ã–π —Ä–µ–∂–∏–º —á–µ—Ä–µ–∑ –º–µ–Ω—é –∏–ª–∏:\n"
            "`/calc_tile –¥–ª–∏–Ω–∞ —à–∏—Ä–∏–Ω–∞ —Ä–∞–∑–º–µ—Ä_–ø–ª–∏—Ç–∫–∏ –∑–∞–ø–∞—Å`\n\n"
            "**–ü—Ä–∏–º–µ—Ä:**\n"
            "`/calc_tile 5 4 0.3 10`\n\n"
            "–ü–∞—Ä–∞–º–µ—Ç—Ä—ã:\n"
            "‚Ä¢ –†–∞–∑–º–µ—Ä—ã –ø–æ–º–µ—â–µ–Ω–∏—è (–º)\n"
            "‚Ä¢ –†–∞–∑–º–µ—Ä –ø–ª–∏—Ç–∫–∏ (–º, –Ω–∞–ø—Ä–∏–º–µ—Ä 0.3 –¥–ª—è 30√ó30 —Å–º)\n"
            "‚Ä¢ –ó–∞–ø–∞—Å –Ω–∞ –ø–æ–¥—Ä–µ–∑–∫—É (%)",
            parse_mode='Markdown'
        )

    elif query.data == "calc_paint":
        await query.edit_message_text(
            "üé® **–ö–ê–õ–¨–ö–£–õ–Ø–¢–û–† –ö–†–ê–°–ö–ò**\n\n"
            "–ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ –ø–æ—à–∞–≥–æ–≤—ã–π —Ä–µ–∂–∏–º —á–µ—Ä–µ–∑ –º–µ–Ω—é –∏–ª–∏:\n"
            "`/calc_paint –ø–ª–æ—â–∞–¥—å —Ç–∏–ø –ø–æ–≤–µ—Ä—Ö–Ω–æ—Å—Ç—å —Å–ª–æ–∏`\n\n"
            "**–ü—Ä–∏–º–µ—Ä:**\n"
            "`/calc_paint 50 water smooth 2`\n\n"
            "–ü–∞—Ä–∞–º–µ—Ç—Ä—ã:\n"
            "‚Ä¢ –ü–ª–æ—â–∞–¥—å (–º¬≤)\n"
            "‚Ä¢ –¢–∏–ø: water, oil, latex\n"
            "‚Ä¢ –ü–æ–≤–µ—Ä—Ö–Ω–æ—Å—Ç—å: smooth, rough, porous\n"
            "‚Ä¢ –ö–æ–ª–∏—á–µ—Å—Ç–≤–æ —Å–ª–æ—ë–≤",
            parse_mode='Markdown'
        )

    elif query.data == "calc_wall_area":
        await query.edit_message_text(
            "üìê **–ö–ê–õ–¨–ö–£–õ–Ø–¢–û–† –ü–õ–û–©–ê–î–ò –°–¢–ï–ù**\n\n"
            "–ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ –ø–æ—à–∞–≥–æ–≤—ã–π —Ä–µ–∂–∏–º —á–µ—Ä–µ–∑ –º–µ–Ω—é –∏–ª–∏:\n"
            "`/calc_wall_area –¥–ª–∏–Ω–∞ —à–∏—Ä–∏–Ω–∞ –≤—ã—Å–æ—Ç–∞ –æ–∫–Ω–∞ –¥–≤–µ—Ä–∏`\n\n"
            "**–ü—Ä–∏–º–µ—Ä:**\n"
            "`/calc_wall_area 5 4 2.7 2 1`\n\n"
            "–ü–∞—Ä–∞–º–µ—Ç—Ä—ã:\n"
            "‚Ä¢ –†–∞–∑–º–µ—Ä—ã –ø–æ–º–µ—â–µ–Ω–∏—è (–º)\n"
            "‚Ä¢ –ö–æ–ª–∏—á–µ—Å—Ç–≤–æ –æ–∫–æ–Ω\n"
            "‚Ä¢ –ö–æ–ª–∏—á–µ—Å—Ç–≤–æ –¥–≤–µ—Ä–µ–π",
            parse_mode='Markdown'
        )

    elif query.data == "calc_roof":
        await query.edit_message_text(
            "üè† **–ö–ê–õ–¨–ö–£–õ–Ø–¢–û–† –ö–†–û–í–õ–ò**\n\n"
            "–ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ –ø–æ—à–∞–≥–æ–≤—ã–π —Ä–µ–∂–∏–º —á–µ—Ä–µ–∑ –º–µ–Ω—é –∏–ª–∏:\n"
            "`/calc_roof –¥–ª–∏–Ω–∞ —à–∏—Ä–∏–Ω–∞ —Ç–∏–ø —É–≥–æ–ª —Å–≤–µ—Å`\n\n"
            "**–ü—Ä–∏–º–µ—Ä:**\n"
            "`/calc_roof 10 8 gable 30 0.5`\n\n"
            "–ü–∞—Ä–∞–º–µ—Ç—Ä—ã:\n"
            "‚Ä¢ –†–∞–∑–º–µ—Ä—ã –∑–¥–∞–Ω–∏—è (–º)\n"
            "‚Ä¢ –¢–∏–ø: flat, gable, hip\n"
            "‚Ä¢ –£–≥–æ–ª –Ω–∞–∫–ª–æ–Ω–∞ (–≥—Ä–∞–¥—É—Å—ã)\n"
            "‚Ä¢ –°–≤–µ—Å (–º)",
            parse_mode='Markdown'
        )

    elif query.data == "calc_plaster":
        await query.edit_message_text(
            "üß± **–ö–ê–õ–¨–ö–£–õ–Ø–¢–û–† –®–¢–£–ö–ê–¢–£–†–ö–ò**\n\n"
            "–ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ –ø–æ—à–∞–≥–æ–≤—ã–π —Ä–µ–∂–∏–º —á–µ—Ä–µ–∑ –º–µ–Ω—é –∏–ª–∏:\n"
            "`/calc_plaster –ø–ª–æ—â–∞–¥—å —Ç–æ–ª—â–∏–Ω–∞ —Ç–∏–ø`\n\n"
            "**–ü—Ä–∏–º–µ—Ä:**\n"
            "`/calc_plaster 100 0.02 cement`\n\n"
            "–ü–∞—Ä–∞–º–µ—Ç—Ä—ã:\n"
            "‚Ä¢ –ü–ª–æ—â–∞–¥—å (–º¬≤)\n"
            "‚Ä¢ –¢–æ–ª—â–∏–Ω–∞ —Å–ª–æ—è (–º, –Ω–∞–ø—Ä–∏–º–µ—Ä 0.02 –¥–ª—è 20 –º–º)\n"
            "‚Ä¢ –¢–∏–ø: cement, gypsum, decorative",
            parse_mode='Markdown'
        )

    elif query.data == "calc_wallpaper":
        await query.edit_message_text(
            "üñºÔ∏è **–ö–ê–õ–¨–ö–£–õ–Ø–¢–û–† –û–ë–û–ï–í**\n\n"
            "–ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ –ø–æ—à–∞–≥–æ–≤—ã–π —Ä–µ–∂–∏–º —á–µ—Ä–µ–∑ –º–µ–Ω—é –∏–ª–∏:\n"
            "`/calc_wallpaper –ø–ª–æ—â–∞–¥—å –¥–ª–∏–Ω–∞_—Ä—É–ª–æ–Ω–∞ —à–∏—Ä–∏–Ω–∞ —Ä–∞–ø–ø–æ—Ä—Ç –∑–∞–ø–∞—Å`\n\n"
            "**–ü—Ä–∏–º–µ—Ä:**\n"
            "`/calc_wallpaper 50 10 0.53 0 10`\n\n"
            "–ü–∞—Ä–∞–º–µ—Ç—Ä—ã:\n"
            "‚Ä¢ –ü–ª–æ—â–∞–¥—å —Å—Ç–µ–Ω (–º¬≤)\n"
            "‚Ä¢ –†–∞–∑–º–µ—Ä—ã —Ä—É–ª–æ–Ω–∞ (–º)\n"
            "‚Ä¢ –†–∞–ø–ø–æ—Ä—Ç —Ä–∏—Å—É–Ω–∫–∞ (–º, 0 –µ—Å–ª–∏ –±–µ–∑ —Ä–∏—Å—É–Ω–∫–∞)\n"
            "‚Ä¢ –ó–∞–ø–∞—Å (%)",
            parse_mode='Markdown'
        )

    elif query.data == "calc_laminate":
        await query.edit_message_text(
            "ü™µ **–ö–ê–õ–¨–ö–£–õ–Ø–¢–û–† –õ–ê–ú–ò–ù–ê–¢–ê**\n\n"
            "–ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ –ø–æ—à–∞–≥–æ–≤—ã–π —Ä–µ–∂–∏–º —á–µ—Ä–µ–∑ –º–µ–Ω—é –∏–ª–∏:\n"
            "`/calc_laminate –¥–ª–∏–Ω–∞ —à–∏—Ä–∏–Ω–∞ –¥–ª–∏–Ω–∞_–¥–æ—Å–∫–∏ —à–∏—Ä–∏–Ω–∞_–¥–æ—Å–∫–∏ –∑–∞–ø–∞—Å`\n\n"
            "**–ü—Ä–∏–º–µ—Ä:**\n"
            "`/calc_laminate 5 4 1.28 0.192 7`\n\n"
            "–ü–∞—Ä–∞–º–µ—Ç—Ä—ã:\n"
            "‚Ä¢ –†–∞–∑–º–µ—Ä—ã –ø–æ–º–µ—â–µ–Ω–∏—è (–º)\n"
            "‚Ä¢ –†–∞–∑–º–µ—Ä—ã –¥–æ—Å–∫–∏ (–º)\n"
            "‚Ä¢ –ó–∞–ø–∞—Å –Ω–∞ –ø–æ–¥—Ä–µ–∑–∫—É (%)",
            parse_mode='Markdown'
        )

    elif query.data == "calc_insulation":
        await query.edit_message_text(
            "üßä **–ö–ê–õ–¨–ö–£–õ–Ø–¢–û–† –£–¢–ï–ü–õ–ò–¢–ï–õ–Ø**\n\n"
            "–ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ –ø–æ—à–∞–≥–æ–≤—ã–π —Ä–µ–∂–∏–º —á–µ—Ä–µ–∑ –º–µ–Ω—é –∏–ª–∏:\n"
            "`/calc_insulation –ø–ª–æ—â–∞–¥—å —Ç–æ–ª—â–∏–Ω–∞ —Ç–∏–ø`\n\n"
            "**–ü—Ä–∏–º–µ—Ä:**\n"
            "`/calc_insulation 100 0.1 mineral_wool`\n\n"
            "–ü–∞—Ä–∞–º–µ—Ç—Ä—ã:\n"
            "‚Ä¢ –ü–ª–æ—â–∞–¥—å (–º¬≤)\n"
            "‚Ä¢ –¢–æ–ª—â–∏–Ω–∞ (–º, –Ω–∞–ø—Ä–∏–º–µ—Ä 0.1 –¥–ª—è 100 –º–º)\n"
            "‚Ä¢ –¢–∏–ø: mineral_wool, polystyrene, penoplex",
            parse_mode='Markdown'
        )

    elif query.data == "calc_foundation":
        await query.edit_message_text(
            "‚öì **–ö–ê–õ–¨–ö–£–õ–Ø–¢–û–† –§–£–ù–î–ê–ú–ï–ù–¢–ê**\n\n"
            "–ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ –ø–æ—à–∞–≥–æ–≤—ã–π —Ä–µ–∂–∏–º —á–µ—Ä–µ–∑ –º–µ–Ω—é –∏–ª–∏:\n"
            "`/calc_foundation –¥–ª–∏–Ω–∞ —à–∏—Ä–∏–Ω–∞ –≥–ª—É–±–∏–Ω–∞ —Ç–∏–ø`\n\n"
            "**–ü—Ä–∏–º–µ—Ä:**\n"
            "`/calc_foundation 10 8 1.5 strip`\n\n"
            "–ü–∞—Ä–∞–º–µ—Ç—Ä—ã:\n"
            "‚Ä¢ –†–∞–∑–º–µ—Ä—ã —Ñ—É–Ω–¥–∞–º–µ–Ω—Ç–∞ (–º)\n"
            "‚Ä¢ –ì–ª—É–±–∏–Ω–∞ –∑–∞–ª–æ–∂–µ–Ω–∏—è (–º)\n"
            "‚Ä¢ –¢–∏–ø: strip, slab, pile",
            parse_mode='Markdown'
        )

    elif query.data == "calc_stairs":
        await query.edit_message_text(
            "ü™ú **–ö–ê–õ–¨–ö–£–õ–Ø–¢–û–† –õ–ï–°–¢–ù–ò–¶–´**\n\n"
            "–ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ –ø–æ—à–∞–≥–æ–≤—ã–π —Ä–µ–∂–∏–º —á–µ—Ä–µ–∑ –º–µ–Ω—é –∏–ª–∏:\n"
            "`/calc_stairs –≤—ã—Å–æ—Ç–∞_—ç—Ç–∞–∂–∞ —à–∏—Ä–∏–Ω–∞_–ø—Ä–æ—Å—Ç—É–ø–∏ –≤—ã—Å–æ—Ç–∞_–ø–æ–¥—Å—Ç—É–ø–µ–Ω–∫–∞ —à–∏—Ä–∏–Ω–∞_–ª–µ—Å—Ç–Ω–∏—Ü—ã`\n\n"
            "**–ü—Ä–∏–º–µ—Ä:**\n"
            "`/calc_stairs 3 0.3 0.15 1`\n\n"
            "–ü–∞—Ä–∞–º–µ—Ç—Ä—ã:\n"
            "‚Ä¢ –í—ã—Å–æ—Ç–∞ —ç—Ç–∞–∂–∞ (–º)\n"
            "‚Ä¢ –®–∏—Ä–∏–Ω–∞ –ø—Ä–æ—Å—Ç—É–ø–∏ (–º)\n"
            "‚Ä¢ –í—ã—Å–æ—Ç–∞ –ø–æ–¥—Å—Ç—É–ø–µ–Ω–∫–∞ (–º)\n"
            "‚Ä¢ –®–∏—Ä–∏–Ω–∞ –ª–µ—Å—Ç–Ω–∏—Ü—ã (–º)",
            parse_mode='Markdown'
        )

    elif query.data == "calc_drywall":
        await query.edit_message_text(
            "üìê **–ö–ê–õ–¨–ö–£–õ–Ø–¢–û–† –ì–ò–ü–°–û–ö–ê–†–¢–û–ù–ê**\n\n"
            "–ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ –ø–æ—à–∞–≥–æ–≤—ã–π —Ä–µ–∂–∏–º —á–µ—Ä–µ–∑ –º–µ–Ω—é –∏–ª–∏:\n"
            "`/calc_drywall –¥–ª–∏–Ω–∞ —à–∏—Ä–∏–Ω–∞ –≤—ã—Å–æ—Ç–∞ –ø–æ—Ç–æ–ª–æ–∫`\n\n"
            "**–ü—Ä–∏–º–µ—Ä:**\n"
            "`/calc_drywall 5 4 2.7 1`\n\n"
            "–ü–∞—Ä–∞–º–µ—Ç—Ä—ã:\n"
            "‚Ä¢ –†–∞–∑–º–µ—Ä—ã –ø–æ–º–µ—â–µ–Ω–∏—è (–º)\n"
            "‚Ä¢ –ü–æ—Ç–æ–ª–æ–∫: 1 - –¥–∞, 0 - –Ω–µ—Ç",
            parse_mode='Markdown'
        )

    elif query.data == "calc_earthwork":
        await query.edit_message_text(
            "üöú **–ö–ê–õ–¨–ö–£–õ–Ø–¢–û–† –ó–ï–ú–õ–Ø–ù–´–• –†–ê–ë–û–¢**\n\n"
            "–ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ –ø–æ—à–∞–≥–æ–≤—ã–π —Ä–µ–∂–∏–º —á–µ—Ä–µ–∑ –º–µ–Ω—é –∏–ª–∏:\n"
            "`/calc_earthwork –¥–ª–∏–Ω–∞ —à–∏—Ä–∏–Ω–∞ –≥–ª—É–±–∏–Ω–∞ —Ç–∏–ø_–≥—Ä—É–Ω—Ç–∞`\n\n"
            "**–ü—Ä–∏–º–µ—Ä:**\n"
            "`/calc_earthwork 20 15 2 clay`\n\n"
            "–ü–∞—Ä–∞–º–µ—Ç—Ä—ã:\n"
            "‚Ä¢ –†–∞–∑–º–µ—Ä—ã –∫–æ—Ç–ª–æ–≤–∞–Ω–∞ (–º)\n"
            "‚Ä¢ –ì–ª—É–±–∏–Ω–∞ (–º)\n"
            "‚Ä¢ –¢–∏–ø –≥—Ä—É–Ω—Ç–∞: clay, sand, rock",
            parse_mode='Markdown'
        )

    elif query.data == "calc_labor":
        await query.edit_message_text(
            "üë∑ **–ö–ê–õ–¨–ö–£–õ–Ø–¢–û–† –¢–†–£–î–û–ó–ê–¢–†–ê–¢**\n\n"
            "–ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ –ø–æ—à–∞–≥–æ–≤—ã–π —Ä–µ–∂–∏–º —á–µ—Ä–µ–∑ –º–µ–Ω—é –∏–ª–∏:\n"
            "`/calc_labor —Ç–∏–ø_—Ä–∞–±–æ—Ç—ã –æ–±—ä—ë–º –µ–¥–∏–Ω–∏—Ü–∞`\n\n"
            "**–ü—Ä–∏–º–µ—Ä:**\n"
            "`/calc_labor concrete 50 m3`\n\n"
            "–ü–∞—Ä–∞–º–µ—Ç—Ä—ã:\n"
            "‚Ä¢ –¢–∏–ø: concrete, masonry, plaster, paint, tile, roof\n"
            "‚Ä¢ –û–±—ä—ë–º —Ä–∞–±–æ—Ç\n"
            "‚Ä¢ –ï–¥–∏–Ω–∏—Ü–∞: m3, m2, m",
            parse_mode='Markdown'
        )

    # –û–±—Ä–∞–±–æ—Ç—á–∏–∫–∏ –∫–∞—Ç–µ–≥–æ—Ä–∏–π –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤
    elif query.data.startswith("reg_cat_"):
        category = query.data.replace("reg_cat_", "")
        category_names = {
            "structures": "üèóÔ∏è –ö–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏",
            "foundations": "‚öì –û—Å–Ω–æ–≤–∞–Ω–∏—è –∏ —Ñ—É–Ω–¥–∞–º–µ–Ω—Ç—ã",
            "enclosures": "üß± –û–≥—Ä–∞–∂–¥–∞—é—â–∏–µ –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏",
            "engineering": "‚ö° –ò–Ω–∂–µ–Ω–µ—Ä–Ω—ã–µ —Å–∏—Å—Ç–µ–º—ã",
            "fire": "üî• –ü–æ–∂–∞—Ä–Ω–∞—è –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç—å",
            "quality": "‚úÖ –ö–æ–Ω—Ç—Ä–æ–ª—å –∫–∞—á–µ—Å—Ç–≤–∞",
            "inspection": "üîç –û–±—Å–ª–µ–¥–æ–≤–∞–Ω–∏–µ",
            "accessibility": "‚ôø –î–æ—Å—Ç—É–ø–Ω–æ—Å—Ç—å"
        }

        if IMPROVEMENTS_V3_AVAILABLE:
            from improvements_v3 import REGULATIONS_CATEGORIES
            cat_name = category_names.get(category, "–ö–∞—Ç–µ–≥–æ—Ä–∏—è")

            if category in ["structures", "foundations", "enclosures", "engineering",
                          "fire", "quality", "inspection", "accessibility"]:
                # –ü—Ä–µ–æ–±—Ä–∞–∑—É–µ–º –∞–Ω–≥–ª–∏–π—Å–∫–æ–µ –Ω–∞–∑–≤–∞–Ω–∏–µ –≤ —Ä—É—Å—Å–∫–æ–µ –¥–ª—è –ø–æ–∏—Å–∫–∞ –≤ —Å–ª–æ–≤–∞—Ä–µ
                rus_cat = {
                    "structures": "–ö–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏",
                    "foundations": "–û—Å–Ω–æ–≤–∞–Ω–∏—è –∏ —Ñ—É–Ω–¥–∞–º–µ–Ω—Ç—ã",
                    "enclosures": "–û–≥—Ä–∞–∂–¥–∞—é—â–∏–µ –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏",
                    "engineering": "–ò–Ω–∂–µ–Ω–µ—Ä–Ω—ã–µ —Å–∏—Å—Ç–µ–º—ã",
                    "fire": "–ü–æ–∂–∞—Ä–Ω–∞—è –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç—å",
                    "quality": "–ö–æ–Ω—Ç—Ä–æ–ª—å –∫–∞—á–µ—Å—Ç–≤–∞",
                    "inspection": "–û–±—Å–ª–µ–¥–æ–≤–∞–Ω–∏–µ",
                    "accessibility": "–î–æ—Å—Ç—É–ø–Ω–æ—Å—Ç—å"
                }

                cat_rus = rus_cat.get(category)
                if cat_rus and cat_rus in REGULATIONS_CATEGORIES:
                    docs = REGULATIONS_CATEGORIES[cat_rus]
                    text = f"üìö **{cat_name}**\n\n"
                    for doc_name, doc_info in docs.items():
                        text += f"**{doc_name}**\n"
                        text += f"_{doc_info['description']}_\n"
                        if 'link' in doc_info:
                            text += f"üîó [–û—Ç–∫—Ä—ã—Ç—å –¥–æ–∫—É–º–µ–Ω—Ç]({doc_info['link']})\n"
                        text += "\n"

                    await query.edit_message_text(text, parse_mode='Markdown', disable_web_page_preview=True)
                else:
                    await query.edit_message_text(f"‚ö†Ô∏è –î–æ–∫—É–º–µ–Ω—Ç—ã –∫–∞—Ç–µ–≥–æ—Ä–∏–∏ '{cat_name}' –Ω–µ –Ω–∞–π–¥–µ–Ω—ã.")
            else:
                await query.edit_message_text("‚ö†Ô∏è –ù–µ–∏–∑–≤–µ—Å—Ç–Ω–∞—è –∫–∞—Ç–µ–≥–æ—Ä–∏—è.")
        else:
            await query.edit_message_text("‚ö†Ô∏è –ú–æ–¥—É–ª—å –∫–∞—Ç–µ–≥–æ—Ä–∏–π –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω.")

    # –û–±—Ä–∞–±–æ—Ç—á–∏–∫–∏ –≤—ã–±–æ—Ä–∞ —Ä–µ–≥–∏–æ–Ω–∞
    elif query.data.startswith("region_"):
        region = query.data.replace("region_", "")
        region_names = {
            "moscow": "–ú–æ—Å–∫–≤–∞ –∏ –ú–û",
            "spb": "–°–∞–Ω–∫—Ç-–ü–µ—Ç–µ—Ä–±—É—Ä–≥",
            "sochi": "–°–æ—á–∏",
            "yakutsk": "–Ø–∫—É—Ç—Å–∫",
            "kamchatka": "–ö–∞–º—á–∞—Ç–∫–∞"
        }

        region_name = region_names.get(region, "–†–µ–≥–∏–æ–Ω")
        # –í –±—É–¥—É—â–µ–º –∑–¥–µ—Å—å –±—É–¥–µ—Ç —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ —Ä–µ–≥–∏–æ–Ω–∞ –≤ –ë–î
        await query.edit_message_text(
            f"‚úÖ **–†–µ–≥–∏–æ–Ω –≤—ã–±—Ä–∞–Ω: {region_name}**\n\n"
            "–¢–µ–ø–µ—Ä—å —è –±—É–¥—É —É—á–∏—Ç—ã–≤–∞—Ç—å —Ä–µ–≥–∏–æ–Ω–∞–ª—å–Ω—ã–µ –æ—Å–æ–±–µ–Ω–Ω–æ—Å—Ç–∏ –≤ —Å–≤–æ–∏—Ö –æ—Ç–≤–µ—Ç–∞—Ö:\n"
            "‚Ä¢ –ö–ª–∏–º–∞—Ç–∏—á–µ—Å–∫–∞—è –∑–æ–Ω–∞\n"
            "‚Ä¢ –ö–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç—ã —Å—Ç–æ–∏–º–æ—Å—Ç–∏ —Ä–∞–±–æ—Ç\n"
            "‚Ä¢ –°–µ–π—Å–º–∏—á–Ω–æ—Å—Ç—å\n"
            "‚Ä¢ –¢–µ–º–ø–µ—Ä–∞—Ç—É—Ä–Ω—ã–µ —É—Å–ª–æ–≤–∏—è\n\n"
            "üí° _–§—É–Ω–∫—Ü–∏—è —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è —Ä–µ–≥–∏–æ–Ω–∞ –±—É–¥–µ—Ç –¥–æ—Å—Ç—É–ø–Ω–∞ –≤ v3.1_",
            parse_mode='Markdown'
        )

    # === –û–ë–†–ê–ë–û–¢–ß–ò–ö–ò v3.2 ===

    # –û–±—Ä–∞–±–æ—Ç—á–∏–∫–∏ —à–∞–±–ª–æ–Ω–æ–≤ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤
    elif query.data.startswith("template_"):
        if TEMPLATES_AVAILABLE:
            template_id = query.data.replace("template_", "")
            await handle_template_selection(update, context, template_id)
        else:
            await query.edit_message_text("‚ö†Ô∏è –ú–æ–¥—É–ª—å —à–∞–±–ª–æ–Ω–æ–≤ –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω.")

    # –û–±—Ä–∞–±–æ—Ç—á–∏–∫–∏ —Å–∫–∞—á–∏–≤–∞–Ω–∏—è –ø—É—Å—Ç–æ–≥–æ —à–∞–±–ª–æ–Ω–∞
    elif query.data.startswith("download_empty_"):
        if TEMPLATES_AVAILABLE:
            template_id = query.data.replace("download_empty_", "")
            await handle_download_empty_template(update, context, template_id)
        else:
            await query.edit_message_text("‚ö†Ô∏è –ú–æ–¥—É–ª—å —à–∞–±–ª–æ–Ω–æ–≤ –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω.")

    # –û–±—Ä–∞–±–æ—Ç—á–∏–∫–∏ –≤—ã–±–æ—Ä–∞ —Ä–æ–ª–∏
    elif query.data.startswith("role_"):
        if ROLES_AVAILABLE:
            role_id = query.data.replace("role_", "")
            await handle_role_selection(update, context, role_id)
        else:
            await query.edit_message_text("‚ö†Ô∏è –ú–æ–¥—É–ª—å —Ä–æ–ª–µ–π –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω.")

    # –û–±—Ä–∞–±–æ—Ç—á–∏–∫–∏ –≤—ã–±–æ—Ä–∞ –ø—Ä–æ–µ–∫—Ç–∞
    elif query.data.startswith("setproj_"):
        if PROJECTS_AVAILABLE:
            project_name = query.data.replace("setproj_", "")
            user_id = update.effective_user.id
            context.user_data["current_project"] = project_name
            project = load_project(user_id, project_name)
            await query.edit_message_text(
                f"‚úÖ –ê–∫—Ç–∏–≤–Ω—ã–π –ø—Ä–æ–µ–∫—Ç —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω: **{project_name}**\n\n"
                f"{project.get_log_summary()}\n\n"
                "üìå –í—Å–µ –≤–∞—à–∏ –≤–æ–ø—Ä–æ—Å—ã –∏ –æ—Ç–≤–µ—Ç—ã —Ç–µ–ø–µ—Ä—å —Å–æ—Ö—Ä–∞–Ω—è—é—Ç—Å—è –≤ —ç—Ç–æ—Ç –ø—Ä–æ–µ–∫—Ç.",
                parse_mode="Markdown"
            )
        else:
            await query.edit_message_text("‚ö†Ô∏è –ú–æ–¥—É–ª—å –ø—Ä–æ–µ–∫—Ç–æ–≤ –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω.")
    elif query.data == "proj_create":
        # –°–æ–∑–¥–∞–Ω–∏–µ –Ω–æ–≤–æ–≥–æ –ø—Ä–æ–µ–∫—Ç–∞ —á–µ—Ä–µ–∑ –¥–∏–∞–ª–æ–≥
        if PROJECTS_AVAILABLE:
            await query.edit_message_text(
                "üìù **–°–û–ó–î–ê–ù–ò–ï –ü–†–û–ï–ö–¢–ê**\n\n"
                "–í–≤–µ–¥–∏—Ç–µ –Ω–∞–∑–≤–∞–Ω–∏–µ –ø—Ä–æ–µ–∫—Ç–∞:\n\n"
                "_–ù–∞–ø—Ä–∏–º–µ—Ä: –†–µ–∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏—è –¢–¶ –ú–µ–≥–∞, –°—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤–æ –∂–∏–ª–æ–≥–æ –¥–æ–º–∞, –†–µ–º–æ–Ω—Ç –º–æ—Å—Ç–∞_",
                parse_mode="Markdown"
            )
            context.user_data["waiting_for_project_name"] = True
        else:
            await query.edit_message_text("‚ö†Ô∏è –ú–æ–¥—É–ª—å –ø—Ä–æ–µ–∫—Ç–æ–≤ –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω.")
    elif query.data.startswith("proj_open_"):
        # –û—Ç–∫—Ä—ã—Ç–∏–µ —Å—É—â–µ—Å—Ç–≤—É—é—â–µ–≥–æ –ø—Ä–æ–µ–∫—Ç–∞
        if PROJECTS_AVAILABLE:
            project_name = query.data.replace("proj_open_", "")
            user_id = update.effective_user.id
            project = load_project(user_id, project_name)

            if project:
                context.user_data["current_project"] = project_name

                # –°–æ–∑–¥–∞—ë–º –º–µ–Ω—é –ø—Ä–æ–µ–∫—Ç–∞
                keyboard = [
                    [InlineKeyboardButton("üìä –ò–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è", callback_data=f"proj_info_{project_name}"),
                     InlineKeyboardButton("üìã –ñ—É—Ä–Ω–∞–ª", callback_data=f"proj_log_{project_name}")],
                    [InlineKeyboardButton("üìÅ –§–∞–π–ª—ã", callback_data=f"proj_files_{project_name}"),
                     InlineKeyboardButton("üìù –î–æ–±–∞–≤–∏—Ç—å –∑–∞–º–µ—Ç–∫—É", callback_data=f"proj_note_{project_name}")],
                    [InlineKeyboardButton("üì¶ –≠–∫—Å–ø–æ—Ä—Ç –ø—Ä–æ–µ–∫—Ç–∞", callback_data=f"proj_export_{project_name}")],
                    [InlineKeyboardButton("¬´ –ù–∞–∑–∞–¥ –∫ –ø—Ä–æ–µ–∫—Ç–∞–º", callback_data="project_menu")]
                ]
                reply_markup = InlineKeyboardMarkup(keyboard)

                info_text = f"üìÅ **–ü–†–û–ï–ö–¢: {project_name}**\n\n"
                info_text += f"‚úÖ –ü—Ä–æ–µ–∫—Ç –∞–∫—Ç–∏–≤–∏—Ä–æ–≤–∞–Ω\n\n"
                info_text += f"{project.get_log_summary()}\n\n"
                info_text += "–ß—Ç–æ –≤—ã —Ö–æ—Ç–∏—Ç–µ —Å–¥–µ–ª–∞—Ç—å?"

                await query.edit_message_text(
                    info_text,
                    reply_markup=reply_markup,
                    parse_mode="Markdown"
                )
            else:
                await query.edit_message_text("‚ùå –û—à–∏–±–∫–∞ –∑–∞–≥—Ä—É–∑–∫–∏ –ø—Ä–æ–µ–∫—Ç–∞")
        else:
            await query.edit_message_text("‚ö†Ô∏è –ú–æ–¥—É–ª—å –ø—Ä–æ–µ–∫—Ç–æ–≤ –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω.")
    elif query.data.startswith("proj_info_"):
        # –ò–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è –æ –ø—Ä–æ–µ–∫—Ç–µ
        if PROJECTS_AVAILABLE:
            project_name = query.data.replace("proj_info_", "")
            user_id = update.effective_user.id
            project = load_project(user_id, project_name)

            if project:
                info = project.get_project_summary()
                log_summary = project.get_log_summary()

                keyboard = [[InlineKeyboardButton("¬´ –ù–∞–∑–∞–¥", callback_data=f"proj_open_{project_name}")]]
                reply_markup = InlineKeyboardMarkup(keyboard)

                await query.edit_message_text(
                    f"{info}\n\nüìä **–ñ—É—Ä–Ω–∞–ª —Ä–∞–±–æ—Ç—ã:** {log_summary}",
                    reply_markup=reply_markup,
                    parse_mode="Markdown"
                )
            else:
                await query.edit_message_text("‚ùå –û—à–∏–±–∫–∞ –∑–∞–≥—Ä—É–∑–∫–∏ –ø—Ä–æ–µ–∫—Ç–∞")
        else:
            await query.edit_message_text("‚ö†Ô∏è –ú–æ–¥—É–ª—å –ø—Ä–æ–µ–∫—Ç–æ–≤ –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω.")
    elif query.data.startswith("proj_log_"):
        # –ñ—É—Ä–Ω–∞–ª –ø—Ä–æ–µ–∫—Ç–∞
        if PROJECTS_AVAILABLE:
            project_name = query.data.replace("proj_log_", "")
            user_id = update.effective_user.id
            project = load_project(user_id, project_name)

            if project:
                log = project.get_conversation_log()

                keyboard = []

                if not log:
                    keyboard.append([InlineKeyboardButton("¬´ –ù–∞–∑–∞–¥", callback_data=f"proj_open_{project_name}")])
                    reply_markup = InlineKeyboardMarkup(keyboard)
                    await query.edit_message_text(
                        "üìã –ñ—É—Ä–Ω–∞–ª —Ä–∞–±–æ—Ç—ã –ø—É—Å—Ç",
                        reply_markup=reply_markup
                    )
                else:
                    # –ü–æ—Å–ª–µ–¥–Ω–∏–µ 10 –∑–∞–ø–∏—Å–µ–π
                    recent_log = log[-10:]
                    response = f"üìã **–ñ–£–†–ù–ê–õ: {project_name}**\n\n"
                    response += f"–í—Å–µ–≥–æ –∑–∞–ø–∏—Å–µ–π: {len(log)}\n"
                    response += f"–ü–æ—Å–ª–µ–¥–Ω–∏–µ {len(recent_log)} –∑–∞–ø–∏—Å–µ–π:\n\n"

                    # –°–æ–∑–¥–∞—ë–º –∫–Ω–æ–ø–∫–∏ –¥–ª—è –∫–∞–∂–¥–æ–π –∑–∞–ø–∏—Å–∏
                    start_index = len(log) - len(recent_log)
                    for i, entry in enumerate(recent_log, start=start_index):
                        timestamp = entry["timestamp"][:16].replace("T", " ")
                        question = entry.get("question", "")[:40]
                        response += f"{i+1}. {timestamp}\n   Q: {question}...\n\n"

                        # –î–æ–±–∞–≤–ª—è–µ–º –∫–Ω–æ–ø–∫—É –¥–ª—è –ø—Ä–æ—Å–º–æ—Ç—Ä–∞ —ç—Ç–æ–π –∑–∞–ø–∏—Å–∏
                        keyboard.append([
                            InlineKeyboardButton(
                                f"üìñ –ó–∞–ø–∏—Å—å #{i+1}",
                                callback_data=f"proj_entry_{project_name}_{i}"
                            )
                        ])

                    keyboard.append([InlineKeyboardButton("¬´ –ù–∞–∑–∞–¥", callback_data=f"proj_open_{project_name}")])
                    reply_markup = InlineKeyboardMarkup(keyboard)

                    await query.edit_message_text(
                        response,
                        reply_markup=reply_markup,
                        parse_mode="Markdown"
                    )
            else:
                await query.edit_message_text("‚ùå –û—à–∏–±–∫–∞ –∑–∞–≥—Ä—É–∑–∫–∏ –ø—Ä–æ–µ–∫—Ç–∞")
        else:
            await query.edit_message_text("‚ö†Ô∏è –ú–æ–¥—É–ª—å –ø—Ä–æ–µ–∫—Ç–æ–≤ –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω.")
    elif query.data.startswith("proj_entry_"):
        # –ü—Ä–æ—Å–º–æ—Ç—Ä –∫–æ–Ω–∫—Ä–µ—Ç–Ω–æ–π –∑–∞–ø–∏—Å–∏ –∏–∑ –∂—É—Ä–Ω–∞–ª–∞
        if PROJECTS_AVAILABLE:
            parts = query.data.replace("proj_entry_", "").rsplit("_", 1)
            project_name = parts[0]
            entry_index = int(parts[1])
            user_id = update.effective_user.id
            project = load_project(user_id, project_name)

            if project:
                log = project.get_conversation_log()
                if 0 <= entry_index < len(log):
                    entry = log[entry_index]
                    timestamp = entry["timestamp"][:19].replace("T", " ")
                    question = entry.get("question", "–ù–µ—Ç –≤–æ–ø—Ä–æ—Å–∞")
                    answer = entry.get("answer", "–ù–µ—Ç –æ—Ç–≤–µ—Ç–∞")
                    entry_type = entry.get("type", "qa")

                    response = f"üìñ **–ó–ê–ü–ò–°–¨ #{entry_index+1}**\n\n"
                    response += f"‚è∞ {timestamp}\n"
                    response += f"üìÇ –¢–∏–ø: {entry_type}\n\n"
                    response += f"**‚ùì –í–æ–ø—Ä–æ—Å:**\n{question}\n\n"
                    response += f"**üí¨ –û—Ç–≤–µ—Ç:**\n{answer[:3000]}"  # –û–≥—Ä–∞–Ω–∏—á–µ–Ω–∏–µ –Ω–∞ –¥–ª–∏–Ω—É

                    if len(answer) > 3000:
                        response += "\n\n_...–æ—Ç–≤–µ—Ç –æ–±—Ä–µ–∑–∞–Ω, –ø–æ–ª–Ω—ã–π —Ç–µ–∫—Å—Ç —Å–æ—Ö—Ä–∞–Ω—ë–Ω –≤ –ø—Ä–æ–µ–∫—Ç–µ_"

                    keyboard = [
                        [InlineKeyboardButton("üîÑ –ü–æ–≤—Ç–æ—Ä–Ω–æ –æ—Ç–ø—Ä–∞–≤–∏—Ç—å", callback_data=f"proj_resend_{project_name}_{entry_index}")],
                        [InlineKeyboardButton("¬´ –ö –∂—É—Ä–Ω–∞–ª—É", callback_data=f"proj_log_{project_name}")]
                    ]
                    reply_markup = InlineKeyboardMarkup(keyboard)

                    await query.edit_message_text(
                        response,
                        reply_markup=reply_markup,
                        parse_mode="Markdown"
                    )
                else:
                    await query.edit_message_text("‚ùå –ó–∞–ø–∏—Å—å –Ω–µ –Ω–∞–π–¥–µ–Ω–∞")
            else:
                await query.edit_message_text("‚ùå –û—à–∏–±–∫–∞ –∑–∞–≥—Ä—É–∑–∫–∏ –ø—Ä–æ–µ–∫—Ç–∞")
        else:
            await query.edit_message_text("‚ö†Ô∏è –ú–æ–¥—É–ª—å –ø—Ä–æ–µ–∫—Ç–æ–≤ –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω.")
    elif query.data.startswith("proj_resend_"):
        # –ü–æ–≤—Ç–æ—Ä–Ω–∞—è –æ—Ç–ø—Ä–∞–≤–∫–∞ –∑–∞–ø–∏—Å–∏ –∏–∑ –∂—É—Ä–Ω–∞–ª–∞
        if PROJECTS_AVAILABLE:
            parts = query.data.replace("proj_resend_", "").rsplit("_", 1)
            project_name = parts[0]
            entry_index = int(parts[1])
            user_id = update.effective_user.id
            project = load_project(user_id, project_name)

            if project:
                log = project.get_conversation_log()
                if 0 <= entry_index < len(log):
                    entry = log[entry_index]
                    question = entry.get("question", "–ù–µ—Ç –≤–æ–ø—Ä–æ—Å–∞")
                    answer = entry.get("answer", "–ù–µ—Ç –æ—Ç–≤–µ—Ç–∞")

                    # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é –∫–∞–∫ –Ω–æ–≤–æ–µ —Å–æ–æ–±—â–µ–Ω–∏–µ
                    await query.answer("üì® –û—Ç–ø—Ä–∞–≤–ª—è—é –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é...")

                    full_message = f"üìÇ **–ò–∑ –ø—Ä–æ–µ–∫—Ç–∞:** {project_name}\n\n"
                    full_message += f"**‚ùì –í–æ–ø—Ä–æ—Å:**\n{question}\n\n"
                    full_message += f"**üí¨ –û—Ç–≤–µ—Ç:**\n{answer}"

                    # –†–∞–∑–±–∏–≤–∞–µ–º –Ω–∞ —á–∞—Å—Ç–∏ –µ—Å–ª–∏ —Å–ª–∏—à–∫–æ–º –¥–ª–∏–Ω–Ω–æ–µ
                    max_length = 4000
                    if len(full_message) > max_length:
                        # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –ø–æ —á–∞—Å—Ç—è–º
                        await query.message.reply_text(
                            f"üìÇ **–ò–∑ –ø—Ä–æ–µ–∫—Ç–∞:** {project_name}\n\n**‚ùì –í–æ–ø—Ä–æ—Å:**\n{question}",
                            parse_mode="Markdown"
                        )

                        # –î–µ–ª–∏–º –æ—Ç–≤–µ—Ç –Ω–∞ —á–∞—Å—Ç–∏
                        chunks = [answer[i:i+max_length] for i in range(0, len(answer), max_length)]
                        for i, chunk in enumerate(chunks):
                            prefix = f"**üí¨ –û—Ç–≤–µ—Ç (—á–∞—Å—Ç—å {i+1}/{len(chunks)}):**\n" if len(chunks) > 1 else "**üí¨ –û—Ç–≤–µ—Ç:**\n"
                            await query.message.reply_text(
                                prefix + chunk,
                                parse_mode="Markdown"
                            )
                    else:
                        await query.message.reply_text(full_message, parse_mode="Markdown")

                    # –í–æ–∑–≤—Ä–∞—â–∞–µ–º—Å—è –∫ –∑–∞–ø–∏—Å–∏
                    keyboard = [
                        [InlineKeyboardButton("¬´ –ö –∂—É—Ä–Ω–∞–ª—É", callback_data=f"proj_log_{project_name}")]
                    ]
                    reply_markup = InlineKeyboardMarkup(keyboard)
                    await query.edit_message_reply_markup(reply_markup=reply_markup)
                else:
                    await query.answer("‚ùå –ó–∞–ø–∏—Å—å –Ω–µ –Ω–∞–π–¥–µ–Ω–∞", show_alert=True)
            else:
                await query.answer("‚ùå –û—à–∏–±–∫–∞ –∑–∞–≥—Ä—É–∑–∫–∏ –ø—Ä–æ–µ–∫—Ç–∞", show_alert=True)
        else:
            await query.answer("‚ö†Ô∏è –ú–æ–¥—É–ª—å –ø—Ä–æ–µ–∫—Ç–æ–≤ –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω", show_alert=True)
    elif query.data.startswith("proj_files_"):
        # –§–∞–π–ª—ã –ø—Ä–æ–µ–∫—Ç–∞
        if PROJECTS_AVAILABLE:
            project_name = query.data.replace("proj_files_", "")
            user_id = update.effective_user.id
            project = load_project(user_id, project_name)

            if project:
                files = project.list_files()

                keyboard = [[InlineKeyboardButton("¬´ –ù–∞–∑–∞–¥", callback_data=f"proj_open_{project_name}")]]
                reply_markup = InlineKeyboardMarkup(keyboard)

                if not files:
                    await query.edit_message_text(
                        "üìÅ –í –ø—Ä–æ–µ–∫—Ç–µ –ø–æ–∫–∞ –Ω–µ—Ç —Ñ–∞–π–ª–æ–≤\n\n"
                        "–û—Ç–ø—Ä–∞–≤—å—Ç–µ —Ñ–∞–π–ª —Å –ø–æ–¥–ø–∏—Å—å—é –¥–ª—è –¥–æ–±–∞–≤–ª–µ–Ω–∏—è –≤ –ø—Ä–æ–µ–∫—Ç",
                        reply_markup=reply_markup
                    )
                else:
                    response = f"üìÅ **–§–ê–ô–õ–´ –ü–†–û–ï–ö–¢–ê: {project_name}**\n\n"
                    response += f"–í—Å–µ–≥–æ —Ñ–∞–π–ª–æ–≤: {len(files)}\n\n"

                    for i, file_info in enumerate(files, 1):
                        name = file_info["original_name"]
                        size_mb = file_info["size_bytes"] / 1024 / 1024
                        file_type = file_info["type"]
                        response += f"{i}. {name}\n"
                        response += f"   –¢–∏–ø: {file_type} | –†–∞–∑–º–µ—Ä: {size_mb:.2f} –ú–ë\n\n"

                    await query.edit_message_text(
                        response,
                        reply_markup=reply_markup,
                        parse_mode="Markdown"
                    )
            else:
                await query.edit_message_text("‚ùå –û—à–∏–±–∫–∞ –∑–∞–≥—Ä—É–∑–∫–∏ –ø—Ä–æ–µ–∫—Ç–∞")
        else:
            await query.edit_message_text("‚ö†Ô∏è –ú–æ–¥—É–ª—å –ø—Ä–æ–µ–∫—Ç–æ–≤ –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω.")
    elif query.data.startswith("proj_note_"):
        # –î–æ–±–∞–≤–ª–µ–Ω–∏–µ –∑–∞–º–µ—Ç–∫–∏
        if PROJECTS_AVAILABLE:
            project_name = query.data.replace("proj_note_", "")
            await query.edit_message_text(
                f"üìù **–î–û–ë–ê–í–õ–ï–ù–ò–ï –ó–ê–ú–ï–¢–ö–ò**\n\n"
                f"–ü—Ä–æ–µ–∫—Ç: {project_name}\n\n"
                "–í–≤–µ–¥–∏—Ç–µ —Ç–µ–∫—Å—Ç –∑–∞–º–µ—Ç–∫–∏:",
                parse_mode="Markdown"
            )
            context.user_data["waiting_for_note"] = project_name
        else:
            await query.edit_message_text("‚ö†Ô∏è –ú–æ–¥—É–ª—å –ø—Ä–æ–µ–∫—Ç–æ–≤ –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω.")
    elif query.data.startswith("proj_export_"):
        # –≠–∫—Å–ø–æ—Ä—Ç –ø—Ä–æ–µ–∫—Ç–∞
        if PROJECTS_AVAILABLE:
            project_name = query.data.replace("proj_export_", "")
            user_id = update.effective_user.id

            await query.edit_message_text("‚è≥ –ü–æ–¥–≥–æ—Ç–∞–≤–ª–∏–≤–∞—é —ç–∫—Å–ø–æ—Ä—Ç –ø—Ä–æ–µ–∫—Ç–∞...")

            try:
                project = load_project(user_id, project_name)
                if project:
                    # –≠–∫—Å–ø–æ—Ä—Ç–∏—Ä—É–µ–º –ø—Ä–æ–µ–∫—Ç –≤ JSON
                    export_data = {
                        "project_name": project_name,
                        "metadata": project.metadata,
                        "exported_at": datetime.now().isoformat()
                    }

                    # –°–æ–∑–¥–∞—ë–º —Ç–µ–∫—Å—Ç–æ–≤—ã–π —Ñ–∞–π–ª –¥–ª—è —ç–∫—Å–ø–æ—Ä—Ç–∞
                    export_text = f"üìÅ –ü–†–û–ï–ö–¢: {project_name}\n"
                    export_text += f"–≠–∫—Å–ø–æ—Ä—Ç–∏—Ä–æ–≤–∞–Ω: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n\n"
                    export_text += "=" * 50 + "\n\n"

                    # –ò–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è –æ –ø—Ä–æ–µ–∫—Ç–µ
                    export_text += f"{project.get_project_summary()}\n\n"
                    export_text += "=" * 50 + "\n\n"

                    # –ñ—É—Ä–Ω–∞–ª —Ä–∞–±–æ—Ç—ã
                    log = project.get_conversation_log()
                    export_text += f"üìã –ñ–£–†–ù–ê–õ –†–ê–ë–û–¢–´ ({len(log)} –∑–∞–ø–∏—Å–µ–π)\n\n"

                    for i, entry in enumerate(log, 1):
                        timestamp = entry["timestamp"][:16].replace("T", " ")
                        question = entry.get("question", "")
                        answer = entry.get("answer", "")
                        export_text += f"‚ïê‚ïê‚ïê –ó–∞–ø–∏—Å—å #{i} ‚ïê‚ïê‚ïê\n"
                        export_text += f"‚è∞ {timestamp}\n\n"
                        export_text += f"‚ùì –í–û–ü–†–û–°:\n{question}\n\n"
                        export_text += f"üí¨ –û–¢–í–ï–¢:\n{answer}\n\n"
                        export_text += "-" * 50 + "\n\n"

                    # –°–ø–∏—Å–æ–∫ —Ñ–∞–π–ª–æ–≤
                    files = project.list_files()
                    if files:
                        export_text += "=" * 50 + "\n\n"
                        export_text += f"üìÅ –§–ê–ô–õ–´ –ü–†–û–ï–ö–¢–ê ({len(files)})\n\n"
                        for file_info in files:
                            export_text += f"‚Ä¢ {file_info['original_name']}\n"
                            export_text += f"  –¢–∏–ø: {file_info['type']}\n"
                            export_text += f"  –†–∞–∑–º–µ—Ä: {file_info['size_bytes'] / 1024 / 1024:.2f} –ú–ë\n"
                            export_text += f"  –î–æ–±–∞–≤–ª–µ–Ω: {file_info['added_at'][:16]}\n\n"

                    # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º —Ñ–∞–π–ª
                    from io import BytesIO
                    buffer = BytesIO(export_text.encode('utf-8'))
                    buffer.seek(0)
                    filename = f"{project_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.txt"

                    await query.message.reply_document(
                        document=buffer,
                        filename=filename,
                        caption=f"üì¶ –≠–∫—Å–ø–æ—Ä—Ç –ø—Ä–æ–µ–∫—Ç–∞: **{project_name}**\n\n"
                                f"–ó–∞–ø–∏—Å–µ–π: {len(log)} | –§–∞–π–ª–æ–≤: {len(files)}",
                        parse_mode="Markdown"
                    )

                    keyboard = [[InlineKeyboardButton("¬´ –ù–∞–∑–∞–¥", callback_data=f"proj_open_{project_name}")]]
                    reply_markup = InlineKeyboardMarkup(keyboard)

                    await query.edit_message_text(
                        "‚úÖ –ü—Ä–æ–µ–∫—Ç —ç–∫—Å–ø–æ—Ä—Ç–∏—Ä–æ–≤–∞–Ω!",
                        reply_markup=reply_markup
                    )
                else:
                    await query.edit_message_text("‚ùå –û—à–∏–±–∫–∞ –∑–∞–≥—Ä—É–∑–∫–∏ –ø—Ä–æ–µ–∫—Ç–∞")
            except Exception as e:
                logger.error(f"–û—à–∏–±–∫–∞ —ç–∫—Å–ø–æ—Ä—Ç–∞ –ø—Ä–æ–µ–∫—Ç–∞: {e}")
                await query.edit_message_text(f"‚ùå –û—à–∏–±–∫–∞ —ç–∫—Å–ø–æ—Ä—Ç–∞: {str(e)}")
        else:
            await query.edit_message_text("‚ö†Ô∏è –ú–æ–¥—É–ª—å –ø—Ä–æ–µ–∫—Ç–æ–≤ –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω.")
    elif query.data == "ignore":
        # –ò–≥–Ω–æ—Ä–∏—Ä—É–µ–º —Ä–∞–∑–¥–µ–ª–∏—Ç–µ–ª—å
        await query.answer()


async def error_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–û–±—Ä–∞–±–æ—Ç–∫–∞ –æ—à–∏–±–æ–∫"""
    logger.error(f"Update {update} caused error {context.error}")


# === –ù–û–í–´–ï –ö–û–ú–ê–ù–î–´ v3.0 ===

async def calculators_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /calculators - –ü–æ–∫–∞–∑–∞—Ç—å –º–µ–Ω—é –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä–æ–≤"""
    if not CALCULATORS_AVAILABLE:
        await update.message.reply_text(
            "‚ö†Ô∏è –ú–æ–¥—É–ª—å –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä–æ–≤ –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω.\n\n–û–±—Ä–∞—Ç–∏—Ç–µ—Å—å –∫ –∞–¥–º–∏–Ω–∏—Å—Ç—Ä–∞—Ç–æ—Ä—É."
        )
        return

    calc_text = """üßÆ **–ö–ê–õ–¨–ö–£–õ–Ø–¢–û–†–´ v5.0**

–í—Å–µ –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä—ã —Å –ø–æ—à–∞–≥–æ–≤—ã–º –≤–≤–æ–¥–æ–º –¥–∞–Ω–Ω—ã—Ö!

**–û—Å–Ω–æ–≤–Ω—ã–µ –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏:**
üèóÔ∏è –ë–µ—Ç–æ–Ω - –æ–±—ä—ë–º, –º–∞—Ä–∫–∞, –º–∞—Ç–µ—Ä–∏–∞–ª—ã
üîß –ê—Ä–º–∞—Ç—É—Ä–∞ - –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ, –≤–µ—Å
üì¶ –û–ø–∞–ª—É–±–∫–∞ - –ø–ª–æ—â–∞–¥—å, –æ–±–æ—Ä–∞—á–∏–≤–∞–µ–º–æ—Å—Ç—å
‚öì –§—É–Ω–¥–∞–º–µ–Ω—Ç - –∫–æ–º–ø–ª–µ–∫—Å–Ω—ã–π —Ä–∞—Å—á–µ—Ç

**–ú–∞—Ç–µ—Ä–∏–∞–ª—ã –æ—Ç–¥–µ–ª–∫–∏:**
üß± –ö–∏—Ä–ø–∏—á/–±–ª–æ–∫–∏ - —Å —É—á–µ—Ç–æ–º –ø—Ä–æ—ë–º–æ–≤
üî≤ –ü–ª–∏—Ç–∫–∞ - –ø–ª–∏—Ç–∫–∞ + –∫–ª–µ–π + –∑–∞—Ç–∏—Ä–∫–∞
üß± –®—Ç—É–∫–∞—Ç—É—Ä–∫–∞ - —Ä–∞—Å—Ö–æ–¥ —Å–º–µ—Å–∏
üé® –ö—Ä–∞—Å–∫–∞ - —Ä–∞—Å—Ö–æ–¥ –ø–æ —Å–ª–æ—è–º
üßä –£—Ç–µ–ø–ª–∏—Ç–µ–ª—å - —Ç–æ–ª—â–∏–Ω–∞ + –ø–ª–æ—â–∞–¥—å
üè† –ö—Ä–æ–≤–ª—è - –ø–ª–æ—â–∞–¥—å —Å–∫–∞—Ç–æ–≤

**–ö–æ–º–º—É–Ω–∏–∫–∞—Ü–∏–∏:**
‚ö° –≠–ª–µ–∫—Ç—Ä–æ—Å–Ω–∞–±–∂–µ–Ω–∏–µ - –º–æ—â–Ω–æ—Å—Ç—å
üíß –í–æ–¥–æ—Å–Ω–∞–±–∂–µ–Ω–∏–µ - —Ä–∞—Å—Ö–æ–¥ –≤–æ–¥—ã
‚ùÑÔ∏è –ó–∏–º–Ω–∏–π –ø—Ä–æ–≥—Ä–µ–≤ - –ø—Ä–æ–≥—Ä–µ–≤ –±–µ—Ç–æ–Ω–∞

–í—ã–±–µ—Ä–∏—Ç–µ –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä –∏–∑ –º–µ–Ω—é –Ω–∏–∂–µ:"""

    if IMPROVEMENTS_V3_AVAILABLE:
        keyboard = create_calculators_menu()
        await update.message.reply_text(calc_text, parse_mode='Markdown', reply_markup=keyboard)
    else:
        await update.message.reply_text(calc_text, parse_mode='Markdown')


async def region_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /region - –í—ã–±–æ—Ä —Ä–µ–≥–∏–æ–Ω–∞"""
    if not IMPROVEMENTS_V3_AVAILABLE:
        await update.message.reply_text(
            "‚ö†Ô∏è –§—É–Ω–∫—Ü–∏—è —Ä–µ–≥–∏–æ–Ω–∞–ª—å–Ω—ã—Ö –Ω–∞—Å—Ç—Ä–æ–µ–∫ –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∞."
        )
        return

    region_text = """üìç **–†–ï–ì–ò–û–ù–ê–õ–¨–ù–´–ï –ù–ê–°–¢–†–û–ô–ö–ò**

–í—ã–±–µ—Ä–∏—Ç–µ –≤–∞—à —Ä–µ–≥–∏–æ–Ω –¥–ª—è —É—á—ë—Ç–∞ –∫–ª–∏–º–∞—Ç–∏—á–µ—Å–∫–∏—Ö –æ—Å–æ–±–µ–Ω–Ω–æ—Å—Ç–µ–π:

‚Ä¢ –ö–ª–∏–º–∞—Ç–∏—á–µ—Å–∫–∏–µ –∫–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç—ã
‚Ä¢ –¢–µ–º–ø–µ—Ä–∞—Ç—É—Ä–Ω—ã–µ –∑–æ–Ω—ã
‚Ä¢ –°–µ–π—Å–º–∏—á–Ω–æ—Å—Ç—å
‚Ä¢ –°—Ç–æ–∏–º–æ—Å—Ç—å —Ä–∞–±–æ—Ç

–í—ã–±–µ—Ä–∏—Ç–µ —Ä–µ–≥–∏–æ–Ω –∏–∑ —Å–ø–∏—Å–∫–∞:"""

    keyboard = create_region_selection_menu()
    await update.message.reply_text(region_text, parse_mode='Markdown', reply_markup=keyboard)


async def generate_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """–ö–æ–º–∞–Ω–¥–∞ /generate - –ì–µ–Ω–µ—Ä–∞—Ü–∏—è –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π —á–µ—Ä–µ–∑ Gemini AI"""
    if not GEMINI_AVAILABLE:
        await update.message.reply_text(
            "‚ö†Ô∏è –§—É–Ω–∫—Ü–∏—è –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∞.\n\n"
            "–ü—Ä–æ–≤–µ—Ä—å—Ç–µ –Ω–∞–ª–∏—á–∏–µ GEMINI_API_KEY –≤ –ø–µ—Ä–µ–º–µ–Ω–Ω—ã—Ö –æ–∫—Ä—É–∂–µ–Ω–∏—è."
        )
        return

    # –ü—Ä–æ–≤–µ—Ä—è–µ–º –∞—Ä–≥—É–º–µ–Ω—Ç—ã –∫–æ–º–∞–Ω–¥—ã
    if not context.args:
        await update.message.reply_text(
            "üé® **–ì–ï–ù–ï–†–ê–¶–ò–Ø –ò–ó–û–ë–†–ê–ñ–ï–ù–ò–ô - Gemini AI**\n\n"
            "–ò—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏–µ:\n"
            "`/generate –æ–ø–∏—Å–∞–Ω–∏–µ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è`\n\n"
            "**–ü—Ä–∏–º–µ—Ä—ã:**\n"
            "‚Ä¢ `/generate —Å—Ö–µ–º–∞ —Ñ—É–Ω–¥–∞–º–µ–Ω—Ç–∞ —Å –∞—Ä–º–∏—Ä–æ–≤–∞–Ω–∏–µ–º`\n"
            "‚Ä¢ `/generate —É–∑–µ–ª —Å–æ–µ–¥–∏–Ω–µ–Ω–∏—è –±–∞–ª–∫–∏ –∏ –∫–æ–ª–æ–Ω–Ω—ã`\n"
            "‚Ä¢ `/generate –∞—Ä–º–∞—Ç—É—Ä–Ω—ã–π –∫–∞—Ä–∫–∞—Å –ø–µ—Ä–µ–∫—Ä—ã—Ç–∏—è`\n"
            "‚Ä¢ `/generate —ç–ª–µ–∫—Ç—Ä–∏—á–µ—Å–∫–∞—è —Å—Ö–µ–º–∞ –∫–≤–∞—Ä—Ç–∏—Ä—ã`\n"
            "‚Ä¢ `/generate —Å—Ö–µ–º–∞ –≤–æ–¥–æ–ø—Ä–æ–≤–æ–¥–∞ –≤ –¥–æ–º–µ`\n\n"
            "–¢–∞–∫–∂–µ –º–æ–∂–Ω–æ –ø—Ä–æ—Å—Ç–æ –Ω–∞–ø–∏—Å–∞—Ç—å:\n"
            "\"–Ω–∞—Ä–∏—Å—É–π —Ç—Ä–µ—â–∏–Ω—É –≤ —Å—Ç–µ–Ω–µ\" - –±–æ—Ç –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏ —Å–≥–µ–Ω–µ—Ä–∏—Ä—É–µ—Ç",
            parse_mode="Markdown"
        )
        return

    # –ü–æ–ª—É—á–∞–µ–º –æ–ø–∏—Å–∞–Ω–∏–µ
    user_request = " ".join(context.args)

    # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º —Å–æ–æ–±—â–µ–Ω–∏–µ –æ –ø—Ä–æ—Ü–µ—Å—Å–µ
    generating_message = await update.message.reply_text(
        "üé® –ì–µ–Ω–µ—Ä–∏—Ä—É—é –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ...\n"
        "–≠—Ç–æ –∑–∞–π–º–µ—Ç 15-30 —Å–µ–∫—É–Ω–¥\n\n"
        "üí° –ò—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è Gemini AI"
    )

    try:
        # –ì–µ–Ω–µ—Ä–∏—Ä—É–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ —á–µ—Ä–µ–∑ Gemini
        result = await generate_construction_image_gemini(user_request)

        if result and result.get("image_data"):
            # –£–¥–∞–ª—è–µ–º —Å–æ–æ–±—â–µ–Ω–∏–µ –æ –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏
            try:
                await generating_message.delete()
            except:
                pass

            # –§–æ—Ä–º–∏—Ä—É–µ–º –ø–æ–¥–ø–∏—Å—å
            caption = f"""üé® **–°–ì–ï–ù–ï–†–ò–†–û–í–ê–ù–ù–û–ï –ò–ó–û–ë–†–ê–ñ–ï–ù–ò–ï**

**–ó–∞–ø—Ä–æ—Å:** {user_request}

{result.get('text', '')[:500] if result.get('text') else ''}

---
ü§ñ _–ú–æ–¥–µ–ª—å: {result.get('model', 'Gemini AI')}_"""

            # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ
            result["image_data"].seek(0)
            await update.message.reply_photo(
                photo=result["image_data"],
                caption=caption[:1024],  # –û–≥—Ä–∞–Ω–∏—á–µ–Ω–∏–µ Telegram
                parse_mode="Markdown"
            )

            logger.info(f"‚úÖ –ò–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ –æ—Ç–ø—Ä–∞–≤–ª–µ–Ω–æ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—é {update.effective_user.id}")
        else:
            await generating_message.edit_text(
                "‚ùå –ù–µ —É–¥–∞–ª–æ—Å—å —Å–æ–∑–¥–∞—Ç—å –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ.\n\n"
                "–í–æ–∑–º–æ–∂–Ω—ã–µ –ø—Ä–∏—á–∏–Ω—ã:\n"
                "‚Ä¢ –ü—Ä–æ–±–ª–µ–º—ã —Å Gemini API\n"
                "‚Ä¢ –ó–∞–ø—Ä–æ—Å –Ω–µ –ø–æ–¥—Ö–æ–¥–∏—Ç –¥–ª—è –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏\n"
                "‚Ä¢ –í—Ä–µ–º–µ–Ω–Ω–∞—è –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–æ—Å—Ç—å —Å–µ—Ä–≤–∏—Å–∞\n\n"
                "–ü–æ–ø—Ä–æ–±—É–π—Ç–µ –∏–∑–º–µ–Ω–∏—Ç—å –æ–ø–∏—Å–∞–Ω–∏–µ –∏–ª–∏ –ø–æ–≤—Ç–æ—Ä–∏—Ç–µ –ø–æ–ø—ã—Ç–∫—É –ø–æ–∑–∂–µ."
            )

    except Exception as e:
        logger.error(f"–û—à–∏–±–∫–∞ –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è: {e}")
        await generating_message.edit_text(
            f"‚ùå –û—à–∏–±–∫–∞ –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏:\n`{str(e)}`\n\n"
            "–ü—Ä–æ–≤–µ—Ä—å—Ç–µ –Ω–∞–ª–∏—á–∏–µ GEMINI_API_KEY –≤ –ø–µ—Ä–µ–º–µ–Ω–Ω—ã—Ö –æ–∫—Ä—É–∂–µ–Ω–∏—è.",
            parse_mode="Markdown"
        )


# === –ì–õ–ê–í–ù–ê–Ø –§–£–ù–ö–¶–ò–Ø ===

async def setup_bot_menu(application):
    """–£—Å—Ç–∞–Ω–æ–≤–∫–∞ –ø–æ—Å—Ç–æ—è–Ω–Ω–æ–≥–æ –º–µ–Ω—é –∫–æ–º–∞–Ω–¥ –±–æ—Ç–∞"""
    commands = [
        BotCommand("start", "üè† –ì–ª–∞–≤–Ω–æ–µ –º–µ–Ω—é"),
        BotCommand("help", "üìñ –°–ø—Ä–∞–≤–∫–∞ –ø–æ –≤—Å–µ–º –∫–æ–º–∞–Ω–¥–∞–º"),
        # BotCommand("generate", "üé® –ì–µ–Ω–µ—Ä–∞—Ü–∏—è —Å—Ö–µ–º (Gemini AI)"),  # –û—Ç–∫–ª—é—á–µ–Ω–æ
        # BotCommand("visualize", "üé® –í–∏–∑—É–∞–ª–∏–∑–∞—Ü–∏—è –¥–µ—Ñ–µ–∫—Ç–æ–≤ (Gemini AI)"),  # –û—Ç–∫–ª—é—á–µ–Ω–æ
        BotCommand("calculators", "üßÆ –ö–∞–ª—å–∫—É–ª—è—Ç–æ—Ä—ã"),
        BotCommand("regulations", "üìö –ù–æ—Ä–º–∞—Ç–∏–≤—ã (27 –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤)"),
        BotCommand("regulations_menu", "üìñ –ö–∞—Ç–µ–≥–æ—Ä–∏–∏ –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤"),
        BotCommand("faq", "‚ùì –ß–∞—Å—Ç—ã–µ –≤–æ–ø—Ä–æ—Å—ã"),
        BotCommand("templates", "üìÑ –®–∞–±–ª–æ–Ω—ã –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤"),
        BotCommand("projects", "üìÅ –ú–æ–∏ –ø—Ä–æ–µ–∫—Ç—ã"),
        BotCommand("role", "üëî –í—ã–±—Ä–∞—Ç—å —Ä–æ–ª—å"),
        BotCommand("history", "üìú –ò—Å—Ç–æ—Ä–∏—è –¥–∏–∞–ª–æ–≥–æ–≤"),
        BotCommand("stats", "üìä –°—Ç–∞—Ç–∏—Å—Ç–∏–∫–∞"),
        BotCommand("hse", "ü¶∫ –û—Ö—Ä–∞–Ω–∞ —Ç—Ä—É–¥–∞"),
        BotCommand("technology", "üèóÔ∏è –¢–µ—Ö–Ω–æ–ª–æ–≥–∏—è —Å—Ç—Ä–æ–∏—Ç–µ–ª—å—Å—Ç–≤–∞"),
        BotCommand("estimating", "üí∞ –°–º–µ—Ç–Ω–æ–µ –¥–µ–ª–æ"),
        BotCommand("legal", "‚öñÔ∏è –Æ—Ä–∏–¥–∏—á–µ—Å–∫–∏–µ –≤–æ–ø—Ä–æ—Å—ã"),
        BotCommand("management", "üìà –£–ø—Ä–∞–≤–ª–µ–Ω–∏–µ –ø—Ä–æ–µ–∫—Ç–∞–º–∏"),
        BotCommand("suggestions", "üí° –ü—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è –ø–æ —É–ª—É—á—à–µ–Ω–∏—é"),
        BotCommand("dev", "üîß –†–µ–∂–∏–º —Ä–∞–∑—Ä–∞–±–æ—Ç—á–∏–∫–∞"),
    ]

    await application.bot.set_my_commands(commands)
    logger.info("‚úÖ –ú–µ–Ω—é –∫–æ–º–∞–Ω–¥ –±–æ—Ç–∞ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω–æ")


def main():
    """–ó–∞–ø—É—Å–∫ –±–æ—Ç–∞"""
    import asyncio

    # –°–æ–∑–¥–∞–µ–º event loop –¥–ª—è Python 3.14+
    try:
        loop = asyncio.get_event_loop()
    except RuntimeError:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)

    # –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä—É–µ–º PostgreSQL –±–∞–∑—É –¥–∞–Ω–Ω—ã—Ö
    if DATABASE_AVAILABLE:
        try:
            loop.run_until_complete(init_db())
        except Exception as e:
            logger.error(f"–û—à–∏–±–∫–∞ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏–∏ PostgreSQL: {e}")
            logger.info("–ü—Ä–æ–¥–æ–ª–∂–∞–µ–º —Ä–∞–±–æ—Ç—É —Å JSON —Ö—Ä–∞–Ω–∏–ª–∏—â–µ–º")

    logger.info("‚úÖ –ë–æ—Ç –°—Ç—Ä–æ–π–ù–∞–¥–∑–æ—ÄAI –∑–∞–ø—É—â–µ–Ω —É—Å–ø–µ—à–Ω–æ!")

    # –°–æ–∑–¥–∞–µ–º –ø—Ä–∏–ª–æ–∂–µ–Ω–∏–µ
    application = Application.builder().token(TELEGRAM_TOKEN).build()

    # –†–µ–≥–∏—Å—Ç—Ä–∏—Ä—É–µ–º –æ–±—Ä–∞–±–æ—Ç—á–∏–∫–∏ –∫–æ–º–∞–Ω–¥
    application.add_handler(CommandHandler("start", start_command))
    application.add_handler(CommandHandler("help", help_command))
    application.add_handler(CommandHandler("regulations", regulations_command))
    application.add_handler(CommandHandler("examples", examples_command))
    application.add_handler(CommandHandler("history", history_command))
    application.add_handler(CommandHandler("stats", stats_command))
    application.add_handler(CommandHandler("clear", clear_command))
    # –ù–æ–≤—ã–µ –∫–æ–º–∞–Ω–¥—ã v2.1
    application.add_handler(CommandHandler("export", export_command))
    application.add_handler(CommandHandler("search", search_command))
    application.add_handler(CommandHandler("recommendations", recommendations_command))
    application.add_handler(CommandHandler("updates", updates_command))
    # –ö–æ–º–∞–Ω–¥—ã –¥–ª—è –∞–∫—Ç—É–∞–ª—å–Ω—ã—Ö —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π 2025
    application.add_handler(CommandHandler("requirements2025", requirements2025_command))
    application.add_handler(CommandHandler("laws", laws_command))
    application.add_handler(CommandHandler("checklist", checklist_command))
    # –ö–æ–º–∞–Ω–¥—ã –¥–ª—è –ø—Ä–∞–∫—Ç–∏—á–µ—Å–∫–∏—Ö –∑–Ω–∞–Ω–∏–π 2025
    application.add_handler(CommandHandler("hse", hse_command))
    application.add_handler(CommandHandler("technology", technology_command))
    application.add_handler(CommandHandler("estimating", estimating_command))
    application.add_handler(CommandHandler("legal", legal_command))
    application.add_handler(CommandHandler("management", management_command))

    # –ù–æ–≤—ã–µ –∫–æ–º–∞–Ω–¥—ã v3.0
    application.add_handler(CommandHandler("calculators", calculators_command))
    application.add_handler(CommandHandler("region", region_command))

    # LLM Council - –°–æ–≤–µ—Ç AI –º–æ–¥–µ–ª–µ–π (Karpathy's approach)
    if LLM_COUNCIL_AVAILABLE:
        application.add_handler(CommandHandler("council", council_command))
        logger.info("‚úÖ –ö–æ–º–∞–Ω–¥–∞ /council –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω–∞")

    # === –ò–ù–¢–ï–†–ê–ö–¢–ò–í–ù–´–ï –ö–ê–õ–¨–ö–£–õ–Ø–¢–û–†–´ v4.0 ===
    # –ò–Ω—Ç–µ—Ä–∞–∫—Ç–∏–≤–Ω—ã–µ –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä—ã –≤—ã–∑—ã–≤–∞—é—Ç—Å—è —á–µ—Ä–µ–∑ –º–µ–Ω—é /calculators –∏ –∫–æ–º–∞–Ω–¥—ã
    if CALCULATOR_HANDLERS_AVAILABLE:
        # –û—Å–Ω–æ–≤–Ω—ã–µ 7 –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä–æ–≤
        application.add_handler(create_concrete_calculator_handler())
        application.add_handler(create_rebar_calculator_handler())
        application.add_handler(create_formwork_calculator_handler())
        application.add_handler(create_electrical_calculator_handler())
        application.add_handler(create_water_calculator_handler())
        application.add_handler(create_winter_calculator_handler())
        application.add_handler(create_math_calculator_handler())

        # –ù–æ–≤—ã–µ 14 –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä–æ–≤
        application.add_handler(create_brick_calculator_handler())
        application.add_handler(create_tile_calculator_handler())
        application.add_handler(create_paint_calculator_handler())
        application.add_handler(create_wall_area_calculator_handler())
        application.add_handler(create_roof_calculator_handler())
        application.add_handler(create_plaster_calculator_handler())
        application.add_handler(create_wallpaper_calculator_handler())
        application.add_handler(create_laminate_calculator_handler())
        application.add_handler(create_insulation_calculator_handler())
        application.add_handler(create_foundation_calculator_handler())
        application.add_handler(create_stairs_calculator_handler())
        application.add_handler(create_drywall_calculator_handler())
        application.add_handler(create_earthwork_calculator_handler())
        application.add_handler(create_labor_calculator_handler())

        # –ë—ã—Å—Ç—Ä—ã–µ –∫–æ–º–∞–Ω–¥—ã –¥–ª—è —Ä–∞—Å—á—ë—Ç–∞ –æ–¥–Ω–æ–π —Å—Ç—Ä–æ–∫–æ–π
        application.add_handler(CommandHandler("calc_concrete", quick_concrete))
        application.add_handler(CommandHandler("calc_math", quick_math))

        logger.info("‚úÖ –ò–Ω—Ç–µ—Ä–∞–∫—Ç–∏–≤–Ω—ã–µ –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä—ã v4.0 –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω—ã (–≤—Å–µ 21 + –±—ã—Å—Ç—Ä—ã–µ –∫–æ–º–∞–Ω–¥—ã)")

    # –£–î–ê–õ–ï–ù–û: –¥—É–±–ª–∏—Ä–æ–≤–∞–Ω–∏–µ –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä–æ–≤ —Å–æ–∑–¥–∞–≤–∞–ª–æ –∫–æ–Ω—Ñ–ª–∏–∫—Ç ConversationHandler
    # –°—Ç–∞—Ä–∞—è —Å–∏—Å—Ç–µ–º–∞ interactive_calculators.py –±–æ–ª—å—à–µ –Ω–µ –∏—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è

    # === –ö–ê–¢–ï–ì–û–†–ò–ó–ê–¶–ò–Ø –ù–û–†–ú–ê–¢–ò–í–û–í v1.0 ===
    if REGULATIONS_CATEGORIES_AVAILABLE:
        async def regulations_menu_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
            """–ö–æ–º–∞–Ω–¥–∞ /regulations_menu - –∫–∞—Ç–µ–≥–æ—Ä–∏–∏ –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤"""
            keyboard = get_categories_keyboard()
            await update.message.reply_text(
                "üìö **–ö–ê–¢–ï–ì–û–†–ò–ò –ù–û–†–ú–ê–¢–ò–í–û–í**\n\n"
                "–í—ã–±–µ—Ä–∏—Ç–µ –∫–∞—Ç–µ–≥–æ—Ä–∏—é –¥–ª—è –ø—Ä–æ—Å–º–æ—Ç—Ä–∞ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤:\n\n"
                "_27 –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ —Ä–∞–∑–¥–µ–ª–µ–Ω—ã –Ω–∞ 7 —É–¥–æ–±–Ω—ã—Ö –∫–∞—Ç–µ–≥–æ—Ä–∏–π_",
                reply_markup=keyboard,
                parse_mode='Markdown'
            )

        application.add_handler(CommandHandler("regulations_menu", regulations_menu_command))

        # Callback handler –¥–ª—è –∫–∞—Ç–µ–≥–æ—Ä–∏–π –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤
        async def handle_regulations_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
            """–û–±—Ä–∞–±–æ—Ç—á–∏–∫ callback –¥–ª—è –∫–∞—Ç–µ–≥–æ—Ä–∏–π –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤"""
            query = update.callback_query
            await query.answer()

            callback_data = query.data

            if callback_data.startswith("cat_"):
                category_id = callback_data.replace("cat_", "")

                if category_id == "all":
                    # –ü–æ–∫–∞–∑–∞—Ç—å –≤—Å–µ –Ω–æ—Ä–º–∞—Ç–∏–≤—ã
                    text = get_all_regulations_text()
                    await query.edit_message_text(text, parse_mode='Markdown')
                else:
                    # –ü–æ–∫–∞–∑–∞—Ç—å –Ω–æ—Ä–º–∞—Ç–∏–≤—ã –∫–∞—Ç–µ–≥–æ—Ä–∏–∏
                    text = get_regulations_by_category(category_id)
                    if text:
                        # –ö–Ω–æ–ø–∫–∞ "–ù–∞–∑–∞–¥"
                        keyboard = [[InlineKeyboardButton("‚¨ÖÔ∏è –ù–∞–∑–∞–¥ –∫ –∫–∞—Ç–µ–≥–æ—Ä–∏—è–º", callback_data="cat_back")]]
                        reply_markup = InlineKeyboardMarkup(keyboard)
                        await query.edit_message_text(text, reply_markup=reply_markup, parse_mode='Markdown')

            elif callback_data == "cat_back":
                # –í–µ—Ä–Ω—É—Ç—å—Å—è –∫ –∫–∞—Ç–µ–≥–æ—Ä–∏—è–º
                keyboard = get_categories_keyboard()
                await query.edit_message_text(
                    "üìö **–ö–ê–¢–ï–ì–û–†–ò–ò –ù–û–†–ú–ê–¢–ò–í–û–í**\n\n"
                    "–í—ã–±–µ—Ä–∏—Ç–µ –∫–∞—Ç–µ–≥–æ—Ä–∏—é –¥–ª—è –ø—Ä–æ—Å–º–æ—Ç—Ä–∞ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤:\n\n"
                    "_27 –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ —Ä–∞–∑–¥–µ–ª–µ–Ω—ã –Ω–∞ 7 —É–¥–æ–±–Ω—ã—Ö –∫–∞—Ç–µ–≥–æ—Ä–∏–π_",
                    reply_markup=keyboard,
                    parse_mode='Markdown'
                )

        application.add_handler(CallbackQueryHandler(handle_regulations_callback, pattern="^cat_"))
        logger.info("‚úÖ –ö–æ–º–∞–Ω–¥–∞ /regulations_menu –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω–∞ (–∫–∞—Ç–µ–≥–æ—Ä–∏–∏ –Ω–æ—Ä–º–∞—Ç–∏–≤–æ–≤)")

    # === –ì–ï–ù–ï–†–ê–¶–ò–Ø –ò–ó–û–ë–†–ê–ñ–ï–ù–ò–ô - Gemini AI ===
    if GEMINI_AVAILABLE:
        application.add_handler(CommandHandler("generate", generate_command))
        logger.info("‚úÖ –ö–æ–º–∞–Ω–¥–∞ /generate –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω–∞ (Gemini AI)")

    # === –í–ò–ó–£–ê–õ–ò–ó–ê–¶–ò–Ø –î–ï–§–ï–ö–¢–û–í - Gemini AI ===
    if GEMINI_AVAILABLE:
        application.add_handler(CommandHandler("visualize", visualize_command))
        logger.info("‚úÖ –ö–æ–º–∞–Ω–¥–∞ /visualize –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω–∞ (Gemini AI)")

    # === –ù–û–í–´–ï –ö–û–ú–ê–ù–î–´ v3.9 ===
    if TEMPLATES_AVAILABLE:
        application.add_handler(CommandHandler("templates", templates_command))
        logger.info("‚úÖ –ö–æ–º–∞–Ω–¥–∞ /templates –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω–∞")

    if PROJECTS_AVAILABLE:
        application.add_handler(CommandHandler("projects", projects_command))
        application.add_handler(CommandHandler("new_project", new_project_command))
        application.add_handler(CommandHandler("set_project", set_project_command))
        application.add_handler(CommandHandler("project_info", project_info_command))
        application.add_handler(CommandHandler("project_log", project_log_command))
        logger.info("‚úÖ –ö–æ–º–∞–Ω–¥—ã —É–ø—Ä–∞–≤–ª–µ–Ω–∏—è –ø—Ä–æ–µ–∫—Ç–∞–º–∏ –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω—ã (5 –∫–æ–º–∞–Ω–¥)")

    if ROLES_AVAILABLE:
        application.add_handler(CommandHandler("role", role_command))
        logger.info("‚úÖ –ö–æ–º–∞–Ω–¥–∞ /role –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω–∞")

    # === –ü–†–ï–î–õ–û–ñ–ï–ù–ò–Ø –ü–û –£–õ–£–ß–®–ï–ù–ò–Æ v1.0 ===
    if SUGGESTIONS_AVAILABLE:
        # –î–æ–±–∞–≤–ª—è–µ–º –∫–æ–º–∞–Ω–¥—É /suggestions –¥–ª—è –±—ã—Å—Ç—Ä–æ–≥–æ –¥–æ—Å—Ç—É–ø–∞
        async def suggestions_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
            """–ö–æ–º–∞–Ω–¥–∞ /suggestions - –æ—Ç–∫—Ä—ã–≤–∞–µ—Ç –º–µ–Ω—é –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π"""
            await suggestions_menu(update, context)

        application.add_handler(CommandHandler("suggestions", suggestions_command))
        logger.info("‚úÖ –ö–æ–º–∞–Ω–¥–∞ /suggestions –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω–∞")

    # === –£–ü–†–ê–í–õ–ï–ù–ò–ï –ò–°–¢–û–†–ò–ï–ô v3.5 ===
    if HISTORY_MANAGER_AVAILABLE:
        application.add_handler(CommandHandler("history", history_command))
        application.add_handler(CommandHandler("stats", stats_command))
        application.add_handler(CommandHandler("search", search_command))
        application.add_handler(CommandHandler("export", export_command))
        application.add_handler(CommandHandler("clear", clear_history_command))
        # –û–±—Ä–∞–±–æ—Ç—á–∏–∫ callback –¥–ª—è —Ä–∞–±–æ—Ç—ã —Å –∏—Å—Ç–æ—Ä–∏–µ–π
        application.add_handler(CallbackQueryHandler(handle_history_callback, pattern="^hist_|^export_|^clear_"))
        logger.info("‚úÖ –£–ø—Ä–∞–≤–ª–µ–Ω–∏–µ –∏—Å—Ç–æ—Ä–∏–µ–π v3.5 –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω–æ (5 –∫–æ–º–∞–Ω–¥)")

    # === –°–û–•–†–ê–ù–Å–ù–ù–´–ï –†–ê–°–ß–Å–¢–´ v3.6 ===
    if SAVED_CALCS_AVAILABLE:
        application.add_handler(CommandHandler("saved", saved_command))
        # –û–±—Ä–∞–±–æ—Ç—á–∏–∫ callback –¥–ª—è —Å–æ—Ö—Ä–∞–Ω—ë–Ω–Ω—ã—Ö —Ä–∞—Å—á—ë—Ç–æ–≤
        application.add_handler(CallbackQueryHandler(handle_saved_callback, pattern="^saved_"))
        logger.info("‚úÖ –°–æ—Ö—Ä–∞–Ω—ë–Ω–Ω—ã–µ —Ä–∞—Å—á—ë—Ç—ã v3.6 –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω—ã")

    # === –ë–ê–ó–ê –ß–ê–°–¢–´–• –í–û–ü–†–û–°–û–í (FAQ) v3.7 ===
    if FAQ_AVAILABLE:
        application.add_handler(CommandHandler("faq", faq_command))
        application.add_handler(CommandHandler("faq_search", faq_search_command))
        # –û–±—Ä–∞–±–æ—Ç—á–∏–∫ callback –¥–ª—è FAQ
        application.add_handler(CallbackQueryHandler(handle_faq_callback, pattern="^faq_"))
        logger.info(f"‚úÖ FAQ v3.7 –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω ({get_total_faq_count()} –≤–æ–ø—Ä–æ—Å–æ–≤)")

    # === –ü–õ–ê–ù–ò–†–û–í–©–ò–ö –†–ê–ë–û–¢ v3.8 ===
    if PLANNER_AVAILABLE:
        application.add_handler(CommandHandler("planner", planner_command))
        application.add_handler(CommandHandler("plan_calc", plan_calc_command))
        # –û–±—Ä–∞–±–æ—Ç—á–∏–∫ callback –¥–ª—è –ø–ª–∞–Ω–∏—Ä–æ–≤—â–∏–∫–∞
        application.add_handler(CallbackQueryHandler(handle_planner_callback, pattern="^plan_"))
        logger.info("‚úÖ Work planner v3.8 –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω")

    # –£–î–ê–õ–ï–ù–û: –¥—É–±–ª–∏—Ä—É—é—â–∞—è—Å—è —Ä–µ–≥–∏—Å—Ç—Ä–∞—Ü–∏—è –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä–æ–≤ v3.5
    # –í—Å–µ 21 –∫–∞–ª—å–∫—É–ª—è—Ç–æ—Ä —É–∂–µ –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω—ã –≤—ã—à–µ –≤ –±–ª–æ–∫–µ v4.0

    # === –ò–ù–¢–ï–†–ê–ö–¢–ò–í–ù–´–ï –û–ë–†–ê–ë–û–¢–ß–ò–ö–ò –î–û–ö–£–ú–ï–ù–¢–û–í v1.0 ===
    if DOCUMENT_HANDLERS_AVAILABLE:
        # ConversationHandler –¥–ª—è –≤—Å–µ—Ö 4 —à–∞–±–ª–æ–Ω–æ–≤ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤
        application.add_handler(create_acceptance_foundation_handler())
        application.add_handler(create_complaint_contractor_handler())
        application.add_handler(create_safety_plan_handler())
        application.add_handler(create_hidden_works_act_handler())
        logger.info("‚úÖ –í—Å–µ 4 –∏–Ω—Ç–µ—Ä–∞–∫—Ç–∏–≤–Ω—ã—Ö –æ–±—Ä–∞–±–æ—Ç—á–∏–∫–∞ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω—ã (v1.0)")

    # === –†–ï–ñ–ò–ú –†–ê–ó–†–ê–ë–û–¢–ß–ò–ö–ê v1.0 ===
    if DEV_MODE_AVAILABLE:
        # ConversationHandler –¥–ª—è —Ä–µ–∂–∏–º–∞ —Ä–∞–∑—Ä–∞–±–æ—Ç—á–∏–∫–∞
        application.add_handler(create_dev_mode_handler())
        logger.info("‚úÖ –†–µ–∂–∏–º —Ä–∞–∑—Ä–∞–±–æ—Ç—á–∏–∫–∞ v1.0 –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω")

    # –†–µ–≥–∏—Å—Ç—Ä–∏—Ä—É–µ–º –æ–±—Ä–∞–±–æ—Ç—á–∏–∫–∏ —Å–æ–æ–±—â–µ–Ω–∏–π
    application.add_handler(MessageHandler(filters.PHOTO, handle_photo))

    # === –ì–û–õ–û–°–û–í–´–ï –°–û–û–ë–©–ï–ù–ò–Ø v3.9 ===
    if VOICE_HANDLER_AVAILABLE:
        application.add_handler(MessageHandler(filters.VOICE, handle_voice))
        logger.info("‚úÖ –û–±—Ä–∞–±–æ—Ç—á–∏–∫ –≥–æ–ª–æ—Å–æ–≤—ã—Ö —Å–æ–æ–±—â–µ–Ω–∏–π –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω")

    # === –ó–ê–ì–†–£–ó–ö–ê –î–û–ö–£–ú–ï–ù–¢–û–í –í –ü–†–û–ï–ö–¢–´ v3.9 ===
    if PROJECTS_AVAILABLE:
        application.add_handler(MessageHandler(filters.Document.ALL, handle_document))
        logger.info("‚úÖ –û–±—Ä–∞–±–æ—Ç—á–∏–∫ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω")

    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))

    # === –ê–í–¢–û–ü–†–ò–ú–ï–ù–ï–ù–ò–ï –ò–ó–ú–ï–ù–ï–ù–ò–ô v1.0 ===
    if AUTO_APPLY_AVAILABLE:
        application.add_handler(CallbackQueryHandler(handle_apply_changes, pattern="^apply_changes"))
        logger.info("‚úÖ –û–±—Ä–∞–±–æ—Ç—á–∏–∫ –∞–≤—Ç–æ–ø—Ä–∏–º–µ–Ω–µ–Ω–∏—è –∏–∑–º–µ–Ω–µ–Ω–∏–π –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω")

    # === –ü–†–ï–î–õ–û–ñ–ï–ù–ò–Ø –ü–û –£–õ–£–ß–®–ï–ù–ò–Æ v1.0 ===
    if SUGGESTIONS_AVAILABLE:
        application.add_handler(create_suggestions_handler())
        logger.info("‚úÖ –û–±—Ä–∞–±–æ—Ç—á–∏–∫ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω")

    # === GEMINI LIVE API - –ì–û–õ–û–°–û–í–û–ô –ê–°–°–ò–°–¢–ï–ù–¢ ===
    if VOICE_ASSISTANT_AVAILABLE:
        try:
            from gemini_live_bot_integration import (
                register_voice_assistant_handlers,
                init_voice_assistant
            )
            # –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è –≥–æ–ª–æ—Å–æ–≤–æ–≥–æ –∞—Å—Å–∏—Å—Ç–µ–Ω—Ç–∞
            init_voice_assistant()
            # –†–µ–≥–∏—Å—Ç—Ä–∞—Ü–∏—è –æ–±—Ä–∞–±–æ—Ç—á–∏–∫–æ–≤ (ConversationHandler –¥–ª—è –≥–æ–ª–æ—Å–æ–≤—ã—Ö —Å–µ—Å—Å–∏–π)
            register_voice_assistant_handlers(application)
            logger.info("‚úÖ Gemini Live API (–≥–æ–ª–æ—Å–æ–≤–æ–π –∞—Å—Å–∏—Å—Ç–µ–Ω—Ç) –∞–∫—Ç–∏–≤–∏—Ä–æ–≤–∞–Ω")
        except Exception as e:
            logger.error(f"‚ùå –û—à–∏–±–∫–∞ –∞–∫—Ç–∏–≤–∞—Ü–∏–∏ –≥–æ–ª–æ—Å–æ–≤–æ–≥–æ –∞—Å—Å–∏—Å—Ç–µ–Ω—Ç–∞: {e}")

    # === OPENAI REALTIME API - –ê–õ–¨–¢–ï–†–ù–ê–¢–ò–í–ù–´–ô –ì–û–õ–û–°–û–í–û–ô –ê–°–°–ò–°–¢–ï–ù–¢ ===
    if OPENAI_REALTIME_AVAILABLE:
        try:
            from openai_realtime_bot_integration import (
                register_realtime_assistant_handlers,
                init_realtime_assistant
            )
            # –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è OpenAI Realtime –∞—Å—Å–∏—Å—Ç–µ–Ω—Ç–∞
            init_realtime_assistant()
            # –†–µ–≥–∏—Å—Ç—Ä–∞—Ü–∏—è –æ–±—Ä–∞–±–æ—Ç—á–∏–∫–æ–≤
            register_realtime_assistant_handlers(application)
            logger.info("‚úÖ OpenAI Realtime API (–≥–æ–ª–æ—Å–æ–≤–æ–π –∞—Å—Å–∏—Å—Ç–µ–Ω—Ç) –∞–∫—Ç–∏–≤–∏—Ä–æ–≤–∞–Ω")
        except Exception as e:
            logger.error(f"‚ùå –û—à–∏–±–∫–∞ –∞–∫—Ç–∏–≤–∞—Ü–∏–∏ OpenAI Realtime: {e}")

    # –†–µ–≥–∏—Å—Ç—Ä–∏—Ä—É–µ–º –æ–±—Ä–∞–±–æ—Ç—á–∏–∫ –∫–Ω–æ–ø–æ–∫
    application.add_handler(CallbackQueryHandler(handle_callback))

    # –†–µ–≥–∏—Å—Ç—Ä–∏—Ä—É–µ–º –æ–±—Ä–∞–±–æ—Ç—á–∏–∫ –æ—à–∏–±–æ–∫
    application.add_error_handler(error_handler)

    # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –º–µ–Ω—é –∫–æ–º–∞–Ω–¥ –±–æ—Ç–∞
    loop.run_until_complete(setup_bot_menu(application))

    # –ó–∞–ø—É—Å–∫–∞–µ–º –±–æ—Ç–∞
    logger.info("Bot is running... Press Ctrl+C to stop")
    application.run_polling(allowed_updates=Update.ALL_TYPES)


if __name__ == "__main__":
    main()