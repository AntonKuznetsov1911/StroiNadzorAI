"""
Модуль для генерации строительных документов по шаблонам
Соответствует ГОСТ, СП и формам КС
"""

import os
import logging
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# Папка для сгенерированных документов
DOCUMENTS_DIR = Path("generated_documents")
DOCUMENTS_DIR.mkdir(exist_ok=True)


# ==========================
# ШАБЛОНЫ ДОКУМЕНТОВ
# ==========================

DOCUMENT_TEMPLATES = {
    "acceptance_foundation": {
        "name": "Акт приёмки фундамента",
        "description": "Акт освидетельствования скрытых работ (фундамент) по форме ОС-3",
        "params": ["act_number", "object_name", "contractor", "customer", "date", "foundation_type", "volume_m3", "concrete_class", "inspector_name", "defects"],
        "params_display": ["Номер акта", "Наименование объекта", "Подрядчик", "Заказчик", "Дата", "Тип фундамента", "Объём бетона (м³)", "Класс бетона", "ФИО инспектора", "Выявленные дефекты"]
    },
    "complaint_contractor": {
        "name": "Претензия подрядчику",
        "description": "Официальная претензия по выявленным дефектам работ",
        "params": ["complaint_number", "date", "contractor_name", "contractor_address", "sender_name", "sender_address", "contract_number", "contract_date", "defect_description", "deadline_days", "penalty_percent"],
        "params_display": ["Номер претензии", "Дата", "Наименование подрядчика", "Адрес подрядчика", "Ваше наименование/ФИО", "Ваш адрес", "Номер договора", "Дата договора", "Описание дефектов", "Срок устранения (дни)", "Размер штрафа (%)"]
    },
    "safety_plan": {
        "name": "План мероприятий по охране труда",
        "description": "План ОТ и ТБ на объекте на 2025 год по СП 12-135-2003",
        "params": ["object_name", "year", "responsible_person", "worker_count", "total_budget"],
        "params_display": ["Наименование объекта", "Год", "ФИО ответственного лица", "Численность работников", "Общий бюджет (руб.)"]
    },
    "hidden_works_act": {
        "name": "Акт освидетельствования скрытых работ",
        "description": "Универсальный акт для скрытых работ по форме КС-3",
        "params": ["act_number", "object_name", "contractor", "customer", "date", "work_type", "volume", "standards", "inspector_name", "project_compliance"],
        "params_display": ["Номер акта", "Наименование объекта", "Подрядчик", "Заказчик", "Дата", "Вид работ", "Объём работ", "Применяемые нормативы", "ФИО инспектора", "Соответствие проекту"]
    },
}


def add_table_border(table):
    """Добавляет границы ко всем ячейкам таблицы"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)

    tblPr.append(tblBorders)


def generate_acceptance_foundation(doc: Document, params: dict):
    """Генерация акта приёмки фундамента по форме ОС-3"""

    # Заголовок
    heading = doc.add_heading('АКТ ПРИЕМКИ ФУНДАМЕНТА', level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Номер и дата
    p = doc.add_paragraph()
    p.add_run(f'№ {params.get("act_number", "______")} от «___» __________ {params.get("date", "2025")} г.').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    # Стороны
    doc.add_paragraph('Мы, нижеподписавшиеся:').runs[0].font.bold = True
    doc.add_paragraph()

    doc.add_paragraph(f'Заказчик: {params.get("customer", "_________________________________________________________")}')
    doc.add_paragraph('(наименование организации, ФИО представителя, должность, реквизиты)')
    doc.add_paragraph()

    doc.add_paragraph(f'Подрядчик: {params.get("contractor", "________________________________________________________")}')
    doc.add_paragraph('(наименование организации, ФИО представителя, должность, реквизиты)')
    doc.add_paragraph()

    # Основной текст
    doc.add_paragraph('составили настоящий акт о том, что выполнены следующие работы:')
    doc.add_paragraph()

    # Информация об объекте
    p = doc.add_paragraph('Объект: ')
    p.add_run(params.get("object_name", "_____________________________________________________"))
    doc.add_paragraph('(адрес, наименование строительного объекта)')
    doc.add_paragraph()

    doc.add_paragraph('Вид работ: Приемка фундамента под здание/сооружение.')
    doc.add_paragraph()

    # Технические параметры
    p = doc.add_paragraph('Технические параметры:')
    p.runs[0].font.bold = True

    doc.add_paragraph(f'• Тип фундамента: {params.get("foundation_type", "ленточный/свайный/плитный")};')
    doc.add_paragraph(f'• Объём бетона: {params.get("volume_m3", "___")} м³;')
    doc.add_paragraph(f'• Класс бетона: {params.get("concrete_class", "В25")};')
    doc.add_paragraph('• Соответствие проектной документации: □ Да / □ Нет;')
    doc.add_paragraph(f'• Выявленные дефекты (при наличии): {params.get("defects", "_____________________________")}')
    doc.add_paragraph()

    # Результат проверки
    p = doc.add_paragraph('Результат проверки:')
    p.runs[0].font.bold = True
    doc.add_paragraph('Фундамент принят/не принят к дальнейшим работам.')
    doc.add_paragraph('Причины отказа (если применимо): _______________________________')
    doc.add_paragraph()

    # Приложения
    p = doc.add_paragraph('Приложения:')
    p.runs[0].font.bold = True
    doc.add_paragraph('• Протоколы испытаний бетона (№______);')
    doc.add_paragraph('• Чертежи с отметками о соответствии (№______).')
    doc.add_paragraph()

    # Нормативные документы
    p = doc.add_paragraph('Нормативные документы:')
    p.runs[0].font.bold = True
    doc.add_paragraph('• СП 48.13330.2024 — «Организация строительства. Актуализированная редакция СНиП 12-01-2004» (порядок ведения исполнительной документации);')
    doc.add_paragraph('• СП 70.13330.2025 — «Несущие и ограждающие конструкции. Актуализированная редакция СНиП 3.03.01-87» (правила приемки бетонных работ);')
    doc.add_paragraph('• СП 50-101-2024 — «Проектирование и устройство оснований и фундаментов зданий и сооружений»;')
    doc.add_paragraph('• СП 22.13330.2025 — «Основания зданий и сооружений. Актуализированная редакция СНиП 2.02.01-83»;')
    doc.add_paragraph('• ГОСТ Р 21.1101-2025 — «Система проектной документации для строительства. Основные требования к рабочей документации»;')
    doc.add_paragraph('• ГОСТ 7473-2023 — «Смеси бетонные. Технические условия»;')
    doc.add_paragraph('• ГОСТ 34329-2024 — «Арматура для железобетонных конструкций»;')
    doc.add_paragraph('• ГОСТ 10180-2023 — «Бетоны. Методы определения прочности по контрольным образцам»;')
    doc.add_paragraph('• Федеральный закон № 384-ФЗ от 30.12.2009 — «Технический регламент о безопасности зданий и сооружений»;')
    doc.add_paragraph('• Приказ Минстроя России № 845/пр от 20.12.2024 — «Об утверждении формы акта освидетельствования скрытых работ».')
    doc.add_paragraph()

    # Примечания о 2025
    p = doc.add_paragraph('Примечание (изменения с 01.01.2025):')
    p.runs[0].font.bold = True
    p.runs[0].font.italic = True
    doc.add_paragraph('• Обязательно ведение электронного журнала работ (п. 5.2 СП 48.13330.2024);')
    doc.add_paragraph('• Обязательное фото-видеофиксирование скрытых работ (п. 5.6 СП 48.13330.2024);')
    doc.add_paragraph('• Цифровые подписи на актах приобретают юридическую силу наравне с «мокрыми» печатями.')
    doc.add_paragraph()

    # Подписи
    p = doc.add_paragraph('Подписи сторон:')
    p.runs[0].font.bold = True
    doc.add_paragraph()

    doc.add_paragraph('Заказчик:')
    doc.add_paragraph('__________________ / __________________')
    doc.add_paragraph()

    doc.add_paragraph('Подрядчик:')
    doc.add_paragraph('__________________ / __________________')
    doc.add_paragraph()

    doc.add_paragraph(f'Технический надзор: {params.get("inspector_name", "__________________")}')
    doc.add_paragraph('__________________ / __________________')


def generate_complaint_contractor(doc: Document, params: dict):
    """Генерация претензии подрядчику"""

    # Заголовок
    p = doc.add_paragraph()
    p.add_run('ПРЕТЕНЗИЯ').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(16)

    doc.add_paragraph()

    # Номер и дата
    p = doc.add_paragraph()
    p.add_run(f'№ {params.get("complaint_number", "______")} от «___» __________ {params.get("date", "2025")} г.').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    doc.add_paragraph()

    # Адресат
    p = doc.add_paragraph('Кому: ')
    p.add_run(f'[{params.get("contractor_name", "Наименование подрядчика")}]').bold = True
    doc.add_paragraph(f'Адрес: [{params.get("contractor_address", "Юридический адрес подрядчика")}]')

    doc.add_paragraph()

    # Отправитель
    p = doc.add_paragraph('От: ')
    p.add_run(f'[{params.get("sender_name", "Ваше наименование/ФИО")}]').bold = True
    doc.add_paragraph(f'Адрес: [{params.get("sender_address", "Ваш юридический адрес")}]')

    doc.add_paragraph()

    # Основание
    contract_num = params.get("contract_number", "____")
    contract_date = params.get("contract_date", "«___» __________ 2025 г.")
    doc.add_paragraph(
        f'На основании договора № {contract_num} от {contract_date} сообщаем о выявленных нарушениях:'
    )

    doc.add_paragraph()

    # Суть претензии
    p = doc.add_paragraph('Суть претензии:')
    p.runs[0].font.bold = True

    defects = params.get("defect_description", "")
    if defects:
        doc.add_paragraph(f'• {defects}')
    else:
        doc.add_paragraph('• Не соблюдены сроки выполнения работ по устройству фундамента (п. ___ договора);')
        doc.add_paragraph('• Выявлены дефекты: трещины в бетоне, отклонение от проектных размеров (см. акт от «___» __________ 2025 г.);')
        doc.add_paragraph('• Отсутствуют документы на материалы (паспорта качества).')

    doc.add_paragraph()

    # Требования
    p = doc.add_paragraph('Требования:')
    p.runs[0].font.bold = True

    deadline = params.get("deadline_days", "___")
    doc.add_paragraph(f'• Устранить недостатки в течение {deadline} рабочих дней;')
    doc.add_paragraph('• Предоставить письменное подтверждение устранения;')

    penalty_percent = params.get("penalty_percent", "___")
    doc.add_paragraph(f'• Оплатить штраф в размере {penalty_percent}% от стоимости работ (п. ___ договора).')

    doc.add_paragraph()

    # Последствия
    p = doc.add_paragraph('Последствия:')
    p.runs[0].font.bold = True
    doc.add_paragraph(
        'В случае невыполнения требований в срок, заказчик вправе расторгнуть договор '
        'и взыскать убытки в судебном порядке.'
    )

    doc.add_paragraph()

    # Приложения
    p = doc.add_paragraph('Приложения:')
    p.runs[0].font.bold = True
    doc.add_paragraph('• Копия акта приемки фундамента;')
    doc.add_paragraph('• Фотоматериалы дефектов.')

    doc.add_paragraph()

    # Нормативные документы
    p = doc.add_paragraph('Правовое обоснование:')
    p.runs[0].font.bold = True
    doc.add_paragraph('• Гражданский кодекс РФ, Глава 37 — «Подряд» (ст. 702 — исполнение договора, ст. 715 — ответственность за качество работ, ст. 723 — устранение недостатков);')
    doc.add_paragraph('• Постановление Пленума Верховного Суда РФ № 14 от 26.03.2025 — «О применении судами законодательства при разрешении споров, возникающих из договоров подряда»;')
    doc.add_paragraph('• Арбитражный процессуальный кодекс РФ (ст. 125, 126) — досудебный порядок урегулирования споров;')
    doc.add_paragraph('• ГОСТ Р 51144-2024 — «Управление качеством. Правила оформления претензий к качеству строительно-монтажных работ»;')
    doc.add_paragraph('• СП 48.13330.2024 (п. 6.4) — ответственность за отступления от проектной документации;')
    doc.add_paragraph('• Приказ Минэкономразвития России № 112 от 15.02.2025 — «Об утверждении порядка претензионной работы в строительстве».')
    doc.add_paragraph()

    # Примечание
    p = doc.add_paragraph('Примечание:')
    p.runs[0].font.bold = True
    p.runs[0].font.italic = True
    doc.add_paragraph('Согласно ст. 125 АПК РФ, претензия должна быть направлена до обращения в суд. Срок ответа на претензию — 30 календарных дней (если иное не установлено договором).')
    doc.add_paragraph()

    # Подпись
    doc.add_paragraph('С уважением,')
    doc.add_paragraph(f'________________ / {params.get("sender_name", "__________________")}')
    doc.add_paragraph('(подпись)             (ФИО, должность)')


def generate_safety_plan(doc: Document, params: dict):
    """Генерация плана мероприятий по охране труда"""

    # Заголовок
    heading = doc.add_heading('ПЛАН МЕРОПРИЯТИЙ ПО ОХРАНЕ ТРУДА', level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subheading = doc.add_paragraph(f'на {params.get("year", "2025")} г.')
    subheading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subheading.runs[0].font.size = Pt(12)

    doc.add_paragraph()

    # Информация об объекте
    doc.add_paragraph(f'Объект: {params.get("object_name", "_________________________")}')
    doc.add_paragraph(f'Ответственное лицо: {params.get("responsible_person", "_________________________")}')
    doc.add_paragraph(f'Численность работников: {params.get("worker_count", "___")} чел.')

    doc.add_paragraph()

    # Таблица мероприятий
    table = doc.add_table(rows=4, cols=5)
    add_table_border(table)
    table.style = 'Table Grid'

    # Заголовки
    headers = ['№ п/п', 'Мероприятие', 'Ответственный', 'Срок', 'Бюджет, руб.']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Мероприятия
    measures = [
        ('1', 'Обучение персонала технике безопасности', 'Иванов А.А.', 'до 15.01.2025', '50 000'),
        ('2', 'Проверка исправности лесов и опалубки', 'Петров С.В.', 'ежемесячно', '20 000'),
        ('3', 'Закупка СИЗ (каски, страховочные пояса)', 'Сидорова Е.М.', 'до 30.11.2025', '150 000'),
    ]

    for i, (num, measure, responsible, deadline, budget) in enumerate(measures, start=1):
        table.rows[i].cells[0].text = num
        table.rows[i].cells[1].text = measure
        table.rows[i].cells[2].text = responsible
        table.rows[i].cells[3].text = deadline
        table.rows[i].cells[4].text = budget
        table.rows[i].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    # Общая сумма
    total_budget = params.get("total_budget", "______")
    p = doc.add_paragraph(f'Общая сумма: {total_budget} руб.')
    p.runs[0].font.bold = True

    doc.add_paragraph()

    # Нормативные документы
    p = doc.add_paragraph('Нормативные документы:')
    p.runs[0].font.bold = True
    doc.add_paragraph('• Трудовой кодекс РФ (ст. 212 — обязанности работодателя по обеспечению безопасных условий труда, ст. 219 — право работника на безопасные условия труда, ст. 223 — санитарно-бытовое обслуживание);')
    doc.add_paragraph('• ГОСТ 12.0.230.01-2025 — «Система стандартов безопасности труда. Системы управления охраной труда (СУОТ). Общие требования»;')
    doc.add_paragraph('• СП 1.13130.2025 — «Системы противопожарной защиты. Эвакуационные пути и выходы»;')
    doc.add_paragraph('• Приказ Минтруда России № 578н от 10.06.2025 — «Об утверждении Правил проведения специальной оценки условий труда (СОУТ)»;')
    doc.add_paragraph('• Приказ Минтруда России № 116н от 28.03.2025 — «Об утверждении Порядка обучения по охране труда и проверки знаний требований охраны труда»;')
    doc.add_paragraph('• Приказ Минтруда России № 154н от 05.05.2025 — «Об утверждении Правил обеспечения работников средствами индивидуальной защиты (СИЗ)»;')
    doc.add_paragraph('• СП 284.1325800.2025 — «Безопасность труда в строительстве. Часть 1. Общие требования (работы на высоте)»;')
    doc.add_paragraph('• Постановление Правительства РФ № 2464 от 24.12.2024 — «О финансовом обеспечении предупредительных мер по сокращению производственного травматизма».')
    doc.add_paragraph()

    # Важные примечания о 2025
    p = doc.add_paragraph('Обязательные требования с 01.01.2025:')
    p.runs[0].font.bold = True
    p.runs[0].font.italic = True
    doc.add_paragraph('• Ежегодная оценка профессиональных рисков (ГОСТ 12.0.230.01-2025, п. 6.3);')
    doc.add_paragraph('• Обучение по охране труда — не реже 1 раза в 3 года для офисных работников, ежегодно для рабочих профессий;')
    doc.add_paragraph('• СОУТ должна быть проведена на всех рабочих местах (за исключением дистанционных);')
    doc.add_paragraph('• Работодатель обязан вести электронный учёт выдачи СИЗ.')
    doc.add_paragraph()

    # Утверждение
    p = doc.add_paragraph('Утверждено:')
    p.runs[0].font.bold = True
    doc.add_paragraph(f'________________ / {params.get("responsible_person", "__________________")}')
    doc.add_paragraph('(руководитель организации)')


def generate_hidden_works_act(doc: Document, params: dict):
    """Генерация универсального акта освидетельствования скрытых работ"""

    # Заголовок
    heading = doc.add_heading('АКТ ОСВИДЕТЕЛЬСТВОВАНИЯ СКРЫТЫХ РАБОТ', level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Номер и дата
    p = doc.add_paragraph()
    p.add_run(f'№ {params.get("act_number", "______")} от «___» __________ {params.get("date", "2025")} г.').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    # Основная информация
    doc.add_paragraph(f'Объект: {params.get("object_name", "_________________________________________________________")}')
    doc.add_paragraph(f'Место проведения работ: __________________________________________')
    doc.add_paragraph(f'Вид работ: {params.get("work_type", "Устройство фундамента (армирование, бетонирование).")}')

    doc.add_paragraph()

    # Описание работ
    p = doc.add_paragraph('Описание работ:')
    p.runs[0].font.bold = True

    doc.add_paragraph(f'• Выполнено армирование монолитного фундамента по чертежу №______;')
    doc.add_paragraph(f'• Уложена бетонная смесь класса {params.get("concrete_class", "В22,5")};')
    standards = params.get("standards", "СП 70.13330.2012")
    doc.add_paragraph(f'• Работы выполнены в соответствии с {standards}.')

    doc.add_paragraph()

    # Проверка
    p = doc.add_paragraph('Проверка:')
    p.runs[0].font.bold = True

    compliance = params.get("project_compliance", "Да")
    if compliance == "Да":
        doc.add_paragraph('☑ Соответствует проекту;')
        doc.add_paragraph('☐ Не соответствует (указать причины): ___________________________')
    else:
        doc.add_paragraph('☐ Соответствует проекту;')
        doc.add_paragraph(f'☑ Не соответствует (указать причины): {compliance}')

    doc.add_paragraph()

    # Решение
    p = doc.add_paragraph('Решение:')
    p.runs[0].font.bold = True
    doc.add_paragraph('Работы приняты/не приняты к закрытию.')

    doc.add_paragraph()

    # Подписи
    p = doc.add_paragraph('Подписи:')
    p.runs[0].font.bold = True
    doc.add_paragraph()

    doc.add_paragraph('Представитель заказчика')
    doc.add_paragraph(f'__________________ / {params.get("customer", "__________________")}')
    doc.add_paragraph()

    doc.add_paragraph('Представитель подрядчика')
    doc.add_paragraph(f'__________________ / {params.get("contractor", "__________________")}')
    doc.add_paragraph()

    doc.add_paragraph('Представитель проектной организации')
    doc.add_paragraph(f'__________________ / {params.get("inspector_name", "__________________")}')

    doc.add_paragraph()

    # Нормативные документы
    p = doc.add_paragraph('Нормативные документы:')
    p.runs[0].font.bold = True
    doc.add_paragraph('• СП 48.13330.2024 (п. 5.5–5.7) — «Организация строительства. Актуализированная редакция СНиП 12-01-2004» (требования к освидетельствованию скрытых работ);')
    doc.add_paragraph('• СП 53-102-2025 — «Общие правила проектирования сборных бетонных и железобетонных конструкций» (арматурные работы);')
    doc.add_paragraph('• СП 70.13330.2025 (п. 8.3–8.5) — «Несущие и ограждающие конструкции» (контроль качества бетонирования);')
    doc.add_paragraph('• ГОСТ Р 21.1101-2025 (п. 4.12) — «Система проектной документации для строительства. Основные требования к рабочей документации» (исполнительная документация);')
    doc.add_paragraph('• ГОСТ 34329-2024 — «Арматура для железобетонных конструкций. Общие технические условия»;')
    doc.add_paragraph('• ГОСТ 10180-2023 — «Бетоны. Методы определения прочности по контрольным образцам»;')
    doc.add_paragraph('• Приказ Ростехнадзора № 350 от 12.09.2024 — «Правила осуществления строительного контроля»;')
    doc.add_paragraph('• РД 11-05-2025 (МЧС России) — «Требования пожарной безопасности к строительным конструкциям».')
    doc.add_paragraph()

    # Важное примечание
    p = doc.add_paragraph('Важно!')
    p.runs[0].font.bold = True
    doc.add_paragraph(
        'Все документы должны быть подписаны уполномоченными лицами с указанием должностей и реквизитов.'
    )
    doc.add_paragraph(
        'Для юридической силы акты часто требуют печати (если организация ее использует).'
    )

    doc.add_paragraph()

    # Дополнительные требования 2025
    p = doc.add_paragraph('Требования 2025 года:')
    p.runs[0].font.bold = True
    p.runs[0].font.italic = True
    doc.add_paragraph('• Обязательная фото-видеофиксация скрытых работ (п. 5.6 СП 48.13330.2024);')
    doc.add_paragraph('• Электронный журнал производства работ обязателен для объектов госзаказа;')
    doc.add_paragraph('• Акт должен быть подписан в течение 24 часов после завершения работ, иначе работы считаются выполненными ненадлежащим образом.')

    doc.add_paragraph()
    doc.add_paragraph(f'Дата актуализации шаблона: 29 ноября 2025 г.')


def generate_document(template_id: str, params: dict) -> dict:
    """
    Генерирует документ по шаблону

    Args:
        template_id: ID шаблона
        params: параметры для заполнения

    Returns:
        dict: {"success": bool, "filepath": str, "error": str}
    """
    try:
        if template_id not in DOCUMENT_TEMPLATES:
            return {
                "success": False,
                "filepath": "",
                "error": "Шаблон не найден"
            }

        doc = Document()

        # Устанавливаем стандартные поля документа
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(0.6)

        # Вызываем соответствующую функцию генерации
        if template_id == "acceptance_foundation":
            generate_acceptance_foundation(doc, params)
        elif template_id == "complaint_contractor":
            generate_complaint_contractor(doc, params)
        elif template_id == "safety_plan":
            generate_safety_plan(doc, params)
        elif template_id == "hidden_works_act":
            generate_hidden_works_act(doc, params)
        else:
            return {
                "success": False,
                "filepath": "",
                "error": "Неизвестный тип шаблона"
            }

        # Сохраняем документ
        filename = f"{template_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = DOCUMENTS_DIR / filename
        doc.save(filepath)

        logger.info(f"✅ Документ создан: {filepath}")
        return {
            "success": True,
            "filepath": str(filepath),
            "error": ""
        }
    except Exception as e:
        logger.error(f"❌ Ошибка генерации документа: {e}")
        return {
            "success": False,
            "filepath": "",
            "error": str(e)
        }
